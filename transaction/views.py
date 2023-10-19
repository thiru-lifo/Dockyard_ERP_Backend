from django.views.decorators.clickjacking import xframe_options_exempt
from urllib import response
from django.shortcuts import render
from functools import partial
from queue import Empty
from unicodedata import name
from rest_framework.response import Response
from rest_framework import status
from rest_framework.views import APIView
#from rest_framework import authentication, AllowAny
#from rest_framework.permissions import Is_Authenticated
from rest_framework.authtoken.views import ObtainAuthToken
from rest_framework.authtoken.models import Token
from rest_framework.response import Response
from django.http import Http404
from functools import reduce
import operator
from django.db.models import Q
from collections import defaultdict
from rest_framework import filters
from configuration.models import Templatestable


from NavyTrials import language,error,settings,common
from access.views import Common

from collections import namedtuple

from django.template import Template, Context
from datetime import datetime
from datetime import date
import time
from django.template import loader
from django.http import HttpResponse
from io import BytesIO
from xhtml2pdf import pisa

from django.template import Context, Template
import requests
import json
import barcode
from barcode.writer import ImageWriter
from django.utils.crypto import get_random_string
from . import models
from . import serializer
from . import serializer as cSerializer
from configuration.models import Approval
from notification.models import NotificationUser,NotificationUserLog
from master import models as masterModels
from master import serializer as masterSerializer
from access import models as accessModels
#from accounts import models as accountsModels

from notification import models as notificationModels
from access import models as accessSerializer
import pandas as pd
import uuid
import os
import csv
from django.core.files.storage import FileSystemStorage

import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

from django.core.mail import send_mail
from django.core.mail import EmailMessage

from django.http import JsonResponse
from django.template.loader import render_to_string

toMailAddress = 'abc@lifotechnologies.com'

class PrimaryRolesViews(APIView):

    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['name','description']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.PrimaryRoles.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.ListPrimaryRolesSerializer(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.PrimaryRoles.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"Primary roles" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.PrimaryRoles.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        serializer = cSerializer.ListPrimaryRolesSerializer(lists, many=True)
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)

class PrimaryRolesDetailViews(APIView):

    def get_object(self,pk):

        try:
            return models.PrimaryRoles.objects.get(pk = pk)
        except PrimaryRoles.DoesNotExist:
            raise Http404


    def post(self,request, pk = None):

        if 'name' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "Name" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        elif 'code' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "Code" +language.context[language.defaultLang]['missing'] },status=status.HTTP_200_OK)
        else: 

            if 'code' in request.data:
                request.data['code']=(request.data['code']).upper()
            if 'sequence' in request.data:
                request.data['sequence']=request.data['sequence'] if(request.data['sequence']!='')  else 0
            if 'id' in request.data:
                pk = request.data['id']
                created_ip = Common.get_client_ip(request)                                       
                request.data['created_ip'] = created_ip  
               
                if not pk: 
                    duplicate_code = models.PrimaryRoles.objects.values('code').filter(code=request.data['code']).exclude(status=3)
                    duplicate_name = models.PrimaryRoles.objects.values('name').filter(name=request.data['name']).exclude(status=3)
                            
                    if duplicate_code:            
                        return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit code'] },status=status.HTTP_200_OK)                    
                    
                    elif duplicate_name:   
                        return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit name'] },status=status.HTTP_200_OK)                    
                    
                    else:
                
                        saveserialize = cSerializer.PrimaryRolesSerializer(data = request.data)
                        if saveserialize.is_valid():
                            saveserialize.save()
                            return Response({"status" :error.context['success_code'], "message":"New primary roles" +language.context[language.defaultLang]['insert'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                        else:
                            return Response({"status" :error.context['error_code'], "message": error.serializerError(saveserialize)}, status=status.HTTP_200_OK)

                else:
                    if request.data['status'] != 3:
                        duplicate_code = models.PrimaryRoles.objects.values('code').filter(code=request.data['code']).exclude(Q(id=request.data['id']) | Q(status=3))
                        duplicate_name = models.PrimaryRoles.objects.values('name').filter(name=request.data['name']).exclude(Q(id=request.data['id']) | Q(status=3))
                        if duplicate_code:            
                            return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit code'] },status=status.HTTP_200_OK)                    
                        elif duplicate_name:   
                            return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit name'] },status=status.HTTP_200_OK)                                            

                    list = self.get_object(pk)

                    saveserialize = cSerializer.PrimaryRolesSerializer(list, data = request.data, partial= True)
                    if saveserialize.is_valid():
                        saveserialize.save()
                        return Response({"status" :error.context['success_code'], "message":"Primary roles" +language.context[language.defaultLang]['update'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                    else:
                        return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)
            else: 
                return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)   


class SecondaryRolesViews(APIView):

    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['name','description']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.SecondaryRoles.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.ListSecondaryRolesSerializer(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.SecondaryRoles.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"Region" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.SecondaryRoles.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        serializer = cSerializer.ListSecondaryRolesSerializer(lists, many=True)
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)

class SecondaryRolesDetailViews(APIView):

    def get_object(self,pk):

        try:
            return models.SecondaryRoles.objects.get(pk = pk)
        except SecondaryRoles.DoesNotExist:
            raise Http404


    def post(self,request, pk = None):

        if 'name' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "Name" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        elif 'code' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "Code" +language.context[language.defaultLang]['missing'] },status=status.HTTP_200_OK)
        else: 

            if 'code' in request.data:
                request.data['code']=(request.data['code']).upper()
            if 'sequence' in request.data:
                request.data['sequence']=request.data['sequence'] if(request.data['sequence']!='')  else 0
            if 'id' in request.data:
                pk = request.data['id']
                created_ip = Common.get_client_ip(request)                                       
                request.data['created_ip'] = created_ip  
               
                if not pk: 
                    duplicate_code = models.SecondaryRoles.objects.values('code').filter(code=request.data['code']).exclude(status=3)
                    duplicate_name = models.SecondaryRoles.objects.values('name').filter(name=request.data['name']).exclude(status=3)
                            
                    if duplicate_code:            
                        return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit code'] },status=status.HTTP_200_OK)                    
                    
                    elif duplicate_name:   
                        return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit name'] },status=status.HTTP_200_OK)                    
                    
                    else:
                
                        saveserialize = cSerializer.SecondaryRolesSerializer(data = request.data)
                        if saveserialize.is_valid():
                            saveserialize.save()
                            return Response({"status" :error.context['success_code'], "message":"New Secondary Roles" +language.context[language.defaultLang]['insert'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                        else:
                            return Response({"status" :error.context['error_code'], "message": error.serializerError(saveserialize)}, status=status.HTTP_200_OK)

                else:
                    if request.data['status'] != 3:
                        duplicate_code = models.SecondaryRoles.objects.values('code').filter(code=request.data['code']).exclude(Q(id=request.data['id']) | Q(status=3))
                        duplicate_name = models.SecondaryRoles.objects.values('name').filter(name=request.data['name']).exclude(Q(id=request.data['id']) | Q(status=3))
                        if duplicate_code:            
                            return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit code'] },status=status.HTTP_200_OK)                    
                        elif duplicate_name:   
                            return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit name'] },status=status.HTTP_200_OK)                                            

                    list = self.get_object(pk)

                    saveserialize = cSerializer.SecondaryRolesSerializer(list, data = request.data, partial= True)
                    if saveserialize.is_valid():
                        saveserialize.save()
                        return Response({"status" :error.context['success_code'], "message":"Secondary roles" +language.context[language.defaultLang]['update'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                    else:
                        return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)
            else: 
                return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)   


class StandardViews(APIView):

    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['name','description']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.Standard.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.ListStandardSerializer(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.Standard.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"Region" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.Standard.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        serializer = cSerializer.ListStandardSerializer(lists, many=True)
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)

class StandardDetailViews(APIView):

    def get_object(self,pk):

        try:
            return models.Standard.objects.get(pk = pk)
        except Standard.DoesNotExist:
            raise Http404


    def post(self,request, pk = None):

        if 'name' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "Name" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        elif 'code' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "Code" +language.context[language.defaultLang]['missing'] },status=status.HTTP_200_OK)
        else: 

            if 'code' in request.data:
                request.data['code']=(request.data['code']).upper()
            if 'sequence' in request.data:
                request.data['sequence']=request.data['sequence'] if(request.data['sequence']!='')  else 0
            if 'id' in request.data:
                pk = request.data['id']
                created_ip = Common.get_client_ip(request)                                       
                request.data['created_ip'] = created_ip  
               
                if not pk: 
                    duplicate_code = models.Standard.objects.values('code').filter(code=request.data['code']).exclude(status=3)
                    duplicate_name = models.Standard.objects.values('name').filter(name=request.data['name']).exclude(status=3)
                            
                    if duplicate_code:            
                        return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit code'] },status=status.HTTP_200_OK)                    
                    
                    elif duplicate_name:   
                        return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit name'] },status=status.HTTP_200_OK)                    
                    
                    else:
                
                        saveserialize = cSerializer.StandardSerializer(data = request.data)
                        if saveserialize.is_valid():
                            saveserialize.save()
                            return Response({"status" :error.context['success_code'], "message":"New Standard" +language.context[language.defaultLang]['insert'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                        else:
                            return Response({"status" :error.context['error_code'], "message": error.serializerError(saveserialize)}, status=status.HTTP_200_OK)

                else:
                    if request.data['status'] != 3:
                        duplicate_code = models.Standard.objects.values('code').filter(code=request.data['code']).exclude(Q(id=request.data['id']) | Q(status=3))
                        duplicate_name = models.Standard.objects.values('name').filter(name=request.data['name']).exclude(Q(id=request.data['id']) | Q(status=3))
                        if duplicate_code:            
                            return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit code'] },status=status.HTTP_200_OK)                    
                        elif duplicate_name:   
                            return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit name'] },status=status.HTTP_200_OK)                                            

                    list = self.get_object(pk)

                    saveserialize = cSerializer.StandardSerializer(list, data = request.data, partial= True)
                    if saveserialize.is_valid():
                        saveserialize.save()
                        return Response({"status" :error.context['success_code'], "message":"Standard" +language.context[language.defaultLang]['update'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                    else:
                        return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)
            else: 
                return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)   


class PlanForManpowerInductionViews(APIView):

    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['name','description']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.PlanForManpowerInduction.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.ListPlanForManpowerInductionSerializer(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.PlanForManpowerInduction.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"Region" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.PlanForManpowerInduction.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        serializer = cSerializer.ListPlanForManpowerInductionSerializer(lists, many=True)
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)

class PlanForManpowerInductionDetailViews(APIView):

    def get_object(self,pk):

        try:
            return models.PlanForManpowerInduction.objects.get(pk = pk)
        except PlanForManpowerInduction.DoesNotExist:
            raise Http404


    def post(self,request, pk = None):

        if 'name' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "Name" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        elif 'code' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "Code" +language.context[language.defaultLang]['missing'] },status=status.HTTP_200_OK)
        else: 

            if 'code' in request.data:
                request.data['code']=(request.data['code']).upper()
            if 'sequence' in request.data:
                request.data['sequence']=request.data['sequence'] if(request.data['sequence']!='')  else 0
            if 'id' in request.data:
                pk = request.data['id']
                created_ip = Common.get_client_ip(request)                                       
                request.data['created_ip'] = created_ip  
               
                if not pk: 
                    duplicate_code = models.PlanForManpowerInduction.objects.values('code').filter(code=request.data['code']).exclude(status=3)
                    duplicate_name = models.PlanForManpowerInduction.objects.values('name').filter(name=request.data['name']).exclude(status=3)
                            
                    if duplicate_code:            
                        return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit code'] },status=status.HTTP_200_OK)                    
                    
                    elif duplicate_name:   
                        return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit name'] },status=status.HTTP_200_OK)                    
                    
                    else:
                
                        saveserialize = cSerializer.PlanForManpowerInductionSerializer(data = request.data)
                        if saveserialize.is_valid():
                            saveserialize.save()
                            return Response({"status" :error.context['success_code'], "message":"New Plan For Manpower Induction" +language.context[language.defaultLang]['insert'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                        else:
                            return Response({"status" :error.context['error_code'], "message": error.serializerError(saveserialize)}, status=status.HTTP_200_OK)

                else:
                    if request.data['status'] != 3:
                        duplicate_code = models.PlanForManpowerInduction.objects.values('code').filter(code=request.data['code']).exclude(Q(id=request.data['id']) | Q(status=3))
                        duplicate_name = models.PlanForManpowerInduction.objects.values('name').filter(name=request.data['name']).exclude(Q(id=request.data['id']) | Q(status=3))
                        if duplicate_code:            
                            return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit code'] },status=status.HTTP_200_OK)                    
                        elif duplicate_name:   
                            return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit name'] },status=status.HTTP_200_OK)                                            

                    list = self.get_object(pk)

                    saveserialize = cSerializer.PlanForManpowerInductionSerializer(list, data = request.data, partial= True)
                    if saveserialize.is_valid():
                        saveserialize.save()
                        return Response({"status" :error.context['success_code'], "message":"Plan For Manpower Induction" +language.context[language.defaultLang]['update'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                    else:
                        return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)
            else: 
                return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)   

class AcquisitionMethodViews(APIView):

    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['name','description']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.AcquisitionMethod.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.ListAcquisitionMethodSerializer(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.AcquisitionMethod.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"Region" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.AcquisitionMethod.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        serializer = cSerializer.ListAcquisitionMethodSerializer(lists, many=True)
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)

class AcquisitionMethodDetailViews(APIView):

    def get_object(self,pk):

        try:
            return models.AcquisitionMethod.objects.get(pk = pk)
        except AcquisitionMethod.DoesNotExist:
            raise Http404


    def post(self,request, pk = None):

        if 'name' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "Name" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        elif 'code' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "Code" +language.context[language.defaultLang]['missing'] },status=status.HTTP_200_OK)
        else: 

            if 'code' in request.data:
                request.data['code']=(request.data['code']).upper()
            if 'sequence' in request.data:
                request.data['sequence']=request.data['sequence'] if(request.data['sequence']!='')  else 0
            if 'id' in request.data:
                pk = request.data['id']
                created_ip = Common.get_client_ip(request)                                       
                request.data['created_ip'] = created_ip  
               
                if not pk: 
                    duplicate_code = models.AcquisitionMethod.objects.values('code').filter(code=request.data['code']).exclude(status=3)
                    duplicate_name = models.AcquisitionMethod.objects.values('name').filter(name=request.data['name']).exclude(status=3)
                            
                    if duplicate_code:            
                        return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit code'] },status=status.HTTP_200_OK)                    
                    
                    elif duplicate_name:   
                        return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit name'] },status=status.HTTP_200_OK)                    
                    
                    else:
                
                        saveserialize = cSerializer.AcquisitionMethodSerializer(data = request.data)
                        if saveserialize.is_valid():
                            saveserialize.save()
                            return Response({"status" :error.context['success_code'], "message":"New Acquisition Method" +language.context[language.defaultLang]['insert'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                        else:
                            return Response({"status" :error.context['error_code'], "message": error.serializerError(saveserialize)}, status=status.HTTP_200_OK)

                else:
                    if request.data['status'] != 3:
                        duplicate_code = models.AcquisitionMethod.objects.values('code').filter(code=request.data['code']).exclude(Q(id=request.data['id']) | Q(status=3))
                        duplicate_name = models.AcquisitionMethod.objects.values('name').filter(name=request.data['name']).exclude(Q(id=request.data['id']) | Q(status=3))
                        if duplicate_code:            
                            return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit code'] },status=status.HTTP_200_OK)                    
                        elif duplicate_name:   
                            return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit name'] },status=status.HTTP_200_OK)                                            

                    list = self.get_object(pk)

                    saveserialize = cSerializer.AcquisitionMethodSerializer(list, data = request.data, partial= True)
                    if saveserialize.is_valid():
                        saveserialize.save()
                        return Response({"status" :error.context['success_code'], "message":"Acquisition Method" +language.context[language.defaultLang]['update'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                    else:
                        return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)
            else: 
                return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)   



class SSSViews(APIView):

    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['name','description']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.SSS.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.ListSSSSerializer(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.SSS.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"Region" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.SSS.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        serializer = cSerializer.ListSSSSerializer(lists, many=True)
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)

class SSSDetailViews(APIView):

    def get_object(self,pk):

        try:
            return models.SSS.objects.get(pk = pk)
        except SSS.DoesNotExist:
            raise Http404


    def post(self,request, pk = None):

        if 'name' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "Name" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        elif 'code' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "Code" +language.context[language.defaultLang]['missing'] },status=status.HTTP_200_OK)
        else: 

            if 'code' in request.data:
                request.data['code']=(request.data['code']).upper()
            if 'sequence' in request.data:
                request.data['sequence']=request.data['sequence'] if(request.data['sequence']!='')  else 0
            if 'id' in request.data:
                pk = request.data['id']
                created_ip = Common.get_client_ip(request)                                       
                request.data['created_ip'] = created_ip  
               
                if not pk: 
                    duplicate_code = models.SSS.objects.values('code').filter(code=request.data['code']).exclude(status=3)
                    duplicate_name = models.SSS.objects.values('name').filter(name=request.data['name']).exclude(status=3)
                            
                    if duplicate_code:            
                        return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit code'] },status=status.HTTP_200_OK)                    
                    
                    elif duplicate_name:   
                        return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit name'] },status=status.HTTP_200_OK)                    
                    
                    else:
                
                        saveserialize = cSerializer.SSSSerializer(data = request.data)
                        if saveserialize.is_valid():
                            saveserialize.save()
                            return Response({"status" :error.context['success_code'], "message":"New PSR SSS" +language.context[language.defaultLang]['insert'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                        else:
                            return Response({"status" :error.context['error_code'], "message": error.serializerError(saveserialize)}, status=status.HTTP_200_OK)

                else:
                    if request.data['status'] != 3:
                        duplicate_code = models.SSS.objects.values('code').filter(code=request.data['code']).exclude(Q(id=request.data['id']) | Q(status=3))
                        duplicate_name = models.SSS.objects.values('name').filter(name=request.data['name']).exclude(Q(id=request.data['id']) | Q(status=3))
                        if duplicate_code:            
                            return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit code'] },status=status.HTTP_200_OK)                    
                        elif duplicate_name:   
                            return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit name'] },status=status.HTTP_200_OK)                                            

                    list = self.get_object(pk)

                    saveserialize = cSerializer.SSSSerializer(list, data = request.data, partial= True)
                    if saveserialize.is_valid():
                        saveserialize.save()
                        return Response({"status" :error.context['success_code'], "message":"PSR SSS" +language.context[language.defaultLang]['update'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                    else:
                        return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)
            else: 
                return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)   


###############

class PSRSubSectionViews(APIView):

    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['name','description']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.PSRSubSection.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.ListPSRSubSectionSerializer(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.PSRSubSection.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"Region" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.PSRSubSection.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        serializer = cSerializer.ListPSRSubSectionSerializer(lists, many=True)
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)

class PSRSubSectionDetailViews(APIView):

    def get_object(self,pk):

        try:
            return models.PSRSubSection.objects.get(pk = pk)
        except PSRSubSection.DoesNotExist:
            raise Http404


    def post(self,request, pk = None):

        if 'name' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "Name" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        elif 'code' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "Code" +language.context[language.defaultLang]['missing'] },status=status.HTTP_200_OK)
        else: 

            if 'code' in request.data:
                request.data['code']=(request.data['code']).upper()
            if 'sequence' in request.data:
                request.data['sequence']=request.data['sequence'] if(request.data['sequence']!='')  else 0
            if 'id' in request.data:
                pk = request.data['id']
                created_ip = Common.get_client_ip(request)                                       
                request.data['created_ip'] = created_ip  
               
                if not pk: 
                    duplicate_code = models.PSRSubSection.objects.values('code').filter(code=request.data['code']).exclude(status=3)
                    duplicate_name = models.PSRSubSection.objects.values('name').filter(name=request.data['name']).exclude(status=3)
                            
                    if duplicate_code:            
                        return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit code'] },status=status.HTTP_200_OK)                    
                    
                    elif duplicate_name:   
                        return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit name'] },status=status.HTTP_200_OK)                    
                    
                    else:
                
                        saveserialize = cSerializer.PSRSubSectionSerializer(data = request.data)
                        if saveserialize.is_valid():
                            saveserialize.save()
                            return Response({"status" :error.context['success_code'], "message":"New PSR Sub Section" +language.context[language.defaultLang]['insert'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                        else:
                            return Response({"status" :error.context['error_code'], "message": error.serializerError(saveserialize)}, status=status.HTTP_200_OK)

                else:
                    if request.data['status'] != 3:
                        duplicate_code = models.PSRSubSection.objects.values('code').filter(code=request.data['code']).exclude(Q(id=request.data['id']) | Q(status=3))
                        duplicate_name = models.PSRSubSection.objects.values('name').filter(name=request.data['name']).exclude(Q(id=request.data['id']) | Q(status=3))
                        if duplicate_code:            
                            return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit code'] },status=status.HTTP_200_OK)                    
                        elif duplicate_name:   
                            return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit name'] },status=status.HTTP_200_OK)                                            

                    list = self.get_object(pk)

                    saveserialize = cSerializer.PSRSubSectionSerializer(list, data = request.data, partial= True)
                    if saveserialize.is_valid():
                        saveserialize.save()
                        return Response({"status" :error.context['success_code'], "message":"PSR Sub Section" +language.context[language.defaultLang]['update'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                    else:
                        return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)
            else: 
                return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)   


#### PSR Section 

class PSRSectionViews(APIView):

    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['name','description']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.PSRSection.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.ListPSRSectionSerializer(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.PSRSection.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"Region" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.PSRSection.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        serializer = cSerializer.ListPSRSectionSerializer(lists, many=True)
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)

class PSRSectionViews2(APIView):

    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['name','description']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.PSRSection.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.ListPSRSectionSerializer2(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.PSRSection.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"Region" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.PSRSection.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        serializer = cSerializer.ListPSRSectionSerializer2(lists, many=True)
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)


class PSRSectionDetailViews(APIView):

    def get_object(self,pk):

        try:
            return models.PSRSection.objects.get(pk = pk)
        except PSRSection.DoesNotExist:
            raise Http404


    def post(self,request, pk = None):

        if 'name' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "Name" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        elif 'code' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "Code" +language.context[language.defaultLang]['missing'] },status=status.HTTP_200_OK)
        else: 

            if 'code' in request.data:
                request.data['code']=(request.data['code']).upper()
            if 'sequence' in request.data:
                request.data['sequence']=request.data['sequence'] if(request.data['sequence']!='')  else 0
            if 'id' in request.data:
                pk = request.data['id']
                created_ip = Common.get_client_ip(request)                                       
                request.data['created_ip'] = created_ip  
               
                if not pk: 
                    duplicate_code = models.PSRSection.objects.values('code').filter(code=request.data['code']).exclude(status=3)
                    duplicate_name = models.PSRSection.objects.values('name').filter(name=request.data['name']).exclude(status=3)
                            
                    if duplicate_code:            
                        return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit code'] },status=status.HTTP_200_OK)                    
                    
                    elif duplicate_name:   
                        return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit name'] },status=status.HTTP_200_OK)                    
                    
                    else:
                
                        saveserialize = cSerializer.PSRSectionSerializer(data = request.data)

                        if saveserialize.is_valid():
                            saveserialize.save()
                            #print(saveserialize.data,"$$$$$$$$$$ ADD")

                            unit_id = request.data['unit_id']
                            running_id = saveserialize.data['id']

                            #print(unit_id, len(unit_id), running_id, "##############")


                            for unit in unit_id:
                                models.PSRSectionUnit.objects.create(
                                    section_id = running_id,
                                    unit_id = unit
                                )


                            return Response({"status" :error.context['success_code'], "message":"New PSR Section" +language.context[language.defaultLang]['insert'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                        else:
                            return Response({"status" :error.context['error_code'], "message": error.serializerError(saveserialize)}, status=status.HTTP_200_OK)

                else:
                    if request.data['status'] != 3:
                        duplicate_code = models.PSRSection.objects.values('code').filter(code=request.data['code']).exclude(Q(id=request.data['id']) | Q(status=3))
                        duplicate_name = models.PSRSection.objects.values('name').filter(name=request.data['name']).exclude(Q(id=request.data['id']) | Q(status=3))
                        if duplicate_code:            
                            return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit code'] },status=status.HTTP_200_OK)                    
                        elif duplicate_name:   
                            return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit name'] },status=status.HTTP_200_OK)                                            

                    list = self.get_object(pk)

                    saveserialize = cSerializer.PSRSectionSerializer(list, data = request.data, partial= True)

                    

                    if saveserialize.is_valid():
                        saveserialize.save()

                        running_id = saveserialize.data['id']
                        unit_id = request.data['unit_id']

                        #print(unit_id, len(unit_id), running_id, "##############")

                        models.PSRSectionUnit.objects.filter(section_id=running_id).delete()
                        for unit in unit_id:
                            models.PSRSectionUnit.objects.create(
                                section_id = running_id,
                                unit_id = unit
                            )

                        return Response({"status" :error.context['success_code'], "message":"PSR Section" +language.context[language.defaultLang]['update'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                    else:
                        return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)
            else: 
                return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)   







# GLS Masters

class DocumentSectionsViews(APIView):

    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['name','description']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.DocumentSections.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.ListDocumentSectionsSerializer(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.DocumentSections.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"Document sections" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.DocumentSections.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        serializer = cSerializer.ListDocumentSectionsSerializer(lists, many=True)
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)

class DocumentSectionsDetailViews(APIView):

    def get_object(self,pk):

        try:
            return models.DocumentSections.objects.get(pk = pk)
        except DocumentSections.DoesNotExist:
            raise Http404


    def post(self,request, pk = None):

        if 'name' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "Name" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        elif 'code' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "Code" +language.context[language.defaultLang]['missing'] },status=status.HTTP_200_OK)
        else: 

            if 'code' in request.data:
                request.data['code']=(request.data['code']).upper()
            if 'sequence' in request.data:
                request.data['sequence']=request.data['sequence'] if(request.data['sequence']!='')  else 0
            if 'id' in request.data:
                pk = request.data['id']
                created_ip = Common.get_client_ip(request)                                       
                request.data['created_ip'] = created_ip  
               
                if not pk: 
                    duplicate_code = models.DocumentSections.objects.values('code').filter(code=request.data['code']).exclude(status=3)
                    duplicate_name = models.DocumentSections.objects.values('name').filter(name=request.data['name']).exclude(status=3)
                            
                    if duplicate_code:            
                        return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit code'] },status=status.HTTP_200_OK)                    
                    
                    elif duplicate_name:   
                        return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit name'] },status=status.HTTP_200_OK)                    
                    
                    else:
                
                        saveserialize = cSerializer.DocumentSectionsSerializer(data = request.data)
                        if saveserialize.is_valid():
                            saveserialize.save()
                            return Response({"status" :error.context['success_code'], "message":"New Document sections" +language.context[language.defaultLang]['insert'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                        else:
                            return Response({"status" :error.context['error_code'], "message": error.serializerError(saveserialize)}, status=status.HTTP_200_OK)

                else:
                    if request.data['status'] != 3:
                        duplicate_code = models.DocumentSections.objects.values('code').filter(code=request.data['code']).exclude(Q(id=request.data['id']) | Q(status=3))
                        duplicate_name = models.DocumentSections.objects.values('name').filter(name=request.data['name']).exclude(Q(id=request.data['id']) | Q(status=3))
                        if duplicate_code:            
                            return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit code'] },status=status.HTTP_200_OK)                    
                        elif duplicate_name:   
                            return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit name'] },status=status.HTTP_200_OK)                                            

                    list = self.get_object(pk)

                    saveserialize = cSerializer.DocumentSectionsSerializer(list, data = request.data, partial= True)
                    if saveserialize.is_valid():
                        saveserialize.save()
                        return Response({"status" :error.context['success_code'], "message":"Document sections" +language.context[language.defaultLang]['update'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                    else:
                        return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)
            else: 
                return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)


class DocumentSubSectionsViews(APIView):

    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['name','description']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.DocumentSubSections.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.ListDocumentSubSectionsSerializer(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.DocumentSubSections.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"Document sub sections" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.DocumentSubSections.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        serializer = cSerializer.ListDocumentSubSectionsSerializer(lists, many=True)
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)

class DocumentSubSectionsDetailViews(APIView):

    def get_object(self,pk):

        try:
            return models.DocumentSubSections.objects.get(pk = pk)
        except DocumentSubSections.DoesNotExist:
            raise Http404


    def post(self,request, pk = None):

        if 'name' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "Name" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        elif 'code' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "Code" +language.context[language.defaultLang]['missing'] },status=status.HTTP_200_OK)
        else: 

            if 'code' in request.data:
                request.data['code']=(request.data['code']).upper()
            if 'sequence' in request.data:
                request.data['sequence']=request.data['sequence'] if(request.data['sequence']!='')  else 0
            if 'id' in request.data:
                pk = request.data['id']
                created_ip = Common.get_client_ip(request)                                       
                request.data['created_ip'] = created_ip  
               
                if not pk: 
                    duplicate_code = models.DocumentSubSections.objects.values('code').filter(code=request.data['code']).exclude(status=3)
                    duplicate_name = models.DocumentSubSections.objects.values('name').filter(name=request.data['name']).exclude(status=3)
                            
                    if duplicate_code:            
                        return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit code'] },status=status.HTTP_200_OK)                    
                    
                    elif duplicate_name:   
                        return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit name'] },status=status.HTTP_200_OK)                    
                    
                    else:
                
                        saveserialize = cSerializer.DocumentSubSectionsSerializer(data = request.data)
                        if saveserialize.is_valid():
                            saveserialize.save()
                            return Response({"status" :error.context['success_code'], "message":"New Document sub sections" +language.context[language.defaultLang]['insert'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                        else:
                            return Response({"status" :error.context['error_code'], "message": error.serializerError(saveserialize)}, status=status.HTTP_200_OK)

                else:
                    if request.data['status'] != 3:
                        duplicate_code = models.DocumentSubSections.objects.values('code').filter(code=request.data['code']).exclude(Q(id=request.data['id']) | Q(status=3))
                        duplicate_name = models.DocumentSubSections.objects.values('name').filter(name=request.data['name']).exclude(Q(id=request.data['id']) | Q(status=3))
                        if duplicate_code:            
                            return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit code'] },status=status.HTTP_200_OK)                    
                        elif duplicate_name:   
                            return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit name'] },status=status.HTTP_200_OK)                                            

                    list = self.get_object(pk)

                    saveserialize = cSerializer.DocumentSubSectionsSerializer(list, data = request.data, partial= True)
                    if saveserialize.is_valid():
                        saveserialize.save()
                        return Response({"status" :error.context['success_code'], "message":"Document sub sections" +language.context[language.defaultLang]['update'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                    else:
                        return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)
            else: 
                return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)



class DocumentSubSections2Views(APIView):

    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['name','description']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.DocumentSubSections2.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.ListDocumentSubSections2Serializer(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.DocumentSubSections2.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"Document sub sections2" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.DocumentSubSections2.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        serializer = cSerializer.ListDocumentSubSections2Serializer(lists, many=True)
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)

class DocumentSubSections2DetailViews(APIView):

    def get_object(self,pk):

        try:
            return models.DocumentSubSections2.objects.get(pk = pk)
        except DocumentSubSections2.DoesNotExist:
            raise Http404


    def post(self,request, pk = None):

        if 'name' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "Name" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        elif 'code' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "Code" +language.context[language.defaultLang]['missing'] },status=status.HTTP_200_OK)
        else: 

            if 'code' in request.data:
                request.data['code']=(request.data['code']).upper()
            if 'sequence' in request.data:
                request.data['sequence']=request.data['sequence'] if(request.data['sequence']!='')  else 0
            if 'id' in request.data:
                pk = request.data['id']
                created_ip = Common.get_client_ip(request)                                       
                request.data['created_ip'] = created_ip  
               
                if not pk: 
                    duplicate_code = models.DocumentSubSections2.objects.values('code').filter(code=request.data['code']).exclude(status=3)
                    duplicate_name = models.DocumentSubSections2.objects.values('name').filter(name=request.data['name']).exclude(status=3)
                            
                    if duplicate_code:            
                        return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit code'] },status=status.HTTP_200_OK)                    
                    
                    elif duplicate_name:   
                        return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit name'] },status=status.HTTP_200_OK)                    
                    
                    else:
                
                        saveserialize = cSerializer.DocumentSubSections2Serializer(data = request.data)
                        if saveserialize.is_valid():
                            saveserialize.save()
                            return Response({"status" :error.context['success_code'], "message":"New Document sub sections2" +language.context[language.defaultLang]['insert'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                        else:
                            return Response({"status" :error.context['error_code'], "message": error.serializerError(saveserialize)}, status=status.HTTP_200_OK)

                else:
                    if request.data['status'] != 3:
                        duplicate_code = models.DocumentSubSections2.objects.values('code').filter(code=request.data['code']).exclude(Q(id=request.data['id']) | Q(status=3))
                        duplicate_name = models.DocumentSubSections2.objects.values('name').filter(name=request.data['name']).exclude(Q(id=request.data['id']) | Q(status=3))
                        if duplicate_code:            
                            return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit code'] },status=status.HTTP_200_OK)                    
                        elif duplicate_name:   
                            return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit name'] },status=status.HTTP_200_OK)                                            

                    list = self.get_object(pk)

                    saveserialize = cSerializer.DocumentSubSections2Serializer(list, data = request.data, partial= True)
                    if saveserialize.is_valid():
                        saveserialize.save()
                        return Response({"status" :error.context['success_code'], "message":"Document sub sections2" +language.context[language.defaultLang]['update'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                    else:
                        return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)
            else: 
                return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)


class AnnexuresViews(APIView):

    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['name','description']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.Annexures.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.ListAnnexuresSerializer(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.Annexures.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"Annexures" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.Annexures.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        serializer = cSerializer.ListAnnexuresSerializer(lists, many=True)
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)

class AnnexuresDetailViews(APIView):

    def get_object(self,pk):

        try:
            return models.Annexures.objects.get(pk = pk)
        except Annexures.DoesNotExist:
            raise Http404


    def post(self,request, pk = None):

        if 'name' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "Name" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        elif 'code' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "Code" +language.context[language.defaultLang]['missing'] },status=status.HTTP_200_OK)
        else: 

            if 'code' in request.data:
                request.data['code']=(request.data['code']).upper()
            if 'sequence' in request.data:
                request.data['sequence']=request.data['sequence'] if(request.data['sequence']!='')  else 0
            if 'id' in request.data:
                pk = request.data['id']
                created_ip = Common.get_client_ip(request)                                       
                request.data['created_ip'] = created_ip  
               
                if not pk: 
                    duplicate_code = models.Annexures.objects.values('code').filter(code=request.data['code']).exclude(status=3)
                    duplicate_name = models.Annexures.objects.values('name').filter(name=request.data['name']).exclude(status=3)
                            
                    if duplicate_code:            
                        return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit code'] },status=status.HTTP_200_OK)                    
                    
                    elif duplicate_name:   
                        return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit name'] },status=status.HTTP_200_OK)                    
                    
                    else:
                
                        saveserialize = cSerializer.AnnexuresSerializer(data = request.data)
                        if saveserialize.is_valid():
                            saveserialize.save()
                            return Response({"status" :error.context['success_code'], "message":"New Annexures" +language.context[language.defaultLang]['insert'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                        else:
                            return Response({"status" :error.context['error_code'], "message": error.serializerError(saveserialize)}, status=status.HTTP_200_OK)

                else:
                    if request.data['status'] != 3:
                        duplicate_code = models.Annexures.objects.values('code').filter(code=request.data['code']).exclude(Q(id=request.data['id']) | Q(status=3))
                        duplicate_name = models.Annexures.objects.values('name').filter(name=request.data['name']).exclude(Q(id=request.data['id']) | Q(status=3))
                        if duplicate_code:            
                            return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit code'] },status=status.HTTP_200_OK)                    
                        elif duplicate_name:   
                            return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit name'] },status=status.HTTP_200_OK)                                            

                    list = self.get_object(pk)

                    saveserialize = cSerializer.AnnexuresSerializer(list, data = request.data, partial= True)
                    if saveserialize.is_valid():
                        saveserialize.save()
                        return Response({"status" :error.context['success_code'], "message":"Annexures" +language.context[language.defaultLang]['update'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                    else:
                        return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)
            else: 
                return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)




# PSR Initiation Notes
class InitiationNotesList(APIView):
    
    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.InitiationNotes.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.ListInitiationNotesSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.InitiationNotes.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Initiation Notes" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        #lists = models.InitiationNotes.objects
        lists = models.InitiationNotes.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListInitiationNotesSerializer(lists, many=True)
        print('serializer serializer',serializer.data)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class InitiationNotesCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.InitiationNotes.objects.get(pk = pk)
            except models.InitiationNotes.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):

        #print(request.data)
        #pass
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']
            if not request.data['project'] and request.data['status'] != 3:
                return Response({"status":error.context['error_code'], "message": "Project is missing"}, status=status.HTTP_200_OK)
            if not pk:


                
            
                request.data['primary_role'] = json.dumps(request.data['primary_role'])
                request.data['secondary_role'] = json.dumps(request.data['secondary_role'])
                request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.InitiationNotesSerializer(data = request.data)
                
                if saveserialize.is_valid():
                    #models.InitiationNotes.objects.filter(trial_id=request.data['trial']).delete()
                    saveserialize.save()
                    '''
                    logData=request.data
                    logData['running_id']=saveserialize.data['id']
                    saveserializelog = cSerializer.InitiationNoteslogSerializer(data = logData)
                    if saveserializelog.is_valid():
                        saveserializelog.save()
                    trials=InitiationNotesListSerializer(models.InitiationNotes.objects.filter(id=saveserialize.data['id']), many=True).data[0]
                    '''
                    return Response({"status" : error.context['success_code'], "message":"Initiation notes" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:

                #request.data['primary_role'] = json.dumps(request.data['primary_role'])
                #request.data['secondary_role'] = json.dumps(request.data['secondary_role'])
                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id

                if request.data['approve']=='no': 
                    request.data['primary_role'] = json.dumps(request.data['primary_role'])
                    request.data['secondary_role'] = json.dumps(request.data['secondary_role'])

                list = self.get_object(pk)
                
                saveserialize = cSerializer.InitiationNotesSerializer(list, data = request.data, partial= True)                
                if saveserialize.is_valid():
                    saveserialize.save()
                    logData=request.data

                    if 'approved_status' in request.data:
                        models.ProjectModuleStatus.objects.filter(project_id=request.data['project'], project_module_master_id=1).delete()

                        models.ProjectModuleStatus.objects.create(
                            project_module_master_id = 1,
                            project_id = request.data['project'],
                            status  = request.data['approved_status']
                        )

                    return Response({"status" :error.context['success_code'], "message":"Initiation notes" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)


# Initiation Notes Document
class InitiationNotesDocumentList(APIView):

    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.InitiationNotesDocument.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.InitiationNotesDocumentListSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.InitiationNotesDocument.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Initiation Notes" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.InitiationNotesDocument.objects
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListInitiationNotesDocumentSerializer(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class InitiationNotesDocumentCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.InitiationNotesDocument.objects.get(pk = pk)
            except models.InitiationNotesDocument.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):

        #print('######@@@@@@@',request.data)
        #pass

        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']  
            if not request.data['document_name']:
                return Response({"status":error.context['error_code'], "message": "Document name is missing"}, status=status.HTTP_200_OK)
            if not request.data['document_remark']:
                return Response({"status":error.context['error_code'], "message": "Document remark is missing"}, status=status.HTTP_200_OK)
            # if not request.data['file_name']:
            #     return Response({"status":error.context['error_code'], "message": "Document file is missing"}, status=status.HTTP_200_OK)
            if not pk:
                request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.InitiationNotesDocumentSerializer(data = request.data)
                
                if saveserialize.is_valid():
                    #models.InitiationNotesDocument.objects.filter(trial_id=request.data['trial']).delete()
                    saveserialize.save()
                    #print(saveserialize.query())
                    return Response({"status" : error.context['success_code'], "message":"Document" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:
                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id 
                list = self.get_object(pk)
                saveserialize = cSerializer.InitiationNotesDocumentSerializer(list, data = request.data, partial= True)
                if saveserialize.is_valid():
                    saveserialize.save()
                    logData=request.data
                    return Response({"status" :error.context['success_code'], "message":"Document" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)


# Initiation Notes Send Mail
class InitiationNotesSendMailList(APIView):

    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.InitiationNotesSendMail.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.InitiationNotesSendMailListSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.InitiationNotesSendMail.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Initiation Notes" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.InitiationNotesSendMail.objects
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListInitiationNotesSendMailSerializer(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class InitiationNotesSendMailCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.InitiationNotesSendMail.objects.get(pk = pk)
            except models.InitiationNotesSendMail.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):

        #print('@@@@', request.data)
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']

            if not request.data['to_email']:
                return Response({"status":error.context['error_code'], "message": "To email address is missing"}, status=status.HTTP_200_OK)
            if not request.data['subject']:
                return Response({"status":error.context['error_code'], "message": "Subject is missing"}, status=status.HTTP_200_OK)
            if not pk:
                request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.InitiationNotesSendMailSerializer(data = request.data)
                
                if saveserialize.is_valid():
                    saveserialize.save()
                    #print(saveserialize.data)
                    #print(saveserialize.query())

                    subject = saveserialize.data['subject']
                    comments = saveserialize.data['comments']
                    file_name = saveserialize.data['file_name']

                    email = toMailAddress
                    attach = request.FILES['file_name']
                    #print(attach)
                    #print(request.FILES['file_name'])
                    fhand = open(f"{settings.BASE_DIR}"+file_name,"r")
                    #print(settings.BASE_DIR)
                    #print(f"{settings.BASE_DIR}"+file_name,"QQQQ")
                    try:
                        mail = EmailMessage(subject, comments, settings.EMAIL_HOST_USER, [email])
                        mail.attach_file(f"{settings.BASE_DIR}"+file_name)
                        mail.send()
                        #print('OK')
                        return Response({"status" : error.context['success_code'], "message":"Send email" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                    except:
                        return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
                        #print('Fail')

                    return Response({"status" : error.context['success_code'], "message":"Send email" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:
                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id 
                list = self.get_object(pk)
                saveserialize = cSerializer.InitiationNotesSendMailSerializer(list, data = request.data, partial= True)
                if saveserialize.is_valid():
                    saveserialize.save()
                    logData=request.data
                    return Response({"status" :error.context['success_code'], "message":"Send email" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)



# PSR Formulation Of Approach Paper
class FormulationOfApproachPaperList(APIView):
    
    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.FormulationOfApproachPaper.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.FormulationOfApproachPaperListSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.FormulationOfApproachPaper.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"PSR Formulation Of Approach Paper" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.FormulationOfApproachPaper.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListFormulationOfApproachPaperSerializer(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class FormulationOfApproachPaperCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.FormulationOfApproachPaper.objects.get(pk = pk)
            except models.FormulationOfApproachPaper.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):



        #project = masterModels.Project.objects.first()
        #project = masterModels.Project.objects.filter(code='DD').first()
        #a1 = masterModels.Project.objects.select_related(models.FormulationOfApproachPaper)
        #print("######@@@",a1.query,"#######")
        #print(request.data)
        #pass
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']
            print("Delete",request.data)

            if not request.data['project'] and request.data['status'] != 3:
                return Response({"status":error.context['error_code'], "message": "Project is missing"}, status=status.HTTP_200_OK)
            if not pk:

                request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.FormulationOfApproachPaperSerializer(data = request.data)
                
                if saveserialize.is_valid():
                    #models.FormulationOfApproachPaper.objects.filter(trial_id=request.data['trial']).delete()
                    saveserialize.save()

                    return Response({"status" : error.context['success_code'], "message":"PSR Formulation Of Approach Paper" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:
                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id 
                list = self.get_object(pk)
                
                saveserialize = cSerializer.FormulationOfApproachPaperSerializer(list, data = request.data, partial= True)                
                if saveserialize.is_valid():
                    saveserialize.save()
                    logData=request.data


                    # Dashboard

                    if 'approved_status' in request.data:
                        models.ProjectModuleStatus.objects.filter(project_id=request.data['project'], project_module_master_id=2).delete()

                        models.ProjectModuleStatus.objects.create(
                            project_module_master_id = 2,
                            project_id = request.data['project'],
                            status  = request.data['approved_status']
                        )

                    return Response({"status" :error.context['success_code'], "message":"PSR Formulation Of Approach Paper" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)


# PSR Formulation Of Approach Paper Document
class FormulationOfApproachPaperDocumentList(APIView):

    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.FormulationOfApproachPaperDocument.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.FormulationOfApproachPaperDocumentListSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.FormulationOfApproachPaperDocument.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Initiation Notes" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.FormulationOfApproachPaperDocument.objects
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListFormulationOfApproachPaperDocumentSerializer(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class FormulationOfApproachPaperDocumentCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.FormulationOfApproachPaperDocument.objects.get(pk = pk)
            except models.FormulationOfApproachPaperDocument.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):

        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']  
            if not request.data['document_name']:
                return Response({"status":error.context['error_code'], "message": "Document name is missing"}, status=status.HTTP_200_OK)
            if not request.data['document_remark']:
                return Response({"status":error.context['error_code'], "message": "Document remark is missing"}, status=status.HTTP_200_OK)
            # if not request.data['file_name']:
            #     return Response({"status":error.context['error_code'], "message": "Document file is missing"}, status=status.HTTP_200_OK)
            if not pk:
                request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.FormulationOfApproachPaperDocumentSerializer(data = request.data)
                
                if saveserialize.is_valid():
                    #models.FormulationOfApproachPaperDocument.objects.filter(trial_id=request.data['trial']).delete()
                    saveserialize.save()
                    #print(saveserialize.query())
                    return Response({"status" : error.context['success_code'], "message":"Document" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:
                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id 
                list = self.get_object(pk)
                saveserialize = cSerializer.FormulationOfApproachPaperDocumentSerializer(list, data = request.data, partial= True)
                if saveserialize.is_valid():
                    saveserialize.save()
                    logData=request.data
                    return Response({"status" :error.context['success_code'], "message":"Document" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)


# PSR Formulation Of Approach Paper Send Mail
class FormulationOfApproachPaperSendMailList(APIView):

    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.FormulationOfApproachPaperSendMail.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.FormulationOfApproachPaperSendMailListSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.FormulationOfApproachPaperSendMail.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Initiation Notes" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.FormulationOfApproachPaperSendMail.objects
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListFormulationOfApproachPaperSendMailSerializer(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class FormulationOfApproachPaperSendMailCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.FormulationOfApproachPaperSendMail.objects.get(pk = pk)
            except models.FormulationOfApproachPaperSendMail.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):

        #print('@@@@', request.data)
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']

            if not request.data['to_email']:
                return Response({"status":error.context['error_code'], "message": "To email address is missing"}, status=status.HTTP_200_OK)
            if not request.data['subject']:
                return Response({"status":error.context['error_code'], "message": "Subject is missing"}, status=status.HTTP_200_OK)
            if not pk:
                request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.FormulationOfApproachPaperSendMailSerializer(data = request.data)
                
                if saveserialize.is_valid():
                    saveserialize.save()
                    #print(saveserialize.data)
                    #print(saveserialize.query())

                    subject = saveserialize.data['subject']
                    comments = saveserialize.data['comments']
                    file_name = saveserialize.data['file_name']

                    email = toMailAddress
                    attach = request.FILES['file_name']
                    fhand = open(f"{settings.BASE_DIR}"+file_name,"r")

                    try:
                        mail = EmailMessage(subject, comments, settings.EMAIL_HOST_USER, [email])
                        mail.attach_file(f"{settings.BASE_DIR}"+file_name)
                        mail.send()
                        
                        return Response({"status" : error.context['success_code'], "message":"Send email" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                    except:
                        return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
                        #print('Fail')

                    return Response({"status" : error.context['success_code'], "message":"Send email" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:
                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id 
                list = self.get_object(pk)
                saveserialize = cSerializer.FormulationOfApproachPaperSendMailSerializer(list, data = request.data, partial= True)
                if saveserialize.is_valid():
                    saveserialize.save()
                    logData=request.data
                    return Response({"status" :error.context['success_code'], "message":"Send email" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)


# PSR Formulation Of Approach Paper Responsibility
class FormulationOfApproachPaperResponsibilityList(APIView):
    
    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.PresentationOfApproachPaper.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.PresentationOfApproachPaperListSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.PresentationOfApproachPaper.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"PSR Formulation Of Approach Paper" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        #lists = models.PresentationOfApproachPaper.objects.exclude(status='3')
        #lists = models.FormulationOfApproachPaperResponsibility.objects.all()
        lists = models.FormulationOfApproachPaperResponsibility.objects.values('id','section','unit','project_id').filter(project_id='40')
        #sectionList = masterModels.Section.objects.exclude(status='3').get()

        #foe

        print(lists)
        #lists.values()
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.FormulationOfApproachPaperResponsibilitySerializer(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class FormulationOfApproachPaperResponsibilityCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.FormulationOfApproachPaperResponsibility.objects.get(pk = pk)
            except models.FormulationOfApproachPaperResponsibility.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):
        created_ip = Common.get_client_ip(request)
        #pass
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']
            print("Delete",request.data)

            if not request.data['project'] and request.data['status'] != 3:
                return Response({"status":error.context['error_code'], "message": "Project is missing"}, status=status.HTTP_200_OK)
            if not pk:


                count = len(request.data['section'])
                #print(len(request.data['section']),'@@@@@@@@@@')
                #print(request.data['compartment'], request.data['equipment'][count:],'@@@@@@@@@@')
                #request.data['compartment'] = json.dumps(request.data['compartment'])
                request.data['section'] = json.dumps(request.data['section'])
                request.data['unit'] = json.dumps(request.data['unit'][count:])
                request.data['compartment'] = json.dumps(request.data['compartment'][count:])
                request.data['system'] = json.dumps(request.data['system'][count:])
                request.data['equipment'] = json.dumps(request.data['equipment'][count:])
                #request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.FormulationOfApproachPaperResponsibilitySerializer(data = request.data)
                
                if saveserialize.is_valid():
                    #models.FormulationOfApproachPaperResponsibility.objects.filter(trial_id=request.data['trial']).delete()
                    saveserialize.save()

                    # print(saveserialize.data)

                    res_id = saveserialize.data['id']

                    models.FormulationOfApproachPaperResponsibilityCompartment.objects.create(
                    formulation_of_approach_paper_responsibility_id = res_id,
                    compartment = request.data['compartment'],
                    created_ip =  created_ip
                    )

                    models.FormulationOfApproachPaperResponsibilitySystem.objects.create(
                    formulation_of_approach_paper_responsibility_id = res_id,
                    system = request.data['system'],
                    created_ip =  created_ip
                    )

                    models.FormulationOfApproachPaperResponsibilityEquipment.objects.create(
                    formulation_of_approach_paper_responsibility_id = res_id,
                    equipment = request.data['equipment'],
                    created_ip =  created_ip
                    )

                    return Response({"status" : error.context['success_code'], "message":"PSR Formulation Of Approach Paper Responsability" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:
                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id 
                list = self.get_object(pk)
                
                saveserialize = cSerializer.FormulationOfApproachPaperResponsibilitySerializer(list, data = request.data, partial= True)                
                if saveserialize.is_valid():
                    saveserialize.save()
                    logData=request.data

                    return Response({"status" :error.context['success_code'], "message":"PSR Formulation Of Approach Paper" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)




# PSR Presentation Of Approach Paper
class PresentationOfApproachPaperList(APIView):
    
    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.PresentationOfApproachPaper.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.PresentationOfApproachPaperListSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.PresentationOfApproachPaper.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"PSR Presentation Of Approach Paper" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.PresentationOfApproachPaper.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListPresentationOfApproachPaperSerializer(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class PresentationOfApproachPaperCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.PresentationOfApproachPaper.objects.get(pk = pk)
            except models.PresentationOfApproachPaper.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):
        created_ip = Common.get_client_ip(request)
        #pass
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']
            print("Delete",request.data)

            if not request.data['project'] and request.data['status'] != 3:
                return Response({"status":error.context['error_code'], "message": "Project is missing"}, status=status.HTTP_200_OK)
            if not pk or pk=='null':

                request.data['created_ip'] = created_ip
                saveserialize = cSerializer.PresentationOfApproachPaperSerializer(data = request.data)
                
                if saveserialize.is_valid():
                    #models.PresentationOfApproachPaper.objects.filter(trial_id=request.data['trial']).delete()
                    saveserialize.save()

                    # print(saveserialize.data)

                    return Response({"status" : error.context['success_code'], "message":"PSR Presentation Of Approach Paper" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:
                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id 
                list = self.get_object(pk)
                
                saveserialize = cSerializer.PresentationOfApproachPaperSerializer(list, data = request.data, partial= True)                
                if saveserialize.is_valid():
                    saveserialize.save()
                    logData=request.data
                    
                    #Dashboard
                    if 'approved_status' in request.data:
                        models.ProjectModuleStatus.objects.filter(project_id=request.data['project'], project_module_master_id=3).delete()

                        models.ProjectModuleStatus.objects.create(
                            project_module_master_id = 3,
                            project_id = request.data['project'],
                            status  = request.data['approved_status']
                        )

                    return Response({"status" :error.context['success_code'], "message":"PSR Presentation Of Approach Paper" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)

# PSR Presentation Of Approach Paper Send Mail
class PresentationOfApproachPaperSendMailList(APIView):

    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.PresentationOfApproachPaperSendMail.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.PresentationOfApproachPaperSendMailListSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.PresentationOfApproachPaperSendMail.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Initiation Notes" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.PresentationOfApproachPaperSendMail.objects
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListPresentationOfApproachPaperSendMailSerializer(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class PresentationOfApproachPaperSendMailCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.PresentationOfApproachPaperSendMail.objects.get(pk = pk)
            except models.PresentationOfApproachPaperSendMail.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):

        #print('@@@@', request.data)
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']

            if not request.data['to_email']:
                return Response({"status":error.context['error_code'], "message": "To email address is missing"}, status=status.HTTP_200_OK)
            if not request.data['subject']:
                return Response({"status":error.context['error_code'], "message": "Subject is missing"}, status=status.HTTP_200_OK)
            if not pk:
                request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.PresentationOfApproachPaperSendMailSerializer(data = request.data)
                
                if saveserialize.is_valid():
                    saveserialize.save()
                    #print(saveserialize.data)
                    #print(saveserialize.query())

                    subject = saveserialize.data['subject']
                    comments = saveserialize.data['comments']
                    file_name = saveserialize.data['file_name']

                    email = toMailAddress
                    attach = request.FILES['file_name']
                    fhand = open(f"{settings.BASE_DIR}"+file_name,"r")

                    try:
                        mail = EmailMessage(subject, comments, settings.EMAIL_HOST_USER, [email])
                        mail.attach_file(f"{settings.BASE_DIR}"+file_name)
                        mail.send()
                        
                        return Response({"status" : error.context['success_code'], "message":"Send email" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                    except:
                        return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
                        #print('Fail')

                    return Response({"status" : error.context['success_code'], "message":"Send email" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:
                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id 
                list = self.get_object(pk)
                saveserialize = cSerializer.PresentationOfApproachPaperSendMailSerializer(list, data = request.data, partial= True)
                if saveserialize.is_valid():
                    saveserialize.save()
                    logData=request.data
                    return Response({"status" :error.context['success_code'], "message":"Send email" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)

# PSR Presentation Of Approach Paper Document
class PresentationOfApproachPaperDocumentList(APIView):

    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.PresentationOfApproachPaperDocument.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.PresentationOfApproachPaperDocumentListSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.PresentationOfApproachPaperDocument.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Initiation Notes" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.PresentationOfApproachPaperDocument.objects
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListPresentationOfApproachPaperDocumentSerializer(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class PresentationOfApproachPaperDocumentCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.PresentationOfApproachPaperDocument.objects.get(pk = pk)
            except models.PresentationOfApproachPaperDocument.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):
        print(request,'######')

        #request.data['created_ip'] = created_ip
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']  
            if not request.data['document_name']:
                return Response({"status":error.context['error_code'], "message": "Document name is missing"}, status=status.HTTP_200_OK)
            if not request.data['document_remark']:
                return Response({"status":error.context['error_code'], "message": "Document remark is missing"}, status=status.HTTP_200_OK)
            # if not request.data['file_name']:
            #     return Response({"status":error.context['error_code'], "message": "Document file is missing"}, status=status.HTTP_200_OK)
            if not pk:
                #request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.PresentationOfApproachPaperDocumentSerializer(data = request.data)
                
                if saveserialize.is_valid():
                    #models.PresentationOfApproachPaperDocument.objects.filter(trial_id=request.data['trial']).delete()
                    saveserialize.save()
                    #print(saveserialize.query())
                    return Response({"status" : error.context['success_code'], "message":"Document" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:
                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id 
                list = self.get_object(pk)
                saveserialize = cSerializer.PresentationOfApproachPaperDocumentSerializer(list, data = request.data, partial= True)
                if saveserialize.is_valid():
                    saveserialize.save()
                    logData=request.data
                    return Response({"status" :error.context['success_code'], "message":"Document" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)


# PSR Inputs For Staff Requirement
class InputsForStaffRequirementList(APIView):
    
    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.InputsForStaffRequirement.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.InputsForStaffRequirementListSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.InputsForStaffRequirement.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Inputs For Staff Requirement" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        #lists = models.InputsForStaffRequirement.objects
        lists = models.InputsForStaffRequirement.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListInputsForStaffRequirementSerializer(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class InputsForStaffRequirementCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.InputsForStaffRequirement.objects.get(pk = pk)
            except models.InputsForStaffRequirement.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):

        #print(request.data)
        #pass
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']
            if not request.data['project'] and request.data['status'] != 3:
                return Response({"status":error.context['error_code'], "message": "Project is missing"}, status=status.HTTP_200_OK)
            if not pk:

                project = request.data['project']
                request.data['created_ip'] = Common.get_client_ip(request)
                created_ip = request.data['created_ip']
                saveserialize = cSerializer.InputsForStaffRequirementSerializer(data = request.data)
                
                if saveserialize.is_valid():
                    #models.InputsForStaffRequirement.objects.filter(trial_id=request.data['trial']).delete()
                    saveserialize.save()


                    ##################
                    models.InputsForStaffRequirementCompartment.objects.create(
                        project_id = project,
                        ser = request.data['ser_compartment'],
                        name = request.data['name_compartment'],
                        numbers = request.data['numbers_compartment'],
                        weight_volume_and_power_consumption = request.data['weight_volume_and_power_consumption_compartment'],
                        location = request.data['location_compartment'],
                        equipment = request.data['equipment_compartment'],
                        features = request.data['features_compartment'],
                        layout = request.data['layout_compartment'],
                        special_requirements = request.data['special_requirements_compartment'],
                        #standards = request.data['standards_compartment'],
                        recommendations = request.data['recommendations_compartment'],
                        created_ip =  created_ip
                    )

                    models.InputsForStaffRequirementEquipment.objects.create(
                        project_id = project,
                        ser = request.data['ser_equipment'],
                        name = request.data['name_equipment'],
                        numbers = request.data['numbers_equipment'],
                        weight_volume_and_power_consumption = request.data['weight_volume_and_power_consumption_equipment'],
                        location = request.data['location_equipment'],
                        equipment = request.data['equipment_equipment'],
                        features = request.data['features_equipment'],
                        layout = request.data['layout_equipment'],
                        special_requirements = request.data['special_requirements_equipment'],
                        #standards = request.data['standards_equipment'],
                        recommendations = request.data['recommendations_equipment'],
                        created_ip =  created_ip
                    )

                    models.InputsForStaffRequirementSystem.objects.create(
                        project_id = project,
                        ser = request.data['ser_system'],
                        name = request.data['name_system'],
                        numbers = request.data['numbers_system'],
                        weight_volume_and_power_consumption = request.data['weight_volume_and_power_consumption_system'],
                        location = request.data['location_system'],
                        equipment = request.data['equipment_system'],
                        features = request.data['features_system'],
                        layout = request.data['layout_system'],
                        special_requirements = request.data['special_requirements_system'],
                        #standards = request.data['standards_system'],
                        recommendations = request.data['recommendations_system'],
                        created_ip =  created_ip
                    )

                    return Response({"status" : error.context['success_code'], "message":"Inputs For Staff Requirement" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:
                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id 
                list = self.get_object(pk)
                
                saveserialize = cSerializer.InputsForStaffRequirementSerializer(list, data = request.data, partial= True)                
                if saveserialize.is_valid():
                    saveserialize.save()
                    logData=request.data


                    #Dashboard
                    if 'approved_status' in request.data:
                        models.ProjectModuleStatus.objects.filter(project_id=request.data['project'], project_module_master_id=4).delete()

                        models.ProjectModuleStatus.objects.create(
                            project_module_master_id = 4,
                            project_id = request.data['project'],
                            status  = request.data['approved_status']
                        )

                    return Response({"status" :error.context['success_code'], "message":"Inputs For Staff Requirement" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)

# PSR Inputs For Staff Requirement Send Mail
class InputsForStaffRequirementSendMailList(APIView):

    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.InputsForStaffRequirementSendMail.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.InputsForStaffRequirementSendMailSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.InputsForStaffRequirementSendMail.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Inputs For Staff Requirement " +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.InputsForStaffRequirementSendMail.objects
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListInputsForStaffRequirementSendMailSerializer(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class InputsForStaffRequirementSendMailCRUD(APIView):
    def get_object(self, pk):
        
            try:
                return models.InputsForStaffRequirementSendMail.objects.get(pk = pk)
            except models.InputsForStaffRequirementSendMail.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):

        #print('@@@@', request.data)
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']

            if not request.data['to_email']:
                return Response({"status":error.context['error_code'], "message": "To email address is missing"}, status=status.HTTP_200_OK)
            if not request.data['subject']:
                return Response({"status":error.context['error_code'], "message": "Subject is missing"}, status=status.HTTP_200_OK)
            if not pk:
                request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.InputsForStaffRequirementSendMailSerializer(data = request.data)
                
                if saveserialize.is_valid():
                    saveserialize.save()
                    #print(saveserialize.data)
                    #print(saveserialize.query())

                    subject = saveserialize.data['subject']
                    comments = saveserialize.data['comments']
                    file_name = saveserialize.data['file_name']

                    email = toMailAddress
                    attach = request.FILES['file_name']
                    fhand = open(f"{settings.BASE_DIR}"+file_name,"r")

                    try:
                        mail = EmailMessage(subject, comments, settings.EMAIL_HOST_USER, [email])
                        mail.attach_file(f"{settings.BASE_DIR}"+file_name)
                        mail.send()
                        
                        return Response({"status" : error.context['success_code'], "message":"Send email" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                    except:
                        return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
                        #print('Fail')

                    return Response({"status" : error.context['success_code'], "message":"Send email" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:
                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id 
                list = self.get_object(pk)
                saveserialize = cSerializer.InputsForStaffRequirementSendMailSerializer(list, data = request.data, partial= True)
                if saveserialize.is_valid():
                    saveserialize.save()
                    logData=request.data
                    return Response({"status" :error.context['success_code'], "message":"Send email" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)

# PSR Inputs For Staff Requirement Document
class InputsForStaffRequirementDocumentList(APIView):

    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.InputsForStaffRequirementDocument.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cserializer.InputsForStaffRequirementDocumentSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.InputsForStaffRequirementDocument.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Inputs For Staff Requirement " +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.InputsForStaffRequirementDocument.objects
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListInputsForStaffRequirementDocumentSerializer(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class InputsForStaffRequirementDocumentCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.InputsForStaffRequirementDocument.objects.get(pk = pk)
            except models.InputsForStaffRequirementDocument.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):
        print(request,'######')

        #request.data['created_ip'] = created_ip
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']  
            if not request.data['document_name']:
                return Response({"status":error.context['error_code'], "message": "Document name is missing"}, status=status.HTTP_200_OK)
            if not request.data['document_remark']:
                return Response({"status":error.context['error_code'], "message": "Document remark is missing"}, status=status.HTTP_200_OK)
            # if not request.data['file_name']:
            #     return Response({"status":error.context['error_code'], "message": "Document file is missing"}, status=status.HTTP_200_OK)
            if not pk:
                #request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.InputsForStaffRequirementDocumentSerializer(data = request.data)
                
                if saveserialize.is_valid():
                    #models.PresentationOfApproachPaperDocument.objects.filter(trial_id=request.data['trial']).delete()
                    saveserialize.save()
                    #print(saveserialize.query())
                    return Response({"status" : error.context['success_code'], "message":"Document" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:
                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id 
                list = self.get_object(pk)
                saveserialize = cSerializer.InputsForStaffRequirementDocumentSerializer(list, data = request.data, partial= True)
                if saveserialize.is_valid():
                    saveserialize.save()
                    logData=request.data
                    return Response({"status" :error.context['success_code'], "message":"Document" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)



# PSR Inputs For Staff Requirement Compartment
class InputsForStaffRequirementCompartmentList(APIView):
    
    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.InputsForStaffRequirementCompartment.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.InputsForStaffRequirementCompartmentListSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.InputsForStaffRequirementCompartment.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Inputs For Staff Requirement" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        #lists = models.InputsForStaffRequirementCompartment.objects
        lists = models.InputsForStaffRequirementCompartment.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.InputsForStaffRequirementCompartment(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

# PSR Inputs For Staff Requirement Equipment
class InputsForStaffRequirementEquipmentList(APIView):
    
    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.InputsForStaffRequirementEquipment.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.InputsForStaffRequirementEquipmentListSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.InputsForStaffRequirementEquipment.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Inputs For Staff Requirement" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        #lists = models.InputsForStaffRequirementEquipment.objects
        lists = models.InputsForStaffRequirementEquipment.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.InputsForStaffRequirementEquipment(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

# PSR Inputs For Staff Requirement System
class InputsForStaffRequirementSystemList(APIView):
    
    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.InputsForStaffRequirementSystem.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.InputsForStaffRequirementSystemListSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.InputsForStaffRequirementSystem.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Inputs For Staff Requirement" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        #lists = models.InputsForStaffRequirementSystem.objects
        lists = models.InputsForStaffRequirementSystem.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.InputsForStaffRequirementSystem(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)


# PSR Concept Design
class ConceptDesignList(APIView):
    
    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.ConceptDesign.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.ConceptDesignListSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.ConceptDesign.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Concept Design" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        #lists = models.ConceptDesign.objects
        lists = models.ConceptDesign.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListConceptDesignSerializer(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class ConceptDesignCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.ConceptDesign.objects.get(pk = pk)
            except models.ConceptDesign.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):

        #print(request.data)
        #pass
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']
            if not request.data['project'] and request.data['status'] != 3:
                return Response({"status":error.context['error_code'], "message": "Project is missing"}, status=status.HTTP_200_OK)
            if not pk:

                project = request.data['project']
                request.data['created_ip'] = Common.get_client_ip(request)
                created_ip = request.data['created_ip']
                saveserialize = cSerializer.ConceptDesignSerializer(data = request.data)
                
                if saveserialize.is_valid():
                    #models.ConceptDesign.objects.filter(trial_id=request.data['trial']).delete()
                    saveserialize.save()


                    ##################
                    # models.ConceptDesignCompartment.objects.create(
                    #     project_id = project,
                    #     ser = request.data['ser_compartment'],
                    #     name = request.data['name_compartment'],
                    #     numbers = request.data['numbers_compartment'],
                    #     weight_volume_and_power_consumption = request.data['weight_volume_and_power_consumption_compartment'],
                    #     location = request.data['location_compartment'],
                    #     equipment = request.data['equipment_compartment'],
                    #     features = request.data['features_compartment'],
                    #     layout = request.data['layout_compartment'],
                    #     special_requirements = request.data['special_requirements_compartment'],
                    #     #standards = request.data['standards_compartment'],
                    #     recommendations = request.data['recommendations_compartment'],
                    #     created_ip =  created_ip
                    # )

                    # models.ConceptDesignEquipment.objects.create(
                    #     project_id = project,
                    #     ser = request.data['ser_equipment'],
                    #     name = request.data['name_equipment'],
                    #     numbers = request.data['numbers_equipment'],
                    #     weight_volume_and_power_consumption = request.data['weight_volume_and_power_consumption_equipment'],
                    #     location = request.data['location_equipment'],
                    #     equipment = request.data['equipment_equipment'],
                    #     features = request.data['features_equipment'],
                    #     layout = request.data['layout_equipment'],
                    #     special_requirements = request.data['special_requirements_equipment'],
                    #     #standards = request.data['standards_equipment'],
                    #     recommendations = request.data['recommendations_equipment'],
                    #     created_ip =  created_ip
                    # )

                    # models.ConceptDesignSystem.objects.create(
                    #     project_id = project,
                    #     ser = request.data['ser_system'],
                    #     name = request.data['name_system'],
                    #     numbers = request.data['numbers_system'],
                    #     weight_volume_and_power_consumption = request.data['weight_volume_and_power_consumption_system'],
                    #     location = request.data['location_system'],
                    #     equipment = request.data['equipment_system'],
                    #     features = request.data['features_system'],
                    #     layout = request.data['layout_system'],
                    #     special_requirements = request.data['special_requirements_system'],
                    #     #standards = request.data['standards_system'],
                    #     recommendations = request.data['recommendations_system'],
                    #     created_ip =  created_ip
                    # )

                    return Response({"status" : error.context['success_code'], "message":"Concept Design" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:
                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id 
                list = self.get_object(pk)
                
                saveserialize = cSerializer.ConceptDesignSerializer(list, data = request.data, partial= True)                
                if saveserialize.is_valid():
                    saveserialize.save()
                    logData=request.data

                    #Dashboard
                    if 'approved_status' in request.data:
                        models.ProjectModuleStatus.objects.filter(project_id=request.data['project'], project_module_master_id=5).delete()

                        models.ProjectModuleStatus.objects.create(
                            project_module_master_id = 5,
                            project_id = request.data['project'],
                            status  = request.data['approved_status']
                        )

                    return Response({"status" :error.context['success_code'], "message":"Concept Design" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)

# Concept Design Document
class ConceptDesignDocumentList(APIView):

    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.ConceptDesignDocument.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.ConceptDesignDocumentSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.ConceptDesignDocument.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Concept Design" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.ConceptDesignDocument.objects
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListConceptDesignDocumentSerializer(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class ConceptDesignDocumentCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.ConceptDesignDocument.objects.get(pk = pk)
            except models.ConceptDesignDocument.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):

        #print('######@@@@@@@',request.data)
        #pass

        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']  
            if not request.data['document_name']:
                return Response({"status":error.context['error_code'], "message": "Document name is missing"}, status=status.HTTP_200_OK)
            if not request.data['document_remark']:
                return Response({"status":error.context['error_code'], "message": "Document remark is missing"}, status=status.HTTP_200_OK)
            # if not request.data['file_name']:
            #     return Response({"status":error.context['error_code'], "message": "Document file is missing"}, status=status.HTTP_200_OK)
            if not pk:
                request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.ConceptDesignDocumentSerializer(data = request.data)
                
                if saveserialize.is_valid():
                    #models.InitiationNotesDocument.objects.filter(trial_id=request.data['trial']).delete()
                    saveserialize.save()
                    #print(saveserialize.query())
                    return Response({"status" : error.context['success_code'], "message":"Document" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:
                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id 
                list = self.get_object(pk)
                saveserialize = cSerializer.ConceptDesignDocumentSerializer(list, data = request.data, partial= True)
                if saveserialize.is_valid():
                    saveserialize.save()
                    logData=request.data
                    return Response({"status" :error.context['success_code'], "message":"Document" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)


# PSR Concept Design Send Mail
class ConceptDesignSendMailList(APIView):

    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.ConceptDesignSendMail.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.ConceptDesignSendMailListSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.ConceptDesignSendMail.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Concept Design" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.ConceptDesignSendMail.objects
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListConceptDesignSendMailSerializer(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class ConceptDesignSendMailCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.ConceptDesignSendMail.objects.get(pk = pk)
            except models.ConceptDesignSendMail.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):

        #print('@@@@', request.data)
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']

            if not request.data['to_email']:
                return Response({"status":error.context['error_code'], "message": "To email address is missing"}, status=status.HTTP_200_OK)
            if not request.data['subject']:
                return Response({"status":error.context['error_code'], "message": "Subject is missing"}, status=status.HTTP_200_OK)
            if not pk:
                request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.ConceptDesignSendMailSerializer(data = request.data)
                
                if saveserialize.is_valid():
                    saveserialize.save()
                    #print(saveserialize.data)
                    #print(saveserialize.query())

                    subject = saveserialize.data['subject']
                    comments = saveserialize.data['comments']
                    file_name = saveserialize.data['file_name']

                    email = toMailAddress
                    attach = request.FILES['file_name']
                    #print(attach)
                    #print(request.FILES['file_name'])
                    fhand = open(f"{settings.BASE_DIR}"+file_name,"r")
                    #print(settings.BASE_DIR)
                    #print(f"{settings.BASE_DIR}"+file_name,"QQQQ")
                    try:
                        mail = EmailMessage(subject, comments, settings.EMAIL_HOST_USER, [email])
                        mail.attach_file(f"{settings.BASE_DIR}"+file_name)
                        mail.send()
                        #print('OK')
                        return Response({"status" : error.context['success_code'], "message":"Send email" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                    except:
                        return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
                        #print('Fail')

                    return Response({"status" : error.context['success_code'], "message":"Send email" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:
                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id 
                list = self.get_object(pk)
                saveserialize = cSerializer.ConceptDesignSendMailSerializer(list, data = request.data, partial= True)
                if saveserialize.is_valid():
                    saveserialize.save()
                    logData=request.data
                    return Response({"status" :error.context['success_code'], "message":"Send email" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)





# PSR Incorporation of Design Inputs
class IncorporationOfDesignInputsList(APIView):
    
    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.IncorporationOfDesignInputs.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.IncorporationOfDesignInputsListSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.IncorporationOfDesignInputs.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Incorporation Of Design Inputs" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        #lists = models.IncorporationOfDesignInputs.objects
        lists = models.IncorporationOfDesignInputs.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListIncorporationOfDesignInputsSerializer(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class IncorporationOfDesignInputsCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.IncorporationOfDesignInputs.objects.get(pk = pk)
            except models.IncorporationOfDesignInputs.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):

        #print(request.data)
        #pass
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']
            if not request.data['project'] and request.data['status'] != 3:
                return Response({"status":error.context['error_code'], "message": "Project is missing"}, status=status.HTTP_200_OK)
            if not pk:

                project = request.data['project']
                request.data['created_ip'] = Common.get_client_ip(request)
                created_ip = request.data['created_ip']
                saveserialize = cSerializer.IncorporationOfDesignInputsSerializer(data = request.data)
                
                if saveserialize.is_valid():
                    #models.IncorporationOfDesignInputs.objects.filter(trial_id=request.data['trial']).delete()
                    saveserialize.save()


                    ##################
                    # models.IncorporationOfDesignInputsCompartment.objects.create(
                    #     project_id = project,
                    #     ser = request.data['ser_compartment'],
                    #     name = request.data['name_compartment'],
                    #     numbers = request.data['numbers_compartment'],
                    #     weight_volume_and_power_consumption = request.data['weight_volume_and_power_consumption_compartment'],
                    #     location = request.data['location_compartment'],
                    #     equipment = request.data['equipment_compartment'],
                    #     features = request.data['features_compartment'],
                    #     layout = request.data['layout_compartment'],
                    #     special_requirements = request.data['special_requirements_compartment'],
                    #     #standards = request.data['standards_compartment'],
                    #     recommendations = request.data['recommendations_compartment'],
                    #     created_ip =  created_ip
                    # )

                    # models.IncorporationOfDesignInputsEquipment.objects.create(
                    #     project_id = project,
                    #     ser = request.data['ser_equipment'],
                    #     name = request.data['name_equipment'],
                    #     numbers = request.data['numbers_equipment'],
                    #     weight_volume_and_power_consumption = request.data['weight_volume_and_power_consumption_equipment'],
                    #     location = request.data['location_equipment'],
                    #     equipment = request.data['equipment_equipment'],
                    #     features = request.data['features_equipment'],
                    #     layout = request.data['layout_equipment'],
                    #     special_requirements = request.data['special_requirements_equipment'],
                    #     #standards = request.data['standards_equipment'],
                    #     recommendations = request.data['recommendations_equipment'],
                    #     created_ip =  created_ip
                    # )

                    # models.IncorporationOfDesignInputsSystem.objects.create(
                    #     project_id = project,
                    #     ser = request.data['ser_system'],
                    #     name = request.data['name_system'],
                    #     numbers = request.data['numbers_system'],
                    #     weight_volume_and_power_consumption = request.data['weight_volume_and_power_consumption_system'],
                    #     location = request.data['location_system'],
                    #     equipment = request.data['equipment_system'],
                    #     features = request.data['features_system'],
                    #     layout = request.data['layout_system'],
                    #     special_requirements = request.data['special_requirements_system'],
                    #     #standards = request.data['standards_system'],
                    #     recommendations = request.data['recommendations_system'],
                    #     created_ip =  created_ip
                    # )

                    return Response({"status" : error.context['success_code'], "message":"Incorporation Of Design Inputs" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:
                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id 
                list = self.get_object(pk)
                
                saveserialize = cSerializer.IncorporationOfDesignInputsSerializer(list, data = request.data, partial= True)                
                if saveserialize.is_valid():
                    saveserialize.save()
                    logData=request.data

                    #Dashboard
                    if 'approved_status' in request.data:
                        models.ProjectModuleStatus.objects.filter(project_id=request.data['project'], project_module_master_id=6).delete()

                        models.ProjectModuleStatus.objects.create(
                            project_module_master_id = 6,
                            project_id = request.data['project'],
                            status  = request.data['approved_status']
                        )

                    return Response({"status" :error.context['success_code'], "message":"Incorporation Of Design Inputs" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)


# Incorporation of Design Inputs Document
class IncorporationOfDesignInputsDocumentList(APIView):

    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.IncorporationOfDesignInputsDocument.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.IncorporationOfDesignInputsDocumentSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.IncorporationOfDesignInputsDocument.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Incorporation Of Design Inputs" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.IncorporationOfDesignInputsDocument.objects
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListIncorporationOfDesignInputsDocumentSerializer(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class IncorporationOfDesignInputsDocumentCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.IncorporationOfDesignInputsDocument.objects.get(pk = pk)
            except models.IncorporationOfDesignInputsDocument.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):

        #print('######@@@@@@@',request.data)
        #pass

        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']  
            if not request.data['document_name']:
                return Response({"status":error.context['error_code'], "message": "Document name is missing"}, status=status.HTTP_200_OK)
            if not request.data['document_remark']:
                return Response({"status":error.context['error_code'], "message": "Document remark is missing"}, status=status.HTTP_200_OK)
            # if not request.data['file_name']:
            #     return Response({"status":error.context['error_code'], "message": "Document file is missing"}, status=status.HTTP_200_OK)
            if not pk:
                request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.IncorporationOfDesignInputsDocumentSerializer(data = request.data)
                
                if saveserialize.is_valid():
                    #models.InitiationNotesDocument.objects.filter(trial_id=request.data['trial']).delete()
                    saveserialize.save()
                    #print(saveserialize.query())
                    return Response({"status" : error.context['success_code'], "message":"Document" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:
                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id 
                list = self.get_object(pk)
                saveserialize = cSerializer.IncorporationOfDesignInputsDocumentSerializer(list, data = request.data, partial= True)
                if saveserialize.is_valid():
                    saveserialize.save()
                    logData=request.data
                    return Response({"status" :error.context['success_code'], "message":"Document" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)


# PSR Incorporation of Design Inputs Send Mail
class IncorporationOfDesignInputsSendMailList(APIView):

    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.IncorporationOfDesignInputsSendMail.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.IncorporationOfDesignInputsSendMailListSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.IncorporationOfDesignInputsSendMail.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Incorporation Of Design Inputs" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.IncorporationOfDesignInputsSendMail.objects
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListIncorporationOfDesignInputsSendMailSerializer(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class IncorporationOfDesignInputsSendMailCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.IncorporationOfDesignInputsSendMail.objects.get(pk = pk)
            except models.IncorporationOfDesignInputsSendMail.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):

        #print('@@@@', request.data)
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']

            if not request.data['to_email']:
                return Response({"status":error.context['error_code'], "message": "To email address is missing"}, status=status.HTTP_200_OK)
            if not request.data['subject']:
                return Response({"status":error.context['error_code'], "message": "Subject is missing"}, status=status.HTTP_200_OK)
            if not pk:
                request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.IncorporationOfDesignInputsSendMailSerializer(data = request.data)
                
                if saveserialize.is_valid():
                    saveserialize.save()
                    #print(saveserialize.data)
                    #print(saveserialize.query())

                    subject = saveserialize.data['subject']
                    comments = saveserialize.data['comments']
                    file_name = saveserialize.data['file_name']

                    email = toMailAddress
                    attach = request.FILES['file_name']
                    #print(attach)
                    #print(request.FILES['file_name'])
                    fhand = open(f"{settings.BASE_DIR}"+file_name,"r")
                    #print(settings.BASE_DIR)
                    #print(f"{settings.BASE_DIR}"+file_name,"QQQQ")
                    try:
                        mail = EmailMessage(subject, comments, settings.EMAIL_HOST_USER, [email])
                        mail.attach_file(f"{settings.BASE_DIR}"+file_name)
                        mail.send()
                        #print('OK')
                        return Response({"status" : error.context['success_code'], "message":"Send email" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                    except:
                        return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
                        #print('Fail')

                    return Response({"status" : error.context['success_code'], "message":"Send email" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:
                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id 
                list = self.get_object(pk)
                saveserialize = cSerializer.IncorporationOfDesignInputsSendMailSerializer(list, data = request.data, partial= True)
                if saveserialize.is_valid():
                    saveserialize.save()
                    logData=request.data
                    return Response({"status" :error.context['success_code'], "message":"Send email" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)





# PSR Receipt Of RFI Responses
class ReceiptOfRFIResponsesList(APIView):
    
    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.ReceiptOfRFIResponses.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.ReceiptOfRFIResponsesListSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.ReceiptOfRFIResponses.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Initiation Notes" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.ReceiptOfRFIResponses.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListReceiptOfRFIResponsesSerializer(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class ReceiptOfRFIResponsesCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.ReceiptOfRFIResponses.objects.get(pk = pk)
            except models.ReceiptOfRFIResponses.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):

        #print(request.data)
        #pass
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']
            if not request.data['project'] and request.data['status'] != 3:
                return Response({"status":error.context['error_code'], "message": "Project is missing"}, status=status.HTTP_200_OK)
            if not pk:
            
                # request.data['primary_role'] = json.dumps(request.data['primary_role'])
                # request.data['secondary_role'] = json.dumps(request.data['secondary_role'])
                request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.ReceiptOfRFIResponsesSerializer(data = request.data)
                
                if saveserialize.is_valid():
                    saveserialize.save()
                    return Response({"status" : error.context['success_code'], "message":"Receipt of RFI Responses" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:
                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id 
                list = self.get_object(pk)
                
                saveserialize = cSerializer.ReceiptOfRFIResponsesSerializer(list, data = request.data, partial= True)                
                if saveserialize.is_valid():
                    saveserialize.save()
                    logData=request.data

                    #Dashboard
                    if 'approved_status' in request.data:
                        models.ProjectModuleStatus.objects.filter(project_id=request.data['project'], project_module_master_id=7).delete()

                        models.ProjectModuleStatus.objects.create(
                            project_module_master_id = 7,
                            project_id = request.data['project'],
                            status  = request.data['approved_status']
                        )

                    return Response({"status" :error.context['success_code'], "message":"Receipt of RFI Responses" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)

# Receipt Of RFI Responses Document
class ReceiptOfRFIResponsesDocumentList(APIView):

    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.ReceiptOfRFIResponsesDocument.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.ReceiptOfRFIResponsesDocumentSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.ReceiptOfRFIResponsesDocument.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Receipt Of RFI Responses" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.ReceiptOfRFIResponsesDocument.objects
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListReceiptOfRFIResponsesDocumentSerializer(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class ReceiptOfRFIResponsesDocumentCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.ReceiptOfRFIResponsesDocument.objects.get(pk = pk)
            except models.ReceiptOfRFIResponsesDocument.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):

        #print('######@@@@@@@',request.data)
        #pass

        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']  
            if not request.data['document_name']:
                return Response({"status":error.context['error_code'], "message": "Document name is missing"}, status=status.HTTP_200_OK)
            if not request.data['document_remark']:
                return Response({"status":error.context['error_code'], "message": "Document remark is missing"}, status=status.HTTP_200_OK)
            # if not request.data['file_name']:
            #     return Response({"status":error.context['error_code'], "message": "Document file is missing"}, status=status.HTTP_200_OK)
            if not pk:
                request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.ReceiptOfRFIResponsesDocumentSerializer(data = request.data)
                
                if saveserialize.is_valid():
                    #models.InitiationNotesDocument.objects.filter(trial_id=request.data['trial']).delete()
                    saveserialize.save()
                    #print(saveserialize.query())
                    return Response({"status" : error.context['success_code'], "message":"Document" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:
                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id 
                list = self.get_object(pk)
                saveserialize = cSerializer.ReceiptOfRFIResponsesDocumentSerializer(list, data = request.data, partial= True)
                if saveserialize.is_valid():
                    saveserialize.save()
                    logData=request.data
                    return Response({"status" :error.context['success_code'], "message":"Document" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)


# PSR Receipt Of RFI Responses Send Mail
class ReceiptOfRFIResponsesSendMailList(APIView):

    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.ReceiptOfRFIResponsesSendMail.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.ReceiptOfRFIResponsesSendMailListSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.ReceiptOfRFIResponsesSendMail.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Receipt Of RFI Responses" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.ReceiptOfRFIResponsesSendMail.objects
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListReceiptOfRFIResponsesSendMailSerializer(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class ReceiptOfRFIResponsesSendMailCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.ReceiptOfRFIResponsesSendMail.objects.get(pk = pk)
            except models.ReceiptOfRFIResponsesSendMail.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):

        #print('@@@@', request.data)
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']

            if not request.data['to_email']:
                return Response({"status":error.context['error_code'], "message": "To email address is missing"}, status=status.HTTP_200_OK)
            if not request.data['subject']:
                return Response({"status":error.context['error_code'], "message": "Subject is missing"}, status=status.HTTP_200_OK)
            if not pk:
                request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.ReceiptOfRFIResponsesSendMailSerializer(data = request.data)
                
                if saveserialize.is_valid():
                    saveserialize.save()
                    #print(saveserialize.data)
                    #print(saveserialize.query())

                    subject = saveserialize.data['subject']
                    comments = saveserialize.data['comments']
                    file_name = saveserialize.data['file_name']

                    email = toMailAddress
                    attach = request.FILES['file_name']
                    #print(attach)
                    #print(request.FILES['file_name'])
                    fhand = open(f"{settings.BASE_DIR}"+file_name,"r")
                    #print(settings.BASE_DIR)
                    #print(f"{settings.BASE_DIR}"+file_name,"QQQQ")
                    try:
                        mail = EmailMessage(subject, comments, settings.EMAIL_HOST_USER, [email])
                        mail.attach_file(f"{settings.BASE_DIR}"+file_name)
                        mail.send()
                        #print('OK')
                        return Response({"status" : error.context['success_code'], "message":"Send email" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                    except:
                        return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
                        #print('Fail')

                    return Response({"status" : error.context['success_code'], "message":"Send email" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:
                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id 
                list = self.get_object(pk)
                saveserialize = cSerializer.ReceiptOfRFIResponsesSendMailSerializer(list, data = request.data, partial= True)
                if saveserialize.is_valid():
                    saveserialize.save()
                    logData=request.data
                    return Response({"status" :error.context['success_code'], "message":"Send email" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)




# GLS Initiation Notes

class InitiationNotesGLSMasterList(APIView):
    
    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.InitiationNotesGLSMaster.objects.filter(pk=pk).exclude(status='3').get()
                
                serializeobj = serializer.ListInitiationNotesGLSMasterSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.InitiationNotesGLSMaster.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"GLS Initiation Notes Master" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.InitiationNotesGLSMaster.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListInitiationNotesGLSMasterSerializer(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)


class InitiationNotesGLSList(APIView):
    
    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try: 
            if pk:
                list = models.InitiationNotesGLS.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.InitiationNotesGLSListSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.InitiationNotesGLS.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"GLS Initiation Notes" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        #lists = models.InitiationNotesGLS.objects

        project_id = request.GET.get('project_id')
        #lists = models.FormulationOfApproachPaperResponsibility.objects.exclude(status='3')

        listMaster = models.InitiationNotesGLSMaster.objects.values('project_id','type_name').filter(project_id=project_id)
        #listsMaster = models.InitiationNotesGLSMaster.objects.filter(project_id=project_id)
        lists = models.InitiationNotesGLS.objects.filter(project_id=project_id)
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListInitiationNotesGLSSerializer(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data,"master":listMaster}, status=status.HTTP_200_OK)

class InitiationNotesGLSCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.InitiationNotesGLS.objects.get(pk = pk)
            except models.InitiationNotesGLS.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):

        print(request.data)
        #pass
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']
            created_ip = Common.get_client_ip(request)                                       
            request.data['created_ip'] = created_ip
            created_by = request.data['created_by']

            if not request.data['project']:
                return Response({"status":error.context['error_code'], "message": "Project is missing"}, status=status.HTTP_200_OK)

            count = models.InitiationNotesGLSMaster.objects.filter(project_id=request.data['project']).count()

            if count > 0:
                approveRes = models.InitiationNotesGLSMaster.objects.values('approved_status','approved_by','approved_on', 'approved_remark','project_id').filter(project_id = request.data['project']).get()

                approved_status = approveRes['approved_status']
                approved_by = approveRes['approved_by']
                approved_on = approveRes['approved_on']
                approved_remark = approveRes['approved_remark']


            models.InitiationNotesGLSMaster.objects.filter(project_id=request.data['project']).delete()
            if not pk:

                if count > 0:
                    master_gls = {
                        'project': request.data['project'],
                        'type_name' : request.data['type_name'],
                        'created_ip': created_ip,
                        'created_by': created_by,
                        "approved_status" : approved_status,
                        "approved_by" : approved_by,
                        "approved_on" : approved_on,
                        "approved_remark" : approved_remark
                    }
                else:
                    master_gls = {
                        'project': request.data['project'],
                        'type_name' : request.data['type_name'],
                        'created_ip': created_ip,
                        'created_by': created_by
                    }

                saveserializelog = cSerializer.InitiationNotesGLSMasterSerializer(data = master_gls)
                if saveserializelog.is_valid():
                    saveserializelog.save()
                    running_id = saveserializelog.data['id']


                #request.data['initiation_notes_gls_master'] = running_id
                #saveserialize = cSerializer.InitiationNotesGLSSerializer(data = request.data)

                for row in request.data['sections1']:
                    request_data = {
                        'initiation_notes_gls_master' : running_id,
                        'project' : request.data['project'],
                        'type_name' : request.data['type_name'],
                        'document_sections' : row['document_sections'],
                        'document_sub_sections' : row['document_sub_sections'],
                        #'document_sub_sections2' : row['document_sub_sections2'],
                        #'annexures' : row['annexures'],
                        'paragraph_title' : row['paragraph_title'],
                        'paragraph' : row['paragraph'],
                        'primary' : row['primary'],
                        'secondary1' : row['secondary1'],
                        'secondary2' : row['secondary2'],
                        'secondary3' : row['secondary3'],                    
                        'created_ip' : created_ip,
                        'created_by' : created_by
                    }
            
                    saveserialize = cSerializer.InitiationNotesGLSSerializer(data = request_data)
                    if saveserialize.is_valid():
                        saveserialize.save()


                        #saveDataFlowGeneral(request.data['project'],request.data)

                
                if saveserializelog.is_valid():
                    #models.InitiationNotesGLS.objects.filter(trial_id=request.data['trial']).delete()
                    #saveserialize.save()
                    if count > 0:
                        return Response({"status" : error.context['success_code'], "message":"GLS Initiation notes" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                    else:
                        return Response({"status" : error.context['success_code'], "message":"GLS Initiation notes" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:
                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id 
                list = self.get_object(pk)
                
                saveserialize = cSerializer.InitiationNotesGLSSerializer(list, data = request.data, partial= True)                
                if saveserialize.is_valid():
                    saveserialize.save()
                    logData=request.data
                    return Response({"status" :error.context['success_code'], "message":"Initiation notes" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)



class InitiationNotesGLSMasterCURD(APIView):
    def get_object(self, pk):
            try:
                return models.InitiationNotesGLSMaster.objects.get(pk = pk)
            except models.InitiationNotesGLSMaster.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):

        #print(request.data)
        #pass
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']
            if not request.data['project'] and request.data['status'] != 3:
                return Response({"status":error.context['error_code'], "message": "Project is missing"}, status=status.HTTP_200_OK)
            if not pk:
            
                request.data['primary_role'] = json.dumps(request.data['primary_role'])
                request.data['secondary_role'] = json.dumps(request.data['secondary_role'])
                request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.InitiationNotesGLSMasterSerializer(data = request.data)
                
                if saveserialize.is_valid():
                    #models.InitiationNotesGLSMaster.objects.filter(trial_id=request.data['trial']).delete()
                    saveserialize.save()
                    '''
                    logData=request.data
                    logData['running_id']=saveserialize.data['id']
                    saveserializelog = cSerializer.InitiationNotesGLSMasterlogSerializer(data = logData)
                    if saveserializelog.is_valid():
                        saveserializelog.save()
                    trials=InitiationNotesGLSMasterListSerializer(models.InitiationNotesGLSMaster.objects.filter(id=saveserialize.data['id']), many=True).data[0]
                    '''
                    return Response({"status" : error.context['success_code'], "message":"Initiation notes" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:
                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id 
                list = self.get_object(pk)
                
                saveserialize = cSerializer.InitiationNotesGLSMasterSerializer(list, data = request.data, partial= True)                
                if saveserialize.is_valid():
                    saveserialize.save()
                    logData=request.data

                    #Dashboard
                    if 'approved_status' in request.data:
                        qs = models.ProjectModuleStatus.objects.filter(project_id=request.data['project'], project_module_master_id=8).delete()

                        #print(qs.query,"My query")
                        models.ProjectModuleStatus.objects.create(
                            project_module_master_id = 8,
                            project_id = request.data['project'],
                            status  = request.data['approved_status']
                        )
                    '''
                    logData['running_id']=saveserialize.data['id']
                    saveserializelog = InitiationNotesGLSMasterlogSerializer(data = logData)
                    if saveserializelog.is_valid():
                        saveserializelog.save()
                    trials=InitiationNotesGLSMasterListSerializer(models.InitiationNotesGLSMaster.objects.filter(id=saveserialize.data['id']), many=True).data[0]
                    '''
                    return Response({"status" :error.context['success_code'], "message":"Initiation notes" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)

class deleteInitiationNotesGLSMaster(APIView):

    def post(self,request, pk = None):
        if 'id' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "id" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        else: 
            if models.InitiationNotesGLSMaster.objects.filter(id=request.data['id']).count()>0:
                models.InitiationNotesGLSMaster.objects.filter(id=request.data['id']).update(**{"status":3})
                return Response({"status" :error.context['success_code'], "message":"GLS Initiation Note deleted successfully"}, status=status.HTTP_200_OK)
            else:
                return Response({"status" :error.context['error_code'], "message":"Incorrect id supplied in the parameter"}, status=status.HTTP_200_OK)

class InitiationNotesBLSMasterCURD(APIView):
    def get_object(self, pk):
            try:
                return models.InitiationNotesBLSMaster.objects.get(pk = pk)
            except models.InitiationNotesBLSMaster.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):

        #print(request.data)
        #pass
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']
            if not request.data['project'] and request.data['status'] != 3:
                return Response({"status":error.context['error_code'], "message": "Project is missing"}, status=status.HTTP_200_OK)
            if not pk:
            
                request.data['primary_role'] = json.dumps(request.data['primary_role'])
                request.data['secondary_role'] = json.dumps(request.data['secondary_role'])
                request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.InitiationNotesBLSMasterSerializer(data = request.data)
                
                if saveserialize.is_valid():
                    #models.InitiationNotesBLSMaster.objects.filter(trial_id=request.data['trial']).delete()
                    saveserialize.save()

                    return Response({"status" : error.context['success_code'], "message":"Initiation notes" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:
                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id 
                list = self.get_object(pk)
                
                saveserialize = cSerializer.InitiationNotesBLSMasterSerializer(list, data = request.data, partial= True)                
                if saveserialize.is_valid():
                    saveserialize.save()
                    logData=request.data

                    #Dashboard
                    if 'approved_status' in request.data:
                        models.ProjectModuleStatus.objects.filter(project_id=request.data['project'], project_module_master_id=9).delete()

                        models.ProjectModuleStatus.objects.create(
                            project_module_master_id = 9,
                            project_id = request.data['project'],
                            status  = request.data['approved_status']
                        )

                    return Response({"status" :error.context['success_code'], "message":"Initiation notes" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)

class deleteInitiationNotesBLSMaster(APIView):

    def post(self,request, pk = None):
        if 'id' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "id" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        else: 
            if models.InitiationNotesBLSMaster.objects.filter(id=request.data['id']).count()>0:
                models.InitiationNotesBLSMaster.objects.filter(id=request.data['id']).update(**{"status":3})
                return Response({"status" :error.context['success_code'], "message":"BLS Initiation Note deleted successfully"}, status=status.HTTP_200_OK)
            else:
                return Response({"status" :error.context['error_code'], "message":"Incorrect id supplied in the parameter"}, status=status.HTTP_200_OK)


class ExportImportExcelGLS(APIView):

    def get(self, request):
        annexures_obj = models.Annexures.objects.all()
        serializer = cSerializer.AnnexuresSerializer(annexures_obj, many=True)
        df = pd.DataFrame(serializer.data)
        print(df)
        df = df.to_csv(f"{settings.BASE_DIR}/media/Excel/GLS/{uuid.uuid4()}.csv", encoding="UTF-8", index=False)
        return Response({'status':200})

    def post(self,request, pk = None): 
        
        created_ip = Common.get_client_ip(request)
        request_file = request.FILES['excel_file_upload']
        project = request.data['project']
        start_date = request.data['start_date']
        close_date = request.data['close_date']
        type_name = request.data['type_name']
        created_by = request.data['created_by']

        dir_storage='static/import_excel'
        fs = FileSystemStorage(location=dir_storage)
        filename = fs.save(request_file.name, request_file)
        if os.path.splitext(request_file.name)[1] == ".xls" or  os.path.splitext(request_file.name)[1] == ".xlsx":
            excel_folder = os.path.join(settings.BASE_DIR, 'media/Excel/GLS/')
            read_file = pd.read_excel(request_file)
            read_file.to_csv(excel_folder +'import_excel_file.csv')
            fhand = open('media/Excel/GLS/import_excel_file.csv')
        else:
             return Response({"status":error.context['error_code'],"message" : "File format not supported (Xls and Xlsx only allowed)" })    
        reader = csv.reader(fhand)
        next(reader)
        #print(reader)

        master_gls = {
            'project': project,
            'type_name' : type_name,
            'created_ip': created_ip,
            'created_by': created_by
        }

        saveserializelog = cSerializer.InitiationNotesGLSMasterSerializer(data = master_gls)
        if saveserializelog.is_valid():
            saveserializelog.save()
            running_id = saveserializelog.data['id']
        else:
            return Response({"status":error.context['error_code'],"message" : "Provide data for all the required fields" })

        
        for row in reader:
            #print(row[1])
            if not project:
                return Response({"status":error.context['error_code'],"message" : "Project Code is required" })

            #project = masterModels.Project.objects.filter(code=row[1]).first()
            document_sections = models.DocumentSections.objects.filter(code=row[1]).first()            
            document_sub_sections = models.DocumentSubSections.objects.filter(code=row[2]).first()
            document_sub_sections2 = models.DocumentSubSections2.objects.filter(code=row[3]).first()

            if not document_sections:
                document_sections = None
            else:
                document_sections = document_sections.id

            if not document_sub_sections:
                document_sub_sections = None
            else:
                document_sub_sections = document_sub_sections.id

            if not document_sub_sections2:
                document_sub_sections2 = None
            else:
                document_sub_sections2 = document_sub_sections2.id

            if not row[4]:
                annexures = None
            else:
                annexures = row[4]


            request_data = {
                'initiation_notes_gls_master': running_id,
                'project' : project,
                'type_name' : type_name,
                'document_sections' : document_sections,
                'document_sub_sections' : document_sub_sections,
                'document_sub_sections2' : document_sub_sections2,
                'annexures' : annexures,
                'paragraph_title' : row[5],
                'paragraph' : row[6],
                'primary' : row[7],
                'secondary1' : row[8],
                'secondary2' : row[9],
                'secondary3' : row[10],
                'start_date' : start_date,
                'close_date' : close_date,
                'created_ip': created_ip,
                'created_by': created_by
            }

            saveserialize = cSerializer.InitiationNotesGLSSerializer(data = request_data)
            if saveserialize.is_valid():
                saveserialize.save()

        exceled_upload_obj = models.ExcelFileUpload.objects.create(
        excel_file_upload = request.data['excel_file_upload'],
        created_ip =  created_ip
        )
        
        #saveserialize = cSerializer.InitiationNotesGLSSerializer(data = request_data)
        if exceled_upload_obj:
            #saveserialize.save()
            #logData=request.data
            return Response({"status" :error.context['success_code'], "message":"File imported successfully"}, status=status.HTTP_200_OK)
        else:
            return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)



# BLS Initiation Notes



class GetGLSInitiationNotesList(APIView):
    
    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.InitiationNotesBLSMaster.objects.filter(pk=pk).exclude(status='3').get()
                
                serializeobj = serializer.ListInitiationNotesBLSMasterSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.InitiationNotesBLSMaster.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"BLS Initiation Notes Master" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.InitiationNotesBLSMaster.objects
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           




        project_id = request.GET.get('project_id')
        gls_res = models.InitiationNotesGLSMaster.objects.values('id','type_name','project_id').filter(project_id=project_id)

        gls_id = gls_res[0]['id']




        gls = models.InitiationNotesGLS.objects.values('initiation_notes_gls_master','project','type_name','document_sections','document_sub_sections','document_sub_sections2','annexures','paragraph_title','paragraph','primary','secondary1','secondary2','secondary3').filter(initiation_notes_gls_master_id=gls_id).exclude(status='3')

        




    # initiation_notes_gls_master = models.ForeignKey(InitiationNotesGLSMaster, on_delete=models.CASCADE,null=True)
    # project = models.ForeignKey(Project, on_delete=models.CASCADE,null=True)
    # type_name = models.SmallIntegerField(choices=((1,'Section'),(2,'Annexure')))

    # document_sections = models.ForeignKey(DocumentSections, on_delete=models.CASCADE,null=True)
    # document_sub_sections = models.ForeignKey(DocumentSubSections, on_delete=models.CASCADE,null=True)
    # document_sub_sections2 = models.ForeignKey(DocumentSubSections2, on_delete=models.CASCADE,null=True)
    # annexures = models.ForeignKey(Annexures, on_delete=models.CASCADE,null=True)
    # paragraph_title = models.CharField(max_length=255, blank=True, null=True)
    # paragraph = models.TextField(null=True, blank=True)
    # primary = models.IntegerField(null=True)
    # secondary1 = models.IntegerField(null=True)
    # secondary2 = models.IntegerField(null=True)
    # secondary3 = models.IntegerField(null=True)




        #serializer = cSerializer.ListInitiationNotesBLSMasterSerializer(lists, many=True)
        #return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)
        return Response({"status":error.context['success_code'], "data": gls}, status=status.HTTP_200_OK)


# GLS Initiation Notes Document
class GLSInitiationNotesDocumentList(APIView):

    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.GLSInitiationNotesDocument.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.GLSInitiationNotesDocumentListSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.GLSInitiationNotesDocument.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"GLS Initiation Notes" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.GLSInitiationNotesDocument.objects
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListGLSInitiationNotesDocumentSerializer(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class GLSInitiationNotesDocumentCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.GLSInitiationNotesDocument.objects.get(pk = pk)
            except models.GLSInitiationNotesDocument.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):

        #print('######@@@@@@@',request.data)
        #pass

        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']  
            if not request.data['document_name']:
                return Response({"status":error.context['error_code'], "message": "Document name is missing"}, status=status.HTTP_200_OK)
            if not request.data['document_remark']:
                return Response({"status":error.context['error_code'], "message": "Document remark is missing"}, status=status.HTTP_200_OK)
            # if not request.data['file_name']:
            #     return Response({"status":error.context['error_code'], "message": "Document file is missing"}, status=status.HTTP_200_OK)
            if not pk:
                request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.GLSInitiationNotesDocumentSerializer(data = request.data)
                
                if saveserialize.is_valid():
                    #models.InitiationNotesDocument.objects.filter(trial_id=request.data['trial']).delete()
                    saveserialize.save()
                    #print(saveserialize.query())
                    return Response({"status" : error.context['success_code'], "message":"Document" +language.context[language.defaultLang]['insert'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:
                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id 
                list = self.get_object(pk)
                saveserialize = cSerializer.GLSInitiationNotesDocumentSerializer(list, data = request.data, partial= True)
                if saveserialize.is_valid():
                    saveserialize.save()
                    logData=request.data
                    return Response({"status" :error.context['success_code'], "message":"Document" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)


# Initiation Notes Send Mail
class GLSInitiationNotesSendMailList(APIView):

    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.GLSInitiationNotesSendMail.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.GLSInitiationNotesSendMailListSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.GLSInitiationNotesSendMail.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"GLS Initiation Notes" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.GLSInitiationNotesSendMail.objects
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListGLSInitiationNotesSendMailSerializer(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class GLSInitiationNotesSendMailCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.GLSInitiationNotesSendMail.objects.get(pk = pk)
            except models.GLSInitiationNotesSendMail.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):

        #print('@@@@', request.data)
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']

            if not request.data['to_email']:
                return Response({"status":error.context['error_code'], "message": "To email address is missing"}, status=status.HTTP_200_OK)
            if not request.data['subject']:
                return Response({"status":error.context['error_code'], "message": "Subject is missing"}, status=status.HTTP_200_OK)
            if not pk:
                request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.GLSInitiationNotesSendMailSerializer(data = request.data)
                
                if saveserialize.is_valid():
                    saveserialize.save()
                    #print(saveserialize.data)
                    #print(saveserialize.query())

                    subject = saveserialize.data['subject']
                    comments = saveserialize.data['comments']
                    file_name = saveserialize.data['file_name']
                    initiation_notes_gls_master=saveserialize.data['initiation_notes_gls_master']

                    email = toMailAddress
                    attach = request.FILES['file_name']
                    #print(attach)
                    #print(request.FILES['file_name'])
                    fhand = open(f"{settings.BASE_DIR}"+file_name,"r")
                    #print(settings.BASE_DIR)
                    #print(f"{settings.BASE_DIR}"+file_name,"QQQQ")
                    try:
                        mail = EmailMessage(subject, comments, settings.EMAIL_HOST_USER, [email])
                        mail.attach_file(f"{settings.BASE_DIR}"+file_name)
                        mail.send()
                        #print('OK')
                        return Response({"status" : error.context['success_code'], "message":"Send email" +language.context[language.defaultLang]['insert'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                    except:
                        return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
                        #print('Fail')

                    return Response({"status" : error.context['success_code'], "message":"Send email" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:
                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id 
                list = self.get_object(pk)
                saveserialize = cSerializer.GLSInitiationNotesSendMailSerializer(list, data = request.data, partial= True)
                if saveserialize.is_valid():
                    saveserialize.save()
                    logData=request.data
                    return Response({"status" :error.context['success_code'], "message":"Send email" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)



class InitiationNotesBLSMasterList(APIView):
    
    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.InitiationNotesBLSMaster.objects.filter(pk=pk).exclude(status='3').get()
                
                serializeobj = serializer.ListInitiationNotesBLSMasterSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.InitiationNotesBLSMaster.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"BLS Initiation Notes Master" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.InitiationNotesBLSMaster.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListInitiationNotesBLSMasterSerializer(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)




class InitiationNotesBLSList(APIView):
    
    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.InitiationNotesBLS.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.InitiationNotesBLSListSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.InitiationNotesBLS.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"BLS Initiation Notes" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.InitiationNotesBLS.objects
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListInitiationNotesBLSSerializer(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class InitiationNotesBLSCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.InitiationNotesBLS.objects.get(pk = pk)
            except models.InitiationNotesBLS.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):

        #print(request.data)
        #pass
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:

            count = models.InitiationNotesBLSMaster.objects.filter(project_id=request.data['project']).count()

            if count > 0:
                approveRes = models.InitiationNotesBLSMaster.objects.values('approved_status','approved_by','approved_on', 'approved_remark','project_id').filter(project_id = request.data['project']).get()

                approved_status = approveRes['approved_status']
                approved_by = approveRes['approved_by']
                approved_on = approveRes['approved_on']
                approved_remark = approveRes['approved_remark']

            models.InitiationNotesBLSMaster.objects.filter(project_id=request.data['project']).delete()

            pk = request.data['id']
            created_ip = Common.get_client_ip(request)                                       
            request.data['created_ip'] = created_ip
            created_by = request.data['created_by']

            if not request.data['project']:
                return Response({"status":error.context['error_code'], "message": "Project is missing"}, status=status.HTTP_200_OK)
            if not pk:

                if count > 0:
                    master_gls = {
                        'project': request.data['project'],
                        'type_name' : request.data['type_name'],
                        'created_ip': created_ip,
                        'created_by':created_by,

                        "approved_status" : approved_status,
                        "approved_by" : approved_by,
                        "approved_on" : approved_on,
                        "approved_remark" : approved_remark
                    }
                else:
                    master_gls = {
                        'project': request.data['project'],
                        'type_name' : request.data['type_name'],
                        'created_ip': created_ip,
                        'created_by':created_by
                    }

                saveserializelog = cSerializer.InitiationNotesBLSMasterSerializer(data = master_gls)
                if saveserializelog.is_valid():
                    saveserializelog.save()
                    running_id = saveserializelog.data['id']


                request.data['initiation_notes_gls_master'] = running_id
                
                #print(request.data['sections1'],"@@@@@@1111111")
                for row in request.data['sections1']:
                    request_data = {
                        'initiation_notes_bls_master' : running_id,
                        'project' : request.data['project'],
                        'type_name' : request.data['type_name'],
                        'document_sections' : row['document_sections'],
                        'document_sub_sections' : row['document_sub_sections'],
                        'document_sub_sections2' : row['document_sub_sections2'],
                        #'annexures' : row['annexures'],
                        'paragraph_title' : row['paragraph_title'],
                        'paragraph' : row['paragraph'],
                        'primary' : row['primary'],
                        'secondary1' : row['secondary1'],
                        'secondary2' : row['secondary2'],
                        'secondary3' : row['secondary3'],                    
                        'created_ip' : created_ip,
                        'created_by' : created_by
                    }
            
                    saveserialize = cSerializer.InitiationNotesBLSSerializer(data = request_data)
                    if saveserialize.is_valid():
                        saveserialize.save()

                #saveserialize = cSerializer.InitiationNotesBLSSerializer(data = request.data)
                
                if saveserializelog.is_valid():
                    #models.InitiationNotesBLS.objects.filter(trial_id=request.data['trial']).delete()
                    #saveserialize.save()
                    return Response({"status" : error.context['success_code'], "message":"BLS Initiation notes" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:
                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id 
                list = self.get_object(pk)
                
                saveserialize = cSerializer.InitiationNotesBLSSerializer(list, data = request.data, partial= True)                
                if saveserialize.is_valid():
                    saveserialize.save()
                    logData=request.data
                    if count > 0:
                        return Response({"status" :error.context['success_code'], "message":"Initiation notes" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                    else:
                        return Response({"status" :error.context['success_code'], "message":"Initiation notes" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)

class ExportImportExcelBLS(APIView):

    def get(self, request):
        annexures_obj = models.Annexures.objects.all()
        serializer = cSerializer.AnnexuresSerializer(annexures_obj, many=True)
        df = pd.DataFrame(serializer.data) 
        print(df)
        df = df.to_csv(f"{settings.BASE_DIR}/media/Excel/BLS/{uuid.uuid4()}.csv", encoding="UTF-8", index=False)
        return Response({'status':200})

    def post(self,request, pk = None): 
        
        created_ip = Common.get_client_ip(request)
        request_file = request.FILES['excel_file_upload']
        project = request.data['project']
        start_date = request.data['start_date']
        close_date = request.data['close_date']
        type_name = request.data['type_name']
        created_by = request.data['created_by']

        dir_storage='static/import_excel'
        fs = FileSystemStorage(location=dir_storage)
        filename = fs.save(request_file.name, request_file)
        if os.path.splitext(request_file.name)[1] == ".xls" or  os.path.splitext(request_file.name)[1] == ".xlsx":
            excel_folder = os.path.join(settings.BASE_DIR, 'media/Excel/BLS/')
            read_file = pd.read_excel(request_file)
            read_file.to_csv(excel_folder +'import_excel_file.csv')
            fhand = open('media/Excel/BLS/import_excel_file.csv')
        else:
             return Response({"status":error.context['error_code'],"message" : "File format not supported (Xls and Xlsx only allowed)" })    
        reader = csv.reader(fhand)
        next(reader)
        #print(reader)

        master_bls = {
            'project': project,
            'type_name' : type_name,
            'created_ip': created_ip,
            'created_by': created_by
        }

        saveserializelog = cSerializer.InitiationNotesBLSMasterSerializer(data = master_bls)
        if saveserializelog.is_valid():
            saveserializelog.save()
            running_id = saveserializelog.data['id']
        else:
            return Response({"status":error.context['error_code'],"message" : "Provide data for all the required fields" })

        
        for row in reader:
            #print(row[1])
            if not project:
                return Response({"status":error.context['error_code'],"message" : "Project Code is required" })

            #project = masterModels.Project.objects.filter(code=row[1]).first()
            document_sections = models.DocumentSections.objects.filter(code=row[1]).first()            
            document_sub_sections = models.DocumentSubSections.objects.filter(code=row[2]).first()
            document_sub_sections2 = models.DocumentSubSections2.objects.filter(code=row[3]).first()

            if not document_sections:
                document_sections = None
            else:
                document_sections = document_sections.id

            if not document_sub_sections:
                document_sub_sections = None
            else:
                document_sub_sections = document_sub_sections.id

            if not document_sub_sections2:
                document_sub_sections2 = None
            else:
                document_sub_sections2 = document_sub_sections2.id

            if not row[4]:
                annexures = None
            else:
                annexures = row[4]


            request_data = {
                'initiation_notes_bls_master': running_id,
                'project' : project,
                'type_name' : type_name,
                'document_sections' : document_sections,
                'document_sub_sections' : document_sub_sections,
                'document_sub_sections2' : document_sub_sections2,
                'annexures' : annexures,
                'paragraph_title' : row[5],
                'paragraph' : row[6],
                'primary' : row[7],
                'secondary1' : row[8],
                'secondary2' : row[9],
                'secondary3' : row[10],
                'start_date' : start_date,
                'close_date' : close_date,
                'created_ip': created_ip,
                'created_by': created_by
            }

            saveserialize = cSerializer.InitiationNotesBLSSerializer(data = request_data)
            if saveserialize.is_valid():
                saveserialize.save()

        exceled_upload_obj = models.ExcelFileUpload.objects.create(
        excel_file_upload = request.data['excel_file_upload'],
        created_ip =  created_ip
        )
        
        #saveserialize = cSerializer.InitiationNotesBLSSerializer(data = request_data)
        if exceled_upload_obj:
            #saveserialize.save()
            #logData=request.data
            return Response({"status" :error.context['success_code'], "message":"File imported successfully"}, status=status.HTTP_200_OK)
        else:
            return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)


# BLS Initiation Notes Document
class BLSInitiationNotesDocumentList(APIView):

    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.BLSInitiationNotesDocument.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.BLSInitiationNotesDocumentListSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.BLSInitiationNotesDocument.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"BLS Initiation Notes" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.BLSInitiationNotesDocument.objects
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListBLSInitiationNotesDocumentSerializer(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class BLSInitiationNotesDocumentCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.BLSInitiationNotesDocument.objects.get(pk = pk)
            except models.BLSInitiationNotesDocument.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):

        #print('######@@@@@@@',request.data)
        #pass

        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']  
            if not request.data['document_name']:
                return Response({"status":error.context['error_code'], "message": "Document name is missing"}, status=status.HTTP_200_OK)
            if not request.data['document_remark']:
                return Response({"status":error.context['error_code'], "message": "Document remark is missing"}, status=status.HTTP_200_OK)
            # if not request.data['file_name']:
            #     return Response({"status":error.context['error_code'], "message": "Document file is missing"}, status=status.HTTP_200_OK)
            if not pk:
                request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.BLSInitiationNotesDocumentSerializer(data = request.data)
                
                if saveserialize.is_valid():
                    #models.InitiationNotesDocument.objects.filter(trial_id=request.data['trial']).delete()
                    saveserialize.save()
                    #print(saveserialize.query())
                    return Response({"status" : error.context['success_code'], "message":"Document" +language.context[language.defaultLang]['insert'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:
                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id 
                list = self.get_object(pk)
                saveserialize = cSerializer.BLSInitiationNotesDocumentSerializer(list, data = request.data, partial= True)
                if saveserialize.is_valid():
                    saveserialize.save()
                    logData=request.data
                    return Response({"status" :error.context['success_code'], "message":"Document" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)


# Initiation Notes Send Mail
class BLSInitiationNotesSendMailList(APIView):

    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.BLSInitiationNotesSendMail.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.BLSInitiationNotesSendMailListSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.BLSInitiationNotesSendMail.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"BLS Initiation Notes" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.BLSInitiationNotesSendMail.objects
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListBLSInitiationNotesSendMailSerializer(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class BLSInitiationNotesSendMailCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.BLSInitiationNotesSendMail.objects.get(pk = pk)
            except models.BLSInitiationNotesSendMail.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):

        #print('@@@@', request.data)
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']

            if not request.data['to_email']:
                return Response({"status":error.context['error_code'], "message": "To email address is missing"}, status=status.HTTP_200_OK)
            if not request.data['subject']:
                return Response({"status":error.context['error_code'], "message": "Subject is missing"}, status=status.HTTP_200_OK)
            if not pk:
                request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.BLSInitiationNotesSendMailSerializer(data = request.data)
                
                if saveserialize.is_valid():
                    saveserialize.save()
                    #print(saveserialize.data)
                    #print(saveserialize.query())

                    subject = saveserialize.data['subject']
                    comments = saveserialize.data['comments']
                    file_name = saveserialize.data['file_name']
                    initiation_notes_bls_master=saveserialize.data['initiation_notes_bls_master']

                    email = toMailAddress
                    attach = request.FILES['file_name']
                    #print(attach)
                    #print(request.FILES['file_name'])
                    fhand = open(f"{settings.BASE_DIR}"+file_name,"r")
                    #print(settings.BASE_DIR)
                    #print(f"{settings.BASE_DIR}"+file_name,"QQQQ")
                    try:
                        mail = EmailMessage(subject, comments, settings.EMAIL_HOST_USER, [email])
                        mail.attach_file(f"{settings.BASE_DIR}"+file_name)
                        mail.send()
                        #print('OK')
                        return Response({"status" : error.context['success_code'], "message":"Send email" +language.context[language.defaultLang]['insert'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                    except:
                        return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
                        #print('Fail')

                    return Response({"status" : error.context['success_code'], "message":"Send email" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:
                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id 
                list = self.get_object(pk)
                saveserialize = cSerializer.BLSInitiationNotesSendMailSerializer(list, data = request.data, partial= True)
                if saveserialize.is_valid():
                    saveserialize.save()
                    logData=request.data
                    return Response({"status" :error.context['success_code'], "message":"Send email" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)





class ListGlobalStatus1(APIView):

    def get(self, request):



        annexures_obj = masterModels.Project.raw('SELECT t1.id, t1.name, t2.project_id FROM "master.project" as t1 INNER JOIN "psr.initiation_notes" as t2 ON t1.id = t2.project_id group by t2.project_id, t1.id')
        serializer = cSerializer.ListGlobalStatusSerializer(annexures_obj, many=True)
        #df = pd.DataFrame(serializer.data)
        print(serializer.data)
        #df = df.to_csv(f"{settings.BASE_DIR}/media/Excel/GLS/{uuid.uuid4()}.csv", encoding="UTF-8", index=False)
        #return Response({'status':200})

        return Response({"status":error.context['success_code'], "data1": serializer.data}, status=status.HTTP_200_OK)










class ListGlobalStatus(APIView):
    
    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                #list = models.InitiationNotes.objects.filter(pk=pk).exclude(status='3').get()
                list = models.InitiationNotes.objects.raw('SELECT t1.id, t1.name, t2.project_id FROM "master.project" as t1 INNER JOIN "psr.initiation_notes" as t2 ON t1.id = t2.project_idgroup by t2.project_id, t1.id')
                serializeobj = serializer.ListGlobalStatusSerializer(list)
                print(serializeobj.query())
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
      
        except models.InitiationNotes.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Initiation Notes" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        #lists = models.InitiationNotes.objects
        #lists = models.InitiationNotes.objects.exclude(status='3')
        #lists = models.InitiationNotes.objects.select_related('master.project').all()
        #lists = models.InitiationNotes.objects.raw('SELECT t1.id, t1.name, t2.project_id FROM "master.project" as t1 INNER JOIN "psr.initiation_notes" as t2 ON t1.id = t2.project_id group by t2.project_id, t1.id')
        #ApplicantPersonal.objects.values('id','first_name','family_name','visa_number','gender__id','is_pregnant').filter(id=applicant_id).first()
        #lists = models.InitiationNotes.objects.values('project__name')
        lists = models.InitiationNotes.objects.values('approved_status','project__id','project__name').filter(approved_status='2')
       # lists = models.InitiationNotes.objects.values('project__id','project__name')
        #lists = models.InitiationNotes.objects.all().select_related('project').filter(approved_status='2')
        
        print(lists,"EEEE")
        # if normal_values:
        #     lists = lists.filter(reduce(operator.and_, 
        #                        (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        # if array_values:
        #     for key,values in array_values.items():
        #         queries= [Q(**{"%s__contains" % key: value }) for value in values]
        #         query=queries.pop()
        #         for item in queries:
        #             query |= item
        #         lists = lists.filter(query)

        # if search_string:
        #     lists = lists.filter(filter_string)

        # if order_type is None: 
        #     if order_column:
        #         lists = lists.order_by(order_column)  

        # elif order_type in 'asc':
        #     if order_column:
        #         lists = lists.order_by(order_column)
        #     else: 
        #         lists = lists.order_by('id')   

        # elif order_type in 'desc':
        #     if order_column:
        #         order_column = '-' + str(order_column)
        #         lists = lists.order_by(order_column)
        #     else: 
        #         lists = lists.order_by('-id') 

        # if limit_start and limit_end:
        #         lists = lists[int(limit_start):int(limit_end)]

        # elif limit_start:
        #         lists = lists[int(limit_start):]

        # elif limit_end:
        #         lists = lists[0:int(limit_end)]           
        
        #serializer = cSerializer.projectSerializer(lists, many=True)
        #serializer = cSerializer.ListInitiationNotesSerializer(lists, many=True)
        serializer = cSerializer.ListGlobalStatusSerializer(lists, many=True)

        
        #print(serializer.query)
        return Response({"status":error.context['success_code'], "data": lists,"test":'ok'}, status=status.HTTP_200_OK)




class ListGlobalStatusInitiationNotes(APIView):

    def get(self,request, pk = None):

        project = masterModels.Project.objects.values('id','name').exclude(status=3)
        #section = masterModels.Project.objects.values('id','name').exclude(status=3)
        lists = models.InitiationNotes.objects.values('approved_status','project__id', 'project_id','project__name').filter(approved_status='2')


        
        print(lists)

        # serializer = cSerializer.ListGlobalStatusSerializer(lists, many=True)

        
        #print(serializer.query)
        return Response({"status":error.context['success_code'], "data": project}, status=status.HTTP_200_OK)




class ListGlobalStatusFormulationOfApproachPaper(APIView):
    
    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                #list = models.FormulationOfApproachPaper.objects.filter(pk=pk).exclude(status='3').get()
                list = models.FormulationOfApproachPaper.objects.raw('SELECT t1.id, t1.name, t2.project_id FROM "master.project" as t1 INNER JOIN "psr.initiation_notes" as t2 ON t1.id = t2.project_idgroup by t2.project_id, t1.id')
                serializeobj = serializer.ListGlobalStatusSerializer(list)
                print(serializeobj.query())
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
      
        except models.FormulationOfApproachPaper.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Initiation Notes" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.FormulationOfApproachPaper.objects.values('approved_status','project__id','project__name').filter(approved_status='2')

        
        print(lists)

        serializer = cSerializer.ListGlobalStatusSerializer(lists, many=True)

        
        #print(serializer.query)
        return Response({"status":error.context['success_code'], "data": lists}, status=status.HTTP_200_OK)


class ListGlobalStatusPresentationOfApproachPaper(APIView):
    
    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                #list = models.PresentationOfApproachPaper.objects.filter(pk=pk).exclude(status='3').get()
                list = models.PresentationOfApproachPaper.objects.raw('SELECT t1.id, t1.name, t2.project_id FROM "master.project" as t1 INNER JOIN "psr.initiation_notes" as t2 ON t1.id = t2.project_idgroup by t2.project_id, t1.id')
                serializeobj = serializer.PresentationOfApproachPaper(list)
                print(serializeobj.query())
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
      
        except models.PresentationOfApproachPaper.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Initiation Notes" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.PresentationOfApproachPaper.objects.values('approved_status','project__id','project__name').filter(approved_status='2')

        
        print(lists)

        #serializer = cSerializer.ListGlobalStatusSerializer(lists, many=True)

        
        #print(serializer.query)
        return Response({"status":error.context['success_code'], "data": lists}, status=status.HTTP_200_OK)


class ListGlobalStatusInputsForStaffRequirement(APIView):
    
    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.InputsForStaffRequirement.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.InputsForStaffRequirement(list)
                print(serializeobj.query())
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
      
        except models.InputsForStaffRequirement.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Inputs For Staff Requirement" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.InputsForStaffRequirement.objects.values('approved_status','project__id','project__name').filter(approved_status='2')

        
        print(lists)

        #serializer = cSerializer.ListGlobalStatusSerializer(lists, many=True)

        
        #print(serializer.query)
        return Response({"status":error.context['success_code'], "data": lists}, status=status.HTTP_200_OK)


class ListGlobalStatusConceptDesign(APIView):
    
    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.ConceptDesign.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.ConceptDesign(list)
                print(serializeobj.query())
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
      
        except models.ConceptDesign.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Inputs For Staff Requirement" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.ConceptDesign.objects.values('approved_status','project__id','project__name').filter(approved_status='2')

        
        print(lists)

        #serializer = cSerializer.ListGlobalStatusSerializer(lists, many=True)

        
        #print(serializer.query)
        return Response({"status":error.context['success_code'], "data": lists}, status=status.HTTP_200_OK)



class ListGlobalStatusIncorporationOfDesignInputs(APIView):
    
    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.IncorporationOfDesignInputs.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.IncorporationOfDesignInputs(list)
                print(serializeobj.query())
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
      
        except models.IncorporationOfDesignInputs.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Incorporation Of Design Inputs" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.IncorporationOfDesignInputs.objects.values('approved_status','project__id','project__name').filter(approved_status='2')

        
        print(lists)

        #serializer = cSerializer.ListGlobalStatusSerializer(lists, many=True)

        
        #print(serializer.query)
        return Response({"status":error.context['success_code'], "data": lists}, status=status.HTTP_200_OK)


class ListGlobalStatusReceiptOfRFIResponse(APIView):
    
    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.ReceiptOfRFIResponses.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.ReceiptOfRFIResponses(list)
                print(serializeobj.query())
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
      
        except models.ReceiptOfRFIResponses.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Incorporation Of Design Inputs" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.ReceiptOfRFIResponses.objects.values('approved_status','project__id','project__name').filter(approved_status='2')

        
        print(lists)

        #serializer = cSerializer.ListGlobalStatusSerializer(lists, many=True)

        
        #print(serializer.query)
        return Response({"status":error.context['success_code'], "data": lists}, status=status.HTTP_200_OK)




class ListGlobalStatusGLSInitiation(APIView):
    
    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.InitiationNotesGLSMaster.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.InitiationNotesGLSMaster(list)
                print(serializeobj.query())
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
      
        except models.InitiationNotesGLSMaster.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Incorporation Of Design Inputs" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.InitiationNotesGLSMaster.objects.values('approved_status','project__id','project__name').filter(approved_status='2')

        
        print(lists)

        #serializer = cSerializer.ListGlobalStatusSerializer(lists, many=True)

        
        #print(serializer.query)
        return Response({"status":error.context['success_code'], "data": lists}, status=status.HTTP_200_OK)



class ResponsibilityList(APIView):
    
    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.FormulationOfApproachPaperResponsibility.objects.filter(pk=pk).exclude(status='3').get()
                print('@@@@@@@@')
                serializeobj = serializer.FormulationOfApproachPaperResponsibilityListSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.FormulationOfApproachPaperResponsibility.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"PSR Formulation Of Approach Paper" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)


        project_id = request.GET.get('project_id')
        #lists = models.FormulationOfApproachPaperResponsibility.objects.exclude(status='3')
        lists = models.FormulationOfApproachPaperResponsibility.objects.values('id','section', 'unit', 'project_id').filter(project_id=project_id).first()

        responsibility_id = lists['id']
        sectionList = lists['section'].strip("[]")
        sectionList = sectionList.replace('"', '').split(',')
        #print(sectionList)

        #print(type(sectionList))
        section = masterModels.Section.objects.values('id','name').filter(id__in=sectionList)

        # Compartment
        compartmentResp = models.FormulationOfApproachPaperResponsibilityCompartment.objects.values('id','compartment', 'formulation_of_approach_paper_responsibility').filter(formulation_of_approach_paper_responsibility=responsibility_id).first()

        #print(compartmentResp)
        compartmentList = compartmentResp['compartment'].strip("[]")
        compartmentList = compartmentList.replace('"', '').split(',')
        #print('Compartment', compartmentList)
        compartment = masterModels.Compartment.objects.values('id','name').filter(id__in=compartmentList)

        # Equipment
        equipmentResp = models.FormulationOfApproachPaperResponsibilityEquipment.objects.values('id','equipment', 'formulation_of_approach_paper_responsibility').filter(formulation_of_approach_paper_responsibility=responsibility_id).first()

        equipmentList = equipmentResp['equipment'].strip("[]")
        equipmentList = equipmentList.replace('"', '').split(',')
        #print('Equipment', equipmentList)
        equipment = masterModels.Equipment.objects.values('id','name').filter(id__in=equipmentList)


        # System 
        systemResp = models.FormulationOfApproachPaperResponsibilitySystem.objects.values('id','system', 'formulation_of_approach_paper_responsibility').filter(formulation_of_approach_paper_responsibility=responsibility_id).first()

        systemList = systemResp['system'].strip("[]")
        systemList = systemList.replace('"', '').split(',')
        #print('system', systemList)
        system = masterModels.System.objects.values('id','name').filter(id__in=systemList)


        #print(section)

        #print(request.GET.get('project_id'),'@@@@@@@@##########')
        # if normal_values:
        #     lists = lists.filter(reduce(operator.and_, 
        #                        (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        # if array_values:
        #     for key,values in array_values.items():
        #         queries= [Q(**{"%s__contains" % key: value }) for value in values]
        #         query=queries.pop()
        #         for item in queries:
        #             query |= item
        #         lists = lists.filter(query)

        # if search_string:
        #     lists = lists.filter(filter_string)

        # if order_type is None: 
        #     if order_column:
        #         lists = lists.order_by(order_column)  

        # elif order_type in 'asc':
        #     if order_column:
        #         lists = lists.order_by(order_column)
        #     else: 
        #         lists = lists.order_by('id')   

        # elif order_type in 'desc':
        #     if order_column:
        #         order_column = '-' + str(order_column)
        #         lists = lists.order_by(order_column)
        #     else: 
        #         lists = lists.order_by('-id') 

        # if limit_start and limit_end:
        #         lists = lists[int(limit_start):int(limit_end)]

        # elif limit_start:
        #         lists = lists[int(limit_start):]

        # elif limit_end:
        #         lists = lists[0:int(limit_end)]           


        #print(lists,"############")                
        #section=''
        #serializer = cSerializer.ListFormulationOfApproachPaperResponsibilitySerializer(lists, many=True)
        return Response({"status":error.context['success_code'], "data": lists, "section": section, "compartment":compartment, "equipment":equipment, "system":system}, status=status.HTTP_200_OK)





class ResponsibilitySectionList(APIView):
    
    def get(self, request, pk = None):
        #print(request.GET.items())
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = masterModels.Section.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = masterSerializer.ListSectionSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except masterModels.Section.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"PSR Formulation Of Approach Paper" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        #lists = masterModels.Section.objects.exclude(status='3')
        arrList = request.GET.get('section_id')
        
        lists = masterModels.Section.objects.filter(id__in=arrList)
        # if normal_values:
        #     lists = lists.filter(reduce(operator.and_, 
        #                        (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        # if array_values:
        #     for key,values in array_values.items():
        #         queries= [Q(**{"%s__contains" % key: value }) for value in values]
        #         query=queries.pop()
        #         for item in queries:
        #             query |= item
        #         lists = lists.filter(query)

        # if search_string:
        #     lists = lists.filter(filter_string)

        # if order_type is None: 
        #     if order_column:
        #         lists = lists.order_by(order_column)  

        # elif order_type in 'asc':
        #     if order_column:
        #         lists = lists.order_by(order_column)
        #     else: 
        #         lists = lists.order_by('id')   

        # elif order_type in 'desc':
        #     if order_column:
        #         order_column = '-' + str(order_column)
        #         lists = lists.order_by(order_column)
        #     else: 
        #         lists = lists.order_by('-id') 

        # if limit_start and limit_end:
        #         lists = lists[int(limit_start):int(limit_end)]

        # elif limit_start:
        #         lists = lists[int(limit_start):]

        # elif limit_end:
        #         lists = lists[0:int(limit_end)]           


        #print(lists,"############")                
        
        serializer = masterSerializer.ListSectionSerializer(lists, many=True)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class saveSSRForInputSR(APIView):

    def post(self,request, pk = None):
        if 'sr_id' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "ssr_id" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        #elif 'specification' not in request.data and request.data['status'] != 3:
            #return Response({"status":error.context['error_code'], "message" : "Specification" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        elif 'description' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "Description" +language.context[language.defaultLang]['missing'] },status=status.HTTP_200_OK)
        else: 
            request.data['created_ip']=Common.get_client_ip(request)
            request.data['created_on']=datetime.now()
            request.data['created_by']=request.user.id

            saveserialize=serializer.InputsSRSSRSerializer(data=request.data,partial=True)
            if saveserialize.is_valid():
                saveserialize.save()
                return Response({"status" :error.context['success_code'], "message":"Single sheet specification " +language.context[language.defaultLang]['insert'], "data":saveserialize.data}, status=status.HTTP_200_OK)
            else:
                return Response({"status" :error.context['error_code'], "message": error.serializerError(saveserialize)}, status=status.HTTP_200_OK)

class ListSSRForInputSR(APIView):

    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['specification','description']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.InputsSRSSR.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.ListInputsSRSSRSerializer(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.InputsSRSSR.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"Primary roles" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.InputsSRSSR.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        serializer = cSerializer.ListInputsSRSSRSerializer(lists, many=True)
        return Response({"status":error.context['success_code'] , "data": serializer.data, 'test':''}, status=status.HTTP_200_OK)

class deleteSSRForInputSR(APIView):

    def post(self,request, pk = None):
        if 'id' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "id" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        else: 
            if models.InputsSRSSR.objects.filter(id=request.data['id']).count()>0:
                models.InputsSRSSR.objects.filter(id=request.data['id']).update(**{"status":3})
                return Response({"status" :error.context['success_code'], "message":"Single sheet specification deleted successfully"}, status=status.HTTP_200_OK)
            else:
                return Response({"status" :error.context['error_code'], "message":"Incorrect id supplied in the parameter"}, status=status.HTTP_200_OK)
                
class ResponsibilityProject(APIView):

    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['specification','description']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.FormulationOfApproachPaperResponsibility.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.reponsibilitySerializer(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.FormulationOfApproachPaperResponsibility.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"Primary roles" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.FormulationOfApproachPaperResponsibility.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        serializer = cSerializer.reponsibilitySerializer(lists, many=True)
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)




class ListSSRForProject(APIView):
    
    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.IncorporationOfDesignInputs.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.IncorporationOfDesignInputs(list)
                print(serializeobj.query())
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
      
        except models.IncorporationOfDesignInputs.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Incorporation Of Design Inputs" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)


        #print(request.GET.get('project_id'),)
        project_id = request.GET.get('project_id')
        staff_res = models.InputsForStaffRequirement.objects.values('id','project_id').filter(project_id=project_id)

        sr_id = staff_res[0]['id']

        #lists = models.InitiationNotes.objects.values('approved_status','project__id','project__name').filter(approved_status='2')
        ssr = models.InputsSRSSR.objects.values('sss__name','description','created_ip').filter(sr_id_id=sr_id).exclude(status='3')
        
        #print(staff_res,staff_res[0]['id'],ssr,"******************")

        #serializer = cSerializer.ListGlobalStatusSerializer(lists, many=True)
        #serializer = cSerializer.ListInputsSRSSRSerializer(ssr, many=True)

        lists=""
        #print(serializer.query)
        return Response({"status":error.context['success_code'], "data_ssr": ssr}, status=status.HTTP_200_OK)

class saveResponsibility(APIView):
    def post(self,request, pk = None):
        if "project_id" not in request.data or request.data['project_id']=='':
            return Response({"status":error.context['error_code'], "message": "Project id is missing"}, status=status.HTTP_200_OK)
        elif "section_unit_id" not in request.data or request.data['section_unit_id']=='':
            return Response({"status":error.context['error_code'], "message": "section_unit_id is missing"}, status=status.HTTP_200_OK)            
        else:
            print(request.data)
            models.PSRUnitCompartments.objects.filter(section_unit_id=request.data['section_unit_id']).delete()
            for i in range(len(request.data['compartments'])):
                models.PSRUnitCompartments.objects.create(**{"project_id":request.data['project_id'],"section_unit_id":request.data['section_unit_id'],"compartment_id":request.data['compartments'][i]})

            models.PSRUnitEquipments.objects.filter(section_unit_id=request.data['section_unit_id']).delete()
            for i in range(len(request.data['equipments'])):
                models.PSRUnitEquipments.objects.create(**{"project_id":request.data['project_id'],"section_unit_id":request.data['section_unit_id'],"equipment_id":request.data['equipments'][i]})

            models.PSRUnitSystems.objects.filter(section_unit_id=request.data['section_unit_id']).delete()
            for i in range(len(request.data['systems'])):
                models.PSRUnitSystems.objects.create(**{"project_id":request.data['project_id'],"section_unit_id":request.data['section_unit_id'],"system_id":request.data['systems'][i]})
            return Response({"status":error.context['success_code'], "message": "Responsiblity saved successfully"}, status=status.HTTP_200_OK)

class getResponsibility(APIView):
    def post(self,request, pk = None):
        if "project_id" not in request.data or request.data['project_id']=='':
            return Response({"status":error.context['error_code'], "message": "Project id is missing"}, status=status.HTTP_200_OK)
        else:
            compartments=models.PSRUnitCompartments.objects.filter(project_id=request.data['project_id'])
            equipments=models.PSRUnitEquipments.objects.filter(project_id=request.data['project_id'])
            systems=models.PSRUnitSystems.objects.filter(project_id=request.data['project_id'])

            if 'section_unit_id' in request.data and request.data['section_unit_id']!='':
                compartments=compartments.filter(section_unit_id=request.data['section_unit_id'])
                equipments=equipments.filter(section_unit_id=request.data['section_unit_id'])
                systems=systems.filter(section_unit_id=request.data['section_unit_id'])

            listCompartments=serializer.PSRUnitCompartmentsSerializer(compartments,many=True).data
            listEquipments=serializer.PSRUnitEquipmentsSerializer(equipments,many=True).data
            listSystems=serializer.PSRUnitSystemsSerializer(systems,many=True).data

            return Response({"status":error.context['success_code'], "compartments": listCompartments,"equipments":listEquipments,"systems":listSystems}, status=status.HTTP_200_OK)

class getResponsibilityOther(APIView):
    def post(self,request, pk = None):
        if "project_id" not in request.data or request.data['project_id']=='':
            return Response({"status":error.context['error_code'], "message": "Project id is missing"}, status=status.HTTP_200_OK)
        else:
            compartments=models.PSRUnitCompartments.objects.filter(project_id=request.data['project_id'])
            equipments=models.PSRUnitEquipments.objects.filter(project_id=request.data['project_id'])
            systems=models.PSRUnitSystems.objects.filter(project_id=request.data['project_id'])

            if 'section_unit_id' in request.data and request.data['section_unit_id']!='':
                compartments=compartments.filter(section_unit_id=request.data['section_unit_id'])
                equipments=equipments.filter(section_unit_id=request.data['section_unit_id'])
                systems=systems.filter(section_unit_id=request.data['section_unit_id'])

            listCompartments=serializer.PSRUnitCompartmentsOtherSerializer(compartments,many=True).data
            listEquipments=serializer.PSRUnitEquipmentsOtherSerializer(equipments,many=True).data
            listSystems=serializer.PSRUnitSystemsOtherSerializer(systems,many=True).data

            return Response({"status":error.context['success_code'], "compartments": listCompartments,"equipments":listEquipments,"systems":listSystems}, status=status.HTTP_200_OK)

class getResponsibilitySection(APIView):
    def post(self,request, pk = None):
        if "project_id" not in request.data or request.data['project_id']=='':
            return Response({"status":error.context['error_code'], "message": "Project id is missing"}, status=status.HTTP_200_OK)
        else:
            sections=models.PSRSection.objects.filter(status=1).exclude(code__in=['GENERAL']).order_by('sequence')
            sectionsList=serializer.ResponsibilityPSRSectionSerializer2(sections,many=True,context={"project_id":request.data['project_id']}).data
            return Response({"status":error.context['success_code'], "sections": sectionsList}, status=status.HTTP_200_OK)

def saveDataFlow(project_id,sections):
    print('sections',len(sections))
    models.PSRDataFlow.objects.filter(project_id=project_id).delete()
    for i in range(0,len(sections)):
        section=sections[i]
        print('project_id',project_id)
        dataToSave={"project":project_id,"psr_section":section['id'],"standards":section['standards'],"whole_ship_features":section['whole_ship_features'],"man_power":section['man_power'],"recommendations":(section['recommendations'] if 'recommendations' in section else '')}
        print('dataToSave',dataToSave)
        saveSer=serializer.PSRDataFlowSerializer(data=dataToSave,partial=True)
        if saveSer.is_valid():
            saveSer.save()
            data_flow_id=saveSer.data['id']
            for j in range(len(section['units'])):
                unit=section['units'][j]
                for k in range(len(unit['compartments'])):
                    compartment=unit['compartments'][k]
                    compartment_id=compartment['compartment_id']
                    for l in range(len(compartment['compartments'])):
                        compart=compartment['compartments'][l]
                        ser=compart['ser']
                        name=compart['name']
                        numbers=compart['numbers']
                        location=compart['location']
                        equipment=compart['equipment']
                        features=compart['features']
                        layout=compart['layout']
                        special_requirements=compart['special_requirements']
                        standards=compart['standards']
                        recommendations=compart['recommendations']
                        if ser!='' and name!='':
                            models.PSRDataFlowCompartment.objects.create(**{"data_flow_id":data_flow_id,"compartment_id":compartment_id,"ser":ser,"name":name,"numbers":numbers,"location":location,"equipment":equipment,"features":features,"layout":layout,"special_requirements":special_requirements,"standards":standards,"recommendations":recommendations})

                for k in range(len(unit['equipments'])):
                    equipment=unit['equipments'][k]
                    equipment_id=equipment['equipment_id']
                    for l in range(len(equipment['equipments'])):
                        equip=equipment['equipments'][l]
                        ser=equip['ser']
                        name=equip['name']
                        numbers=equip['numbers']
                        capabilities=equip['capabilities']
                        weight=equip['weight']
                        location=equip['location']
                        interface=equip['interface']
                        procurement=equip['procurement']
                        vendor=equip['vendor']
                        cost=equip['cost']
                        standards=equip['standards']
                        sustence=equip['sustence']
                        recommendations=equip['recommendations']
                        if ser!='' and name!='':
                            models.PSRDataFlowEquipments.objects.create(**{"data_flow_id":data_flow_id,"equipment_id":equipment_id,"ser":ser,"name":name,"numbers":numbers,"capabilities":capabilities,"weight":weight,"location":location,"interface":interface,"procurement":procurement,"vendor":vendor,"cost":cost,"standards":standards,"sustence":sustence,"recommendations":recommendations})
                for k in range(len(unit['systems'])):
                    system=unit['systems'][k]
                    system_id=system['system_id']
                    for l in range(len(system['systems'])):
                        sys=system['systems'][l]
                        ser=sys['ser']
                        name=sys['name']
                        numbers=sys['numbers']
                        capabilities=sys['capabilities']
                        weight=sys['weight']
                        location=sys['location']
                        interface=sys['interface']
                        procurement=sys['procurement']
                        vendor=sys['vendor']
                        cost=sys['cost']
                        standards=sys['standards']
                        sustence=sys['sustence']
                        recommendations=sys['recommendations']
                        if ser!='' and name!='':
                            models.PSRDataFlowsystems.objects.create(**{"data_flow_id":data_flow_id,"system_id":system_id,"ser":ser,"name":name,"numbers":numbers,"capabilities":capabilities,"weight":weight,"location":location,"interface":interface,"procurement":procurement,"vendor":vendor,"cost":cost,"standards":standards,"sustence":sustence,"recommendations":recommendations})
        else:
            return {"status" :error.context['error_code'], "message": error.serializerError(saveSer)}
    return {"status" :error.context['success_code'], "message": "Data saved successfully"}



def saveDataFlowGeneral(project_id,general):
    #print(project_id,"###",general.data['role'])


    models.PSRDataFlowGeneralSection.objects.filter(project_id=project_id).delete()
    models.PSRDataFlowGeneralSection.objects.create(
        project_id = project_id,
        roles = general.data['roles'],
        critical_design_drivers = general.data['critical_design_drivers'],
        operating_philosophy = general.data['operating_philosophy'],
        area_of_operations = general.data['area_of_operations'],
        rules_and_regulations = general.data['rules_and_regulations'],
        general_remarks = general.data['general_remarks'],
        displacement = general.data['displacement'],
        dimensions = general.data['dimensions'],
        speed = general.data['speed'],
        endurance_and_range = general.data['endurance_and_range'],
        sea_worthiness = general.data['sea_worthiness'],
        propulsion = general.data['propulsion'],
        operating_conditions = general.data['operating_conditions'],
        design_and_construction_standards = general.data['design_and_construction_standards'],
        stealth = general.data['stealth'],
        ergonomics = general.data['ergonomics'],
        complement = general.data['complement'],
        cots_technology = general.data['cots_technology'],
        protection = general.data['protection'],
        unrep = general.data['unrep'],
        boats_and_usvs = general.data['boats_and_usvs'],
        noise_reduction = general.data['noise_reduction'],
        op_logistic_management_information_system = general.data['op_logistic_management_information_system'],
        ipms = general.data['ipms'],
        surveillance_and_security_arrangement = general.data['surveillance_and_security_arrangement'],
        cim = general.data['cim'],
        green_warship = general.data['green_warship'],
        construction = general.data['construction'],
        automation_and_redundancy = general.data['automation_and_redundancy'],
        workshops = general.data['workshops'],
    )

    return {"status" :error.context['success_code'], "message": "Data saved successfully"}



class saveInputSR(APIView):
    def post(self,request, pk = None):
        if "project_id" not in request.data or request.data['project_id']=='':
            return Response({"status":error.context['error_code'], "message": "Project id is missing"}, status=status.HTTP_200_OK)
        else:

            count = models.InputsForStaffRequirement.objects.filter(project_id=request.data['project_id']).count()

            if count > 0:
                approveRes = models.InputsForStaffRequirement.objects.values('approved_status','approved_by','approved_on', 'approved_remark','project_id').filter(project_id = request.data['project_id']).get()

                approved_status = approveRes['approved_status']
                approved_by = approveRes['approved_by']
                approved_on = approveRes['approved_on']
                approved_remark = approveRes['approved_remark']

                models.InputsForStaffRequirement.objects.filter(project_id=request.data['project_id']).delete()

                models.InputsForStaffRequirement.objects.create(**{
                    "project_id":request.data['project_id'],
                    "created_ip":Common.get_client_ip(request),
                    "created_by_id":request.user.id,
                    "approved_status" : approved_status,
                    "approved_by" : approved_by,
                    "approved_on" : approved_on,
                    "approved_remark" : approved_remark
                    })
            else:

                models.InputsForStaffRequirement.objects.filter(project_id=request.data['project_id']).delete()

                models.InputsForStaffRequirement.objects.create(**{
                    "project_id":request.data['project_id'],
                    "created_ip":Common.get_client_ip(request),
                    "created_by_id":request.user.id
                    })


            
            saveDataFlowGeneral(request.data['project_id'],request)

            response = saveDataFlow(request.data['project_id'],request.data['sections'])
            return Response(response, status=status.HTTP_200_OK)

class saveConceptDesign(APIView):
    def post(self,request, pk = None):
        if "project_id" not in request.data or request.data['project_id']=='':
            return Response({"status":error.context['error_code'], "message": "Project id is missing"}, status=status.HTTP_200_OK)
        elif "refinement_of_values" not in request.data or request.data['refinement_of_values']=='':
            return Response({"status":error.context['error_code'], "message": "refinement_of_values is missing"}, status=status.HTTP_200_OK)
        elif "powering_requirements" not in request.data or request.data['powering_requirements']=='':
            return Response({"status":error.context['error_code'], "message": "powering_requirements is missing"}, status=status.HTTP_200_OK)
        elif "total_heat_load_calculation" not in request.data or request.data['total_heat_load_calculation']=='':
            return Response({"status":error.context['error_code'], "message": "total_heat_load_calculation is missing"}, status=status.HTTP_200_OK)
        else:

            count = models.ConceptDesign.objects.filter(project_id=request.data['project_id']).count()
            if count > 0:
                approveRes = models.ConceptDesign.objects.values('approved_status','approved_by','approved_on', 'approved_remark','project_id').filter(project_id = request.data['project_id']).get()

                approved_status = approveRes['approved_status']
                approved_by = approveRes['approved_by']
                approved_on = approveRes['approved_on']
                approved_remark = approveRes['approved_remark']

                models.ConceptDesign.objects.filter(project_id=request.data['project_id']).delete()

                models.ConceptDesign.objects.create(**{
                    "project_id":request.data['project_id'],
                    "refinement_of_values":request.data['refinement_of_values'],
                    "powering_requirements":request.data['powering_requirements'],
                    "powering_requirements":request.data['powering_requirements'],
                    "total_heat_load_calculation":request.data['total_heat_load_calculation'],
                    "created_ip":Common.get_client_ip(request),
                    "created_by_id":request.user.id,

                    "approved_status" : approved_status,
                    "approved_by" : approved_by,
                    "approved_on" : approved_on,
                    "approved_remark" : approved_remark

                    })
            else:
                models.ConceptDesign.objects.filter(project_id=request.data['project_id']).delete()

                models.ConceptDesign.objects.create(**{
                    "project_id":request.data['project_id'],
                    "refinement_of_values":request.data['refinement_of_values'],
                    "powering_requirements":request.data['powering_requirements'],
                    "powering_requirements":request.data['powering_requirements'],
                    "total_heat_load_calculation":request.data['total_heat_load_calculation'],
                    "created_ip":Common.get_client_ip(request),
                    "created_by_id":request.user.id,
                    })


            saveDataFlowGeneral(request.data['project_id'],request)
            response=saveDataFlow(request.data['project_id'],request.data['sections'])
            return Response(response, status=status.HTTP_200_OK)

class saveInCorporation(APIView):
    def post(self,request, pk = None):
        if "project_id" not in request.data or request.data['project_id']=='':
            return Response({"status":error.context['error_code'], "message": "Project id is missing"}, status=status.HTTP_200_OK)
        elif "cost_estimation" not in request.data or request.data['cost_estimation']=='':
            return Response({"status":error.context['error_code'], "message": "cost_estimation is missing"}, status=status.HTTP_200_OK)
        else:

            count = models.IncorporationOfDesignInputs.objects.filter(project_id=request.data['project_id']).count()
            if count > 0:
                approveRes = models.IncorporationOfDesignInputs.objects.values('approved_status','approved_by','approved_on', 'approved_remark','project_id').filter(project_id = request.data['project_id']).get()

                approved_status = approveRes['approved_status']
                approved_by = approveRes['approved_by']
                approved_on = approveRes['approved_on']
                approved_remark = approveRes['approved_remark']

                models.IncorporationOfDesignInputs.objects.filter(project_id=request.data['project_id']).delete()

                models.IncorporationOfDesignInputs.objects.create(**{
                    "project_id":request.data['project_id'],
                    "cost_estimation":request.data['cost_estimation'],
                    "created_ip":Common.get_client_ip(request),
                    "created_by_id":request.user.id,

                    "approved_status" : approved_status,
                    "approved_by" : approved_by,
                    "approved_on" : approved_on,
                    "approved_remark" : approved_remark

                    })
            else:

                models.IncorporationOfDesignInputs.objects.filter(project_id=request.data['project_id']).delete()

                models.IncorporationOfDesignInputs.objects.create(**{
                    "project_id":request.data['project_id'],
                    "cost_estimation":request.data['cost_estimation'],
                    "created_ip":Common.get_client_ip(request),
                    "created_by_id":request.user.id,
                    })

            saveDataFlowGeneral(request.data['project_id'],request)
            response=saveDataFlow(request.data['project_id'],request.data['sections'])
            return Response(response, status=status.HTTP_200_OK)

class saveRFI(APIView):
    def post(self,request, pk = None):
        if "project" not in request.data or request.data['project']=='':
            return Response({"status":error.context['error_code'], "message": "Project id is missing"}, status=status.HTTP_200_OK)
        else:

            count = models.ReceiptOfRFIResponses.objects.filter(project_id=request.data['project']).count()
            if count > 0:
                approveRes = models.ReceiptOfRFIResponses.objects.values('approved_status','approved_by','approved_on', 'approved_remark','project_id').filter(project_id = request.data['project']).get()

                approved_status = approveRes['approved_status']
                approved_by = approveRes['approved_by']
                approved_on = approveRes['approved_on']
                approved_remark = approveRes['approved_remark']

                #print(approved_status,"PPPPPPPPPPPPp")

                models.ReceiptOfRFIResponses.objects.filter(project_id=request.data['project']).delete()
                request.data._mutable = True
                request.data['created_ip'] = Common.get_client_ip(request)

                request.data['approved_status'] = approved_status
                request.data['approved_by'] = approved_by
                request.data['approved_on'] = approved_on
                request.data['approved_remark'] = approved_remark
            else:

                models.ReceiptOfRFIResponses.objects.filter(project_id=request.data['project']).delete()
                request.data._mutable = True
                request.data['created_ip'] = Common.get_client_ip(request)



            saveserialize = cSerializer.ReceiptOfRFIResponsesSerializer(data = request.data,partial=True)
                
            if saveserialize.is_valid():
                saveserialize.save()
                return Response({"status" :error.context['success_code'],"message":"RFI saved successfully"}, status=status.HTTP_200_OK)
            else:
                return Response({"status" :error.context['error_code'], "message": error.serializerError(saveserialize)}, status=status.HTTP_200_OK)

class saveRFISections(APIView):
    def post(self,request, pk = None):
        if "project_id" not in request.data or request.data['project_id']=='':
            return Response({"status":error.context['error_code'], "message": "Project id is missing"}, status=status.HTTP_200_OK)
        elif "sections" not in request.data or request.data['sections']=='':
            return Response({"status":error.context['error_code'], "message": "sections are missing"}, status=status.HTTP_200_OK)
        else:
            saveDataFlowGeneral(request.data['project_id'],request)
            response=saveDataFlow(request.data['project_id'],request.data['sections'])
            return Response(response, status=status.HTTP_200_OK)


class getCompletePSR(APIView):
    def get(self,request, pk = None):
        project_id = request.GET['project_id']
        response={}
        if project_id == '':
            return Response({"status":error.context['error_code'], "message": "Project id is missing"}, status=status.HTTP_200_OK)
        else:
            #sections = models.InitiationNotesGLSMaster.objects.filter(project_id=project_id)
            #sections = {project_id:project_id};
            #sectionsList=serializer.CompletePSRSerializer(sections,many=True,context={"project_id":request.data['project_id']}).data
            #serializer = cSerializer.CompletePSRSerializer(sections,many=True).data
            #serializer = cSerializer.ListPSRSectionSerializer2(lists, many=True)

            # SSS
            staff_res = models.InputsForStaffRequirement.objects.values('id','project_id').filter(project_id=project_id)
            sr_id = staff_res[0]['id']

            staff_res = models.InputsForStaffRequirement.objects.values('id','project_id').filter(project_id=project_id)
            sr_id = staff_res[0]['id']

            response['sss'] = models.InputsSRSSR.objects.values('description','sss__name','created_ip').filter(sr_id_id=sr_id).exclude(status='3')

            # Initiation Notes

            in_id = models.InitiationNotes.objects.values('project','id').filter(project_id=project_id).exclude(status='3')

            in_list = models.InitiationNotes.objects.filter(pk=in_id[0]['id']).exclude(status='3').get()
            in_obj = serializer.ListInitiationNotesSerializer(in_list)

            response['initiation_notes'] = in_obj.data

            # Formulation Of Approach
            for_id = models.FormulationOfApproachPaper.objects.values('project','id').filter(project_id=project_id).exclude(status='3')

            for_list = models.FormulationOfApproachPaper.objects.filter(pk=for_id[0]['id']).exclude(status='3').get()
            for_obj = serializer.FormulationOfApproachPaperSerializer(for_list)
            response['formulation_of_approach'] = for_obj.data

            # Presentation Of Approach Paper
            pre_id = models.PresentationOfApproachPaper.objects.values('project','id').filter(project_id=project_id).exclude(status='3')

            pre_list = models.PresentationOfApproachPaper.objects.filter(pk=pre_id[0]['id']).exclude(status='3').get()
            pre_obj = serializer.ListPresentationOfApproachPaperSerializer(pre_list)
            response['presentation_of_approach'] = pre_obj.data


            # Concept Design
            con_id = models.ConceptDesign.objects.values('project','id').filter(project_id=project_id).exclude(status='3')

            con_list = models.ConceptDesign.objects.filter(pk=con_id[0]['id']).exclude(status='3').get()
            con_obj = serializer.ListConceptDesignSerializer(con_list)
            response['concept_design'] = con_obj.data

            # Incorporation of Design Inputs

            inc_id = models.IncorporationOfDesignInputs.objects.values('project','id').filter(project_id=project_id).exclude(status='3')

            inc_list = models.IncorporationOfDesignInputs.objects.filter(pk=inc_id[0]['id']).exclude(status='3').get()
            inc_obj = serializer.ListIncorporationOfDesignInputsSerializer(inc_list)
            response['incorporation_of_design_input'] = inc_obj.data

            # Receipt of RFI Responses

            rec_id = models.ReceiptOfRFIResponses.objects.values('project','id').filter(project_id=project_id).exclude(status='3')

            rec_list = models.ReceiptOfRFIResponses.objects.filter(pk=rec_id[0]['id']).exclude(status='3').get()
            rec_obj = serializer.ListReceiptOfRFIResponsesSerializer(rec_list)
            response['receipt_of_rfi_responses'] = rec_obj.data

            # Inputs for Staff Requirements

            #isr_id = models.InputsForStaffRequirement.objects.values('project','id').filter(project_id=project_id).exclude(status='3')


            #sections = models.PSRSection.objects.filter(status=1).exclude(code__in=['GENERAL']).order_by('sequence')
            #sectionsList=serializer.ResponsibilityPSRSectionSerializer2(sections,many=True,context={"project_id":request.data['project_id']}).data

            #isr_list = models.InputsSRSSR.objects.filter(pk=isr_id[0]['id']).exclude(status='3').get()
            #isr_obj = serializer.ListInputsSRSSRSerializer(isr_list)
            #response['inputs_for_staff_requirements'] = sections

            return Response({"status":error.context['success_code'], "data": response}, status=status.HTTP_200_OK)


class getGeneralRefForInputSR(APIView):
    def get(self,request, pk = None):
        project_id = request.GET['project_id']
        response={}
        if project_id == '':
            return Response({"status":error.context['error_code'], "message": "Project id is missing"}, status=status.HTTP_200_OK)
        else:


            # SSS
            # staff_res = models.InputsForStaffRequirement.objects.values('id','project_id').filter(project_id=project_id)
            # sr_id = staff_res[0]['id']

            # staff_res = models.InputsForStaffRequirement.objects.values('id','project_id').filter(project_id=project_id)
            # sr_id = staff_res[0]['id']

            # response['sss'] = models.InputsSRSSR.objects.values('description','sss__name','created_ip').filter(sr_id_id=sr_id).exclude(status='3')

            # Initiation Notes

            in_id = models.InitiationNotes.objects.values('project','id').filter(project_id=project_id).exclude(status='3')

            in_list = models.InitiationNotes.objects.filter(pk=in_id[0]['id']).exclude(status='3').get()
            in_obj = serializer.ListInitiationNotesSerializer(in_list)

            response['initiation_notes'] = in_obj.data

            # Formulation Of Approach
            for_id = models.FormulationOfApproachPaper.objects.values('project','id').filter(project_id=project_id).exclude(status='3')

            for_list = models.FormulationOfApproachPaper.objects.filter(pk=for_id[0]['id']).exclude(status='3').get()
            for_obj = serializer.FormulationOfApproachPaperSerializer(for_list)
            response['formulation_of_approach'] = for_obj.data



            return Response({"status":error.context['success_code'], "data": response}, status=status.HTTP_200_OK)


class getSectionGeneral(APIView):
    def get(self,request, pk = None):
        project_id = request.GET['project_id']
        response={}
        if project_id == '':
            return Response({"status":error.context['error_code'], "message": "Project id is missing"}, status=status.HTTP_200_OK)
        else:

            for_list = models.PSRDataFlowGeneralSection.objects.filter(project_id=project_id).get()
            for_obj = serializer.PSRDataFlowGeneralSectionSerializer(for_list)
            response['general_section'] = for_obj.data


            return Response({"status":error.context['success_code'], "data": response}, status=status.HTTP_200_OK)


class getProjectList(APIView):

    def get(self,request, pk = None):

        lists = models.GlobalTransaction.objects.exclude(status='3')
        serializer = cSerializer.GlobalTransactionsSerializer(lists, many=True)
        #print('serializer serializer',serializer.data)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class getAllProjectList(APIView):

    def get(self,request, pk = None):

        response = {}
        lists = models.GlobalTransaction.objects.exclude(status='3').distinct('project_id')
        serializer = cSerializer.DashboardSerializer(lists, many=True)
        response['all_project_list'] = serializer.data

        # PSR Active Count
        psr_count = models.ProjectModuleStatus.objects.filter( project_module_master = 7,status=2).count()

        response['all_psr_active_count'] = psr_count
        #Initiation Count
        # ini_count = 0 
        ini = 0
        ini_count = models.ProjectModuleStatus.objects.filter( project_module_master = 1,status=2).count()
        forp_count = models.ProjectModuleStatus.objects.filter( project_module_master = 2,status=2).count()
        Present_count = models.ProjectModuleStatus.objects.filter( project_module_master = 3,status=2).count()
        Input_count = models.ProjectModuleStatus.objects.filter( project_module_master = 4,status=2).count()
        concept_count = models.ProjectModuleStatus.objects.filter( project_module_master = 5,status=2).count()
        incorp_count = models.ProjectModuleStatus.objects.filter( project_module_master = 6,status=2).count()
        rfi_count = models.ProjectModuleStatus.objects.filter( project_module_master = 7,status=2).count()
        
        
        
        gls_count = models.ProjectModuleStatus.objects.filter( project_module_master = 8,status=2).count()

        response['all_gls_active_count'] = gls_count

        # BLS Active Count
        bls_count = models.ProjectModuleStatus.objects.filter( project_module_master = 9,status=2).count()

        response['all_bls_active_count'] = bls_count

        

        response['psr_ini_active_count'] = ini_count
        response['psr_formulation_active_count'] = forp_count
        response['psr_presentation_active_count'] = Present_count
        response['psr_input_active_count'] = Input_count
        response['psr_concept_active_count'] = concept_count
        response['psr_incorporation_active_count'] = incorp_count
        response['psr_receipt_active_count'] = rfi_count


        return Response({"status":error.context['success_code'], "data": response}, status=status.HTTP_200_OK)


def activeCount(project_list, module_id):
    #print('sections', project_list)

    count = 0
    k = 0
    for x in project_list:
        project_id = project_list[k]['project'][0]['id'];
        cnt = models.ProjectModuleStatus.objects.filter(project = project_id,project_module_master = module_id, status = 2).count()

        if cnt < 1:
            count += 1
            #print(psr_ini_count,"psr_ini_count")
        k += 1

    return count


# def activeCount_(project_list, module_id):
#     count = 0
#     k = 0
#     for x in project_list:
#         project_id = project_list[k]['project'][0]['id'];
#         cnt = models.ProjectModuleStatus.objects.filter(project = project_id,project_module_master = module_id).exclude(status = 2).count()

#         if cnt == 1:
#             count += 1
#         k += 1

#     return count


class getProjectDetail(APIView):

    def get(self,request, pk = None):
        project_id = request.GET['project_id']
        response={}
        if project_id == '':
            return Response({"status":error.context['error_code'], "message": "Project id is missing"}, status=status.HTTP_200_OK)

        else:
            project = masterModels.Project.objects.filter(id=project_id)
            pro_obj = masterSerializer.projectSerializer(project, many=True)
            response['project'] = pro_obj.data

            lists = models.ProjectModuleStatus.objects.filter(project=project_id).exclude(project_module_master__in=[8,9]).order_by('project_module_master_id')
            serializer = cSerializer.ListProjectModuleStatusSerializer(lists, many=True)
            response['psr_list'] = serializer.data

            # PSR
            psr_count = models.ProjectModuleStatus.objects.filter(project=project_id, status=2).exclude(project_module_master__in=[8,9]).count()

            psr_res = (psr_count/7)*100;
            response['psr'] = {'percent':round(psr_res)}


            # Count
            # Pending
            psr_pending_count = models.ProjectModuleStatus.objects.filter(project=project_id, status=8).exclude(project_module_master__in=[8,9]).count()

            # Recommended
            psr_recommended_count = models.ProjectModuleStatus.objects.filter(project=project_id, status=1).exclude(project_module_master__in=[8,9]).count()

            response['psr_approved_count'] = {'psr_approved_count':psr_count}
            response['psr_pending_count'] = {'psr_pending_count':psr_pending_count}
            response['psr_recommended_count'] = {'psr_recommended_count':psr_recommended_count}



            # GLS
            gls_count = models.ProjectModuleStatus.objects.filter(project = project_id, status = 2, project_module_master = 8).count()

            gls_res = (gls_count/1)*100;
            response['gls'] = {'percent':round(gls_res)}

            # Count
            # Pending
            gls_pending_count = models.ProjectModuleStatus.objects.filter(project=project_id, status=8, project_module_master = 8).count()            

            # Recommended
            gls_recommended_count = models.ProjectModuleStatus.objects.filter(project=project_id, status=1, project_module_master = 8).count()

            response['gls_approved_count'] = {'gls_approved_count':gls_count}
            response['gls_pending_count'] = {'gls_pending_count':gls_pending_count}
            response['gls_recommended_count'] = {'gls_recommended_count':gls_recommended_count}


            # BLS
            bls_count = models.ProjectModuleStatus.objects.filter(project = project_id, status = 2, project_module_master = 9).count()

            bls_res = (bls_count/1)*100;
            response['bls'] = {'percent':round(bls_res)}

            # Pie
            # Pending
            bls_pending_count = models.ProjectModuleStatus.objects.filter(project=project_id, status=8, project_module_master = 9).count()

            # Recommended
            bls_recommended_count = models.ProjectModuleStatus.objects.filter(project=project_id, status=1, project_module_master = 9).count() 

            response['bls_approved_count'] = {'bls_approved_count':bls_count}
            response['bls_pending_count'] = {'bls_pending_count':bls_pending_count}
            response['bls_recommended_count'] = {'bls_recommended_count':bls_recommended_count}


            # All Modules List

            all_lists = models.ProjectModuleMaster.objects.exclude(id__in=[8,9]).order_by('id')
            all_serializer = cSerializer.ProjectModuleMasterSerializer(all_lists, many=True)
            response['all_psr_list'] = all_serializer.data


            # project_status = models.ProjectModuleStatus.objects.filter(project=project_id)
            # project = models.ProjectModuleStatus.objects.filter(project=project_id)
            # serializer = cSerializer.ListProjectModuleStatusSerializer(lists, many=True)
            # print('serializer serializer',serializer.data)
            return Response({"status":error.context['success_code'], "data": response}, status=status.HTTP_200_OK)



class GlobalTransactionViews(APIView):

    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['code','name']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.GlobalTransaction.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.ListGlobalTransactionSerializer(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.GlobalTransaction.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"Section" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.GlobalTransaction.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        serializer = cSerializer.ListGlobalTransactionSerializer(lists, many=True)
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)




class GlobalTransactionDetailsViews(APIView):

    def get_object(self,pk):

        try:
            return models.GlobalTransactionDetails.objects.get(pk = pk)
        except GlobalTransactionDetails.DoesNotExist:
            raise Http404


    def post(self,request, pk = None):

        #print(request.data,"YYYYYYYYYYY")
        if 'id' not in request.data:
            return Response({"status":error.context['error_code'], "message" : "id" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        else: 

            if 'id' in request.data:
                pk = request.data['id']
                created_ip = Common.get_client_ip(request)
                created_by = request.data['created_by']
                project = request.data['project_id']
                sub_module = request.data['sub_module']

                if not pk:

                    global_transaction = {
                        'module': request.data['module'],
                        'form': request.data['form'] if request.data['form'] else None,
                        'sub_module': 1,
                        'project': project,
                        'created_ip': created_ip,
                        'created_by': created_by,
                        'status':1
                    }

                    #print(global_transaction,"HHHHHHHHHHHHHHHhh")

                    saveserializelog = cSerializer.GlobalTransactionSerializer(data = global_transaction)
                    if saveserializelog.is_valid():
                        saveserializelog.save()
                        running_id = saveserializelog.data['id']


                        for row in (request.data['psr']):

                            models.GlobalTransactionDetails.objects.create(
                                global_transaction_id = running_id,
                                sub_module_id = row['sub_module'],
                                section_id = row['id'],
                                #sub_section = row['sub_section'],
                                paragraph = row['paragraph'] if 'paragraph' in row else '',
                                project_id = project,
                                status = 1,                                
                                created_ip = created_ip,
                                created_by_id = created_by,
                            )

                            if len(row['subsections'])>0:
                                for sub_row in (row['subsections']):
                                    models.GlobalTransactionDetails.objects.create(
                                        global_transaction_id = running_id,
                                        sub_module_id = sub_row['sub_module'],
                                        section_id = row['id'],
                                        sub_section_id = sub_row['id'],
                                        paragraph = sub_row['paragraph'] if 'paragraph' in sub_row else '',              
                                        project_id = project,
                                        status = 1,
                                        created_ip = created_ip,
                                        created_by_id = created_by,
                                    )

                                    if len(sub_row['subsubsections'])>0:
                                        for sub_sub_row in (sub_row['subsubsections']):
                                            models.GlobalTransactionDetails.objects.create(
                                                global_transaction_id = running_id,
                                                sub_module_id = sub_sub_row['sub_module'],
                                                section_id = row['id'],
                                                sub_section_id = sub_row['id'],
                                                sub_sub_section_id = sub_sub_row['id'],
                                                paragraph = sub_sub_row['paragraph'] if 'paragraph' in sub_sub_row else '',
                                                project_id = project,
                                                status = 1,
                                                created_ip = created_ip,
                                                created_by_id = created_by,
                                            )

                        return Response({"status" :error.context['success_code'],"message": "Global transaction"+language.context[language.defaultLang]['insert'], "data":saveserializelog.data}, status=status.HTTP_200_OK)
                    # else:
                    #     return Response({"status" :error.context['error_code'],"message": error.serializerError(saveserializelog)}, status=status.HTTP_200_OK)



                    # arr = [{'sub_module' : 2,
                    #         'section' :2,
                    #         'sub_section' : 2,
                    #         'sub_sub_section' : 3,
                    #         'paragraph' : 'Para1'},
                    #         {'sub_module' : 2,
                    #         'section' : 2,
                    #         'sub_section' : 2,
                    #         'sub_sub_section' : 4,
                    #         'paragraph' : 'Para2'}]

                    # for row in arr:
                    #     request_data = {
                    #         'global_transaction' : running_id,
                    #         'sub_module' : row['sub_module'],
                    #         'section' : row['section'],
                    #         'sub_section' : row['sub_section'],
                    #         'sub_sub_section' : row['sub_sub_section'],
                    #         'paragraph' : row['paragraph'],
                    #         'created_ip' : created_ip,
                    #         'created_by' : created_by,
                    #         'status': 1
                    #     }
                
                    #     saveserialize = cSerializer.GlobalTransactionDetailsSerializer(data = request_data)
                    #     if saveserialize.is_valid():
                    #         saveserialize.save()
                    #     else:
                    #         return Response({"status" :error.context['error_code'], "message": error.serializerError(saveserialize)}, status=status.HTTP_200_OK)    

                    # if saveserializelog.is_valid():
                    #     return Response({"status" :error.context['success_code'], "message":"Global transaction" +language.context[language.defaultLang]['insert'], "data":saveserializelog.data}, status=status.HTTP_200_OK)
                    # else:
                    #     return Response({"status" :error.context['error_code'], "message": error.serializerError(saveserializelog)}, status=status.HTTP_200_OK)


                else:
                    return Response({"status@" : {"id@" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)

                # else:
                #     if request.data['status'] != 3:
                #         # duplicate_name = models.GlobalTransactionDetails.objects.values('name').filter(name=request.data['name']).exclude(Q(id=request.data['id']) | Q(status=3))
                #         # if duplicate_name:   
                #         #      return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit name'] },status=status.HTTP_200_OK)                                            

                #     list = self.get_object(pk)

                #     saveserialize = cSerializer.GlobalTransactionDetailsSerializer(list, data = request.data, partial= True)
                #     if saveserialize.is_valid():
                #         saveserialize.save()
                #         return Response({"status" :error.context['success_code'], "message":"Global transaction" +language.context[language.defaultLang]['update'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                #     else:
                #         return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)
            else: 
                return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)




class GlobalTransactionDetailsViewsEdit(APIView):

    def get_object(self,pk):

        try:
            return models.GlobalTransactionDetails.objects.get(pk = pk)
        except GlobalTransactionDetails.DoesNotExist:
            raise Http404


    def post(self,request, pk = None):


        #print(request.data,"FFFFFF")
        #print(request.data['psr'],"EEEEEEE")
        #print(request.data['psr'][8]['section_para'][0]['paragraph'],"GGGGGGGGGGGGGgg")
        if 'id' not in request.data:
            return Response({"status":error.context['error_code'], "message" : "id" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        else: 

            if 'id' in request.data:
                pk = request.data['id']
                created_ip = Common.get_client_ip(request)
                created_by = request.data['created_by']
                module = request.data['module']
                sub_module = request.data['sub_module']
                project = request.data['project_id']
                form = request.data['form']

                ### Mapping ###
                if 'mapping' in request.data:
                    sub_module_mapping = request.data['sub_module_mapping']
                    mapping = request.data['mapping']

                    count = models.GlobalTransaction.objects.filter(module_id = request.data['module'], sub_module_id = sub_module_mapping, project_id = project).count()

                    if count > 0:
                        approveRes = models.GlobalTransaction.objects.values('approved_status','approved_by','approved_on','approved_remark','project_id').filter(module_id = request.data['module'], sub_module_id = sub_module_mapping, project_id = project).get()
                        approved_status = approveRes['approved_status']
                        approved_by = approveRes['approved_by']
                        approved_on = approveRes['approved_on']
                        approved_remark = approveRes['approved_remark']


                    models.GlobalTransaction.objects.filter(module_id = request.data['module'], sub_module_id = sub_module_mapping, project_id = project).delete()

                    global_transaction = {
                        'module': request.data['module'],
                        'sub_module': sub_module_mapping,
                        'project': project,
                        'created_ip': created_ip,
                        'created_by': created_by,
                        "approved_status" : approved_status,
                        "approved_by" : approved_by,
                        "approved_on" : approved_on,
                        "approved_remark" : approved_remark,                        
                        'status':1
                    }

                    saveserializelog = cSerializer.GlobalTransactionSerializer(data = global_transaction)
                    if saveserializelog.is_valid():
                        saveserializelog.save()
                        running_id = saveserializelog.data['id']

                    for row in (mapping):

                        models.GlobalTransactionDetails.objects.create(
                            global_transaction_id = running_id,
                            sub_module_id = row['sub_module'],
                            section_id = row['id'],
                            paragraph = row['section_para'][0]['paragraph'],
                            project_id = project,
                            status = 1,
                            created_ip = created_ip,
                            created_by_id = created_by,
                        )

                        if len(row['sub_section'])>0:
                            for sub_row in (row['sub_section']):
                                models.GlobalTransactionDetails.objects.create(
                                    global_transaction_id = running_id,
                                    sub_module_id = row['sub_module'],
                                    section_id = row['id'],
                                    sub_section_id = sub_row['id'],
                                    paragraph = sub_row['sub_section_para'][0]['paragraph'],
                                    project_id = project,
                                    status = 1,
                                    created_ip = created_ip,
                                    created_by_id = created_by,
                                )

                                if len(sub_row['sub_sub_section'])>0:
                                    for sub_sub_row in (sub_row['sub_sub_section']):
                                        models.GlobalTransactionDetails.objects.create(
                                            global_transaction_id = running_id,
                                            sub_module_id = row['sub_module'],
                                            section_id = row['id'],
                                            sub_section_id = sub_row['id'],
                                            sub_sub_section_id = sub_sub_row['id'],
                                            paragraph = sub_sub_row['sub_sub_section_para'][0]['paragraph'],
                                            project_id = project,
                                            status = 1,
                                            created_ip = created_ip,
                                            created_by_id = created_by,
                                        )


                ###
                count = models.GlobalTransaction.objects.filter(id = pk, project_id = project).count()

                if count > 0:
                    approveRes = models.GlobalTransaction.objects.values('approved_status','approved_by','approved_on','approved_remark','project_id').filter(id = pk, project_id = project).get()
                    approved_status = approveRes['approved_status']
                    approved_by = approveRes['approved_by']
                    approved_on = approveRes['approved_on']
                    approved_remark = approveRes['approved_remark']


                models.GlobalTransaction.objects.filter( id = pk).delete()
                
                if pk:
                    if count > 0:
                        global_transaction = {
                            'module': module,
                            'sub_module': sub_module,
                            'form':form,
                            'created_ip': created_ip,
                            'created_by': created_by,
                            'project': project,
                            "approved_status" : approved_status,
                            "approved_by" : approved_by,
                            "approved_on" : approved_on,
                            "approved_remark" : approved_remark,
                            'status':1
                        }
                    else:
                        global_transaction = {
                            'module': module,
                            'form':form,
                            'sub_module': sub_module,
                            'created_ip': created_ip,
                            'created_by': created_by,
                            'project': project,
                            'status':1
                        }

                    saveserializelog = cSerializer.GlobalTransactionSerializer(data = global_transaction)
                    if saveserializelog.is_valid():
                        saveserializelog.save()
                        running_id = saveserializelog.data['id']

                        for row in (request.data['psr']):

                            # print(row['section_para'],'gghhh--->')
                            # if 'paragraph' in row['section_para']:
                            #     print(row['section_para']['paragraph'],'gg')
                            # else:
                            #     print(row['section_para'][0]['paragraph'],'jj')

                            #print(row['section_para'][0]['paragraph'],"PARRRRRRR_11")
                            #print(row['section_para']['paragraph'],"PARRRRRRR_22")
                            #print(row['section_para'].paragraph,'zzzzzzzz')
                            #print(row['section_para'][0]['paragraph'],'xxxxxxxxxxxxx')
                            models.GlobalTransactionDetails.objects.create(
                                global_transaction_id = running_id,
                                sub_module_id = row['sub_module'],
                                section_id = row['id'],
                                #sub_section = row['sub_section'],
                                #paragraph = row['paragraph'] if 'paragraph' in row else '',
                                #paragraph = row['section_para']['paragraph'] if 'paragraph' in row['section_para'] else row['section_para'][0]['paragraph'],
                                paragraph = row['section_para'][0]['paragraph'],
                                project_id = project,
                                status = 1,
                                created_ip = created_ip,
                                created_by_id = created_by,
                            )

                            if len(row['sub_section'])>0:
                                for sub_row in (row['sub_section']):
                                    models.GlobalTransactionDetails.objects.create(
                                        global_transaction_id = running_id,
                                        sub_module_id = row['sub_module'],
                                        section_id = row['id'],
                                        sub_section_id = sub_row['id'],
                                        #paragraph = sub_row['paragraph'] if 'paragraph' in sub_row else '',
                                        paragraph = sub_row['sub_section_para'][0]['paragraph'],
                                        project_id = project,
                                        status = 1,
                                        created_ip = created_ip,
                                        created_by_id = created_by,
                                    )

                                    if len(sub_row['sub_sub_section'])>0:
                                        for sub_sub_row in (sub_row['sub_sub_section']):
                                            models.GlobalTransactionDetails.objects.create(
                                                global_transaction_id = running_id,
                                                sub_module_id = row['sub_module'],
                                                section_id = row['id'],
                                                sub_section_id = sub_row['id'],
                                                sub_sub_section_id = sub_sub_row['id'],
                                                #paragraph = sub_sub_row['paragraph'] if 'paragraph' in sub_sub_row else '',
                                                paragraph = sub_sub_row['sub_sub_section_para'][0]['paragraph'] if sub_sub_row['sub_sub_section_para'] else '',
                                                project_id = project,
                                                status = 1,
                                                created_ip = created_ip,
                                                created_by_id = created_by,
                                            )

                        return Response({"status" :error.context['success_code'],"message": "Global transaction"+language.context[language.defaultLang]['update'], "data":saveserializelog.data}, status=status.HTTP_200_OK)
                    


                else:
                    return Response({"status@" : {"id@" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)

                
            else: 
                return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)


class GlobalTransactionDetailsViewsDelete(APIView):
    def get_object(self, pk):
            try:
                return models.GlobalTransaction.objects.get(pk = pk)
            except models.GlobalTransaction.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):

        #print(request.data)
        #pass
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']
            if not request.data['project']:
                return Response({"status":error.context['error_code'], "message": "Project is missing"}, status=status.HTTP_200_OK)

            else:

                list = self.get_object(pk)
                
                saveserialize = cSerializer.GlobalTransactionSerializer(list, data = request.data, partial= True)                
                if saveserialize.is_valid():
                    saveserialize.save()

                    return Response({"status" :error.context['success_code'], "message":"Global transaction" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)



class GlobalTransactionDetailsViews_Dummy(APIView):

    def get_object(self,pk):

        try:
            return models.GlobalTransactionDetails.objects.get(pk = pk)
        except GlobalTransactionDetails.DoesNotExist:
            raise Http404


    def post(self,request, pk = None):

        #print(request.data.module,"YYYYYYYYYYY")
        if 'id' not in request.data:
            return Response({"status":error.context['error_code'], "message" : "id" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        else: 

            if 'id' in request.data:
                pk = request.data['id']
                created_ip = Common.get_client_ip(request)
                created_by = request.data['created_by']
                #status = request.data['status']

                ###
                #models.GlobalTransactionDetails.objects.filter(id=request.data['global_transaction']).delete()
                if not pk:

                    global_transaction = {
                        'module': request.data['module'],
                        'created_ip': created_ip,
                        'created_by': created_by,
                        'status':1
                    }

                    print(global_transaction,"HHHHHHHHHHHHHHHhh")

                    saveserializelog = cSerializer.GlobalTransactionSerializer(data = global_transaction)
                    if saveserializelog.is_valid():
                        saveserializelog.save()
                        running_id = saveserializelog.data['id']


                    arr = [{'sub_module' : 2,
                            'section' :2,
                            'sub_section' : 2,
                            'sub_sub_section' : 3,
                            'paragraph' : 'Para1'},
                            {'sub_module' : 2,
                            'section' : 2,
                            'sub_section' : 2,
                            'sub_sub_section' : 4,
                            'paragraph' : 'Para2'}]


                    # for row in request.data['sections1']:
                    #     request_data = {
                    #         'global_transaction' : running_id,
                    #         'sub_module' : request.data['sub_module'],
                    #         'section' : request.data['section'],
                    #         'sub_section' : request.data['sub_section'],
                    #         'sub_sub_section' : request.data['sub_sub_section'],
                    #         'paragraph' : request.data['paragraph'],
                    #         'created_ip' : created_ip,
                    #         'created_by' : created_by,
                    #         'status': 1
                    #     }

                    for row in arr:
                        request_data = {
                            'global_transaction' : running_id,
                            'sub_module' : row['sub_module'],
                            'section' : row['section'],
                            'sub_section' : row['sub_section'],
                            'sub_sub_section' : row['sub_sub_section'],
                            'paragraph' : row['paragraph'],
                            'created_ip' : created_ip,
                            'created_by' : created_by,
                            'status': 1
                        }
                
                        saveserialize = cSerializer.GlobalTransactionDetailsSerializer(data = request_data)
                        if saveserialize.is_valid():
                            saveserialize.save()
                        else:
                            return Response({"status" :error.context['error_code'], "message": error.serializerError(saveserialize)}, status=status.HTTP_200_OK)    

                    if saveserializelog.is_valid():
                        return Response({"status" :error.context['success_code'], "message":"Global transaction" +language.context[language.defaultLang]['insert'], "data":saveserializelog.data}, status=status.HTTP_200_OK)
                    else:
                        return Response({"status" :error.context['error_code'], "message": error.serializerError(saveserializelog)}, status=status.HTTP_200_OK)


                else:
                    return Response({"status@" : {"id@" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)

                # else:
                #     if request.data['status'] != 3:
                #         # duplicate_name = models.GlobalTransactionDetails.objects.values('name').filter(name=request.data['name']).exclude(Q(id=request.data['id']) | Q(status=3))
                #         # if duplicate_name:   
                #         #      return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit name'] },status=status.HTTP_200_OK)                                            

                #     list = self.get_object(pk)

                #     saveserialize = cSerializer.GlobalTransactionDetailsSerializer(list, data = request.data, partial= True)
                #     if saveserialize.is_valid():
                #         saveserialize.save()
                #         return Response({"status" :error.context['success_code'], "message":"Global transaction" +language.context[language.defaultLang]['update'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                #     else:
                #         return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)
            else: 
                return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)



class GlobalTransactionEditViews(APIView):

    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['code','name']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.GlobalTransaction.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.ListGlobalTransactionEditSerializer(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.GlobalTransaction.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"Section" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.GlobalTransaction.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        serializer = cSerializer.ListGlobalTransactionEditSerializer(lists, many=True)
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)



class GlobalTransactionEdit(APIView):

    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['code','name']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.GlobalTransaction.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.ListGlobalTransactionAllEditSerializer(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.GlobalTransaction.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"Section" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.GlobalTransaction.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        serializer = cSerializer.ListGlobalTransactionAllEditSerializer(lists, many=True)
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)



class GlobalTransactionEdit_1(APIView):

    def get(self, request, pk = None):
        #print(request.GET.get('module'),"WWWWWWWWWWWWW")
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['code','name']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.GlobalTransaction.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.ListGlobalTransactionAllEditSerializer_1(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.GlobalTransaction.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"Section" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.GlobalTransaction.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        #print(lists,"FFFFFFFFFFf")
        serializer = cSerializer.ListGlobalTransactionAllEditSerializer_1(lists, many=True)
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)



###################
from docx import *
from .serializer import ListGlobalTransactionAllEditSerializer_1
from docx.shared import Inches
from io import BytesIO
from django.http import StreamingHttpResponse
from django.template import Context, Template
from django.template import loader
from docxtpl import DocxTemplate
####################

def global_transaction(self, id=None):  
    globaltransaction = models.GlobalTransaction.objects.values('id','project__name','module__name','module_id','project_id').filter(id=id).first()
    moduleDet=masterModels.Module.objects.values('id','name').filter(id=globaltransaction['module_id']).first()
    globaltransactionname=getDataByFormModule(module_id=globaltransaction['module_id'],project_id=globaltransaction['project_id'])
    globaltransactionsss=ListGlobalTransactionAllEditSerializer_1(models.GlobalTransaction.objects.filter(id=id),many=True).data
    sytemdata=getDataBySystem(module_id=globaltransaction['module_id'],project_id=globaltransaction['project_id'])
    #print(sytemdata,'sytemdata')
    # for globaltransactionnames in globaltransactionname:
    #     for globaltransactionn in globaltransactionnames.sub_module:
    #         print('5656',globaltransactionn )
    doc = DocxTemplate("templates/Test.docx")
    context = { 'globaltransaction':globaltransaction,'globaltransactionname':globaltransactionname ,"moduleDet":moduleDet,"globaltransactionsss":globaltransactionsss,"sytemdata":sytemdata}
    #print('sytemdata',sytemdata)
    # print('globaltransactionname',globaltransactionname)

   
    doc.render(context)
    # doc.save("generated_doc.docx")
    doc_io = BytesIO() # create a file-like object
    doc.save(doc_io) # save data to file-like object
    doc_io.seek(0) # go to the beginning of the file-like object

    response = HttpResponse(doc_io.read())

    # Content-Disposition header makes a file downloadable
    response["Content-Disposition"] = "attachment; filename="+(globaltransaction['project__name']+" - "+globaltransaction['module__name']+" Document")+".docx"

    # Set the appropriate Content-Type for docx file
    response["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    if response:
        #convert_pdf_word(id)
        return response

def global_transaction_submodule(self, id=None, form_id=None):  
    print(form_id,"GGGGGGGGGGG")
    globaltransaction = models.GlobalTransaction.objects.values('id','project__name','module__name','form__name','module_id','project_id').filter(id=id).first()
    moduleDet=masterModels.Module.objects.values('id','name').filter(id=globaltransaction['module_id']).first()
    globaltransactionname=getSubmoduleDataByFormModule(module_id=globaltransaction['module_id'],project_id=globaltransaction['project_id'],form_id=form_id)
    systemdata=getDataBySystemSubmodule(module_id=globaltransaction['module_id'],project_id=globaltransaction['project_id'],form_id=form_id)
    # globaltransactionsss=ListGlobalTransactionAllEditSerializer_1(models.GlobalTransaction.objects.filter(id=id),many=True).data
    # for globaltransactionnames in globaltransactionname:
    #     for globaltransactionn in globaltransactionnames.sub_module:
    #         print('5656',globaltransactionn )
    doc = DocxTemplate("templates/submodule.docx")
    context = { 'globaltransaction':globaltransaction,'globaltransactionname':globaltransactionname ,"moduleDet":moduleDet,"systemdata":systemdata}
    print('systemdata',systemdata)
    # print('globaltransactionname',globaltransactionname)

   
    doc.render(context)
    # doc.save("generated_doc.docx")
    doc_io = BytesIO() # create a file-like object
    doc.save(doc_io) # save data to file-like object
    doc_io.seek(0) # go to the beginning of the file-like object

    response = HttpResponse(doc_io.read())

    # Content-Disposition header makes a file downloadable
    response["Content-Disposition"] = "attachment; filename="+(globaltransaction['project__name']+" - "+globaltransaction['form__name']+" Document")+".docx"

    # Set the appropriate Content-Type for docx file
    response["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    return response


def global_transaction_gls(self, id=None):  
    globaltransaction = models.GlobalTransaction.objects.values('id','project__name','module__name','module_id','project_id').filter(id=id).first()
    moduleDet=masterModels.Module.objects.values('id','name').filter(id=globaltransaction['module_id']).first()
    globaltransactionname=getDataByFormgls(module_id=globaltransaction['module_id'],project_id=globaltransaction['project_id'])
    sytemdata=getDataBySystemgls(module_id=globaltransaction['module_id'],project_id=globaltransaction['project_id'])
   
    # for globaltransactionnames in globaltransactionname:
    #     for globaltransactionn in globaltransactionnames.sub_module:
    #         print('5656',globaltransactionn )
    doc = DocxTemplate("templates/gls.docx")
    context = { 'globaltransaction':globaltransaction,'globaltransactionname':globaltransactionname ,'moduleDet':moduleDet,'sytemdata':sytemdata }
    # print('date',date)
    # print('context',context)

   
    doc.render(context)
    # doc.save("generated_doc.docx")
    doc_io = BytesIO() # create a file-like object
    doc.save(doc_io) # save data to file-like object
    doc_io.seek(0) # go to the beginning of the file-like object

    response = HttpResponse(doc_io.read())

    # Content-Disposition header makes a file downloadable
    response["Content-Disposition"] = "attachment; filename="+(globaltransaction['project__name']+" - "+globaltransaction['module__name']+" Document")+".docx"

    # Set the appropriate Content-Type for docx file
    response["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    return response
def global_transaction_bls(self, id=None):  
    globaltransaction = models.GlobalTransaction.objects.values('id','project__name','module__name','module_id','project_id').filter(id=id).first()
    moduleDet=masterModels.Module.objects.values('id','name').filter(id=globaltransaction['module_id']).first()
    globaltransactionname=getDataByFormbls(module_id=globaltransaction['module_id'],project_id=globaltransaction['project_id'])
    sytemdata=getDataBySystembls(module_id=globaltransaction['module_id'],project_id=globaltransaction['project_id'])
    doc = DocxTemplate("templates/gls.docx")
    context = { 'globaltransaction':globaltransaction,'globaltransactionname':globaltransactionname ,"moduleDet":moduleDet,'sytemdata':sytemdata}
    # print('date',date)
    # print('context',context)

   
    doc.render(context)
    # doc.save("generated_doc.docx")
    doc_io = BytesIO() # create a file-like object
    doc.save(doc_io) # save data to file-like object
    doc_io.seek(0) # go to the beginning of the file-like object

    response = HttpResponse(doc_io.read())

    # Content-Disposition header makes a file downloadable
    response["Content-Disposition"] = "attachment; filename="+(globaltransaction['project__name']+" - "+globaltransaction['module__name']+" Document")+".docx"

    # Set the appropriate Content-Type for docx file
    response["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    return response



def version_transaction(self, id=None):  
    globaltransaction = models.Version.objects.values('id','project__name','module__name','module_id','project_id').filter(id=id).first()
    moduleDet=masterModels.Module.objects.values('id','name').filter(id=globaltransaction['module_id']).first()
    globaltransactionname = getVersionDataByFormModule(module_id=globaltransaction['module_id'],project_id=globaltransaction['project_id'], version_id=id)
    systemdata = getSystemVersionDataByFormModule(module_id=globaltransaction['module_id'],project_id=globaltransaction['project_id'], version_id=id)
    #globaltransactionsss=ListGlobalTransactionAllEditSerializer_1(models.VersionTransaction.objects.filter(id=id),many=True).data
   # print('111',globaltransactionname)
    # for globaltransactionnames in globaltransactionname:
    #     for globaltransactionn in globaltransactionnames.sub_module:
    #         print('5656',globaltransactionn )
    doc = DocxTemplate("templates/Test_Version.docx")
    #context = { 'globaltransaction':globaltransaction,'globaltransactionname':globaltransactionname ,"moduleDet":moduleDet,"globaltransactionsss":globaltransactionsss}

    context = { 'globaltransaction':globaltransaction,'globaltransactionname':globaltransactionname ,"moduleDet":moduleDet,"systemdata":systemdata}
    # print('date',date)
    # print('globaltransactionname',globaltransactionname)

   
    doc.render(context)
    # doc.save("generated_doc.docx")
    doc_io = BytesIO() # create a file-like object
    doc.save(doc_io) # save data to file-like object
    doc_io.seek(0) # go to the beginning of the file-like object

    response = HttpResponse(doc_io.read())

    # Content-Disposition header makes a file downloadable
    response["Content-Disposition"] = "attachment; filename="+(globaltransaction['project__name']+" - "+globaltransaction['module__name']+" Document")+".docx"

    # Set the appropriate Content-Type for docx file
    response["Content-Type"] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

    return response


######## SSS Mapping ###################

class SSSMappingViews(APIView):

    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['name','description']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.SSSMapping.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.ListSSSMappingSerializer(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.SSSMapping.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"SSS mapping" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.SSSMapping.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        serializer = cSerializer.ListSSSMappingSerializer(lists, many=True)
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)

class SSSMappingDetailViews(APIView):

    def get_object(self,pk):

        try:
            return models.SSSMapping.objects.get(pk = pk)
        except SSSMapping.DoesNotExist:
            raise Http404


    def post(self,request, pk = None):

        # if 'name' not in request.data and request.data['status'] != 3:
        #     return Response({"status":error.context['error_code'], "message" : "Name" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        if 'code' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "Code" +language.context[language.defaultLang]['missing'] },status=status.HTTP_200_OK)
        else: 

            if 'code' in request.data:
                request.data['code']=(request.data['code']).upper()
            # if 'sequence' in request.data:
            #     request.data['sequence']=request.data['sequence'] if(request.data['sequence']!='')  else 0
            if 'id' in request.data:
                pk = request.data['id']
                created_ip = Common.get_client_ip(request)                                       
                request.data['created_ip'] = created_ip  
               
                if not pk: 
                    duplicate_code = models.SSSMapping.objects.values('code').filter(code=request.data['code']).exclude(status=3)
                    #duplicate_name = models.SSSMapping.objects.values('name').filter(name=request.data['name']).exclude(status=3)
                            
                    if duplicate_code:            
                        return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit code'] },status=status.HTTP_200_OK)                    
                    
                    # elif duplicate_name:   
                    #     return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit name'] },status=status.HTTP_200_OK)                    
                    
                    else:
                
                        saveserialize = cSerializer.SSSMappingSerializer(data = request.data)
                        if saveserialize.is_valid():
                            saveserialize.save()
                            return Response({"status" :error.context['success_code'], "message":"New SSS mapping" +language.context[language.defaultLang]['insert'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                        else:
                            return Response({"status" :error.context['error_code'], "message": error.serializerError(saveserialize)}, status=status.HTTP_200_OK)

                else:
                    if request.data['status'] != 3:
                        duplicate_code = models.SSSMapping.objects.values('code').filter(code=request.data['code']).exclude(Q(id=request.data['id']) | Q(status=3))
                        #duplicate_name = models.SSSMapping.objects.values('name').filter(name=request.data['name']).exclude(Q(id=request.data['id']) | Q(status=3))
                        if duplicate_code:            
                             return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit code'] },status=status.HTTP_200_OK)                    
                        #if duplicate_name:   
                            #return Response({"status" :error.context['error_code'], "message":language.context[language.defaultLang]['exit name'] },status=status.HTTP_200_OK)                                            

                    list = self.get_object(pk)

                    saveserialize = cSerializer.SSSMappingSerializer(list, data = request.data, partial= True)
                    if saveserialize.is_valid():
                        saveserialize.save()
                        return Response({"status" :error.context['success_code'], "message":"SSS mapping" +language.context[language.defaultLang]['update'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                    else:
                        return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)
            else: 
                return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)   



########## Global Transaction Approve Status ##################

class GlobalTransactionApproveStatus(APIView):
    def get_object(self, pk):
            try:
                return models.GlobalTransaction.objects.get(pk = pk)
            except models.GlobalTransaction.DoesNotExist:
                raise Http404
    
    def post(self,request, pk = None):

        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']
            if not request.data['project']:
                return Response({"status":error.context['error_code'], "message": "Project is missing"}, status=status.HTTP_200_OK)
            else:
                request.data['id'] = request.data['id']
                request.data['project'] = request.data['project']
                created_ip = Common.get_client_ip(request)
                print('ssd',request.data['approved_status'])
                list = self.get_object(pk)
                
                saveserialize = cSerializer.GlobalTransactionSerializer(list, data = request.data, partial= True)                
                if saveserialize.is_valid():
                    saveserialize.save()
                    if 'approved_status' in request.data:
                        print('ddfasf')
                        models.ProjectModuleStatus.objects.filter(project_id=request.data['project'], project_module_master_id=request.data['form']).delete()
                        
                        models.ProjectModuleStatus.objects.create(
                            project_module_master_id = request.data['form'],
                            project_id = request.data['project'],
                            status  = request.data['approved_status']
                        )
                        approved_status = masterSerializer.GlobalStatusSerializer(masterModels.GlobalStatus.objects.filter(id = request.data['approved_status']),many=True).data
                        project_name=  masterSerializer.projectSerializer(masterModels.Project.objects.filter(id = request.data['project']),many=True).data
                        sub_mod=masterSerializer.SubModuleSerializer(masterModels.SubModule.objects.filter(id = request.data['form']),many=True).data
                        # print(project_name['name'],'fdsf')
                        for projectn in project_name:
                            for sub in sub_mod:
                                print(projectn['name'],'fdsf')
                                for status_data in approved_status:
                                    models.ProjectLog.objects.create(project_id=request.data['project'],sub_module_id=request.data['form'],created_on=datetime.now(),created_by_id=request.user.id,status=1,created_ip=created_ip,msg=projectn['name'] + ' has been  '+str(status_data['name']) +' Sucessfully in ' +str(sub['name']))
                                
                    return Response({"status" :error.context['success_code'], "message":"Initiation notes" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)



class GTListGlobalStatus(APIView):
    
    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                #list = models.InitiationNotes.objects.filter(pk=pk).exclude(status='3').get()
                list = models.InitiationNotes.objects.raw('SELECT t1.id, t1.name, t2.project_id FROM "master.project" as t1 INNER JOIN "psr.initiation_notes" as t2 ON t1.id = t2.project_idgroup by t2.project_id, t1.id')
                serializeobj = serializer.ListGlobalStatusSerializer(list)
                print(serializeobj.query())
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
      
        except models.GlobalTransaction.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Initiation Notes" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.GlobalTransaction.objects.values('approved_status','project__id','project__name').filter(module_id=1, sub_module_id=1,approved_status='2')

        
        # print(lists,"EEEEE")
        # serializer = cSerializer.ListGlobalStatusSerializer(lists, many=True)

        
        #print(serializer.query)
        return Response({"status":error.context['success_code'], "data": lists,"test":'ok'}, status=status.HTTP_200_OK)





class GlobalTransactionDownload(APIView):

    def get(self, request, pk = None):
        #print(request.GET.get('module'),"WWWWWWWWWWWWW")
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['code','name']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.GlobalTransaction.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.ListGlobalTransactionAllEditSerializer_1(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.GlobalTransaction.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"Section" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.GlobalTransaction.objects.order_by('project_id').distinct('project_id')


        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        # if order_type is None: 
        #     if order_column:
        #         lists = lists.order_by(order_column)  

        # elif order_type in 'asc':
        #     if order_column:
        #         lists = lists.order_by(order_column)
        #     else: 
        #         lists = lists.order_by('id')   

        # elif order_type in 'desc':
        #     if order_column:
        #         order_column = '-' + str(order_column)
        #         lists = lists.order_by(order_column)
        #     else: 
        #         lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        #lists = {project:{id:13}}
        serializer = cSerializer.ListGlobalTransactionAllEditSerializer_1(lists, many=True)
        #return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)

        return Response({"status" :error.context['success_code'], "message":"Global transaction" +language.context[language.defaultLang]['update'], "data":serializer.data}, status=status.HTTP_200_OK)


############### GT Internal Mapping ##############

class GTInternalMappingViews(APIView):

    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['name','description']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.SSSMapping.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.ListSSSMappingSerializer(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.SSSMapping.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"SSS mapping" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.GlobalTransactionDetails.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        serializer = cSerializer.GlobalTransactionDetailsSerializer(lists, many=True)
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)



##### Forms Mapping #########


class FormsViews(APIView):

    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['code','name']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.Forms.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.ListFormsSerializer(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.Forms.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"Section" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.Forms.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        serializer = cSerializer.ListFormsSerializer(lists, many=True)
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)

class FormsWOTMappingViews(APIView):

    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['code','name']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.Forms.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.ListFormsWOTMappingSerializer(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.Forms.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"Section" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.Forms.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        serializer = cSerializer.ListFormsWOTMappingSerializer(lists, many=True)
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)


class FormsMappingViews(APIView):

    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['code','name']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.FormsMapping.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.ListFormsMappingSerializer(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.FormsMapping.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"Section" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.FormsMapping.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        #print(lists,"####s")
        serializer = cSerializer.ListFormsMappingSerializer(lists, many=True)
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)
class MappingListViews(APIView):

    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['code','name']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.FormsMapping.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.FormsmapSerializer(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.FormsMapping.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"Section" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.FormsMapping.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        serializer = cSerializer.FormsmapSerializer(lists, many=True)
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)


class FormsMappingDetailsViews3333(APIView):

    def get_object(self,pk):

        try:
            return models.FormsMapping.objects.get(pk = pk)
        except FormsMapping.DoesNotExist:
            raise Http404


    def post(self,request, pk = None):

        if 'form' not in request.data:
            return Response({"status":error.context['error_code'], "message" : "Form" +language.context[language.defaultLang]['missing'] },status=status.HTTP_200_OK)
        else: 

            if 'id' in request.data:
                pk = request.data['id']
                created_ip = Common.get_client_ip(request)
                request.data['created_ip'] = created_ip
               
                if not pk:
                    #print(request.data,"GGGGGGGGGGg")
                    saveserialize = cSerializer.FormsMappingSerializer(data = request.data)
                    if saveserialize.is_valid():
                        saveserialize.save()
                        return Response({"status" :error.context['success_code'], "message":"Form mapping" +language.context[language.defaultLang]['insert'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                    else:
                        return Response({"status" :error.context['error_code'], "message": error.serializerError(saveserialize)}, status=status.HTTP_200_OK)

                else:

                    list = self.get_object(pk)

                    saveserialize = cSerializer.FormsMappingSerializer(list, data = request.data, partial= True)
                    if saveserialize.is_valid():
                        saveserialize.save()
                        return Response({"status" :error.context['success_code'], "message":"Form mapping" +language.context[language.defaultLang]['update'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                    else:
                        return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)
            else: 
                return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)  




class FormsMappingDetailsViews(APIView):

    def get_object(self,pk):

        try:
            return models.Forms.objects.get(pk = pk)
        except models.Forms.DoesNotExist:
            raise Http404


    def post(self,request, pk = None):

        #print(request.data)
        if 'form' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "Template" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        else: 
            if 'id' in request.data:
                pk = request.data['id']
                created_ip = Common.get_client_ip(request)
                request.data['created_ip'] = created_ip  


                count = models.FormsMapping.objects.filter(form_id = request.data['form']).count()
               
                models.FormsMapping.objects.filter(form_id = request.data['form']).delete()

                #print(pk,"PK")
                pk = None

                if not pk:

                    if request.data['form']!=9:
                        for index, access in enumerate((request.data['formmapping'])):

                            saveserialize = cSerializer.FormsMappingSerializer(data={
                                'form' : request.data['form'],
                                #"order": access['order'] if 'order' in access else None,
                                "order": index,
                                "Class": access['Class'],
                                "module" : access['module'],
                                "sub_module" : access['sub_module'],
                                "section" : access['section'],
                                "sub_section" : access['sub_section'],
                                "sub_sub_section" : access['sub_sub_section'],
                                "status" : request.data['status'],
                                "created_ip" : created_ip

                                })

                            if saveserialize.is_valid():
                                saveserialize.save()
                    else:

                        if count > 0:

                            for index, access in enumerate((request.data['formmapping'])):
                                saveserialize = cSerializer.FormsMappingSerializer(data={
                                    'form' : request.data['form'],
                                    #"order": access['order'] if 'order' in access else None,
                                    "order": index,
                                    "Class": access['Class'],
                                    "module" : access['module'],
                                    "sub_module" : access['sub_module'],
                                    "section" : access['section'],
                                    "sub_section" : access['sub_section'],
                                    "sub_sub_section" : access['sub_sub_section'],
                                    "status" : request.data['status'],
                                    "created_ip" : created_ip

                                    })

                                if saveserialize.is_valid():
                                    saveserialize.save()
                        else:

                            ss = cSerializer.AllFormsMappingSerializer(models.FormsMapping.objects.filter(form_id=8), many=True)

                            for index, access in enumerate((ss.data)):
                                saveserialize = cSerializer.FormsMappingSerializer(data={
                                    'form' : request.data['form'],
                                    "order": access['order'],
                                    "Class": access['Class'],
                                    "module" : access['module'],
                                    "sub_module" : access['sub_module'],
                                    "section" : access['section'],
                                    "sub_section" : access['sub_section'],
                                    "sub_sub_section" : access['sub_sub_section'],
                                    "status" : access['status'],
                                    'created_on' : access['created_on'],
                                    'created_by' : access['created_by'],
                                    'created_ip' : access['created_ip'],
                                    'modified_on' : access['modified_on'],
                                    'modified_by' : access['modified_by'],
                                    'modified_ip' : access['modified_ip']

                                    })

                                if saveserialize.is_valid():
                                    saveserialize.save()
                        

                    if count > 0:    
                        return Response({"status" :error.context['success_code'], "message":"Form mapping" +language.context[language.defaultLang]['update'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                    else:
                        return Response({"status" :error.context['success_code'], "message":"Form mapping" +language.context[language.defaultLang]['insert'], "data":saveserialize.data}, status=status.HTTP_200_OK)

                else:                                       

                    list = self.get_object(pk)

                    saveserialize = cSerializer.TemplateConfigSerializer(list, data = request.data, partial= True)
                    if saveserialize.is_valid():
                        saveserialize.save()
                        return Response({"status" :error.context['success_code'], "message":"Template config" +language.context[language.defaultLang]['update'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                    else:
                        return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)
            else: 
                return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)







class GenerateTemplateViews(APIView):

    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['code','name']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.FormsMapping.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.TemplateGenerateSerializer(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.FormsMapping.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"Section" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.FormsMapping.objects.exclude(status='3').distinct('form_id')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        serializer = cSerializer.TemplateGenerateSerializer(lists, many=True)
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)



class getFormMapping(APIView):
    authentication_classes = [] #disables authentication
    permission_classes = [] #disables permission
    def post(self,request, pk = None):
        if 'module_id' not in request.data or request.data['module_id'] == '':
            return Response({"status":error.context['error_code'], "message" : "Module id " +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        elif 'form_id' not in request.data or request.data['form_id'] == '':
            return Response({"status":error.context['error_code'], "message" : "Form id " +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        else:
            #print(request.data['directorate_id'],"RRRRRRRDirectorate_id")
            project_id=request.data['project_id'] if 'project_id' in request.data and request.data['project_id'] else ''
            #ss=cSerializer.getFormMappingSerializer(models.FormsMapping.objects.filter(form_id=request.data['form_id']).distinct('sub_module_id'),many=True,context={"project_id":project_id,"directorate_id":request.data['directorate_id']})


            aa_ids = models.FormsMapping.objects.filter(form_id=request.data['form_id']).distinct('sub_module_id')

            print(aa_ids.query, "Suqery")

            #aa_list = models.FormsMapping.objects.filter(id__in=aa_ids).order_by('id')

            #print(aa_list.query, "TTTTT")

            # res = models.FormsMapping.objects.filter(form_id=request.data['form_id']).distinct('sub_module_id')
            # print(res.query,"GGGGG")

            ss=cSerializer.getFormMappingSerializer(models.FormsMapping.objects.filter(form_id=request.data['form_id']).distinct('sub_module_id'),many=True,context={"project_id":project_id,"directorate_id":request.data['directorate_id']})

            return Response({"status":error.context['success_code'] , "data": ss.data}, status=status.HTTP_200_OK)


class getFormMappingAllPSR(APIView):
    authentication_classes = [] #disables authentication
    permission_classes = [] #disables permission
    def post(self,request, pk = None):
        if 'module_id' not in request.data or request.data['module_id'] == '':
            return Response({"status":error.context['error_code'], "message" : "Module id " +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        elif 'form_id' not in request.data or request.data['form_id'] == '':
            return Response({"status":error.context['error_code'], "message" : "Form id " +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        else:
            project_id=request.data['project_id'] if 'project_id' in request.data and request.data['project_id'] else ''
            ss=cSerializer.getFormMappingSerializer(models.FormsMapping.objects.filter(module_id=1).distinct('sub_module_id'),many=True,context={"project_id":project_id,"directorate_id":request.data['directorate_id']})
            return Response({"status":error.context['success_code'] , "data": ss.data}, status=status.HTTP_200_OK)



class getVersion(APIView):
    authentication_classes = [] #disables authentication
    permission_classes = [] #disables permission
    def post(self,request, pk = None):
        if 'module_id' not in request.data or request.data['module_id'] == '':
            return Response({"status":error.context['error_code'], "message" : "Module id " +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        elif 'form_id' not in request.data or request.data['form_id'] == '':
            return Response({"status":error.context['error_code'], "message" : "Form id " +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        else:
            project_id=request.data['project_id'] if 'project_id' in request.data and request.data['project_id'] else ''
            version_id=request.data['version_id']
            ss=cSerializer.getVersionMappingSerializer(models.FormsMapping.objects.filter(module_id=1).distinct('sub_module_id'),many=True,context={"project_id":project_id,"version_id":version_id})
            return Response({"status":error.context['success_code'] , "data": ss.data}, status=status.HTTP_200_OK)



class addVersion(APIView):
    authentication_classes = [] #disables authentication
    permission_classes = [] #disables permission
    def post(self,request, pk = None):
        if 'module_id' not in request.data or request.data['module_id'] == '':
            return Response({"status":error.context['error_code'], "message" : "Module id " +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        elif 'form_id' not in request.data or request.data['form_id'] == '':
            return Response({"status":error.context['error_code'], "message" : "Form id " +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        else:
            project_id=request.data['project_id'] if 'project_id' in request.data and request.data['project_id'] else ''


            gt_id = request.data['gt_id']
            module_id = request.data['module_id']
            form_id = request.data['form_id']
            project_id = request.data['project_id']
            created_by = request.data['created_by']
            created_ip = Common.get_client_ip(request)

            #print(request.data,"########")
            #pass

            ss = cSerializer.getFormMappingSerializerVersion(models.FormsMapping.objects.filter(module_id=1).distinct('sub_module_id'),many=True,context={"project_id":project_id})

            data = saveVerionDetails(module_id, form_id, project_id, created_by, created_ip, ss.data, gt_id, request.data)
            return Response({"status":error.context['success_code'], "message":"Version" +language.context[language.defaultLang]['insert'], "data": ss.data}, status=status.HTTP_200_OK)


def saveVerionDetails(module_id, form_id, project_id, created_by, created_ip, data, gt_id, request):


    #print(request,"HHHJJJJ")
    #print(module_id, form_id, project_id, created_by, created_ip, data, gt_id)
    if module_id and form_id and project_id and data:

        gt_id = gt_id
        module = module_id
        form = form_id
        project = project_id
        created_by = created_by
        created_ip = created_ip

        #print(project,"project_id")
        #pass
        #print(module, form, project, data[0]['section'], "Hiiiiiiiiiiii")
        #return cSerializer.getFormMappingReportSerializer(models.FormsMapping.objects.filter(module_id=module_id).distinct('sub_module_id'),many=True,context={"project_id":project_id}).data

        #count = models.Version.objects.filter(module_id = module, form_id = form, project_id = project).count()

        count = models.Version.objects.filter(module_id = module, project_id = project).count()

        if count == 0:
            versionCount = '0.0.1'
        else:
            versionCount = '0.0.'+str(count+1)

        # version = {
        #     'module': module,
        #     'form': form,
        #     'project': project,
        #     'created_ip': created_ip,
        #     'created_ip': created_by,
        #     'version':versionCount
        # }

        # print(version, "Version")
        # #return []
        # #pass


        version = models.Version.objects.create(
            module_id = module,
            form_id = form,
            project_id = project,
            created_ip = created_ip,
            created_by_id = created_by,
            version = versionCount,
            status = 1,
        )
        version_id = version.id;

        # saveserializelog = cSerializer.VersionSerializer(data = version)
        # if saveserializelog.is_valid():
        #     saveserializelog.save()
        #     version_id = saveserializelog.data['id']

        if version_id:
            for row in (data[0]['section']):

                models.VersionTransaction.objects.create(
                    version_id = version_id,
                    global_transaction_id = gt_id,
                    sub_module_id = row['sub_module'],
                    section_id = row['id'],
                    paragraph = row['paragraph'] if 'paragraph' in row else '',
                    view = row['view'] if 'view' in row else '',
                    project_id = project,
                    status = 1,                                
                    created_ip = created_ip,
                    created_by_id = created_by,
                )

                if len(row['subsections'])>0:
                    for sub_row in (row['subsections']):

                        models.VersionTransaction.objects.create(
                            version_id = version_id,
                            global_transaction_id = gt_id,
                            sub_module_id = sub_row['sub_module'],
                            section_id = row['id'],
                            sub_section_id = sub_row['id'],
                            paragraph = sub_row['paragraph'] if 'paragraph' in sub_row else '',              
                            view = sub_row['view'] if 'view' in sub_row else '',
                            project_id = project,
                            status = 1,
                            created_ip = created_ip,
                            created_by_id = created_by,
                        )

                        if len(sub_row['subsubsections'])>0:
                            for sub_sub_row in (sub_row['subsubsections']):

                                models.VersionTransaction.objects.create(
                                    version_id = version_id,
                                    global_transaction_id = gt_id,
                                    sub_module_id = sub_sub_row['sub_module'],
                                    section_id = row['id'],
                                    sub_section_id = sub_row['id'],
                                    sub_sub_section_id = sub_sub_row['id'],
                                    paragraph = sub_sub_row['paragraph'] if 'paragraph' in sub_sub_row else '',
                                    view = sub_sub_row['view'] if 'view' in sub_row else '',
                                    project_id = project,
                                    status = 1,
                                    created_ip = created_ip,
                                    created_by_id = created_by,
                                )

            #return Response({"status" :error.context['success_code'],"message": formDet['name']+' '+('updated' if request.data['id'] and request.data['id'] else 'created')+' successfully'}, status=status.HTTP_200_OK)
            return Response({"status":error.context['success_code'] , "data": version_id}, status = status.HTTP_200_OK)



            #return data
            #return Response({"status1":error.context['success_code'] , "data1": saveserializelog.data,"gg":data.section}, status=status.HTTP_200_OK)

    else:
        return []



class getGlobalTransactionsOLD(APIView):

    def get(self, request, pk = None):
        #print(request.GET.get('module'),"WWWWWWWWWWWWW")
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['code','name']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.GlobalTransaction.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.GlobalTransactionsSerializer(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.GlobalTransaction.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"Section" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.GlobalTransaction.objects.exclude(status='3')

        if 'Authorized-Role' in request.headers:
            role_code=request.headers['Authorized-Role']
            user_role_id=request.headers['Authorized-By']
            process_id=request.user.process_id if request.user.process_id else None
            # notificationsDet=notificationsDet.exclude(id__in=models.NotificationUserLog.objects.values('notification_id').filter(user_id=request.user.id,user_role_id=user_role_id))
            
            if role_code!='admin':
                # notificationsDet=notificationsDet.filter(Q(user_role_id=user_role_id) & Q(process_id=process_id))
                # modules=masterModels.DataAccess.objects.values('module_id').filter(user_id=request.user.id)
                # sub_modules=masterModels.DataAccessSubModule.objects.values('sub_module_id').filter(data_access__user_id=request.user.id)
                # lists=lists.filter(module_id__in=modules)
                form_ids=models.DataAccessForms.objects.values('form_id').filter(user_id=request.user.id)
                lists=lists.filter(form_id__in=form_ids)

        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        #print(lists,"FFFFFFFFFFf")
        serializer = cSerializer.GlobalTransactionsSerializer(lists, many=True)
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)


class getGlobalTransactionsCurrent(APIView):

    def get(self, request, pk = None):
        #print(request.GET.get('module'),"WWWWWWWWWWWWW")
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['code','name']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  
            approved_level = request.GET.get("approved_level")

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')
            if approved_level is not None:
                normal_values.pop("approved_level")

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.GlobalTransaction.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.GlobalTransactionsSerializer(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.GlobalTransaction.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"Section" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.GlobalTransaction.objects.exclude(status='3')

        if 'Authorized-Role' in request.headers:
            role_code=request.headers['Authorized-Role']
            user_role_id=request.headers['Authorized-By']
            process_id=request.user.process_id if request.user.process_id else None
            # notificationsDet=notificationsDet.exclude(id__in=models.NotificationUserLog.objects.values('notification_id').filter(user_id=request.user.id,user_role_id=user_role_id))
            
            if role_code!='admin':
                # notificationsDet=notificationsDet.filter(Q(user_role_id=user_role_id) & Q(process_id=process_id))
                # modules=masterModels.DataAccess.objects.values('module_id').filter(user_id=request.user.id)
                # sub_modules=masterModels.DataAccessSubModule.objects.values('sub_module_id').filter(data_access__user_id=request.user.id)
                # lists=lists.filter(module_id__in=modules)
                form_ids=models.DataAccessForms.objects.values('form_id').filter(user_id=request.user.id)
                lists=lists.filter(form_id__in=form_ids)

        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        print(lists.query, "AAAAAAAAAAAa")
        if approved_level:
            if approved_level == "-1111":
                lists = lists.filter(approved_level=-1)
            elif approved_level == "-2111":
                lists = lists.exclude(approved_level=-1)
                role_code = request.headers["Authorized-Role"]
                role_id = int(request.headers["Authorized-By"])
                process_id = int(request.user.process_id)

                print(lists.query, "BBBBBBBBBB")
                print(role_id,"role_id123",process_id,"process_id123")
                if process_id == 1 and role_id == 2:
                    lists = lists.filter(
                        (Q(recommender=2) | Q(recommender=None))
                        & Q(initiater=1)
                    )
                    #print(lists, "role_id*******", role_id)
                elif process_id == 1 and role_id == 3:
                    lists = lists.filter(
                        (
                            (Q(approver=2) | Q(approver=None))
                            & Q(recommender=1)
                        )
                        | (Q(initiater=1) & Q(recommender=2))
                        | (Q(recommender=None))
                    )
                    #print(lists.get(), "role_id*******",role_id)
                elif process_id == 1 and role_id == 4:
                    lists = lists.filter(
                        (
                            (Q(approver=1) | Q(approver=None))
                            & Q(approver=1)
                        )
                        | (Q(recommender=1) & Q(approver=2))
                        | (Q(approver=None))
                    )
                    #print(lists.get(), "role_id*******", role_id)

                # elif process_id == 3 and role_id == 3:
                #     lists = lists.filter(
                #         (
                #             (Q(trial_recommender=2) | Q(trial_recommender=None))
                #             & Q(trial_initiater=1)
                #         )
                #         | (Q(approver=1) & Q(trial_initiater=2))
                #         | (Q(trial_initiater=None))
                #     )
                # elif process_id == 3 and role_id == 4:
                #     lists = lists.filter(
                #         (
                #             (Q(trial_approver=2) | Q(trial_approver=None))
                #             & Q(trial_recommender=1)
                #         )
                #         | (Q(trial_initiater=1) & Q(trial_recommender=2))
                #         | (Q(trial_recommender=None))
                #     )
                # elif process_id == 3 and role_id == 5:
                #     lists = lists.filter(
                #         (Q(trial_recommender=1) & Q(trial_approver=None))
                #         | (Q(trial_recommender=1) & Q(trial_approver=2))
                #    )

            elif approved_level == "-3111":
                role_code = request.headers["Authorized-Role"]
                role_id = int(request.headers["Authorized-By"])
                process_id = int(request.user.process_id)
                if process_id == 2 and role_id == 3:
                    # lists = lists.exclude(approved_level=-1)
                    lists = lists.filter(recommender=1, initiater=1)
                elif process_id == 2 and role_id == 4:
                    lists = lists.filter(approver=1, recommender=1)
                elif process_id == 2 and role_id == 5:
                    lists = lists.filter(trial_initiater=1, approver=1)
                # elif process_id == 3 and role_id == 3:
                #     lists = lists.filter(trial_recommender=1, trial_initiater=1)
                # elif process_id == 3 and role_id == 4:
                #     lists = lists.filter(trial_approver=1, trial_recommender=1)
                # elif process_id == 3 and role_id == 5:
                #     lists = lists.filter(trial_approver=1, trial_recommender=1)



        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        #print(lists.query,"FFFFFFFFFFf")
        serializer = cSerializer.GlobalTransactionsSerializer(lists, many=True, context={"request": request})
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)


class getGlobalTransactions(APIView):

    def get(self, request, pk = None):
        #print(request.GET.get('module'),"WWWWWWWWWWWWW")
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['code','name']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  
            approved_level = request.GET.get("approved_level")

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')
            if approved_level is not None:
                normal_values.pop("approved_level")

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.GlobalTransaction.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.GlobalTransactionsSerializer(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.GlobalTransaction.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"Section" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.GlobalTransaction.objects.exclude(status='3')

        if 'Authorized-Role' in request.headers:
            role_code=request.headers['Authorized-Role']
            user_role_id=request.headers['Authorized-By']
            process_id=request.user.process_id if request.user.process_id else None
            # notificationsDet=notificationsDet.exclude(id__in=models.NotificationUserLog.objects.values('notification_id').filter(user_id=request.user.id,user_role_id=user_role_id))
            
            if role_code!='admin':
                # notificationsDet=notificationsDet.filter(Q(user_role_id=user_role_id) & Q(process_id=process_id))
                # modules=masterModels.DataAccess.objects.values('module_id').filter(user_id=request.user.id)
                # sub_modules=masterModels.DataAccessSubModule.objects.values('sub_module_id').filter(data_access__user_id=request.user.id)
                # lists=lists.filter(module_id__in=modules)
                form_ids=models.DataAccessForms.objects.values('form_id').filter(user_id=request.user.id)
                lists=lists.filter(form_id__in=form_ids)

        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 


        if approved_level:
            lists = lists.filter(approved_level=approved_level)

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        print(lists.query,"FFFFFFFFFFf")
        serializer = cSerializer.GlobalTransactionsSerializer(lists, many=True, context={"request": request})
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)


def saveNotifications_OLD(data,level):
    formDet=models.Forms.objects.values('id','name').filter(id=data["form_id"]).first()
    projectDet=models.GlobalTransaction.objects.values('id','project__name','created_by__first_name','created_by__last_name').filter(id=data["transaction_id"]).first()
    data['message']=''
    if level==2:
        data['message']='Project '+projectDet['project__name']+' is initiated by '+projectDet['created_by__first_name']+' '+projectDet['created_by__last_name']+' in '+formDet['name']
    elif level==3:
        data['message']='Project '+projectDet['project__name']+' is recommended by '+projectDet['created_by__first_name']+' '+projectDet['created_by__last_name']+' in '+formDet['name']
    
    #print(data,"Star Data")
    notificationModels.NotificationUser.objects.create(**data)


    return {"status":error.context['success_code'],"message":"Notification sent successfully"}


def saveNotifications(data,level):
    formDet=models.Forms.objects.values('id','name').filter(id=data["form_id"]).first()
    projectDet=models.GlobalTransaction.objects.values('id','project__name','created_by__first_name','created_by__last_name').filter(id=data["transaction_id"]).first()

    data['message']=''
    if level==2:
        data['message']='Project '+projectDet['project__name']+' is initiated by '+projectDet['created_by__first_name']+' '+projectDet['created_by__last_name']+' in '+formDet['name']
    elif level==3:
        data['message']='Project '+projectDet['project__name']+' is recommended by '+projectDet['created_by__first_name']+' '+projectDet['created_by__last_name']+' in '+formDet['name']
    
    #print(data,"Star Data")
    notificationModels.NotificationUser.objects.create(**data)


    return {"status":error.context['success_code'],"message":"Notification sent successfully"}



class saveGlobalTransactionDetails(APIView):

    def post(self,request, pk = None):
        if 'id' not in request.data:
            return Response({"status":error.context['error_code'], "message" : "id" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        else: 
                user_role_id=''
                if 'Authorized-Role' in request.headers:
                    user_role_id=request.headers['Authorized-By']
                

                secondLevel=accessModels.ProcessFlow.objects.values('process_id','user_role_id').filter(level=2).first()

                
                # print("request.user",request.user.process_id)
                # print("user_role_id",user_role_id)
                # print("secondLevel",secondLevel)

                # return Response({"status":error.context['error_code'], "message" : "id" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
                pk = request.data['id']
                created_ip = Common.get_client_ip(request)
                created_by = request.user.id
                project = request.data['project_id']
                form = request.data['form']
                formDet=models.Forms.objects.values('name','code').filter(id=form).first()


                #### System . ###
                section_temp = models.SystemEquipmentCompartmentTemp.objects.values("id","project_id","section_id","sub_section_id","sub_sub_section_id","type","ser","name","numbers","capabilities_feature","weight_volume_power_consumption","location","interface","procurement_router","vendor","cost","standards","sustenance","equipment","features","layout","special_requirements","created_ip","created_by_id").filter(project_id=project, section_id__isnull=False, sub_section_id__isnull=True, sub_sub_section_id__isnull=True)

                if section_temp:
                    # models.SystemEquipmentCompartment.objects.filter(project_id=section_temp[0]['project_id'],section_id=section_temp[0]['section_id'],sub_section_id__isnull=True,sub_sub_section_id__isnull=True, type=section_temp[0]['type']).delete()

                    for section in section_temp: # Section
                        models.SystemEquipmentCompartment.objects.filter(project_id=section['project_id'],section_id=section['section_id'],sub_section_id__isnull=True,sub_sub_section_id__isnull=True, type=section['type']).delete()

                sub_section_temp = models.SystemEquipmentCompartmentTemp.objects.values("id","project_id","section_id","sub_section_id","sub_sub_section_id","type","ser","name","numbers","capabilities_feature","weight_volume_power_consumption","location","interface","procurement_router","vendor","cost","standards","sustenance","equipment","features","layout","special_requirements","created_ip","created_by_id").filter(project_id=project, section_id__isnull=False, sub_section_id__isnull=False, sub_sub_section_id__isnull=True)

                if sub_section_temp:
                    # models.SystemEquipmentCompartment.objects.filter(project_id=sub_section_temp[0]['project_id'],section_id=sub_section_temp[0]['section_id'],sub_section_id=sub_section_temp[0]['sub_section_id'],sub_sub_section_id__isnull=True, type=sub_section_temp[0]['type']).delete()

                    for sub_section in sub_section_temp: # Sub Section
                        models.SystemEquipmentCompartment.objects.filter(project_id=sub_section['project_id'],section_id=sub_section['section_id'],sub_section_id=sub_section_temp[0]['sub_section_id'],sub_sub_section_id__isnull=True, type=sub_section['type']).delete()

                sub_sub_section_temp = models.SystemEquipmentCompartmentTemp.objects.values("id","project_id","section_id","sub_section_id","sub_sub_section_id","type","ser","name","numbers","capabilities_feature","weight_volume_power_consumption","location","interface","procurement_router","vendor","cost","standards","sustenance","equipment","features","layout","special_requirements","created_ip","created_by_id").filter(project_id=project, section_id__isnull=False, sub_section_id__isnull=False, sub_sub_section_id__isnull=False)

                if sub_sub_section_temp:
                    # models.SystemEquipmentCompartment.objects.filter(project_id=sub_sub_section_temp[0]['project_id'],section_id=sub_sub_section_temp[0]['section_id'],sub_section_id=sub_sub_section_temp[0]['sub_section_id'],sub_sub_section_id=sub_sub_section_temp[0]['sub_sub_section_id'], type=sub_sub_section_temp[0]['type']).delete()

                    for sub_sub_section in sub_sub_section_temp: # Sub Sub Section
                        models.SystemEquipmentCompartment.objects.filter(project_id=sub_sub_section['project_id'],section_id=sub_sub_section['section_id'],sub_section_id=sub_sub_section['sub_section_id'],sub_sub_section_id=sub_sub_section['sub_sub_section_id'], type=sub_sub_section['type']).delete()

                #print(tempRes, "EEEEEEEEE")

                for section in section_temp: # Section

                    if section['section_id']!='' and section['sub_section_id']==None and section['sub_sub_section_id']==None:
                        models.SystemEquipmentCompartment.objects.create(
                            project_id = section['project_id'],
                            section_id = section['section_id'],
                            sub_section_id = section['sub_section_id'] if 'sub_section_id' in section else '',
                            sub_sub_section_id = section['sub_sub_section_id'] if 'sub_sub_section_id' in section else '',
                            type = section['type'] if 'type' in section else '',
                            ser = section['ser'] if 'ser' in section else '',
                            name = section['name'] if 'name' in section else '',
                            numbers = section['numbers'] if 'numbers' in section else '',
                            capabilities_feature = section['capabilities_feature'] if 'capabilities_feature' in section else '',
                            weight_volume_power_consumption = section['weight_volume_power_consumption'] if 'weight_volume_power_consumption' in section else '`',
                            location = section['location'] if 'location' in section else '',
                            interface = section['interface'] if 'interface' in section else '',
                            procurement_router = section['procurement_router'] if 'procurement_router' in section else '',
                            vendor = section['vendor'] if 'vendor' in section else '',
                            cost = section['cost'] if 'cost' in section else '',
                            standards = section['standards'] if 'standards' in section else '',
                            sustenance = section['sustenance'] if 'sustenance' in section else '',
                            equipment = section['equipment'] if 'equipment' in section else '',
                            features = section['features'] if 'features' in section else '',
                            layout = section['layout'] if 'layout' in section else '',
                            special_requirements = section['special_requirements'] if 'special_requirements' in section else '',
                            created_ip = created_ip,
                            created_by_id = created_by
                        )

                for sub_section in sub_section_temp: # Sub Section

                    if sub_section['section_id']!='' and sub_section['sub_section_id']!='' and sub_section['sub_sub_section_id']==None:
                        models.SystemEquipmentCompartment.objects.create(
                            project_id = sub_section['project_id'],
                            section_id = sub_section['section_id'],
                            sub_section_id = sub_section['sub_section_id'] if 'sub_section_id' in sub_section else '',
                            sub_sub_section_id = sub_section['sub_sub_section_id'] if 'sub_sub_section_id' in sub_section else '',
                            type = sub_section['type'] if 'type' in sub_section else '',
                            ser = sub_section['ser'] if 'ser' in sub_section else '',
                            name = sub_section['name'] if 'name' in sub_section else '',
                            numbers = sub_section['numbers'] if 'numbers' in sub_section else '',
                            capabilities_feature = sub_section['capabilities_feature'] if 'capabilities_feature' in sub_section else '',
                            weight_volume_power_consumption = sub_section['weight_volume_power_consumption'] if 'weight_volume_power_consumption' in sub_section else '`',
                            location = sub_section['location'] if 'location' in sub_section else '',
                            interface = sub_section['interface'] if 'interface' in sub_section else '',
                            procurement_router = sub_section['procurement_router'] if 'procurement_router' in sub_section else '',
                            vendor = sub_section['vendor'] if 'vendor' in sub_section else '',
                            cost = sub_section['cost'] if 'cost' in sub_section else '',
                            standards = sub_section['standards'] if 'standards' in sub_section else '',
                            sustenance = sub_section['sustenance'] if 'sustenance' in sub_section else '',
                            equipment = sub_section['equipment'] if 'equipment' in sub_section else '',
                            features = sub_section['features'] if 'features' in sub_section else '',
                            layout = sub_section['layout'] if 'layout' in sub_section else '',
                            special_requirements = sub_section['special_requirements'] if 'special_requirements' in sub_section else '',
                            created_ip = created_ip,
                            created_by_id = created_by
                        )

                for sub_sub_section in sub_sub_section_temp: # Sub Sub Section

                    if sub_sub_section['section_id']!='' and sub_sub_section['sub_section_id']!='' and sub_sub_section['sub_sub_section_id']!='':
                        models.SystemEquipmentCompartment.objects.create(
                            project_id = sub_sub_section['project_id'],
                            section_id = sub_sub_section['section_id'],
                            sub_section_id = sub_sub_section['sub_section_id'] if 'sub_section_id' in sub_sub_section else '',
                            sub_sub_section_id = sub_sub_section['sub_sub_section_id'] if 'sub_sub_section_id' in sub_sub_section else '',
                            type = sub_sub_section['type'] if 'type' in sub_sub_section else '',
                            ser = sub_sub_section['ser'] if 'ser' in sub_sub_section else '',
                            name = sub_sub_section['name'] if 'name' in sub_sub_section else '',
                            numbers = sub_sub_section['numbers'] if 'numbers' in sub_sub_section else '',
                            capabilities_feature = sub_sub_section['capabilities_feature'] if 'capabilities_feature' in sub_sub_section else '',
                            weight_volume_power_consumption = sub_sub_section['weight_volume_power_consumption'] if 'weight_volume_power_consumption' in sub_sub_section else '`',
                            location = sub_sub_section['location'] if 'location' in sub_sub_section else '',
                            interface = sub_sub_section['interface'] if 'interface' in sub_sub_section else '',
                            procurement_router = sub_sub_section['procurement_router'] if 'procurement_router' in sub_sub_section else '',
                            vendor = sub_sub_section['vendor'] if 'vendor' in sub_sub_section else '',
                            cost = sub_sub_section['cost'] if 'cost' in sub_sub_section else '',
                            standards = sub_sub_section['standards'] if 'standards' in sub_sub_section else '',
                            sustenance = sub_sub_section['sustenance'] if 'sustenance' in sub_sub_section else '',
                            equipment = sub_sub_section['equipment'] if 'equipment' in sub_sub_section else '',
                            features = sub_sub_section['features'] if 'features' in sub_sub_section else '',
                            layout = sub_sub_section['layout'] if 'layout' in sub_sub_section else '',
                            special_requirements = sub_sub_section['special_requirements'] if 'special_requirements' in sub_sub_section else '',
                            created_ip = created_ip,
                            created_by_id = created_by
                        )

                models.SystemEquipmentCompartmentTemp.objects.filter(project_id=project).delete()
                ### System ###

                
                if request.data['id'] and request.data['id']!='':
                    project_name=  masterSerializer.projectSerializer(masterModels.Project.objects.filter(id = request.data['project_id']),many=True).data
                    for projectn in project_name:
                        if request.data['form']=='1' or request.data['form']==1:
                            models.ProjectLog.objects.create(project_id=request.data['project_id'],module_id=request.data['module'],sub_module_id=request.data['psr'][0]['sub_module'],created_on=datetime.now(),created_by_id=created_by,status=1,created_ip=created_ip,msg=projectn['name'] + '-Initation Notes has been updated sucessfully')
                        if request.data['form']=='2' or request.data['form']==2:
                            models.ProjectLog.objects.create(project_id=request.data['project_id'],module_id=request.data['module'],sub_module_id=request.data['psr'][0]['sub_module'],created_on=datetime.now(),created_by_id=created_by,status=1,created_ip=created_ip,msg=projectn['name'] + '-Formulation of Approach Paper has been updated sucessfully')  
                        if request.data['form']=='3' or request.data['form']==3:
                            models.ProjectLog.objects.create(project_id=request.data['project_id'],module_id=request.data['module'],sub_module_id=request.data['psr'][0]['sub_module'],created_on=datetime.now(),created_by_id=created_by,status=1,created_ip=created_ip,msg=projectn['name'] + '-Presentation of Approach Paper has been updated sucessfully')
                        if request.data['form']=='4' or request.data['form']==4:
                            models.ProjectLog.objects.create(project_id=request.data['project_id'],module_id=request.data['module'],sub_module_id=request.data['psr'][0]['sub_module'],created_on=datetime.now(),created_by_id=created_by,status=1,created_ip=created_ip,msg=projectn['name'] + '-Input of Staff Requirement has been updated sucessfully') 
                        if request.data['form']=='5' or request.data['form']==5:
                            models.ProjectLog.objects.create(project_id=request.data['project_id'],module_id=request.data['module'],sub_module_id=request.data['psr'][0]['sub_module'],created_on=datetime.now(),created_by_id=created_by,status=1,created_ip=created_ip,msg=projectn['name'] + '-Concept Design has been updated sucessfully')
                        if request.data['form']=='6' or request.data['form']==6:
                            models.ProjectLog.objects.create(project_id=request.data['project_id'],module_id=request.data['module'],sub_module_id=request.data['psr'][0]['sub_module'],created_on=datetime.now(),created_by_id=created_by,status=1,created_ip=created_ip,msg=projectn['name'] + '-Incorporation of Design Inputs has been updated sucessfully')
                        if request.data['form']=='7' or request.data['form']==7:
                            models.ProjectLog.objects.create(project_id=request.data['project_id'],module_id=request.data['module'],sub_module_id=request.data['psr'][0]['sub_module'],created_on=datetime.now(),created_by_id=created_by,status=1,created_ip=created_ip,msg=projectn['name'] + 'Receipt of RFI Responses has been updated sucessfully')
                        if request.data['form']=='8' or request.data['form']==8:
                            models.ProjectLog.objects.create(project_id=request.data['project_id'],module_id=request.data['module'],sub_module_id=request.data['psr'][0]['sub_module'],created_on=datetime.now(),created_by_id=created_by,status=1,created_ip=created_ip,msg=projectn['name'] + 'GLS Init has been updated sucessfully')
                        if request.data['form']=='7' or request.data['form']==7:
                            models.ProjectLog.objects.create(project_id=request.data['project_id'],module_id=request.data['module'],sub_module_id=request.data['psr'][0]['sub_module'],created_on=datetime.now(),created_by_id=created_by,status=1,created_ip=created_ip,msg=projectn['name'] + 'Receipt of RFI Responses has been updated sucessfully')
                        if request.data['form']=='8' or request.data['form']==8:
                            models.ProjectLog.objects.create(project_id=request.data['project_id'],module_id=request.data['module'],sub_module_id=request.data['psr'][0]['sub_module'],created_on=datetime.now(),created_by_id=created_by,status=1,created_ip=created_ip,msg=projectn['name'] + 'GLS Initiation has been updated sucessfully')
                        if request.data['form']=='9' or request.data['form']==9:
                            models.ProjectLog.objects.create(project_id=request.data['project_id'],module_id=request.data['module'],sub_module_id=request.data['psr'][0]['sub_module'],created_on=datetime.now(),created_by_id=created_by,status=1,created_ip=created_ip,msg=projectn['name'] + 'BLS Initiation has been updated sucessfully')
                        models.GlobalTransactionDetails.objects.filter(global_transaction_id=request.data['id']).delete()
                        running_id=request.data['id']
                else:
                    global_transaction = {
                        'module': request.data['module'],
                        'form': request.data['form'] if request.data['form'] else None,
                        'project': project,
                        'created_ip': created_ip,
                        'created_by': created_by,
                        'status': 1,
                        'initiater': 1,
                        'approved_level': 1 # Added
                    }

                    saveserializelog = cSerializer.GlobalTransactionSerializer(data = global_transaction)
                    if saveserializelog.is_valid():
                        saveserializelog.save()
                        #print('111',saveserializelog.data['id'])
                        running_id = saveserializelog.data['id']

                        

                        #print(request.headers, "request.headers")
                        # role_id = request.headers["Authorized-By"]
                        # models.GlobalTransactionApproval.objects.create(
                        #     approved_level = 1,
                        #     approved_ip = Common.get_client_ip(request),
                        #     approved_by_id = request.user.id,
                        #     transaction_id = running_id,
                        #     comments = "Initiated",
                        #     status = 1,
                        #     approved_role_id = role_id,
                        #     type=-1,
                        # )

                        # # Getting Recommender Level 1 at Form Level Hierarchy

                        # rec_level_1 = models.FormLevelRecommenderHierarchy.objects.values('recommender_id').filter(form_id=request.data['form'], recommender_level=1)

                        # if len(rec_level_1)>0:
                        #     saveNotifications({"form_id":request.data['form'],"user_role_id":secondLevel['user_role_id'],"process_id":secondLevel['process_id'],"transaction_id":running_id},2,1)
                        # else:
                        #     rec_level_1 = models.ProjectLevelRecommenderHierarchy.objects.values('recommender_id').filter(form_id=request.data['form'],project_id=request.data['project_id'], recommender_level=1)

                        #     if len(rec_level_1)>0:
                        #         saveNotifications({"form_id":request.data['form'],"user_role_id":secondLevel['user_role_id'],"process_id":secondLevel['process_id'],"transaction_id":running_id},2,1)
                        #     else:
                        #         pass



                    saveserializelogg = cSerializer.GlobalTransactionlogSerializer(data =global_transaction)                        
                    if saveserializelogg.is_valid():
                        saveserializelogg.save()
                        if request.data['id'] == '' or request.data['id'] is not None or request.data['id'] == 'null':
                            project_name=  masterSerializer.projectSerializer(masterModels.Project.objects.filter(id = request.data['project_id']),many=True).data
                            for projectn in project_name:
                                if request.data['form']=='1' or request.data['form']==1:
                                    models.ProjectLog.objects.create(project_id=request.data['project_id'],module_id=request.data['module'],sub_module_id=request.data['psr'][0]['sub_module'],created_on=datetime.now(),created_by_id=created_by,status=1,created_ip=created_ip,msg=projectn['name'] + '-Initation Notes has been added sucessfully')
                                if request.data['form']=='2' or request.data['form']==2:
                                    models.ProjectLog.objects.create(project_id=request.data['project_id'],module_id=request.data['module'],sub_module_id=request.data['psr'][0]['sub_module'],created_on=datetime.now(),created_by_id=created_by,status=1,created_ip=created_ip,msg=projectn['name'] + '-Formulation of Approach Paper has been added sucessfully')   
                                if request.data['form']=='3' or request.data['form']==3:
                                    models.ProjectLog.objects.create(project_id=request.data['project_id'],module_id=request.data['module'],sub_module_id=request.data['psr'][0]['sub_module'],created_on=datetime.now(),created_by_id=created_by,status=1,created_ip=created_ip,msg=projectn['name'] + '-Presentation of Approach Paper has been added sucessfully') 
                                if request.data['form']=='4' or request.data['form']==4:
                                    models.ProjectLog.objects.create(project_id=request.data['project_id'],module_id=request.data['module'],sub_module_id=request.data['psr'][0]['sub_module'],created_on=datetime.now(),created_by_id=created_by,status=1,created_ip=created_ip,msg=projectn['name'] + '-Input of Staff Requirement has been added sucessfully') 
                                if request.data['form']=='5' or request.data['form']==5:
                                    models.ProjectLog.objects.create(project_id=request.data['project_id'],module_id=request.data['module'],sub_module_id=request.data['psr'][0]['sub_module'],created_on=datetime.now(),created_by_id=created_by,status=1,created_ip=created_ip,msg=projectn['name'] + '-Concept Design has been added sucessfully') 
                                if request.data['form']=='6' or request.data['form']==6:
                                    models.ProjectLog.objects.create(project_id=request.data['project_id'],module_id=request.data['module'],sub_module_id=request.data['psr'][0]['sub_module'],created_on=datetime.now(),created_by_id=created_by,status=1,created_ip=created_ip,msg=projectn['name'] + '-Incorporation of Design Inputs has been added sucessfully') 
                                if request.data['form']=='7' or request.data['form']==7:
                                    models.ProjectLog.objects.create(project_id=request.data['project_id'],module_id=request.data['module'],sub_module_id=request.data['psr'][0]['sub_module'],created_on=datetime.now(),created_by_id=created_by,status=1,created_ip=created_ip,msg=projectn['name'] + '-Receipt of RFI Responses has been added sucessfully') 
                                if request.data['form']=='8' or request.data['form']==8:
                                    models.ProjectLog.objects.create(project_id=request.data['project_id'],module_id=request.data['module'],sub_module_id=request.data['psr'][0]['sub_module'],created_on=datetime.now(),created_by_id=created_by,status=1,created_ip=created_ip,msg=projectn['name'] + 'GLS Initiation has been added sucessfully') 
                                if request.data['form']=='9' or request.data['form']==9:
                                    models.ProjectLog.objects.create(project_id=request.data['project_id'],module_id=request.data['module'],sub_module_id=request.data['psr'][0]['sub_module'],created_on=datetime.now(),created_by_id=created_by,status=1,created_ip=created_ip,msg=projectn['name'] + 'BLS Initiation has been added sucessfully') 
                        running_idd = saveserializelogg.data['id']


                if running_id:
                    # saveNotifications({"form_id":request.data['form'],"user_role_id":secondLevel['user_role_id'],"process_id":secondLevel['process_id'],"transaction_id":running_id},2,1)
                    for row in (request.data['psr']):
                        if models.GlobalTransactionDetails.objects.filter(sub_module_id = row['sub_module'],project_id = project,section_id = row['id'],sub_section_id__isnull=True,sub_sub_section_id__isnull=True).count()>0:
                            models.GlobalTransactionDetails.objects.filter(id__in=(models.GlobalTransactionDetails.objects.values('id').filter(sub_module_id = row['sub_module'],project_id = project,section_id = row['id'],sub_section_id__isnull=True,sub_sub_section_id__isnull=True))).update(**{"paragraph":(row['paragraph'] if 'paragraph' in row else ''), "para_no":(row['para_no'] if 'para_no' in row else '')})
                        else:
                            models.GlobalTransactionDetails.objects.create(
                                global_transaction_id = running_id,
                                sub_module_id = row['sub_module'],
                                section_id = row['id'],
                                #sub_section = row['sub_section'],
                                paragraph = row['paragraph'] if 'paragraph' in row else '',
                                para_no = row['para_no'] if 'para_no' in row else '',
                                view = row['view'] if 'view' in row else '',
                                project_id = project,
                                status = 1,                                
                                created_ip = created_ip,
                                created_by_id = created_by,
                            )

                        if len(row['subsections'])>0:
                            for sub_row in (row['subsections']):
                                if models.GlobalTransactionDetails.objects.filter(sub_module_id = row['sub_module'],project_id = project,section_id = row['id'],sub_section_id=sub_row['id'],sub_sub_section_id__isnull=True).count()>0:
                                    models.GlobalTransactionDetails.objects.filter(id__in=(models.GlobalTransactionDetails.objects.values('id').filter(sub_module_id = row['sub_module'],project_id = project,section_id = row['id'],sub_section_id=sub_row['id'],sub_sub_section_id__isnull=True))).update(**{"paragraph":(sub_row['paragraph'] if 'paragraph' in sub_row else ''), "para_no":(sub_row['para_no'] if 'para_no' in sub_row else '')})
                                else:
                                    models.GlobalTransactionDetails.objects.create(
                                        global_transaction_id = running_id,
                                        sub_module_id = sub_row['sub_module'],
                                        section_id = row['id'],
                                        sub_section_id = sub_row['id'],
                                        paragraph = sub_row['paragraph'] if 'paragraph' in sub_row else '',
                                        para_no = sub_row['para_no'] if 'para_no' in sub_row else '',
                                        view = sub_row['view'] if 'view' in sub_row else '',
                                        project_id = project,
                                        status = 1,
                                        created_ip = created_ip,
                                        created_by_id = created_by,
                                    )

                                if len(sub_row['subsubsections'])>0:
                                    for sub_sub_row in (sub_row['subsubsections']):
                                        if models.GlobalTransactionDetails.objects.filter(sub_module_id = row['sub_module'],project_id = project,section_id = row['id'],sub_section_id=sub_row['id'],sub_sub_section_id=sub_sub_row['id']).count()>0:
                                            models.GlobalTransactionDetails.objects.filter(id__in=(models.GlobalTransactionDetails.objects.values('id').filter(sub_module_id = row['sub_module'],project_id = project,section_id = row['id'],sub_section_id=sub_row['id'],sub_sub_section_id=sub_sub_row['id']))).update(**{"paragraph":(sub_sub_row['paragraph'] if 'paragraph' in sub_sub_row else ''), "para_no":(sub_sub_row['para_no'] if 'para_no' in sub_sub_row else '')})
                                        else:
                                            models.GlobalTransactionDetails.objects.create(
                                                global_transaction_id = running_id,
                                                sub_module_id = sub_sub_row['sub_module'],
                                                section_id = row['id'],
                                                sub_section_id = sub_row['id'],
                                                sub_sub_section_id = sub_sub_row['id'],
                                                paragraph = sub_sub_row['paragraph'] if 'paragraph' in sub_sub_row else '',
                                                para_no = sub_sub_row['para_no'] if 'para_no' in sub_sub_row else '',
                                                view = sub_sub_row['view'] if 'view' in sub_row else '',
                                                project_id = project,
                                                status = 1,
                                                created_ip = created_ip,
                                                created_by_id = created_by,
                                )

                #  Save Log information
                for row in (request.data['psr']):
                            # print('111')
                            models.GlobalTransactionDetailsLog.objects.create(
                                global_transaction_id = running_id,
                                sub_module_id = row['sub_module'],
                                section_id = row['id'],
                                #sub_section = row['sub_section'],
                                paragraph = row['paragraph'] if 'paragraph' in row else '',
                                para_no = row['para_no'] if 'para_no' in row else '',
                                project_id = project,
                                status = 1,                                
                                created_ip = created_ip,
                                created_by_id = created_by,
                            )

                            if len(row['subsections'])>0:
                                for sub_row in (row['subsections']):
                                    models.GlobalTransactionDetailsLog.objects.create(
                                        global_transaction_id = running_id,
                                        sub_module_id = sub_row['sub_module'],
                                        section_id = row['id'],
                                        sub_section_id = sub_row['id'],
                                        paragraph = sub_row['paragraph'] if 'paragraph' in sub_row else '',
                                        para_no = sub_row['para_no'] if 'para_no' in sub_row else '',
                                        project_id = project,
                                        status = 1,
                                        created_ip = created_ip,
                                        created_by_id = created_by,
                                    )

                                    if len(sub_row['subsubsections'])>0:
                                        for sub_sub_row in (sub_row['subsubsections']):
                                            models.GlobalTransactionDetailsLog.objects.create(
                                                global_transaction_id = running_id,
                                                sub_module_id = sub_sub_row['sub_module'],
                                                section_id = row['id'],
                                                sub_section_id = sub_row['id'],
                                                sub_sub_section_id = sub_sub_row['id'],
                                                paragraph = sub_sub_row['paragraph'] if 'paragraph' in sub_sub_row else '',
                                                para_no = sub_sub_row['para_no'] if 'para_no' in sub_sub_row else '',
                                                project_id = project,
                                                status = 1,
                                                created_ip = created_ip,
                                                created_by_id = created_by,
                                            )

                return Response({"status" :error.context['success_code'],"message": formDet['name']+' '+('updated' if request.data['id'] and request.data['id'] else 'created')+' successfully'}, status=status.HTTP_200_OK)

class deleteGlobalTransactionDetails(APIView):
  
    def post(self,request, pk = None):
        if 'id' not in request.data or request.data['id']=='':
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']
            models.GlobalTransaction.objects.filter(id=pk).update(**{"status":3})
            det=models.GlobalTransaction.objects.values('form__name').filter(id=pk).first()
            return Response({"status" :error.context['success_code'], "message":det['form__name']+' deleted successfully'}, status=status.HTTP_200_OK)

class getFormName(APIView):
  
    def post(self,request, pk = None):
        if 'id' not in request.data or request.data['id']=='':
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']
            det=models.Forms.objects.values('name').filter(id=pk).first()
            return Response({"status" :error.context['success_code'], "name":det['name']}, status=status.HTTP_200_OK)

class sampleTest(APIView):
    authentication_classes = [] #disables authentication
    permission_classes = [] #disables permission
    def get(self,request, pk = None):
        data=getDataByFormModule(1,8)
        # inputSRs=models.FormsMapping.objects.values().filter(form_id=4)
        # for inputSR in inputSRs:
        #     inputSR['id']=None
        #     inputSR['form_id']=8
        #     models.FormsMapping.objects.create(**inputSR)

        return Response({"status":error.context['success_code'] , "data": data}, status=status.HTTP_200_OK)

def getDataByFormModule(module_id,project_id):
    if module_id and project_id:
        return cSerializer.getFormMappingReportSerializer(models.FormsMapping.objects.filter(module_id=module_id).distinct('sub_module_id'),many=True,context={"project_id":project_id}).data
    else:
        return []
def getDataBySystem(module_id,project_id):
    if module_id and project_id:
        return cSerializer.getFormMappingSystemEquipementCompartmentReportSerializer(models.FormsMapping.objects.filter(module_id=module_id).distinct('sub_module_id'),many=True,context={"project_id":project_id}).data
    else:
        return []

def getDataByFormgls(module_id,project_id):
    if module_id and project_id:
        return cSerializer.getFormMappingReportSerializer(models.FormsMapping.objects.filter(form_id=8).distinct('sub_module_id'),many=True,context={"project_id":project_id}).data
    else:
        return []
def getDataBySystemgls(module_id,project_id):
    if module_id and project_id:
        return cSerializer.getFormMappingSystemEquipementCompartmentReportSerializer(models.FormsMapping.objects.filter(form_id=8).distinct('sub_module_id'),many=True,context={"project_id":project_id}).data
    else:
        return []

def getDataByFormbls(module_id,project_id):
    if module_id and project_id:
        return cSerializer.getFormMappingReportSerializer(models.FormsMapping.objects.filter(form_id=9).distinct('sub_module_id'),many=True,context={"project_id":project_id}).data
    else:
        return []

def getDataBySystembls(module_id,project_id):
    if module_id and project_id:
        return cSerializer.getFormMappingSystemEquipementCompartmentReportSerializer(models.FormsMapping.objects.filter(form_id=9).distinct('sub_module_id'),many=True,context={"project_id":project_id}).data
    else:
        return []

def getVersionDataByFormModule(module_id,project_id,version_id):
    #print(version_id,"version_id@@@@")
    if module_id and project_id:
        return cSerializer.getVersionFormMappingReportSerializer(models.FormsMapping.objects.filter(module_id=module_id).distinct('sub_module_id'),many=True,context={"project_id":project_id,"version_id":version_id}).data
    else:
        return []
def getSystemVersionDataByFormModule(module_id,project_id,version_id):
    #print(version_id,"version_id@@@@")
    if module_id and project_id:
        return cSerializer.getFormMappingSystemEquipementCompartmentReportSerializer(models.FormsMapping.objects.filter(module_id=module_id).distinct('sub_module_id'),many=True,context={"project_id":project_id,"version_id":version_id}).data
    else:
        return []
def getSubmoduleDataByFormModule(module_id,project_id,form_id):

    if module_id and project_id:
        return cSerializer.getFormMappingReportSerializer(models.FormsMapping.objects.filter(form_id=form_id).distinct('sub_module_id'),many=True,context={"project_id":project_id}).data
    else:
        return []
def getDataBySystemSubmodule(module_id,project_id,form_id):
    if module_id and project_id:
         return cSerializer.getFormMappingSystemEquipementCompartmentReportSerializer(models.FormsMapping.objects.filter(form_id=form_id).distinct('sub_module_id'),many=True,context={"project_id":project_id}).data
        
    else:
        return []

class getSSS(APIView):
    authentication_classes = [] #disables authentication
    permission_classes = [] #disables permission
    def get(self,request, pk = None):
        lists = models.SSSMapping.objects.exclude(status='3')
        serializer = cSerializer.SSSMappingSerializerNew(lists, many=True, context={"project_id":14})
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)
        return Response({"status":error.context['success_code'] , "data": data}, status=status.HTTP_200_OK)


class Version(APIView):

    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['code','name']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.Version.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.VersionSerializer(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.Version.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"Section" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.Version.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        serializer = cSerializer.VersionSerializer(lists, many=True)
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)


class saveTransactionResponsibility(APIView):

    def post(self,request, pk = None):
        if 'id' not in request.data:
            return Response({"status":error.context['error_code'], "message" : "id" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        elif 'directorate' not in request.data:
            return Response({"status":error.context['error_code'], "message" : "Directorate" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        elif 'module' not in request.data:
            return Response({"status":error.context['error_code'], "message" : "Module" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        #elif 'sub_module' not in request.data:
            #return Response({"status":error.context['error_code'], "message" : "Sub Module" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        #elif 'global_section' not in request.data:
            #return Response({"status":error.context['error_code'], "message" : "Global Section" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        elif 'project' not in request.data:
            return Response({"status":error.context['error_code'], "message" : "Project" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)

        else: 
                id = request.data['id']
                directorate = request.data['directorate']
                module = request.data['module']
                #sub_module = request.data['sub_module']
                responsibility = request.data['formmapping']
                project = request.data['project']
                created_ip = Common.get_client_ip(request)
                created_by = request.data['created_by']

                # if id!="":
                #     count = models.GlobalTransactionResponsibility.objects.filter(id=id).count()
                #     if count > 0:
                #         models.GlobalTransactionResponsibility.objects.filter(id=id).delete()
                print(request.data,'999')
                count = models.GlobalTransactionResponsibility.objects.filter(directorate_id=directorate, project_id=project).count()

                if count > 0:
                     models.GlobalTransactionResponsibility.objects.filter(directorate_id=directorate, project_id=project).delete()

                gt_responsibility = models.GlobalTransactionResponsibility.objects.create(
                    directorate_id = directorate,
                    project_id = project,
                    created_ip = created_ip,
                    created_by_id = created_by,
                    status = 1,
                )
                direct_name=  masterSerializer.UnitSerializer(masterModels.Unit.objects.filter(id = request.data['directorate']),many=True).data
                project_name=  masterSerializer.projectSerializer(masterModels.Project.objects.filter(id = project),many=True).data
                        # print(project_name['name'],'fdsf')
                sub_mod= masterSerializer.SubModuleSerializer(masterModels.SubModule.objects.filter(id = request.data['formmapping'][0]['sub_module']),many=True).data       
                for projectn in project_name:
                    for sub in sub_mod:
                        for dn in direct_name:
                            models.ProjectLog.objects.create(
                            project_id=project,module_id=request.data['module'],sub_module_id=request.data['formmapping'][0]['sub_module'],created_on=datetime.now(),created_by_id=created_by,status=1,created_ip=created_ip,msg=projectn['name'] +'  has been assigned to directorate  ' + str(dn['name']) + ' in '+str(sub['name']))
                        
                # print('ds', request.data['directorate'])
                # test=masterSerializer.UnitSerializer(masterModels.Unit.objects.values('name').filter(id=request.data['directorate']))
                # print('fg',test['name'])
                # if request.data['form']=='1' or request.data['form']==1:
                #     models.ProjectLog.objects.create(project_id=request.data['project_id'],module_id=request.data['module'],sub_module_id=request.data['psr'][0]['sub_module'],created_on=datetime.now(),created_by_id=created_by,status=1,created_ip=created_ip,msg='Initation Notes has been added ' + str(datetime.now()))
                # models.ProjectLog.objects.create(project_id=project,module_id=module,created_on=datetime.now(),created_by_id=created_by,status=1,created_ip=created_ip,msg= + str(datetime.now()))
                 

                gt_responsibility_id = gt_responsibility.id;

                if gt_responsibility_id:
                    for row in (responsibility):
                        models.GlobalTransactionResponsibilityDetail.objects.create(
                            gt_responsibility_id = gt_responsibility_id,
                            directorate_id = directorate,
                            module_id = module,
                            sub_module_id = row['sub_module'],
                            section_id = row['section'],
                            sub_section_id = row['sub_section'],
                            sub_sub_section_id = row['sub_sub_section'],
                            project_id = project,
                            status = 1,
                            created_ip = created_ip,
                            created_by_id = created_by,
                        )

                return Response({"status":error.context['success_code'] , "data": gt_responsibility_id, "message":"Global transaction responsibility" +language.context[language.defaultLang]['insert']}, status=status.HTTP_200_OK)
                # if count > 0:
                #     return Response({"status" :error.context['success_code'], "message":"Global transaction responsibility" +language.context[language.defaultLang]['update'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                # else:
                #     return Response({"status" :error.context['success_code'], "message":"Global transaction responsibility" +language.context[language.defaultLang]['insert'], "data":saveserialize.data}, status=status.HTTP_200_OK)


class deleteTransactionResponsibility(APIView):
  
    def post(self,request, pk = None):
        if 'id' not in request.data or request.data['id']=='':
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']
            models.GlobalTransactionResponsibility.objects.filter(id=pk).delete()

            return Response({"status" :error.context['success_code'], "message":'Responsibility deleted successfully'}, status=status.HTTP_200_OK)

class GlobalTransactionResponsibility(APIView):

    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['code','name']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.GlobalTransactionResponsibility.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.GlobalTransactionResponsibilitySerializer(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.GlobalTransactionResponsibility.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"Section" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.GlobalTransactionResponsibility.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        serializer = cSerializer.GlobalTransactionResponsibilitySerializer(lists, many=True)
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)



class saveGlobalTransactionCommentsDetails(APIView):


    def post(self,request, pk = None):
        if 'module_id' not in request.data or request.data['module_id'] == '':
            return Response({"status":error.context['error_code'], "message" : "Module id " +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        elif 'sub_module_id' not in request.data or request.data['sub_module_id'] == '':
            return Response({"status":error.context['error_code'], "message" : "Sub Module id " +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        elif 'project_id' not in request.data or request.data['project_id'] == '':
            return Response({"status":error.context['error_code'], "message" : "Project id " +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        #elif 'global_transaction_id' not in request.data or request.data['global_transaction_id'] == '':
            return Response({"status":error.context['error_code'], "message" : "Global transaction id" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)

        else:
            project_id=request.data['project_id'] if 'project_id' in request.data and request.data['project_id'] else ''
            id = request.data['id']

            if id!="":
                count = models.GlobalTransactionComments.objects.filter(id=id).count()
                if count > 0:
                     models.GlobalTransactionComments.objects.filter(id=id).delete()
            else:
                count = 0

            created_by = request.data['created_by']
            created_ip = Common.get_client_ip(request)


            global_transaction_comments = {
                'project' : request.data['project_id'] if request.data['project_id'] else None,
                'module' : request.data['module_id'] if request.data['module_id'] else None,
                'sub_module' : request.data['sub_module_id'] if request.data['sub_module_id'] else None,
                'section' : request.data['section_id'] if request.data['section_id'] else None,
                'sub_section' : request.data['sub_section_id'] if request.data['sub_section_id'] else None,
                'sub_sub_section' : request.data['sub_sub_section_id'] if request.data['sub_sub_section_id'] else None,
                #'global_transaction' : request.data['global_transaction_id'],
                #'global_transaction_detail_id' : request.data['global_transaction_detail_id'],
                'comments' : request.data['comments'],
                'created_ip': created_ip,
                'created_by': created_by,
                #'created_on': datetime.now(),
                'status':1,
            }


            saveserializelog = cSerializer.GlobalTransactionCommentsSerializer(data = global_transaction_comments)
            if saveserializelog.is_valid():
                saveserializelog.save()

            if count > 0:
                return Response({"status":error.context['success_code'], "message":"Global transaction comments" +language.context[language.defaultLang]['update'], "data": saveserializelog.data}, status=status.HTTP_200_OK)
            else:
                return Response({"status":error.context['success_code'], "message":"Global transaction comments" +language.context[language.defaultLang]['insert'], "data": saveserializelog.data}, status=status.HTTP_200_OK)

class deleteGlobalTransactionCommentsDetails(APIView):

    def post(self,request, pk = None):
        if 'id' not in request.data or request.data['id']=='':
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']
            models.GlobalTransactionComments.objects.filter(id=pk).delete()

            return Response({"status" :error.context['success_code'], "message":'Comments deleted successfully'}, status=status.HTTP_200_OK)
@xframe_options_exempt
def global_transactionpdf(self,id=None):    

    globaltransaction = models.GlobalTransaction.objects.values('id','project__name','module__name','module_id','project_id').filter(id=id).first()
    moduleDet=masterModels.Module.objects.values('id','name').filter(id=globaltransaction['module_id']).first()
    globaltransactionname=getDataByFormModule(module_id=globaltransaction['module_id'],project_id=globaltransaction['project_id'])
    globaltransactionsss=ListGlobalTransactionAllEditSerializer_1(models.GlobalTransaction.objects.filter(id=id),many=True).data
    sytemdata=getDataBySystem(module_id=globaltransaction['module_id'],project_id=globaltransaction['project_id'])

    context = { 'globaltransaction':globaltransaction,'globaltransactionname':globaltransactionname ,"moduleDet":moduleDet,"globaltransactionsss":globaltransactionsss,"sytemdata":sytemdata}
    # print('globaltransactionname',globaltransactionsss)


    html = loader.render_to_string('service/psr.html', context=context)
    import pdfkit
    import platform
    if platform.system()=='Windows':
        pdf = pdfkit.from_string(html, False, verbose=True,configuration=pdfkit.configuration(wkhtmltopdf=common.WKHTML_PATH))
    else:
        pdf = pdfkit.from_string(html, False, verbose=True)

    f = open('media/PDF/psr/psr_'+id+'.pdf','wb')
    f.write(pdf)

    response = HttpResponse(pdf, content_type='application/pdf')
    if response:
        convert_pdf_word(id)
        return response  # returns the response.

    #return HttpResponse(html, content_type='text/html')   

@xframe_options_exempt
def global_transaction_glspdf(self,id=None):    

    globaltransaction = models.GlobalTransaction.objects.values('id','project__name','module__name','module_id','project_id').filter(id=id).first()
    moduleDet=masterModels.Module.objects.values('id','name').filter(id=globaltransaction['module_id']).first()
    globaltransactionname=getDataByFormgls(module_id=globaltransaction['module_id'],project_id=globaltransaction['project_id'])    
    sytemdata=getDataBySystemgls(module_id=globaltransaction['module_id'],project_id=globaltransaction['project_id'])
    context = { 'globaltransaction':globaltransaction,'globaltransactionname':globaltransactionname ,"moduleDet":moduleDet,"sytemdata":sytemdata}

   
    

    html = loader.render_to_string('service/gls.html', context=context)
    import pdfkit
    import platform
    if platform.system()=='Windows':
        pdf = pdfkit.from_string(html, False, verbose=True,configuration=pdfkit.configuration(wkhtmltopdf=common.WKHTML_PATH))
    else:
        pdf = pdfkit.from_string(html, False, verbose=True)
    response = HttpResponse(pdf, content_type='application/pdf')
    if response:
        return response  # returns the response.

    return HttpResponse(html, content_type='text/html')   
@xframe_options_exempt
def global_transaction_blspdf(self,id=None):    

    globaltransaction = models.GlobalTransaction.objects.values('id','project__name','module__name','module_id','project_id').filter(id=id).first()
    moduleDet=masterModels.Module.objects.values('id','name').filter(id=globaltransaction['module_id']).first()
    globaltransactionname=getDataByFormbls(module_id=globaltransaction['module_id'],project_id=globaltransaction['project_id'])

    sytemdata=getDataBySystembls(module_id=globaltransaction['module_id'],project_id=globaltransaction['project_id'])

    context = { 'globaltransaction':globaltransaction,'globaltransactionname':globaltransactionname ,"moduleDet":moduleDet, "sytemdata":sytemdata}
   

    html = loader.render_to_string('service/bls.html', context=context)
    import pdfkit
    import platform
    if platform.system()=='Windows':
        pdf = pdfkit.from_string(html, False, verbose=True,configuration=pdfkit.configuration(wkhtmltopdf=common.WKHTML_PATH))
    else:
        pdf = pdfkit.from_string(html, False, verbose=True)
    response = HttpResponse(pdf, content_type='application/pdf')
    if response:
        return response  # returns the response.

    return HttpResponse(html, content_type='text/html')   


# PDF Edit.
@xframe_options_exempt
def global_transaction_get_pdf(self, id=None):    

    #print(id)
    globaltransaction = models.GlobalTransaction.objects.values('id','project__name','module__name','module_id','project_id').filter(id=id).first()
    moduleDet=masterModels.Module.objects.values('id','name').filter(id=globaltransaction['module_id']).first()
    globaltransactionname=getDataByFormModule(module_id=globaltransaction['module_id'],project_id=globaltransaction['project_id'])
    globaltransactionsss=ListGlobalTransactionAllEditSerializer_1(models.GlobalTransaction.objects.filter(id=id),many=True).data
    sytemdata=getDataBySystem(module_id=globaltransaction['module_id'],project_id=globaltransaction['project_id'])

    context = { 'globaltransaction':globaltransaction,'globaltransactionname':globaltransactionname ,"moduleDet":moduleDet,"globaltransactionsss":globaltransactionsss,"sytemdata":sytemdata}
    #print('globaltransactionname',globaltransactionsss)

    html = loader.render_to_string('service/psr.html', context=context)

    return JsonResponse({
      'html': render_to_string('service/psr.html', context=context)},
      status=status.HTTP_200_OK)

@xframe_options_exempt
def global_transaction_get_gls_pdf(self, id=None):    

    globaltransaction = models.GlobalTransaction.objects.values('id','project__name','module__name','module_id','project_id').filter(id=id).first()
    moduleDet=masterModels.Module.objects.values('id','name').filter(id=globaltransaction['module_id']).first()
    globaltransactionname=getDataByFormgls(module_id=globaltransaction['module_id'],project_id=globaltransaction['project_id'])    
    sytemdata=getDataBySystemgls(module_id=globaltransaction['module_id'],project_id=globaltransaction['project_id'])
    context = { 'globaltransaction':globaltransaction,'globaltransactionname':globaltransactionname ,"moduleDet":moduleDet,"sytemdata":sytemdata}

    html = loader.render_to_string('service/gls.html', context=context)

    return JsonResponse({
      'html': render_to_string('service/gls.html', context=context),'status': '1'},
      status=status.HTTP_200_OK)

@xframe_options_exempt
def global_transaction_get_bls_pdf(self, id=None):    

    globaltransaction = models.GlobalTransaction.objects.values('id','project__name','module__name','module_id','project_id').filter(id=id).first()
    moduleDet=masterModels.Module.objects.values('id','name').filter(id=globaltransaction['module_id']).first()
    globaltransactionname=getDataByFormbls(module_id=globaltransaction['module_id'],project_id=globaltransaction['project_id'])

    sytemdata=getDataBySystembls(module_id=globaltransaction['module_id'],project_id=globaltransaction['project_id'])

    context = { 'globaltransaction':globaltransaction,'globaltransactionname':globaltransactionname ,"moduleDet":moduleDet, "sytemdata":sytemdata}

    html = loader.render_to_string('service/bls.html', context=context)
    return JsonResponse({
      'html': render_to_string('service/bls.html', context=context),'status': '1'},
      status=status.HTTP_200_OK)


class globalTransactionEditPDF(APIView):

    def post(self,request, pk = None):

        if "html" not in request.data:
            return Response(
                {
                    "status": error.context["error_code"],
                    "message": "PDF content"
                    + language.context[language.defaultLang]["missing"],
                },
                status=status.HTTP_200_OK,
            )
        elif "project_id" not in request.data:
            return Response(
                {
                    "status": error.context["error_code"],
                    "message": "Project ID"
                    + language.context[language.defaultLang]["missing"],
                },
                status=status.HTTP_200_OK,
            )
        else:

            models.GlobalTransactionEditPdf.objects.create(
                project_id = request.data["project_id"],
                html = request.data["html"],
                form_id = request.data["form_id"],
                created_by_id = request.user.id,
                created_ip = Common.get_client_ip(request),
                status = 1,
            )

            content = request.data["html"];

            # res = models.GlobalTransactionEditPdf.objects.values('id','html').filter(id=32).first()
            # content1 = res['html']

            f = open("templates/service/edit_pdf.html","w")
            f.write(content)
            f.close()

        return Response({"status" :error.context['success_code'], "message":'PDF file updated successfully'}, status=status.HTTP_200_OK)    

@xframe_options_exempt
def global_transaction_gen_edit_pdf(self, id=None):

    html = loader.render_to_string('service/edit_pdf.html')
    import pdfkit
    import platform
    if platform.system()=='Windows':
        pdf = pdfkit.from_string(html, False, verbose=True,configuration=pdfkit.configuration(wkhtmltopdf=common.WKHTML_PATH))
    else:
        pdf = pdfkit.from_string(html, False, verbose=True)
    response = HttpResponse(pdf, content_type='application/pdf')
    if response:
        return response  # returns the response.

    return HttpResponse(html, content_type='text/html')




#@xframe_options_exempt
def convert_pdf_word(id=None):

    #print(id,"ssssss")
    from pdf2docx import parse, Converter
    # from wsgiref.util import FileWrapper
    # import mimetypes
    # from django.utils.encoding import smart_str


    #from urllib.request
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    #print(BASE_DIR,"UUUU")
    pdf_file = BASE_DIR + '/media/PDF/psr/psr_'+id+'.pdf'
    docx_file = BASE_DIR + '/media/PDF/psr/psr_'+id+'.docx'
    #download_url = BASE_DIR + '/media/Excel'
    #res = parse(pdf_file=pdf_file,docx_file=docx_file,start=0,end=None)
    #urllib.request.urlretrieve(BASE_DIR + '/media/Excel/', "dummy.docx")
    #urlretrieve(download_url, 'dummy.docx')
    #urllib.request.urlretrieve(download_url, "dummy.docx")

    cv = Converter(pdf_file)
    cv.convert(docx_file, start=0, end=None)
    #cv.close()
    if cv:
        cv.close()
        #parse(pdf_file, docx_file)
        return HttpResponse({"status" :error.context['success_code'], "message":'PDF file updated successfully'}, status=status.HTTP_200_OK)



def convert_pdf_word2(id=None):

    import PyPDF2
    from docx import Document

    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    #pdf_file = BASE_DIR + '/media/PDF/psr/psr_'+id+'.pdf'
    pdf_file = BASE_DIR + '/media/PDF/psr/dummy.pdf'
    docx_file = BASE_DIR + '/media/PDF/psr/dummy.docx'

    with open(pdf_file,'rb') as file:
        pdf = PyPDF2.PdfFileReader(file)
        document = Document()

        for page in range(pdf.getNumPages()):
            text = pdf.getPage(page).extractText()
            document.add_paragraph(text)
    document.save(docx_file)


def convert_pdf_word3rdParty_ASPOSE(id=555):
    import aspose.words as aw

    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    pdf_file = BASE_DIR + '/media/PDF/psr/psr_555.pdf'
    docx_file = BASE_DIR + '/media/PDF/psr/psr_555.docx'

    # load the PDF file
    doc = aw.Document(pdf_file)

    # convert PDF to Word DOCX format
    doc.save(docx_file)


def convert_pdf_word_123(id=None):

    from pdf2docx import parse, Converter

    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    pdf_file = BASE_DIR + '/media/PDF/psr/psr_'+id+'.pdf'
    docx_file = BASE_DIR + '/media/PDF/psr/psr_'+id+'.docx'

    cv = Converter(pdf_file)
    tables = cv.extract_tables(docx_file, start=0, end=None)
    cv.close()
    if tables:
        for table in tables:
            print(table)

def convert_pdf_wordWIN(id=None):

    import win32com.client
    word = win32com.client.Dispatch("Word.Application")
    word.visible = 1
    pdfdoc = BASE_DIR + '/media/PDF/psr/psr_'+id+'.pdf'
    todocx = BASE_DIR + '/media/PDF/psr/psr_'+id+'.docx'
    wb1 = word.Documents.Open(pdfdoc)
    wb1.SaveAs(todocx, FileFormat=16)  # file format for docx
    wb1.Close()
    word.Quit()

def convert_pdf_word4(id=None):
    from pdf2docx import parse, Converter
    from typing import Tuple
    pages: Tuple = None

    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    input_file = BASE_DIR + '/media/PDF/psr/psr_'+id+'.pdf'
    output_file = BASE_DIR + '/media/PDF/psr/psr_'+id+'.docx'

    if pages:
        pages = [int(i) for i in list(pages) if i.isnumeric()]
    result = parse(pdf_file=input_file,
                   docx_file= output_file, pages=pages)

    summary = {
        "File" : input_file, "Pages": str(pages), "Output File": output_file
    }

    print("## Summary #########################################################")
    print("\n".join("{}:{}".format(i, j) for i , j in summary.items()))
    print("#####################################################################")
    return result


def convert_pdf_word_subprocess(id=None):

    import os
    import subprocess

    print('GGGGGGGG123@123')
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    pdf_file = BASE_DIR + '/media/PDF/psr'

    result = subprocess.run(["dir"], shell=True, capture_output=True, text=True)
    res = subprocess.call('lowriter --invisible --convert-to doc:draw_pdf_export "{}"' .format(pdf_file), shell=True)

    #print(pdf_file)
    print(result)
    print(res)


    #pdf_file = BASE_DIR + '/media/PDF/psr/psr_'+id+'.pdf'
    #pdf_file = BASE_DIR + '/media/PDF/psr/dummy.pdf'
    #docx_file = BASE_DIR + '/media/PDF/psr/dummy.docx'

    # for top, dirs, files in os.walk(BASE_DIR + '/media/PDF/psr'):
    #     for filename in files:
    #         if filename.endswith('.pdf'):
    #             abspath = os.path.join(top, filename)

    #             print(abspath, "abspath")
    #             # subprocess.call('lowriter --invisible --convert-to doc "{}"'
    #             #                 .format(abspath), shell=True)
    #             #subprocess.call('lowriter --invisible --convert-to doc:writer_pdf_export "{}"' .format(abspath), shell=True)
    #             subprocess.call('lowriter --invisible --convert-to doc:draw_pdf_export "{}"' .format(abspath), shell=True)


def convert_pdf_word_comtypes(id=None):

    import os
    #import os import sys
    import comtypes.client

    wdFormatPDF = 17

    #def covx_to_pdf(infile, outfile): """Convert a Word .docx to PDF"""

    infile = BASE_DIR + '/media/PDF/psr/psr_547.pdf'
    outfile = BASE_DIR + '/media/PDF/psr/psr_547.docx'

    word = comtypes.client.CreateObject('Word.Application')
    doc = word.Documents.Open(infile)
    doc.SaveAs(outfile, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()




class ApprovalHistory(APIView):
    def post(self, request, pk=None):
        if "transaction_id" in request.data:
            history = (
                models.GlobalTransactionApproval.objects.values(
                    "id",
                    "status",
                    "comments",
                    "approved_on",
                    "approved_by__first_name",
                    "approved_by__last_name",
                    #"approved_by__process__name",
                    "approved_role__name",
                )
                .filter(transaction_id=request.data["transaction_id"])
                .order_by("id")
            )
            return Response(
                {"status": error.context["success_code"], "data": history},
                status=status.HTTP_200_OK,
            )
        else:
            return Response(
                {"status": error.context["success_code"], "data": "No data"},
                status=status.HTTP_200_OK,
            )


# GlobalApproval
class GlobalApproval_OLD(APIView):
    def post(self, request, pk=None):
        if "approved_level" not in request.data:
            return Response(
                {
                    "status": error.context["error_code"],
                    "message": "Level"
                    + language.context[language.defaultLang]["missing"],
                },
                status=status.HTTP_200_OK,
            )
        elif "form_id" not in request.data:
            return Response(
                {
                    "status": error.context["error_code"],
                    "message": "Form ID"
                    + language.context[language.defaultLang]["missing"],
                },
                status=status.HTTP_200_OK,
            )
        elif "comments" not in request.data:
            return Response(
                {
                    "status": error.context["error_code"],
                    "message": "Comment"
                    + language.context[language.defaultLang]["missing"],
                },
                status=status.HTTP_200_OK,
            )
        elif "status" not in request.data:
            return Response(
                {
                    "status": error.context["error_code"],
                    "message": "Status"
                    + language.context[language.defaultLang]["missing"],
                },
                status=status.HTTP_200_OK,
            )
        elif "approved_role_id" not in request.data:
            return Response(
                {
                    "status": error.context["error_code"],
                    "message": "Approved Role ID"
                    + language.context[language.defaultLang]["missing"],
                },
                status=status.HTTP_200_OK,
            )
        else:
            try:
                if int(request.data["approved_level"]) == 2:
                    models.GlobalTransaction.objects.filter(id=request.data["transaction_id"]).update(
                        recommender=request.data["status"]
                    )
                elif int(request.data["approved_level"]) == 3:
                    models.GlobalTransaction.objects.filter(id=request.data["transaction_id"]).update(
                        approver=request.data["status"]
                    )

                if request.data["status"] == 1:
                    approvalData = (
                        accessModels.ProcessFlow.objects.values()
                        .filter(level__gt=int(request.data["approved_level"]))
                        .order_by("id")
                        .first()
                    )

                    if approvalData:
                        # Next level approval notification
                        globalTrans = cSerializer.ListGlobalTransactionSerializer(
                            models.GlobalTransaction.objects.filter(id=request.data["transaction_id"]),
                            many=True,
                            context={"request": request},
                        ).data[0]

                        #print(globalTrans['form'],"GlobalTrans: GlobalTrans")
                        print(approvalData["level"],"Level: Level")

                        formDet = models.Forms.objects.values('id','name').filter(id=request.data["form_id"]).first()
                        projectDet = models.GlobalTransaction.objects.values('id','project__name','created_by__first_name','created_by__last_name').filter(id=request.data["transaction_id"]).first()

                        #print(projectDet)
                        # models.trialStatus.objects.create(
                        #     trial_id=trials["id"],
                        #     process_flow_id=approvalData["id"],
                        #     created_by_id=request.user.id,
                        #     created_ip=Common.get_client_ip(request),
                        # )

                        if approvalData["level"] == 2:
                            notificationMessage = (
                                "Project "+projectDet['project__name']
                                + " request "
                                + projectDet['created_by__first_name']+" "+projectDet['created_by__last_name']+" in "+formDet['name']
                                + ") is re initiated again and waiting for your approval "
                            )
                        elif approvalData["level"] == 3:
                            notificationMessage = (
                                "Project "+projectDet['project__name']
                                + " is recommended by "
                                + projectDet['created_by__first_name']+" "+projectDet['created_by__last_name']+" in "+formDet['name']
                                + " is waiting for your approval "
                            )

                        #print(notificationMessage,"notificationMessage")
                        # elif approvalData["level"] == 4:
                        #     notificationMessage = (
                        #         (trials["trial_type"]["type"])
                        #         + " requestion ("
                        #         + trials["trial_number"]
                        #         + ") is approved and waiting for you to submit"
                        #     )
                        # elif approvalData["level"] == 5:
                        #     notificationMessage = (
                        #         (trials["trial_type"]["type"])
                        #         + " ("
                        #         + trials["trial_number"]
                        #         + ") has submitted and waiting for your recommendation"
                        #     )
                        # elif approvalData["level"] == 6:
                        #     notificationMessage = (
                        #         (trials["trial_type"]["type"])
                        #         + " ("
                        #         + trials["trial_number"]
                        #         + ") has recommended for your approval"
                        #     )

                        notificationModels.NotificationUser.objects.create(
                            form_id = globalTrans['form'],
                            transaction_id = globalTrans["id"],
                            message = notificationMessage,
                            user_role_id = approvalData["user_role_id"],
                            process_id = approvalData["process_id"],
                            hierarchy_type = 2
                        )

                        models.GlobalTransaction.objects.filter(
                            id = request.data["transaction_id"]
                        ).update(approved_level=approvalData["level"])

                        # Form Level
                        res = models.FormLevelRecommenderHierarchy.objects.filter(
                            form_id = globalTrans['form'],
                            recommender_id = request.user.id
                        ).update(recommender_level_status = 1)

                        print(res.query)

                    else:

                        globalTrans = cSerializer.ListGlobalTransactionSerializer(
                            models.GlobalTransaction.objects.filter(id=request.data["transaction_id"]),
                            many=True,
                            context={"request": request},
                        ).data[0]

                        formDet = models.Forms.objects.values('id','name').filter(id=request.data["form_id"]).first()
                        projectDet = models.GlobalTransaction.objects.values('id','project__name','created_by__first_name','created_by__last_name').filter(id=request.data["transaction_id"]).first()

                        
                        notificationMessage = (
                            "Project "+projectDet['project__name']
                            + " request "
                            + projectDet['created_by__first_name']+" "+projectDet['created_by__last_name']+" in "+formDet['name']
                            + ") has been approved and report has been generated"
                        )


                        # NotificationUser.objects.create(
                        #     trial_unit_id=trials["trial_unit"]["id"],
                        #     satellite_unit_id=trials["satellite_unit"]["id"],
                        #     trial_id=trials["id"],
                        #     message=notificationMessage,
                        # )

                        # models.Trials.objects.filter(
                        #     id=request.data["trial_id"]
                        # ).update(approved_level="-1")


                        notificationModels.NotificationUser.objects.create(
                            form_id = globalTrans['form'],
                            transaction_id = globalTrans["id"],
                            message = notificationMessage,
                            #user_role_id = approvalData["user_role_id"],
                            #process_id = approvalData["process_id"],
                            hierarchy_type = 2
                        )

                        models.GlobalTransaction.objects.filter(
                            id = request.data["transaction_id"]
                        ).update(approved_level=-1)


                    # # models.trialApproval.objects.filter(approved_by_id=request.user.id,trial_id=request.data['trial_id'],approved_role_id=request.data['approved_role_id']).delete()
                    # models.trialApproval.objects.create(
                    #     approved_level=request.data["approved_level"],
                    #     approved_ip=Common.get_client_ip(request),
                    #     approved_by_id=request.user.id,
                    #     trial_id=request.data["trial_id"],
                    #     comments=request.data["comments"],
                    #     status=request.data["status"],
                    #     approved_role_id=request.data["approved_role_id"],
                    #     type=request.data["type"],
                    # )

                    models.GlobalTransactionApproval.objects.create(
                        approved_level = request.data["approved_level"],
                        approved_ip = Common.get_client_ip(request),
                        approved_by_id = request.user.id,
                        transaction_id = request.data["transaction_id"],
                        comments = request.data["comments"],
                        status = request.data["status"],
                        approved_role_id = request.data["approved_role_id"],
                        type = request.data["type"],
                    )
                    return Response(
                        {
                            "status": error.context["success_code"],
                            "message": "Approval updated successfully",
                        },
                        status=status.HTTP_200_OK,
                    )
                elif request.data["status"] == 2:
                    #print(request.data["status"],"GGGGGGg$$$$")
                    approvalData = (
                        accessModels.ProcessFlow.objects.values()
                        .filter(level__lt=int(request.data["approved_level"]))
                        .order_by("-id")
                        .first()
                    )
                    if approvalData:
                        # trials = TrialListSerializer(
                        #     models.Trials.objects.filter(id=request.data["trial_id"]),
                        #     many=True,
                        #     context={"request": request},
                        # ).data[0]

                        globalTrans = cSerializer.ListGlobalTransactionSerializer(
                            models.GlobalTransaction.objects.filter(id=request.data["transaction_id"]),
                            many=True,
                            context={"request": request},
                        ).data[0]

                        formDet = models.Forms.objects.values('id','name').filter(id=request.data["form_id"]).first()
                        projectDet = models.GlobalTransaction.objects.values('id','project__name','created_by__first_name','created_by__last_name').filter(id=request.data["transaction_id"]).first()

                        # Previous level approval notification
                        if approvalData["level"] == 1 or approvalData["level"] == 4:
                            notificationMessage = (
                                "Project "+projectDet['project__name']
                                + " request "
                                + projectDet['created_by__first_name']+" "+projectDet['created_by__last_name']+" in "+formDet['name']
                                + " is returned from recommender "
                            )
                        elif approvalData["level"] == 2 or approvalData["level"] == 5:
                            notificationMessage = (
                                "Project "+projectDet['project__name']
                                + " requestion "
                                + projectDet['created_by__first_name']+" "+projectDet['created_by__last_name']+" in "+formDet['name']
                                + " is returned from approver"
                            )
                        elif approvalData["level"] == 3:
                            notificationMessage = (
                                "Project "+projectDet['project__name']
                                + " requestion "
                                + projectDet['created_by__first_name']+" "+projectDet['created_by__last_name']+" in "+formDet['name']
                                + " is returned from initiater of trial unit"
                            )

                        if approvalData["level"]==1:
                            hierarchy_type = 1
                        elif approvalData["level"]==2:
                            hierarchy_type = 2

                        notificationModels.NotificationUser.objects.create(
                            form_id = globalTrans['form'],
                            transaction_id = globalTrans["id"],
                            message = notificationMessage,
                            user_role_id = approvalData["user_role_id"],
                            process_id = approvalData["process_id"],
                            hierarchy_type = hierarchy_type
                        )

                        print(globalTrans['form'], globalTrans["id"],notificationMessage,approvalData["user_role_id"],approvalData["process_id"],"||||||||||||")

                        # models.trialStatus.objects.create(
                        #     trial_id=trials["id"],
                        #     process_flow_id=approvalData["id"],
                        #     created_by_id=request.user.id,
                        #     created_ip=Common.get_client_ip(request),
                        # )

                        models.GlobalTransactionApproval.objects.create(
                            approved_level=request.data["approved_level"],
                            approved_ip=Common.get_client_ip(request),
                            approved_by_id=request.user.id,
                            transaction_id=request.data["transaction_id"],
                            comments=request.data["comments"],
                            status=request.data["status"],
                            approved_role_id=request.data["approved_role_id"],
                            type=request.data["type"],
                        )

                        models.GlobalTransaction.objects.filter(
                            id = request.data["transaction_id"]
                        ).update(approved_level=approvalData["level"])

                    return Response(
                        {
                            "status": error.context["success_code"],
                            "message": "Returned back to previous level successfully",
                        },
                        status=status.HTTP_200_OK,
                    )

            except:
                return Response(
                    {
                        "status": error.context["error_code"],
                        "message": "Failed to perform this action",
                    },
                    status=status.HTTP_200_OK,
                )


# GlobalApproval
class GlobalApproval_Running_Code_Form_Project(APIView):
    def post(self, request, pk=None):
        if "approved_level" not in request.data:
            return Response(
                {
                    "status": error.context["error_code"],
                    "message": "Level"
                    + language.context[language.defaultLang]["missing"],
                },
                status=status.HTTP_200_OK,
            )
        elif "form_id" not in request.data:
            return Response(
                {
                    "status": error.context["error_code"],
                    "message": "Form ID"
                    + language.context[language.defaultLang]["missing"],
                },
                status=status.HTTP_200_OK,
            )
        elif "comments" not in request.data:
            return Response(
                {
                    "status": error.context["error_code"],
                    "message": "Comment"
                    + language.context[language.defaultLang]["missing"],
                },
                status=status.HTTP_200_OK,
            )
        elif "status" not in request.data:
            return Response(
                {
                    "status": error.context["error_code"],
                    "message": "Status"
                    + language.context[language.defaultLang]["missing"],
                },
                status=status.HTTP_200_OK,
            )
        elif "approved_role_id" not in request.data:
            return Response(
                {
                    "status": error.context["error_code"],
                    "message": "Approved Role ID"
                    + language.context[language.defaultLang]["missing"],
                },
                status=status.HTTP_200_OK,
            )
        else:
        #try:
            if int(request.data["approved_level"]) == 2:
                models.GlobalTransaction.objects.filter(id=request.data["transaction_id"]).update(
                    recommender=request.data["status"]
                )
            elif int(request.data["approved_level"]) == 3:
                models.GlobalTransaction.objects.filter(id=request.data["transaction_id"]).update(
                    approver=request.data["status"]
                )

            
            form_id = request.data["form_id"]
            project_id = request.data["project_id"]
            type = request.data["type"]

            # Check the hierarchy type (form level or project level).
            if 'project_id' not in request.data or request.data["project_id"]=="": # Form Level
                rec_level = 1
            else: # Project Level                    
                project_id = request.data["project_id"]
                rec_level = 2

            # Notifiation message.
            formDet = models.Forms.objects.values('id','name').filter(id=request.data["form_id"]).first()
            projectDet = models.GlobalTransaction.objects.values('id','project__name','created_by__first_name','created_by__last_name').filter(id=request.data["transaction_id"]).first()

            userInfo = models.FormLevelRecommenderHierarchy.objects.values('id','recommender_id__first_name','recommender_id__last_name','recommender_level').filter(
                        form_id = form_id,
                        recommender_id = request.user.id
                        ).first()

            #print(userInfo,"userInfo.query")
            #pass
            #sys.exit(1)

            #     accessModels.ProcessFlow.objects.values()

            #print(rec_level, "rec_levelrec_level")

            if request.data["status"] == 1:

                
                if rec_level==1: # Form level

                    if type==1: # Initiator

                        # Get first recommender info.
                        first_level_rec = models.FormLevelRecommenderHierarchy.objects.values('recommender_id','recommender_level').filter(
                        form_id = form_id,
                        recommender_level = 1
                        )

                        # First recommender update
                        first_rec = models.FormLevelRecommenderHierarchy.objects.filter(
                        form_id = form_id,
                        recommender_level = 1
                        ).update(recommender_level_status=None, current_reject=None)

                        notificationMessage = (
                            "Project "+projectDet['project__name']
                            + " request "
                            + request.user.first_name+" "+request.user.last_name+" in "+formDet['name']
                            + " is reinitiated from initiater"
                            )

                        notificationModels.NotificationUser.objects.create(
                            form_id = request.data["form_id"],
                            transaction_id = request.data["transaction_id"],
                            message = notificationMessage,
                            user_role_id = 2,
                            process_id = 1,
                            hierarchy_type = 2, # Recommender
                            from_user_id = request.user.id,
                            to_user_id = first_level_rec[0]['recommender_id'],
                        )

                        models.GlobalTransactionApproval.objects.create(
                            approved_level = request.data["approved_level"],
                            approved_ip = Common.get_client_ip(request),
                            approved_by_id = request.user.id,
                            transaction_id = request.data["transaction_id"],
                            comments = request.data["comments"],
                            status = request.data["status"],
                            approved_role_id = request.data["approved_role_id"],
                            #type = request.data["type"],
                            type = 1,
                        )

                        return Response(
                            {
                                "status": error.context["success_code"],
                                "message": "initiator reinitiated successfully",
                            },
                            status=status.HTTP_200_OK,
                        )

                    if type==2: # Recommendation
                        rec_count = models.FormLevelRecommenderHierarchy.objects.filter(form_id = form_id).count()
                        if rec_count>1:

                            # Current recommender.
                            current_level_rec = models.FormLevelRecommenderHierarchy.objects.values('recommender_id','recommender_level').filter(
                            form_id = form_id,
                            recommender_id = request.user.id
                            )

                            # Current record update.
                            models.FormLevelRecommenderHierarchy.objects.filter(
                            form_id = form_id,
                            recommender_id = request.user.id
                            ).update(recommender_level_status=1, current_reject=None)

                            # Next recommender.
                            next_level_rec = models.FormLevelRecommenderHierarchy.objects.values('recommender_id','recommender_level').filter(
                            form_id = form_id,
                            recommender_level = current_level_rec[0]['recommender_level']+1
                            )

                            if rec_count==current_level_rec[0]['recommender_level']:

                                models.GlobalTransaction.objects.filter(id = request.data["transaction_id"]).update(approved_level=3)

                                userInfoApprover = models.FormLevelApproverHierarchy.objects.values('id','approver_id','approver_id__first_name','approver_id__last_name','approver_level').filter(form_id = form_id).first()

                                current_rec_count = 1
                            else:
                                current_rec_count = 2


                            # Next record update.
                            if current_rec_count==2:
                            
                                models.FormLevelRecommenderHierarchy.objects.filter(
                                form_id = form_id,
                                recommender_level = current_level_rec[0]['recommender_level']+1
                                ).update(recommender_level_status=None)
                            
                            elif current_rec_count==1:

                                #pass
                                models.FormLevelApproverHierarchy.objects.filter(
                                form_id = form_id,
                                approver_level = 1
                                ).update(approver_level_status=-1)

                            if current_rec_count==2:
                                notificationMessage = (
                                    "Project "+projectDet['project__name']
                                    + " is recommended by "
                                    + userInfo['recommender_id__first_name']+" "+userInfo['recommender_id__last_name']+" in "+formDet['name']
                                    + " is waiting for your next level recommendations "
                                )
                            elif current_rec_count==1:
                                notificationMessage = (
                                    "Project "+projectDet['project__name']
                                    + " is recommended by "
                                    + userInfo['recommender_id__first_name']+" "+userInfo['recommender_id__last_name']+" in "+formDet['name']
                                    + " is waiting for your approval"
                                )


                            res1 = notificationModels.NotificationUser.objects.create(
                                form_id = request.data["form_id"],
                                transaction_id = request.data["transaction_id"],
                                project_id = request.data["project_id"],
                                message = notificationMessage,
                                user_role_id = 3 if current_rec_count==2 else 4,
                                process_id = 1,
                                hierarchy_type = 2 if current_rec_count==2 else 3,
                                from_user_id = request.user.id,
                                to_user_id = next_level_rec[0]['recommender_id'] if current_rec_count==2 else userInfoApprover['approver_id'],
                            )

                            res2 = models.GlobalTransactionApproval.objects.create(
                                approved_level = request.data["approved_level"] if current_rec_count==2 else 3,
                                approved_ip = Common.get_client_ip(request),
                                approved_by_id = request.user.id,
                                transaction_id = request.data["transaction_id"],
                                comments = request.data["comments"],
                                status = request.data["status"],
                                approved_role_id = request.data["approved_role_id"],
                                type = request.data["type"] if current_rec_count==2 else 3
                            )

                            return Response(
                                {
                                    "status": error.context["success_code"],
                                    "message": "Recommendation accepted successfully",
                                },
                                status=status.HTTP_200_OK,
                            )
            
                    elif type==3: # Approval

                        # Approver.
                        approver = models.FormLevelApproverHierarchy.objects.filter(
                        form_id = form_id,
                        approver_id = request.user.id
                        ).update(approver_level_status=1)

                        res2 = models.GlobalTransactionApproval.objects.create(
                            approved_level = request.data["approved_level"],
                            approved_ip = Common.get_client_ip(request),
                            approved_by_id = request.user.id,
                            transaction_id = request.data["transaction_id"],
                            comments = request.data["comments"],
                            status = request.data["status"],
                            approved_role_id = request.data["approved_role_id"],
                            type = request.data["type"]
                        )

                        # # Current record update.
                        # models.FormLevelApproverHierarchy.objects.filter(
                        # form_id = form_id,
                        # approver_id = request.user.id
                        # ).update(approver_level_status=1)

                        return Response(
                            {
                                "status": error.context["success_code"],
                                "message": "Recommendation accepted successfully",
                            },
                            status=status.HTTP_200_OK,
                        )




                elif rec_level==2: # Project level


                    curUserInfo = models.ProjectLevelRecommenderHierarchy.objects.values('id','recommender_id','recommender_id__first_name','recommender_id__last_name','recommender_level').filter(form_id = form_id, project_id = project_id, recommender_id = request.user.id).first()

                    if type==1: # Initiator

                        # Get first recommender info.
                        first_level_rec = models.ProjectLevelRecommenderHierarchy.objects.values('recommender_id','recommender_level').filter(
                        form_id = form_id,
                        project_id = project_id,
                        recommender_level = 1
                        )

                        # First recommender update
                        first_rec = models.ProjectLevelRecommenderHierarchy.objects.filter(
                        form_id = form_id,
                        project_id = project_id,
                        recommender_level = 1
                        ).update(recommender_level_status=None, current_reject=None)

                        notificationMessage = (
                            "Project "+projectDet['project__name']
                            + " request "
                            + request.user.first_name+" "+request.user.last_name+" in "+formDet['name']
                            + " is reinitiated from initiater"
                            )

                        notificationModels.NotificationUser.objects.create(
                            form_id = request.data["form_id"],
                            transaction_id = request.data["transaction_id"],
                            project_id = request.data["project_id"],
                            message = notificationMessage,
                            user_role_id = 2,
                            process_id = 1,
                            hierarchy_type = 2, # Recommender
                            from_user_id = request.user.id,
                            to_user_id = first_level_rec[0]['recommender_id'],
                        )

                        models.GlobalTransactionApproval.objects.create(
                            approved_level = request.data["approved_level"],
                            approved_ip = Common.get_client_ip(request),
                            approved_by_id = request.user.id,
                            transaction_id = request.data["transaction_id"],
                            comments = request.data["comments"],
                            status = request.data["status"],
                            approved_role_id = request.data["approved_role_id"],
                            #type = request.data["type"],
                            type = 1,
                        )

                        return Response(
                            {
                                "status": error.context["success_code"],
                                "message": "initiator reinitiated successfully",
                            },
                            status=status.HTTP_200_OK,
                        )

                    if type==2: # Recommendation

                        rec_count = models.ProjectLevelRecommenderHierarchy.objects.filter(form_id = form_id, project_id = project_id).count()


                        print(rec_count, "rec_count")
                        if rec_count==1:
                            print(rec_count, "rec_countAA")
                            # Current record update.
                            models.ProjectLevelRecommenderHierarchy.objects.filter(
                            form_id = form_id,
                            project_id = project_id,
                            recommender_id = request.user.id
                            ).update(recommender_level_status=1, current_reject=None)


                            # Current user info
                            # curUserInfo = models.ProjectLevelRecommenderHierarchy.objects.values('id','recommender_id','recommender_id__first_name','recommender_id__last_name','recommender_level').filter(form_id = form_id, project_id = project_id).first()


                            # Next update the approver.
                            models.ProjectLevelApproverHierarchy.objects.filter(
                            form_id = form_id,
                            project_id = project_id,
                            approver_level = 1
                            ).update(approver_level_status=-1)


                            # Next approver user info
                            nextApproverInfo = models.ProjectLevelApproverHierarchy.objects.values('id','approver_id','approver_level').filter(form_id = form_id, project_id = project_id).first()

                            # Update the global transaction table.
                            models.GlobalTransaction.objects.filter(id = request.data["transaction_id"]).update(approved_level=3)

                            notificationMessage = (
                                "Project "+projectDet['project__name']
                                + " is recommended by "
                                + curUserInfo['recommender_id__first_name']+" "+curUserInfo['recommender_id__last_name']+" in "+formDet['name']
                                + " is waiting for your approval"
                            )


                            res1 = notificationModels.NotificationUser.objects.create(
                                form_id = request.data["form_id"],
                                transaction_id = request.data["transaction_id"],
                                project_id = request.data["project_id"],
                                message = notificationMessage,
                                user_role_id = 3,
                                process_id = 1,
                                hierarchy_type = 3,
                                from_user_id = request.user.id,
                                to_user_id = nextApproverInfo['approver_id']
                            )

                            res2 = models.GlobalTransactionApproval.objects.create(
                                approved_level = request.data["approved_level"],
                                approved_ip = Common.get_client_ip(request),
                                approved_by_id = request.user.id,
                                transaction_id = request.data["transaction_id"],
                                comments = request.data["comments"],
                                status = request.data["status"],
                                approved_role_id = request.data["approved_role_id"],
                                type = request.data["type"]
                            )

                            return Response(
                                {
                                    "status": error.context["success_code"],
                                    "message": "Recommendation accepted successfully",
                                },
                                status=status.HTTP_200_OK,
                            )


                        if rec_count>1:

                            print(rec_count, "rec_countBB")
                            # Current recommender.
                            current_level_rec = models.ProjectLevelRecommenderHierarchy.objects.values('recommender_id','recommender_level').filter(
                            form_id = form_id,
                            recommender_id = request.user.id
                            )

                            # Current record update.
                            models.ProjectLevelRecommenderHierarchy.objects.filter(
                            form_id = form_id,
                            recommender_id = request.user.id
                            ).update(recommender_level_status=1, current_reject=None)

                            # Next recommender.
                            next_level_rec = models.ProjectLevelRecommenderHierarchy.objects.values('recommender_id','recommender_level').filter(
                            form_id = form_id,
                            recommender_level = current_level_rec[0]['recommender_level']+1
                            )

                            if rec_count==current_level_rec[0]['recommender_level']:

                                models.GlobalTransaction.objects.filter(id = request.data["transaction_id"]).update(approved_level=3)

                                userInfoApprover = models.ProjectLevelApproverHierarchy.objects.values('id','approver_id','approver_id__first_name','approver_id__last_name','approver_level').filter(form_id = form_id, project_id = project_id).first()

                                current_rec_count = 1
                            else:
                                current_rec_count = 2


                            # Next record update.
                            if current_rec_count==2:

                                models.ProjectLevelRecommenderHierarchy.objects.filter(
                                form_id = form_id,
                                recommender_level = current_level_rec[0]['recommender_level']+1
                                ).update(recommender_level_status=None)
                            
                            elif current_rec_count==1:
                            
                                #pass
                                models.ProjectLevelApproverHierarchy.objects.filter(
                                form_id = form_id,
                                project_id = project_id,
                                approver_level = 1
                                ).update(approver_level_status=-1)


                            # Current user info
                            # curUserInfo = models.ProjectLevelRecommenderHierarchy.objects.values('id','recommender_id','recommender_id__first_name','recommender_id__last_name','recommender_level').filter(form_id = form_id, project_id = project_id, recommender_id = request.user.id).first()


                            if current_rec_count==2:
                                notificationMessage = (
                                    "Project "+projectDet['project__name']
                                    + " is recommended by "
                                    + curUserInfo['recommender_id__first_name']+" "+curUserInfo['recommender_id__last_name']+" in "+formDet['name']
                                    + " is waiting for your next level recommendations "
                                )
                            elif current_rec_count==1:
                                notificationMessage = (
                                    "Project "+projectDet['project__name']
                                    + " is recommended by "
                                    + curUserInfo['recommender_id__first_name']+" "+curUserInfo['recommender_id__last_name']+" in "+formDet['name']
                                    + " is waiting for your approval"
                                )


                            res1 = notificationModels.NotificationUser.objects.create(
                                form_id = request.data["form_id"],
                                transaction_id = request.data["transaction_id"],
                                project_id = request.data["project_id"],
                                message = notificationMessage,
                                user_role_id = 3 if current_rec_count==2 else 4,
                                process_id = 1,
                                hierarchy_type = 2 if current_rec_count==2 else 3,
                                from_user_id = request.user.id,
                                to_user_id = next_level_rec[0]['recommender_id'] if current_rec_count==2 else userInfoApprover['approver_id'],
                            )

                            res2 = models.GlobalTransactionApproval.objects.create(
                                approved_level = request.data["approved_level"] if current_rec_count==2 else 3,
                                approved_ip = Common.get_client_ip(request),
                                approved_by_id = request.user.id,
                                transaction_id = request.data["transaction_id"],
                                comments = request.data["comments"],
                                status = request.data["status"],
                                approved_role_id = request.data["approved_role_id"],
                                type = request.data["type"] if current_rec_count==2 else 3
                            )

                            return Response(
                                {
                                    "status": error.context["success_code"],
                                    "message": "Recommendation accepted successfully",
                                },
                                status=status.HTTP_200_OK,
                            )


                    elif type==3: # Approver
                        # Approver.
                        approver = models.ProjectLevelApproverHierarchy.objects.filter(
                        form_id = form_id,
                        project_id = project_id,
                        approver_id = request.user.id
                        ).update(approver_level_status=1)

                        res2 = models.GlobalTransactionApproval.objects.create(
                            approved_level = request.data["approved_level"],
                            approved_ip = Common.get_client_ip(request),
                            approved_by_id = request.user.id,
                            transaction_id = request.data["transaction_id"],
                            comments = request.data["comments"],
                            status = request.data["status"],
                            approved_role_id = request.data["approved_role_id"],
                            type = request.data["type"]
                        )

                        models.GlobalTransaction.objects.filter(id=request.data["transaction_id"]).update(
                        approved_level=4)

                        return Response(
                            {
                                "status": error.context["success_code"],
                                "message": "Recommendation accepted successfully",
                            },
                            status=status.HTTP_200_OK,
                        )
                    

            elif request.data["status"] == 2:

                # Check the hierarchy type (form level or project level).
                # Get the current logged in user id.

                formDet = models.Forms.objects.values('id','name').filter(id=request.data["form_id"]).first()
                projectDet = models.GlobalTransaction.objects.values('id','project__name','created_by__first_name','created_by__last_name').filter(id=request.data["transaction_id"]).first()

                if rec_level==1: # Form level

                    # Get the current level and user.
                    current_level_rec = models.FormLevelRecommenderHierarchy.objects.values('recommender_id','recommender_level').filter(
                    form_id = form_id,
                    recommender_id = request.user.id
                    )

                    if request.data["type"]==2: # Reject from recommender.

                        if current_level_rec[0]['recommender_level']==1: # Initiator

                            # Get Initiator details

                            first_rec = models.FormLevelRecommenderHierarchy.objects.filter(
                            form_id = form_id,
                            recommender_level = 1
                            ).update(recommender_level_status=2, current_reject=1)

                            notificationMessage = (
                                "Project "+projectDet['project__name']
                                + " request "
                                + userInfo['recommender_id__first_name']+" "+userInfo['recommender_id__last_name']+" in "+formDet['name']
                                + " is returned from recommender at level 1"
                                )

                            notificationModels.NotificationUser.objects.create(
                                form_id = request.data["form_id"],
                                transaction_id = request.data["transaction_id"],
                                message = notificationMessage,
                                user_role_id = 2,
                                process_id = 1,
                                hierarchy_type = 1, # Initiator
                                from_user_id = request.user.id,
                                #to_user_id = next_level_rec[0]['recommender_id'] if current_rec_count==2 else userInfoApprover['approver_id'],
                            )

                            models.GlobalTransactionApproval.objects.create(
                                approved_level = request.data["approved_level"],
                                approved_ip = Common.get_client_ip(request),
                                approved_by_id = request.user.id,
                                transaction_id = request.data["transaction_id"],
                                comments = request.data["comments"],
                                status = request.data["status"],
                                approved_role_id = request.data["approved_role_id"],
                                #type = request.data["type"],
                                type = 1,
                            )

                        elif current_level_rec[0]['recommender_level']>=1: # Previous recommender

                            # Previous recommender.
                            pre_level_rec = models.FormLevelRecommenderHierarchy.objects.values('recommender_id','recommender_level').filter(
                            form_id = form_id,
                            recommender_level = current_level_rec[0]['recommender_level']-1
                            )

                            # Previous record update.
                            pre_rec = models.FormLevelRecommenderHierarchy.objects.filter(
                            form_id = form_id,
                            recommender_level = current_level_rec[0]['recommender_level']-1
                            ).update(recommender_level_status=None, current_reject=1)

                            # Current record update.
                            cur_rec = models.FormLevelRecommenderHierarchy.objects.filter(
                            form_id = form_id,
                            recommender_id = request.user.id
                            ).update(recommender_level_status=2)


                            notificationMessage = (
                                "Project "+projectDet['project__name']
                                + " request "
                                + userInfo['recommender_id__first_name']+" "+userInfo['recommender_id__last_name']+" in "+formDet['name']
                                + " is returned from recommender"
                            )

                            res1 = notificationModels.NotificationUser.objects.create(
                                form_id = request.data["form_id"],
                                transaction_id = request.data["transaction_id"],
                                message = notificationMessage,
                                user_role_id = 3,
                                process_id = 1,
                                hierarchy_type = 2,
                                from_user_id = request.user.id,
                                to_user_id = pre_level_rec[0]['recommender_id'],
                            )

                            res2 = models.GlobalTransactionApproval.objects.create(
                                approved_level = request.data["approved_level"],
                                approved_ip = Common.get_client_ip(request),
                                approved_by_id = request.user.id,
                                transaction_id = request.data["transaction_id"],
                                comments = request.data["comments"],
                                status = request.data["status"],
                                approved_role_id = request.data["approved_role_id"],
                                type = request.data["type"]
                            )

                    if request.data["type"]==3: # Reject from approver.

                        last_rec = models.FormLevelRecommenderHierarchy.objects.values('recommender_id','recommender_level').order_by('-recommender_level').filter(form_id = form_id).first()


                        userInfoApprover = models.FormLevelApproverHierarchy.objects.values('id','approver_id','approver_id__first_name','approver_id__last_name','approver_level').filter(form_id = form_id).first()
                        #print(last_rec,"last_rec", last_rec['recommender_level'])

                        models.FormLevelRecommenderHierarchy.objects.filter(
                        form_id = form_id,
                        recommender_id = last_rec['recommender_id'],
                        recommender_level = last_rec['recommender_level'],
                        ).update(recommender_level_status=None, current_reject = 1)

                        # Approver.
                        approver = models.FormLevelApproverHierarchy.objects.filter(
                        form_id = form_id,
                        approver_id = request.user.id
                        ).update(approver_level_status=2)


                        notificationMessage = (
                            "Project "+projectDet['project__name']
                            + " request "
                            + userInfoApprover['approver_id__first_name']+" "+userInfoApprover['approver_id__last_name']+" in "+formDet['name']
                            + " is returned from approver"
                        )

                        res1 = notificationModels.NotificationUser.objects.create(
                            form_id = request.data["form_id"],
                            transaction_id = request.data["transaction_id"],
                            message = notificationMessage,
                            user_role_id = 3,
                            process_id = 1,
                            hierarchy_type = 2,
                            from_user_id = request.user.id,
                            to_user_id = last_rec['recommender_id'],
                        )

                        res2 = models.GlobalTransactionApproval.objects.create(
                            approved_level = request.data["approved_level"],
                            approved_ip = Common.get_client_ip(request),
                            approved_by_id = request.user.id,
                            transaction_id = request.data["transaction_id"],
                            comments = request.data["comments"],
                            status = request.data["status"],
                            approved_role_id = request.data["approved_role_id"],
                            type = request.data["type"]
                        )

                else:
                    # Disaaprove from approve level.
                    pass



                if rec_level==2: # Project level

                    # Get the current level and user.
                    current_level_rec = models.ProjectLevelRecommenderHierarchy.objects.values('recommender_id','recommender_level').filter(
                    form_id = form_id,
                    project_id = project_id,
                    recommender_id = request.user.id
                    )

                    # Get Current User Info
                    curUserInfo = models.ProjectLevelRecommenderHierarchy.objects.values('id','recommender_id','recommender_id__first_name','recommender_id__last_name','recommender_level').filter(form_id = form_id, project_id = project_id, recommender_id = request.user.id).first()

                    if request.data["type"]==2: # Reject from recommender.

                        if current_level_rec[0]['recommender_level']==1: # Initiator

                            # Get Initiator details

                            first_rec = models.ProjectLevelRecommenderHierarchy.objects.filter(
                            form_id = form_id,
                            project_id = project_id,
                            recommender_level = 1
                            ).update(recommender_level_status=2, current_reject=1)

                            # # Get Current User Info
                            # curUserInfo = models.ProjectLevelRecommenderHierarchy.objects.values('id','recommender_id','recommender_id__first_name','recommender_id__last_name','recommender_level').filter(form_id = form_id, project_id = project_id, recommender_id = request.user.id).first()


                            notificationMessage = (
                                "Project "+projectDet['project__name']
                                + " request "
                                + curUserInfo['recommender_id__first_name']+" "+curUserInfo['recommender_id__last_name']+" in "+formDet['name']
                                + " is returned from recommender at level 1"
                                )

                            notificationModels.NotificationUser.objects.create(
                                form_id = request.data["form_id"],
                                transaction_id = request.data["transaction_id"],
                                project_id = request.data["project_id"],
                                message = notificationMessage,
                                user_role_id = 2,
                                process_id = 1,
                                hierarchy_type = 1, # Initiator
                                from_user_id = request.user.id,
                                #to_user_id = next_level_rec[0]['recommender_id'] if current_rec_count==2 else userInfoApprover['approver_id'],
                            )

                            models.GlobalTransactionApproval.objects.create(
                                approved_level = request.data["approved_level"],
                                approved_ip = Common.get_client_ip(request),
                                approved_by_id = request.user.id,
                                transaction_id = request.data["transaction_id"],
                                comments = request.data["comments"],
                                status = request.data["status"],
                                approved_role_id = request.data["approved_role_id"],
                                #type = request.data["type"],
                                type = 1,
                            )

                        elif current_level_rec[0]['recommender_level']>=1: # Previous recommender

                            # Previous recommender.
                            pre_level_rec = models.ProjectLevelRecommenderHierarchy.objects.values('recommender_id','recommender_level').filter(
                            form_id = form_id,
                            project_id = project_id,
                            recommender_level = current_level_rec[0]['recommender_level']-1
                            )

                            # Previous record update.
                            pre_rec = models.ProjectLevelRecommenderHierarchy.objects.filter(
                            form_id = form_id,
                            project_id = project_id,
                            recommender_level = current_level_rec[0]['recommender_level']-1
                            ).update(recommender_level_status=None, current_reject=1)

                            # Current record update.
                            cur_rec = models.ProjectLevelRecommenderHierarchy.objects.filter(
                            form_id = form_id,
                            project_id = project_id,
                            recommender_id = request.user.id
                            ).update(recommender_level_status=2)

                            # # Current user info
                            # curUserInfo = models.ProjectLevelRecommenderHierarchy.objects.values('id','recommender_id','recommender_id__first_name','recommender_id__last_name','recommender_level').filter(form_id = form_id, project_id = project_id, recommender_id = request.user.id).first()


                            notificationMessage = (
                                "Project "+projectDet['project__name']
                                + " request "
                                + curUserInfo['recommender_id__first_name']+" "+curUserInfo['recommender_id__last_name']+" in "+formDet['name']
                                + " is returned from recommender"
                            )

                            res1 = notificationModels.NotificationUser.objects.create(
                                form_id = request.data["form_id"],
                                transaction_id = request.data["transaction_id"],
                                project_id = request.data["project_id"],
                                message = notificationMessage,
                                user_role_id = 3,
                                process_id = 1,
                                hierarchy_type = 2,
                                from_user_id = request.user.id,
                                to_user_id = pre_level_rec[0]['recommender_id'],
                            )

                            res2 = models.GlobalTransactionApproval.objects.create(
                                approved_level = request.data["approved_level"],
                                approved_ip = Common.get_client_ip(request),
                                approved_by_id = request.user.id,
                                transaction_id = request.data["transaction_id"],
                                comments = request.data["comments"],
                                status = request.data["status"],
                                approved_role_id = request.data["approved_role_id"],
                                type = request.data["type"]
                            )

                    if request.data["type"]==3: # Reject from approver.

                        last_rec = models.ProjectLevelRecommenderHierarchy.objects.values('recommender_id','recommender_level').order_by('-recommender_level').filter(form_id = form_id, project_id = project_id).first()


                        userInfoApprover = models.ProjectLevelApproverHierarchy.objects.values('id','approver_id','approver_id__first_name','approver_id__last_name','approver_level').filter(form_id = form_id, project_id = project_id, approver_id = request.user.id).first()
                        #print(last_rec,"last_rec", last_rec['recommender_level'])

                        models.ProjectLevelRecommenderHierarchy.objects.filter(
                        form_id = form_id,
                        project_id = project_id,
                        recommender_id = last_rec['recommender_id'],
                        recommender_level = last_rec['recommender_level'],
                        ).update(recommender_level_status=None, current_reject = 1)

                        # Approver.
                        approver = models.ProjectLevelApproverHierarchy.objects.filter(
                        form_id = form_id,
                        project_id = project_id,
                        approver_id = request.user.id
                        ).update(approver_level_status=2)


                        notificationMessage = (
                            "Project "+projectDet['project__name']
                            + " request "
                            + userInfoApprover['approver_id__first_name']+" "+userInfoApprover['approver_id__last_name']+" in "+formDet['name']
                            + " is returned from approver"
                        )

                        res1 = notificationModels.NotificationUser.objects.create(
                            form_id = request.data["form_id"],
                            transaction_id = request.data["transaction_id"],
                            project_id = request.data["project_id"],
                            message = notificationMessage,
                            user_role_id = 3,
                            process_id = 1,
                            hierarchy_type = 2,
                            from_user_id = request.user.id,
                            to_user_id = last_rec['recommender_id'],
                        )

                        res2 = models.GlobalTransactionApproval.objects.create(
                            approved_level = request.data["approved_level"],
                            approved_ip = Common.get_client_ip(request),
                            approved_by_id = request.user.id,
                            transaction_id = request.data["transaction_id"],
                            comments = request.data["comments"],
                            status = request.data["status"],
                            approved_role_id = request.data["approved_role_id"],
                            type = request.data["type"]
                        )

                else:
                    # Disaaprove from approve level.
                    pass


                #print(form_rec_level_count,"# Project Level",form_rec_level_approve_count)
                return Response(
                    {
                        "status": error.context["success_code"],
                        "message": "Returned back to previous level successfully",
                    },
                    status=status.HTTP_200_OK,
                )

            # except:
            #     return Response(
            #         {
            #             "status": error.context["error_code"],
            #             "message": "Failed to perform this action",
            #         },
            #         status=status.HTTP_200_OK,
            #     )


# GlobalApproval
class GlobalApproval_OLD_ONE(APIView):
    def post(self, request, pk=None):
        if "approved_level" not in request.data:
            return Response(
                {
                    "status": error.context["error_code"],
                    "message": "Level"
                    + language.context[language.defaultLang]["missing"],
                },
                status=status.HTTP_200_OK,
            )
        elif "form_id" not in request.data:
            return Response(
                {
                    "status": error.context["error_code"],
                    "message": "Form ID"
                    + language.context[language.defaultLang]["missing"],
                },
                status=status.HTTP_200_OK,
            )
        elif "comments" not in request.data:
            return Response(
                {
                    "status": error.context["error_code"],
                    "message": "Comment"
                    + language.context[language.defaultLang]["missing"],
                },
                status=status.HTTP_200_OK,
            )
        elif "status" not in request.data:
            return Response(
                {
                    "status": error.context["error_code"],
                    "message": "Status"
                    + language.context[language.defaultLang]["missing"],
                },
                status=status.HTTP_200_OK,
            )
        elif "approved_role_id" not in request.data:
            return Response(
                {
                    "status": error.context["error_code"],
                    "message": "Approved Role ID"
                    + language.context[language.defaultLang]["missing"],
                },
                status=status.HTTP_200_OK,
            )
        else:
            try:
                if int(request.data["approved_level"]) == 2:
                    models.GlobalTransaction.objects.filter(id=request.data["transaction_id"]).update(
                        recommender=request.data["status"]
                    )
                elif int(request.data["approved_level"]) == 3:
                    models.GlobalTransaction.objects.filter(id=request.data["transaction_id"]).update(
                        approver=request.data["status"]
                    )

                # Check the hierarchy type (form level or project level).
                form_id = request.data["form_id"]
                
                if 'project_id' not in request.data: # Form Level
                    rec_level = 1
                else: # Project Level                    
                    project_id = request.data["project_id"]
                    rec_level = 2


                if request.data["status"] == 1:

                    ## ##
                    res = models.FormLevelRecommenderHierarchy.objects.values('recommender_id','recommender_level').filter(
                    form_id = form_id,
                    recommender_id = request.user.id
                    )

                    # Next record
                    next_level = models.FormLevelRecommenderHierarchy.objects.values('recommender_id','recommender_level','recommender_level_status').filter(
                    form_id = form_id,
                    recommender_level = res[0]['recommender_level']+1
                    )

                    if next_level[0]['recommender_level_status']==2:
                        next_level = models.FormLevelRecommenderHierarchy.objects.filter(
                        form_id = form_id,
                        recommender_level = res[0]['recommender_level']+1
                        ).update(recommender_level_status=None)

                    # Current record.
                    cur = models.FormLevelRecommenderHierarchy.objects.filter(
                    form_id = form_id,
                    recommender_id = request.user.id
                    ).update(recommender_level_status=1)
                    ## ##


                    approvalData = (
                        accessModels.ProcessFlow.objects.values()
                        .filter(level__gt=int(request.data["approved_level"]))
                        .order_by("id")
                        .first()
                    )

                    if approvalData:
                        # Next level approval notification
                        globalTrans = cSerializer.ListGlobalTransactionSerializer(
                            models.GlobalTransaction.objects.filter(id=request.data["transaction_id"]),
                            many=True,
                            context={"request": request},
                        ).data[0]

                        #print(globalTrans['form'],"GlobalTrans: GlobalTrans")
                        print(approvalData["level"],"Level: Level")

                        formDet = models.Forms.objects.values('id','name').filter(id=request.data["form_id"]).first()
                        projectDet = models.GlobalTransaction.objects.values('id','project__name','created_by__first_name','created_by__last_name').filter(id=request.data["transaction_id"]).first()

                        #print(projectDet)
                        # models.trialStatus.objects.create(
                        #     trial_id=trials["id"],
                        #     process_flow_id=approvalData["id"],
                        #     created_by_id=request.user.id,
                        #     created_ip=Common.get_client_ip(request),
                        # )

                        if approvalData["level"] == 2:
                            notificationMessage = (
                                "Project "+projectDet['project__name']
                                + " request "
                                + projectDet['created_by__first_name']+" "+projectDet['created_by__last_name']+" in "+formDet['name']
                                + ") is re initiated again and waiting for your approval "
                            )
                        elif approvalData["level"] == 3:
                            notificationMessage = (
                                "Project "+projectDet['project__name']
                                + " is recommended by "
                                + projectDet['created_by__first_name']+" "+projectDet['created_by__last_name']+" in "+formDet['name']
                                + " is waiting for your approval "
                            )

                        #print(notificationMessage,"notificationMessage")
                        # elif approvalData["level"] == 4:
                        #     notificationMessage = (
                        #         (trials["trial_type"]["type"])
                        #         + " requestion ("
                        #         + trials["trial_number"]
                        #         + ") is approved and waiting for you to submit"
                        #     )
                        # elif approvalData["level"] == 5:
                        #     notificationMessage = (
                        #         (trials["trial_type"]["type"])
                        #         + " ("
                        #         + trials["trial_number"]
                        #         + ") has submitted and waiting for your recommendation"
                        #     )
                        # elif approvalData["level"] == 6:
                        #     notificationMessage = (
                        #         (trials["trial_type"]["type"])
                        #         + " ("
                        #         + trials["trial_number"]
                        #         + ") has recommended for your approval"
                        #     )

                        notificationModels.NotificationUser.objects.create(
                            form_id = globalTrans['form'],
                            transaction_id = globalTrans["id"],
                            message = notificationMessage,
                            user_role_id = approvalData["user_role_id"],
                            process_id = approvalData["process_id"],
                            hierarchy_type = 2
                        )

                        # models.GlobalTransaction.objects.filter(
                        #     id = request.data["transaction_id"]
                        # ).update(approved_level=approvalData["level"])

                        # Form Level
                        res = models.FormLevelRecommenderHierarchy.objects.filter(
                            form_id = globalTrans['form'],
                            recommender_id = request.user.id
                        ).update(recommender_level_status=1)


                        # Count 
                        rec_level_total = models.FormLevelRecommenderHierarchy.objects.values('recommender_id').filter(form_id = globalTrans['form']).count()

                        rec_level_current = models.FormLevelRecommenderHierarchy.objects.values('recommender_id', 'recommender_level').filter(form_id = globalTrans['form'],recommender_id = request.user.id)

                        if rec_level_total==rec_level_current[0]['recommender_level']:
                            models.GlobalTransaction.objects.filter(
                                id = globalTrans["id"]
                            ).update(approved_level=3)


                    else:

                        globalTrans = cSerializer.ListGlobalTransactionSerializer(
                            models.GlobalTransaction.objects.filter(id=request.data["transaction_id"]),
                            many=True,
                            context={"request": request},
                        ).data[0]

                        formDet = models.Forms.objects.values('id','name').filter(id=request.data["form_id"]).first()
                        projectDet = models.GlobalTransaction.objects.values('id','project__name','created_by__first_name','created_by__last_name').filter(id=request.data["transaction_id"]).first()

                        
                        notificationMessage = (
                            "Project "+projectDet['project__name']
                            + " request "
                            + projectDet['created_by__first_name']+" "+projectDet['created_by__last_name']+" in "+formDet['name']
                            + ") has been approved and report has been generated"
                        )


                        # NotificationUser.objects.create(
                        #     trial_unit_id=trials["trial_unit"]["id"],
                        #     satellite_unit_id=trials["satellite_unit"]["id"],
                        #     trial_id=trials["id"],
                        #     message=notificationMessage,
                        # )

                        # models.Trials.objects.filter(
                        #     id=request.data["trial_id"]
                        # ).update(approved_level="-1")


                        notificationModels.NotificationUser.objects.create(
                            form_id = globalTrans['form'],
                            transaction_id = globalTrans["id"],
                            message = notificationMessage,
                            #user_role_id = approvalData["user_role_id"],
                            #process_id = approvalData["process_id"],
                            hierarchy_type = 2
                        )

                        # models.GlobalTransaction.objects.filter(
                        #     id = request.data["transaction_id"]
                        # ).update(approved_level=-1)



                        # Form Level
                        res = models.FormLevelApproverHierarchy.objects.filter(
                            form_id = globalTrans['form'],
                            approver_id = request.user.id
                        ).update(approver_level_status=1)

                        # models.GlobalTransaction.objects.filter(
                        #     id = globalTrans["id"]
                        # ).update(approved_level=3)


                        # Count 
                        # rec_level_total = models.FormLevelRecommenderHierarchy.objects.values('recommender_id').filter(form_id = globalTrans['form']).count()

                        # rec_level_current = models.FormLevelRecommenderHierarchy.objects.values('recommender_id', 'recommender_level').filter(form_id = globalTrans['form'],recommender_id = request.user.id)

                        # if rec_level_total==rec_level_current[0]['recommender_level']:
                        #     models.GlobalTransaction.objects.filter(
                        #         id = globalTrans["id"]
                        #     ).update(approved_level=3)



                    # # models.trialApproval.objects.filter(approved_by_id=request.user.id,trial_id=request.data['trial_id'],approved_role_id=request.data['approved_role_id']).delete()
                    # models.trialApproval.objects.create(
                    #     approved_level=request.data["approved_level"],
                    #     approved_ip=Common.get_client_ip(request),
                    #     approved_by_id=request.user.id,
                    #     trial_id=request.data["trial_id"],
                    #     comments=request.data["comments"],
                    #     status=request.data["status"],
                    #     approved_role_id=request.data["approved_role_id"],
                    #     type=request.data["type"],
                    # )

                    models.GlobalTransactionApproval.objects.create(
                        approved_level = request.data["approved_level"],
                        approved_ip = Common.get_client_ip(request),
                        approved_by_id = request.user.id,
                        transaction_id = request.data["transaction_id"],
                        comments = request.data["comments"],
                        status = request.data["status"],
                        approved_role_id = request.data["approved_role_id"],
                        type = request.data["type"],
                    )
                    return Response(
                        {
                            "status": error.context["success_code"],
                            "message": "Approval updated successfully",
                        },
                        status=status.HTTP_200_OK,
                    )
                elif request.data["status"] == 2:

                    # Check the hierarchy type (form level or project level).
                    # Get the current logged in user id.

                    if rec_level==1: # Form level

                        # # Get total count
                        # form_rec_level_count = models.FormLevelRecommenderHierarchy.objects.filter(form_id = form_id).count()

                        # # Accept count
                        # form_rec_level_approve_count = models.FormLevelRecommenderHierarchy.objects.filter(form_id = form_id, recommender_level_status = 1).count()


                        # if form_rec_level_count!=form_rec_level_approve_count:

                        #     # Get the current level and user.
                        #     res = models.FormLevelRecommenderHierarchy.objects.values('recommender_id','recommender_level').filter(
                        #         form_id = form_id,
                        #         recommender_id = request.user.id
                        #     )

                        #     a = models.FormLevelRecommenderHierarchy.objects.filter(
                        #         recommender_level__lte = 2
                        #     ).update(recommender_level_status=2)


                        #     models.FormLevelRecommenderHierarchy.objects.filter(
                        #         form_id = form_id,
                        #         recommender_id = request.user.id
                        #     ).update(current_reject=1)

                        #     #print(res,"resresressresres")


                        # Get the current level and user.
                        res = models.FormLevelRecommenderHierarchy.objects.values('recommender_id','recommender_level').filter(
                        form_id = form_id,
                        recommender_id = request.user.id
                        )


                        if res[0]['recommender_level']==1:
                            pass
                        elif res[0]['recommender_level']>=1:

                            # Previous record
                            pre = models.FormLevelRecommenderHierarchy.objects.filter(
                            form_id = form_id,
                            recommender_level = res[0]['recommender_level']-1
                            ).update(recommender_level_status=None)

                            # Current record.
                            cur = models.FormLevelRecommenderHierarchy.objects.filter(
                            form_id = form_id,
                            recommender_id = request.user.id
                            ).update(recommender_level_status=2)

                        else:
                            pass

                    else:
                        # Disaaprove from approve level.
                        pass

                    #print(form_rec_level_count,"# Project Level",form_rec_level_approve_count)
                    return Response(
                        {
                            "status": error.context["success_code"],
                            "message": "Returned back to previous level successfully",
                        },
                        status=status.HTTP_200_OK,
                    )


                    #print(request.data["status"],"GGGGGGg$$$$")
                    # approvalData = (
                    #     accessModels.ProcessFlow.objects.values()
                    #     .filter(level__lt=int(request.data["approved_level"]))
                    #     .order_by("-id")
                    #     .first()
                    # )
                    # if approvalData:
                    #     # trials = TrialListSerializer(
                    #     #     models.Trials.objects.filter(id=request.data["trial_id"]),
                    #     #     many=True,
                    #     #     context={"request": request},
                    #     # ).data[0]

                    #     globalTrans = cSerializer.ListGlobalTransactionSerializer(
                    #         models.GlobalTransaction.objects.filter(id=request.data["transaction_id"]),
                    #         many=True,
                    #         context={"request": request},
                    #     ).data[0]

                    #     formDet = models.Forms.objects.values('id','name').filter(id=request.data["form_id"]).first()
                    #     projectDet = models.GlobalTransaction.objects.values('id','project__name','created_by__first_name','created_by__last_name').filter(id=request.data["transaction_id"]).first()

                    #     # Previous level approval notification
                    #     if approvalData["level"] == 1 or approvalData["level"] == 4:
                    #         notificationMessage = (
                    #             "Project "+projectDet['project__name']
                    #             + " request "
                    #             + projectDet['created_by__first_name']+" "+projectDet['created_by__last_name']+" in "+formDet['name']
                    #             + " is returned from recommender "
                    #         )
                    #     elif approvalData["level"] == 2 or approvalData["level"] == 5:
                    #         notificationMessage = (
                    #             "Project "+projectDet['project__name']
                    #             + " requestion "
                    #             + projectDet['created_by__first_name']+" "+projectDet['created_by__last_name']+" in "+formDet['name']
                    #             + " is returned from approver"
                    #         )
                    #     elif approvalData["level"] == 3:
                    #         notificationMessage = (
                    #             "Project "+projectDet['project__name']
                    #             + " requestion "
                    #             + projectDet['created_by__first_name']+" "+projectDet['created_by__last_name']+" in "+formDet['name']
                    #             + " is returned from initiater of trial unit"
                    #         )

                    #     if approvalData["level"]==1:
                    #         hierarchy_type = 1
                    #     elif approvalData["level"]==2:
                    #         hierarchy_type = 2

                    #     notificationModels.NotificationUser.objects.create(
                    #         form_id = globalTrans['form'],
                    #         transaction_id = globalTrans["id"],
                    #         message = notificationMessage,
                    #         user_role_id = approvalData["user_role_id"],
                    #         process_id = approvalData["process_id"],
                    #         hierarchy_type = hierarchy_type
                    #     )

                    #     print(globalTrans['form'], globalTrans["id"],notificationMessage,approvalData["user_role_id"],approvalData["process_id"],"||||||||||||")

                    #     # models.trialStatus.objects.create(
                    #     #     trial_id=trials["id"],
                    #     #     process_flow_id=approvalData["id"],
                    #     #     created_by_id=request.user.id,
                    #     #     created_ip=Common.get_client_ip(request),
                    #     # )

                    #     models.GlobalTransactionApproval.objects.create(
                    #         approved_level=request.data["approved_level"],
                    #         approved_ip=Common.get_client_ip(request),
                    #         approved_by_id=request.user.id,
                    #         transaction_id=request.data["transaction_id"],
                    #         comments=request.data["comments"],
                    #         status=request.data["status"],
                    #         approved_role_id=request.data["approved_role_id"],
                    #         type=request.data["type"],
                    #     )

                    #     models.GlobalTransaction.objects.filter(
                    #         id = request.data["transaction_id"]
                    #     ).update(approved_level=approvalData["level"])

                    # return Response(
                    #     {
                    #         "status": error.context["success_code"],
                    #         "message": "Returned back to previous level successfully",
                    #     },
                    #     status=status.HTTP_200_OK,
                    # )

            except:
                return Response(
                    {
                        "status": error.context["error_code"],
                        "message": "Failed to perform this action",
                    },
                    status=status.HTTP_200_OK,
                )


# GlobalApproval
class GlobalApproval(APIView):
    def post(self, request, pk=None):
        if "approved_level" not in request.data:
            return Response(
                {
                    "status": error.context["error_code"],
                    "message": "Level"
                    + language.context[language.defaultLang]["missing"],
                },
                status=status.HTTP_200_OK,
            )
        elif "form_id" not in request.data:
            return Response(
                {
                    "status": error.context["error_code"],
                    "message": "Form ID"
                    + language.context[language.defaultLang]["missing"],
                },
                status=status.HTTP_200_OK,
            )
        elif "comments" not in request.data:
            return Response(
                {
                    "status": error.context["error_code"],
                    "message": "Comment"
                    + language.context[language.defaultLang]["missing"],
                },
                status=status.HTTP_200_OK,
            )
        elif "status" not in request.data:
            return Response(
                {
                    "status": error.context["error_code"],
                    "message": "Status"
                    + language.context[language.defaultLang]["missing"],
                },
                status=status.HTTP_200_OK,
            )
        elif "approved_role_id" not in request.data:
            return Response(
                {
                    "status": error.context["error_code"],
                    "message": "Approved Role ID"
                    + language.context[language.defaultLang]["missing"],
                },
                status=status.HTTP_200_OK,
            )
        else:
        #try:
            if int(request.data["approved_level"]) == 2:
                models.GlobalTransaction.objects.filter(id=request.data["transaction_id"]).update(
                    recommender=request.data["status"]
                )
            elif int(request.data["approved_level"]) == 3:
                models.GlobalTransaction.objects.filter(id=request.data["transaction_id"]).update(
                    approver=request.data["status"]
                )

            
            form_id = request.data["form_id"]
            project_id = request.data["project_id"]
            type = request.data["type"]

            # Check the hierarchy type (form level or project level).
            if 'project_id' not in request.data or request.data["project_id"]=="": # Form Level
                rec_level = 1
            else: # Project Level                    
                project_id = request.data["project_id"]
                rec_level = 2

            # Notifiation message.
            formDet = models.Forms.objects.values('id','name').filter(id=request.data["form_id"]).first()
            projectDet = models.GlobalTransaction.objects.values('id','project__name','created_by__first_name','created_by__last_name').filter(id=request.data["transaction_id"]).first()

            userInfo = models.FormLevelRecommenderHierarchy.objects.values('id','recommender_id__first_name','recommender_id__last_name','recommender_level').filter(
                        form_id = form_id,
                        recommender_id = request.user.id
                        ).first()



            if request.data["status"] == 1:


                if rec_level==2: # Project level


                    curUserInfo = models.ProjectLevelRecommenderHierarchy.objects.values('id','recommender_id','recommender_id__first_name','recommender_id__last_name','recommender_level').filter(form_id = form_id, project_id = project_id, recommender_id = request.user.id).first()

                    if type==1: # Initiator

                        # Get first recommender info.
                        first_level_rec = models.ProjectLevelRecommenderHierarchy.objects.values('recommender_id','recommender_level').filter(
                        form_id = form_id,
                        project_id = project_id,
                        recommender_level = 1
                        )

                        # First recommender update
                        first_rec = models.ProjectLevelRecommenderHierarchy.objects.filter(
                        form_id = form_id,
                        project_id = project_id,
                        recommender_level = 1
                        ).update(recommender_level_status=None, current_reject=None)

                        notificationMessage = (
                            "Project "+projectDet['project__name']
                            + " request "
                            + request.user.first_name+" "+request.user.last_name+" in "+formDet['name']
                            + " is reinitiated from initiater"
                            )

                        notificationModels.NotificationUser.objects.create(
                            form_id = request.data["form_id"],
                            transaction_id = request.data["transaction_id"],
                            project_id = request.data["project_id"],
                            message = notificationMessage,
                            user_role_id = 2,
                            process_id = 1,
                            hierarchy_type = 2, # Recommender
                            from_user_id = request.user.id,
                            to_user_id = first_level_rec[0]['recommender_id'],
                        )

                        models.GlobalTransactionApproval.objects.create(
                            approved_level = request.data["approved_level"],
                            approved_ip = Common.get_client_ip(request),
                            approved_by_id = request.user.id,
                            transaction_id = request.data["transaction_id"],
                            comments = request.data["comments"],
                            status = request.data["status"],
                            approved_role_id = request.data["approved_role_id"],
                            #type = request.data["type"],
                            type = 1,
                        )

                        return Response(
                            {
                                "status": error.context["success_code"],
                                "message": "initiator reinitiated successfully",
                            },
                            status=status.HTTP_200_OK,
                        )

                    if type==2: # Recommendation

                        rec_count = models.ProjectLevelRecommenderHierarchy.objects.filter(form_id = form_id, project_id = project_id).count()


                        print(rec_count, "rec_count")
                        if rec_count==1:
                            print(rec_count, "rec_countAA")
                            # Current record update.
                            models.ProjectLevelRecommenderHierarchy.objects.filter(
                            form_id = form_id,
                            project_id = project_id,
                            recommender_id = request.user.id
                            ).update(recommender_level_status=1, current_reject=None)


                            # Current user info
                            # curUserInfo = models.ProjectLevelRecommenderHierarchy.objects.values('id','recommender_id','recommender_id__first_name','recommender_id__last_name','recommender_level').filter(form_id = form_id, project_id = project_id).first()


                            # Next update the approver.
                            models.ProjectLevelApproverHierarchy.objects.filter(
                            form_id = form_id,
                            project_id = project_id,
                            approver_level = 1
                            ).update(approver_level_status=-1)


                            # Next approver user info
                            nextApproverInfo = models.ProjectLevelApproverHierarchy.objects.values('id','approver_id','approver_level').filter(form_id = form_id, project_id = project_id).first()

                            # Update the global transaction table.
                            models.GlobalTransaction.objects.filter(id = request.data["transaction_id"]).update(approved_level=3)

                            notificationMessage = (
                                "Project "+projectDet['project__name']
                                + " is recommended by "
                                + curUserInfo['recommender_id__first_name']+" "+curUserInfo['recommender_id__last_name']+" in "+formDet['name']
                                + " is waiting for your approval"
                            )


                            res1 = notificationModels.NotificationUser.objects.create(
                                form_id = request.data["form_id"],
                                transaction_id = request.data["transaction_id"],
                                project_id = request.data["project_id"],
                                message = notificationMessage,
                                user_role_id = 3,
                                process_id = 1,
                                hierarchy_type = 3,
                                from_user_id = request.user.id,
                                to_user_id = nextApproverInfo['approver_id']
                            )

                            res2 = models.GlobalTransactionApproval.objects.create(
                                approved_level = request.data["approved_level"],
                                approved_ip = Common.get_client_ip(request),
                                approved_by_id = request.user.id,
                                transaction_id = request.data["transaction_id"],
                                comments = request.data["comments"],
                                status = request.data["status"],
                                approved_role_id = request.data["approved_role_id"],
                                type = request.data["type"]
                            )

                            return Response(
                                {
                                    "status": error.context["success_code"],
                                    "message": "Recommendation accepted successfully",
                                },
                                status=status.HTTP_200_OK,
                            )


                        if rec_count>1:

                            #print(rec_count, "rec_countBB")
                            # Current recommender.
                            current_level_rec = models.ProjectLevelRecommenderHierarchy.objects.values('recommender_id','recommender_level').filter(
                            form_id = form_id,
                            recommender_id = request.user.id
                            )

                            # Current record update.
                            models.ProjectLevelRecommenderHierarchy.objects.filter(
                            form_id = form_id,
                            recommender_id = request.user.id
                            ).update(recommender_level_status=1, current_reject=None)

                            # Next recommender.
                            next_level_rec = models.ProjectLevelRecommenderHierarchy.objects.values('recommender_id','recommender_level').filter(
                            form_id = form_id,
                            recommender_level = current_level_rec[0]['recommender_level']+1
                            )

                            if rec_count==current_level_rec[0]['recommender_level']:

                                models.GlobalTransaction.objects.filter(id = request.data["transaction_id"]).update(approved_level=3)

                                userInfoApprover = models.ProjectLevelApproverHierarchy.objects.values('id','approver_id','approver_id__first_name','approver_id__last_name','approver_level').filter(form_id = form_id, project_id = project_id).first()

                                current_rec_count = 1
                            else:
                                current_rec_count = 2


                            # Next record update.
                            if current_rec_count==2:

                                models.ProjectLevelRecommenderHierarchy.objects.filter(
                                form_id = form_id,
                                recommender_level = current_level_rec[0]['recommender_level']+1
                                ).update(recommender_level_status=None)
                            
                            elif current_rec_count==1:
                            
                                #pass
                                models.ProjectLevelApproverHierarchy.objects.filter(
                                form_id = form_id,
                                project_id = project_id,
                                approver_level = 1
                                ).update(approver_level_status=-1)


                            # Current user info
                            # curUserInfo = models.ProjectLevelRecommenderHierarchy.objects.values('id','recommender_id','recommender_id__first_name','recommender_id__last_name','recommender_level').filter(form_id = form_id, project_id = project_id, recommender_id = request.user.id).first()


                            if current_rec_count==2:
                                notificationMessage = (
                                    "Project "+projectDet['project__name']
                                    + " is recommended by "
                                    + curUserInfo['recommender_id__first_name']+" "+curUserInfo['recommender_id__last_name']+" in "+formDet['name']
                                    + " is waiting for your next level recommendations "
                                )
                            elif current_rec_count==1:
                                notificationMessage = (
                                    "Project "+projectDet['project__name']
                                    + " is recommended by "
                                    + curUserInfo['recommender_id__first_name']+" "+curUserInfo['recommender_id__last_name']+" in "+formDet['name']
                                    + " is waiting for your approval"
                                )


                            res1 = notificationModels.NotificationUser.objects.create(
                                form_id = request.data["form_id"],
                                transaction_id = request.data["transaction_id"],
                                project_id = request.data["project_id"],
                                message = notificationMessage,
                                user_role_id = 3 if current_rec_count==2 else 4,
                                process_id = 1,
                                hierarchy_type = 2 if current_rec_count==2 else 3,
                                from_user_id = request.user.id,
                                to_user_id = next_level_rec[0]['recommender_id'] if current_rec_count==2 else userInfoApprover['approver_id'],
                            )

                            res2 = models.GlobalTransactionApproval.objects.create(
                                approved_level = request.data["approved_level"] if current_rec_count==2 else 3,
                                approved_ip = Common.get_client_ip(request),
                                approved_by_id = request.user.id,
                                transaction_id = request.data["transaction_id"],
                                comments = request.data["comments"],
                                status = request.data["status"],
                                approved_role_id = request.data["approved_role_id"],
                                type = request.data["type"] if current_rec_count==2 else 3
                            )

                            return Response(
                                {
                                    "status": error.context["success_code"],
                                    "message": "Recommendation accepted successfully",
                                },
                                status=status.HTTP_200_OK,
                            )


                    elif type==3: # Approver
                        # Approver.
                        approver = models.ProjectLevelApproverHierarchy.objects.filter(
                        form_id = form_id,
                        project_id = project_id,
                        approver_id = request.user.id
                        ).update(approver_level_status=1)

                        res2 = models.GlobalTransactionApproval.objects.create(
                            approved_level = request.data["approved_level"],
                            approved_ip = Common.get_client_ip(request),
                            approved_by_id = request.user.id,
                            transaction_id = request.data["transaction_id"],
                            comments = request.data["comments"],
                            status = request.data["status"],
                            approved_role_id = request.data["approved_role_id"],
                            type = request.data["type"]
                        )

                        models.GlobalTransaction.objects.filter(id=request.data["transaction_id"]).update(
                        approved_level=4)

                        
                        # Notification
                        userInfoApprover = models.ProjectLevelApproverHierarchy.objects.values('id','approver_id','approver_id__first_name','approver_id__last_name','approver_level').filter(form_id = form_id, project_id = project_id, approver_id = request.user.id).first()

                        allRecommender = models.ProjectLevelRecommenderHierarchy.objects.values('id','recommender_id').filter(form_id = form_id, project_id = project_id)

                        # print(allRecommender,"allRecommenderallRecommenderallRecommender")
                        # print(allRecommender.query,"allRecommenderallRecommenderallRecommender")

                        # Send notification to all the recommender.
                        for access in allRecommender:

                            notificationMessage = (
                                "Project "+projectDet['project__name']
                                + " request "
                                + userInfoApprover['approver_id__first_name']+" "+userInfoApprover['approver_id__last_name']+" in "+formDet['name']
                                + " is approved from approver successfully"
                            )

                            res1 = notificationModels.NotificationUser.objects.create(
                                form_id = request.data["form_id"],
                                transaction_id = request.data["transaction_id"],
                                project_id = request.data["project_id"],
                                message = notificationMessage,
                                user_role_id = 3,
                                process_id = 1,
                                hierarchy_type = 2,
                                from_user_id = request.user.id,
                                to_user_id = access['recommender_id'],
                            )

                        # Notification to Initiator
                        userInfoInitiator = notificationModels.NotificationUser.objects.values('id','from_user_id').filter(form_id = form_id, project_id = project_id, user_role_id=2).first()

                        initNotificationMessage = (
                            "Project "+projectDet['project__name']
                            + " request "
                            + userInfoApprover['approver_id__first_name']+" "+userInfoApprover['approver_id__last_name']+" in "+formDet['name']
                            + " is approved from approver successfully"
                        )

                        res1 = notificationModels.NotificationUser.objects.create(
                            form_id = request.data["form_id"],
                            transaction_id = request.data["transaction_id"],
                            project_id = request.data["project_id"],
                            message = initNotificationMessage,
                            user_role_id = 3,
                            process_id = 1,
                            hierarchy_type = 1,  # Initiator
                            from_user_id = request.user.id,
                            to_user_id = userInfoInitiator['from_user_id'],
                        )


                        return Response(
                            {
                                "status": error.context["success_code"],
                                "message": "Recommendation accepted successfully",
                            },
                            status=status.HTTP_200_OK,
                        )
                    

            elif request.data["status"] == 2:

                # Check the hierarchy type (form level or project level).
                # Get the current logged in user id.

                formDet = models.Forms.objects.values('id','name').filter(id=request.data["form_id"]).first()
                projectDet = models.GlobalTransaction.objects.values('id','project__name','created_by__first_name','created_by__last_name').filter(id=request.data["transaction_id"]).first()

                if rec_level==2: # Project level

                    # Get the current level and user.
                    current_level_rec = models.ProjectLevelRecommenderHierarchy.objects.values('recommender_id','recommender_level').filter(
                    form_id = form_id,
                    project_id = project_id,
                    recommender_id = request.user.id
                    )

                    # Get Current User Info
                    curUserInfo = models.ProjectLevelRecommenderHierarchy.objects.values('id','recommender_id','recommender_id__first_name','recommender_id__last_name','recommender_level').filter(form_id = form_id, project_id = project_id, recommender_id = request.user.id).first()

                    if request.data["type"]==2: # Reject from recommender.

                        if current_level_rec[0]['recommender_level']==1: # Initiator

                            # Get Initiator details

                            first_rec = models.ProjectLevelRecommenderHierarchy.objects.filter(
                            form_id = form_id,
                            project_id = project_id,
                            recommender_level = 1
                            ).update(recommender_level_status=2, current_reject=1)

                            # # Get Current User Info
                            # curUserInfo = models.ProjectLevelRecommenderHierarchy.objects.values('id','recommender_id','recommender_id__first_name','recommender_id__last_name','recommender_level').filter(form_id = form_id, project_id = project_id, recommender_id = request.user.id).first()


                            notificationMessage = (
                                "Project "+projectDet['project__name']
                                + " request "
                                + curUserInfo['recommender_id__first_name']+" "+curUserInfo['recommender_id__last_name']+" in "+formDet['name']
                                + " is returned from recommender"
                                )

                            # Notification to Initiator
                            userInfoInitiator = notificationModels.NotificationUser.objects.values('id','from_user_id','from_user_id__first_name','from_user_id__last_name').filter(form_id = form_id, project_id = project_id, user_role_id=2).first()


                            notificationModels.NotificationUser.objects.create(
                                form_id = request.data["form_id"],
                                transaction_id = request.data["transaction_id"],
                                project_id = request.data["project_id"],
                                message = notificationMessage,
                                user_role_id = 2,
                                process_id = 1,
                                hierarchy_type = 1, # Initiator
                                from_user_id = request.user.id,
                                to_user_id = userInfoInitiator['from_user_id'],
                            )

                            models.GlobalTransactionApproval.objects.create(
                                approved_level = request.data["approved_level"],
                                approved_ip = Common.get_client_ip(request),
                                approved_by_id = request.user.id,
                                transaction_id = request.data["transaction_id"],
                                comments = request.data["comments"],
                                status = request.data["status"],
                                approved_role_id = request.data["approved_role_id"],
                                #type = request.data["type"],
                                type = 1,
                            )

                        elif current_level_rec[0]['recommender_level']>=1: # Previous recommender

                            # Previous recommender.
                            pre_level_rec = models.ProjectLevelRecommenderHierarchy.objects.values('recommender_id','recommender_level').filter(
                            form_id = form_id,
                            project_id = project_id,
                            recommender_level = current_level_rec[0]['recommender_level']-1
                            )

                            # Previous record update.
                            pre_rec = models.ProjectLevelRecommenderHierarchy.objects.filter(
                            form_id = form_id,
                            project_id = project_id,
                            recommender_level = current_level_rec[0]['recommender_level']-1
                            ).update(recommender_level_status=None, current_reject=1)

                            # Current record update.
                            cur_rec = models.ProjectLevelRecommenderHierarchy.objects.filter(
                            form_id = form_id,
                            project_id = project_id,
                            recommender_id = request.user.id
                            ).update(recommender_level_status=2)

                            # # Current user info
                            # curUserInfo = models.ProjectLevelRecommenderHierarchy.objects.values('id','recommender_id','recommender_id__first_name','recommender_id__last_name','recommender_level').filter(form_id = form_id, project_id = project_id, recommender_id = request.user.id).first()


                            notificationMessage = (
                                "Project "+projectDet['project__name']
                                + " request "
                                + curUserInfo['recommender_id__first_name']+" "+curUserInfo['recommender_id__last_name']+" in "+formDet['name']
                                + " is returned from recommender"
                            )

                            res1 = notificationModels.NotificationUser.objects.create(
                                form_id = request.data["form_id"],
                                transaction_id = request.data["transaction_id"],
                                project_id = request.data["project_id"],
                                message = notificationMessage,
                                user_role_id = 3,
                                process_id = 1,
                                hierarchy_type = 2,
                                from_user_id = request.user.id,
                                to_user_id = pre_level_rec[0]['recommender_id'],
                            )

                            res2 = models.GlobalTransactionApproval.objects.create(
                                approved_level = request.data["approved_level"],
                                approved_ip = Common.get_client_ip(request),
                                approved_by_id = request.user.id,
                                transaction_id = request.data["transaction_id"],
                                comments = request.data["comments"],
                                status = request.data["status"],
                                approved_role_id = request.data["approved_role_id"],
                                type = request.data["type"]
                            )

                    if request.data["type"]==3: # Reject from approver.

                        last_rec = models.ProjectLevelRecommenderHierarchy.objects.values('recommender_id','recommender_level').order_by('-recommender_level').filter(form_id = form_id, project_id = project_id).first()


                        userInfoApprover = models.ProjectLevelApproverHierarchy.objects.values('id','approver_id','approver_id__first_name','approver_id__last_name','approver_level').filter(form_id = form_id, project_id = project_id, approver_id = request.user.id).first()
                        #print(last_rec,"last_rec", last_rec['recommender_level'])

                        models.ProjectLevelRecommenderHierarchy.objects.filter(
                        form_id = form_id,
                        project_id = project_id,
                        recommender_id = last_rec['recommender_id'],
                        recommender_level = last_rec['recommender_level'],
                        ).update(recommender_level_status=None, current_reject = 1)

                        # Approver.
                        approver = models.ProjectLevelApproverHierarchy.objects.filter(
                        form_id = form_id,
                        project_id = project_id,
                        approver_id = request.user.id
                        ).update(approver_level_status=2)


                        notificationMessage = (
                            "Project "+projectDet['project__name']
                            + " request "
                            + userInfoApprover['approver_id__first_name']+" "+userInfoApprover['approver_id__last_name']+" in "+formDet['name']
                            + " is returned from approver"
                        )

                        res1 = notificationModels.NotificationUser.objects.create(
                            form_id = request.data["form_id"],
                            transaction_id = request.data["transaction_id"],
                            project_id = request.data["project_id"],
                            message = notificationMessage,
                            user_role_id = 3,
                            process_id = 1,
                            hierarchy_type = 2,
                            from_user_id = request.user.id,
                            to_user_id = last_rec['recommender_id'],
                        )

                        res2 = models.GlobalTransactionApproval.objects.create(
                            approved_level = request.data["approved_level"],
                            approved_ip = Common.get_client_ip(request),
                            approved_by_id = request.user.id,
                            transaction_id = request.data["transaction_id"],
                            comments = request.data["comments"],
                            status = request.data["status"],
                            approved_role_id = request.data["approved_role_id"],
                            type = request.data["type"]
                        )

                else:
                    # Disaaprove from approve level.
                    pass


                #print(form_rec_level_count,"# Project Level",form_rec_level_approve_count)
                return Response(
                    {
                        "status": error.context["success_code"],
                        "message": "Returned back to previous level successfully",
                    },
                    status=status.HTTP_200_OK,
                )

            # except:
            #     return Response(
            #         {
            #             "status": error.context["error_code"],
            #             "message": "Failed to perform this action",
            #         },
            #         status=status.HTTP_200_OK,
            #     )



class ProjectLog(APIView):

    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['msg']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try: 

            if pk:
                list = models.ProjectLog.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.ProjectLogSerializer(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.ProjectLog.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"ProjectLog" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.ProjectLog.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        serializer = cSerializer.ProjectLogSerializer(lists, many=True)
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)



class saveSystemEquipmentCompartmentTemp(APIView):


    def post(self,request, pk = None):
        if request.data['system_equ_comp']:

            #print(request.data['system_equ_comp'])
            created_ip = Common.get_client_ip(request)
            created_by = request.user.id
            p_type = request.data['p_type']

            project_id = request.data['system_equ_comp'][0]['project_id']
            section_id = request.data['system_equ_comp'][0]['global_section_id']
            sub_section_id = request.data['system_equ_comp'][0]['global_sub_section_id']
            sub_sub_section_id = request.data['system_equ_comp'][0]['global_sub_sub_section_id']

            models.SystemEquipmentCompartmentTemp.objects.filter(project_id = project_id, section_id = section_id, sub_section_id__isnull=True, sub_sub_section_id__isnull=True, type=p_type).delete()

            models.SystemEquipmentCompartmentTemp.objects.filter(project_id = project_id, section_id = section_id, sub_section_id = sub_section_id, sub_sub_section_id__isnull=True, type=p_type).delete()

            models.SystemEquipmentCompartmentTemp.objects.filter(project_id = project_id, section_id = section_id, sub_section_id = sub_section_id, sub_sub_section_id=sub_sub_section_id, type=p_type).delete()

            for item in (request.data['system_equ_comp']):

                models.SystemEquipmentCompartmentTemp.objects.create(
                project_id = item['project_id'], 
                section_id = item['global_section_id'],
                sub_section_id = item['global_sub_section_id'] if 'global_sub_section_id' in item else '',
                sub_sub_section_id = item['global_sub_sub_section_id'] if 'global_sub_sub_section_id' in item else '',
                type = p_type,
                ser = item['s_ser'] if 's_ser' in item else '',
                name = item['s_name'] if 's_name' in item else '',
                numbers = item['s_numbers'] if 's_numbers' in item else '',
                capabilities_feature = item['s_capabilities_feature'] if 's_capabilities_feature' in item else '',
                weight_volume_power_consumption = item['s_weight_volume_power_consumption'] if 's_weight_volume_power_consumption' in item else '',
                location = item['s_location'] if 's_location' in item else '',
                interface = item['s_interface'] if 's_interface' in item else '',
                procurement_router = item['s_procurement_router'] if 's_procurement_router' in item else '',
                vendor = item['s_vendor'] if 's_vendor' in item else '',
                cost = item['s_cost'] if 's_cost' in item else '',
                standards = item['s_standards'] if 's_standards' in item else '',
                sustenance = item['s_sustenance'] if 's_sustenance' in item else '',
                equipment = item['s_equipment'] if 's_equipment' in item else '',
                features = item['s_features'] if 's_features' in item else '',
                layout = item['s_layout'] if 's_layout' in item else '',
                special_requirements = item['s_special_requirements'] if 's_special_requirements' in item else '',
                created_ip = created_ip,
                created_by_id = created_by
                )

            return Response({"status":"Ok"})

class getModulewiseTimeline(APIView):
    def post(self, request, pk=None):
        from django.db.models import F,ExpressionWrapper
        from django.db.models.functions import ExtractMonth,ExtractYear
        if "project_id" not in request.data or request.data['project_id']=='':
           return Response({"status": error.context["error_code"],"message": "Project id is missing"},status=status.HTTP_200_OK)
        elif "module_id" not in request.data or request.data['module_id']=='':
           return Response({"status": error.context["error_code"],"message": "Module id is missing"},status=status.HTTP_200_OK)
        else:
            request.data['project_id']=int(request.data['project_id'])
            endResponse=[]
            submodules=masterModels.SubModule.objects.values("id","name","targetted_month_moderate","targetted_month_complex").filter(module_id=request.data['module_id'])
            for submodule in submodules:
                sProjDet=models.GlobalTransaction.objects.values('id','project__project_type__code','created_on','id','approved_status','approved_on').filter(project__id=request.data['project_id'],form__id=submodule['id']).annotate(months=(ExtractYear("approved_on") - ExtractYear("created_on")) * 12 + (ExtractMonth("approved_on") - ExtractMonth("created_on"))).first()

                completionStatus="On Time"
                ProjectType=sProjDet['project__project_type__code'] if sProjDet else "MODERATE"
                tentatvieMonths=0
                if ProjectType=='COMPLEX':
                    tentatvieMonths=int(submodule['targetted_month_complex']) if submodule['targetted_month_complex'] else 0
                elif ProjectType=='MODERATE':
                    tentatvieMonths=int(submodule['targetted_month_moderate']) if submodule['targetted_month_moderate'] else 0

                completionStatus=(str(tentatvieMonths-int(sProjDet['months']))+' months Delay' if sProjDet['months']>tentatvieMonths else "On Time") if sProjDet  is not None and tentatvieMonths is not None else "Yet to complete"
                warningStatus=('bg-danger' if sProjDet['months']>tentatvieMonths else 'bg-success') if sProjDet  is not None and tentatvieMonths is not None else 'bg-danger'



                endResponse.append({"submodule":submodule,"name":submodule['name'],'targetted_month_moderate':submodule['targetted_month_moderate'],'targetted_month_complex':submodule['targetted_month_complex'],'sProjDet':sProjDet,"completionStatus":completionStatus,"warningStatus":warningStatus})

            return Response({"status": error.context["success_code"],"message": "Test","endResponse":endResponse},status=status.HTTP_200_OK)



### ----- ###

class AllFormsView(APIView):

    def get(self, request, pk = None):
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = ['code','name']
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  

            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type') 
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')      

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:

            if pk:
                list = models.Forms.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = cSerializer.ListFormsSerializer(list)
                return Response({"status": error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.Forms.DoesNotExist:
            return Response({"status" : error.context['error_code'], "message":"Section" +language.context[language.defaultLang]['dataNotFound'] }, status = status.HTTP_200_OK)

        lists = models.Forms.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]          
    
        serializer = cSerializer.AllFormsListSerializer(lists, many=True)
        return Response({"status":error.context['success_code'] , "data": serializer.data}, status=status.HTTP_200_OK)



class GetFormsView(APIView):

    def get(self, request, pk = None):

        #if form_id in request.GET:
        form_id = request.GET.get('form_id')         
        #lists = models.FormsMapping.objects.values('id', 'form_id','Class_id','module_id','sub_module_id','section_id','sub_section_id','sub_sub_section_id','order').filter(form_id=form_id).exclude(status='3', order=None).order_by('sub_module_id','section_id','sub_section_id','sub_sub_section_id','order')
        #print(lists,"####s")
        #serializer = cSerializer.AllFormsMappingSerializer(lists, many=True)


        lists = models.FormsMapping.objects.values('id', 'form_id','Class_id','module_id','sub_module_id','section_id','sub_section_id','sub_sub_section_id','order').filter(form_id=form_id).exclude(status='3', order=None).order_by('sub_module_id', 'order')

        return Response({"status":error.context['success_code'] , "data": lists}, status=status.HTTP_200_OK)



class getDashBoardModule1(APIView):

    def get(self,request, pk = None):

        dashboard = []
        lists = models.GlobalTransaction.objects.values('form_id','approved_status','project_id').filter(form_id=7, approved_status=2).exclude(status='3').distinct('project_id')

        #print(len(lists),"Length")

        for item in lists:

            #print(item['project_id'],"ProID")
            start = models.GlobalTransaction.objects.values('form_id','approved_status','project_id','created_on','approved_on','project__name').filter(form_id=1, approved_status=2, project_id=item['project_id']).exclude(status='3')

            end = models.GlobalTransaction.objects.values('form_id','approved_status','project_id','created_on','approved_on','project__name').filter(form_id=7, approved_status=2, project_id=item['project_id']).exclude(status='3')

            (dashboard).append({
                "project_id":start[0]['project_id'],
                "project":start[0]['project__name'],
                'start': start[0]['created_on'],
                'start_year':int(start[0]['created_on'].year),
                'start_month':int(start[0]['created_on'].month),
                'start_day':int(start[0]['created_on'].day),
                'end':end[0]['approved_on'],
                'end_year':int(end[0]['approved_on'].year),
                'end_month':int(end[0]['approved_on'].month),
                'end_day':int(end[0]['approved_on'].day),
                })

        return Response({"status":error.context['success_code'], "data": dashboard}, status=status.HTTP_200_OK)



class getDashBoardModule(APIView):

    def get(self,request, pk = None):
        # import pytz
        # from django.utils import timezone
        # timezone.activate(pytz.timezone(user_time_zone))

        response = {}
        psr = []
        gls = []
        bs = []
        allPSR = models.GlobalTransaction.objects.values('form_id','approved_status','project_id').filter(form_id=7, approved_status=2).exclude(status='3').distinct('project_id')

        allGLS = models.GlobalTransaction.objects.values('form_id','approved_status','project_id').filter(form_id=8, approved_status=2).exclude(status='3').distinct('project_id')

        allBS = models.GlobalTransaction.objects.values('form_id','approved_status','project_id').filter(form_id=9, approved_status=2).exclude(status='3').distinct('project_id')

        #print(len(lists),"Length")

        for item in allPSR:

            #print(item['project_id'],"ProID")
            start = models.GlobalTransaction.objects.values('id','form_id','approved_status','project_id','created_on','approved_on','project__name').filter(form_id=1, approved_status=2, project_id=item['project_id']).exclude(status='3')

            end = models.GlobalTransaction.objects.values('id','form_id','approved_status','project_id','created_on','approved_on','project__name').filter(form_id=7, approved_status=2, project_id=item['project_id']).exclude(status='3')

            (psr).append({
                'start_id':start[0]['id'],
                "project_id":start[0]['project_id'],
                "project":start[0]['project__name'],
                'start': start[0]['created_on'],
                'start_year':int(start[0]['created_on'].year),
                'start_month':int(start[0]['created_on'].month),
                'start_day':int(start[0]['created_on'].day),
                'end_id':end[0]['id'],
                'end':end[0]['approved_on'],
                'end_year':int(end[0]['approved_on'].year),
                'end_month':int(end[0]['approved_on'].month),
                'end_day':int(end[0]['approved_on'].day),
                })

            response['psr'] = psr


        for item in allGLS:

            result = models.GlobalTransaction.objects.values('id','form_id','approved_status','project_id','created_on','approved_on','project__name').filter(form_id=8, approved_status=2, project_id=item['project_id']).exclude(status='3')

            #print(result.query,"HHHHHHHHhh")

            (gls).append({
                "id":result[0]['id'],
                "project_id":result[0]['project_id'],
                "project":result[0]['project__name'],
                'start': result[0]['created_on'],
                'start_year':result[0]['created_on'].year,
                'start_month':result[0]['created_on'].month,
                'start_day':result[0]['created_on'].day,
                'end':result[0]['approved_on'],
                'end_year':result[0]['approved_on'].year,
                'end_month':result[0]['approved_on'].month,
                'end_day':result[0]['approved_on'].day,
                })

            response['gls'] = gls


        for item in allBS:

            result = models.GlobalTransaction.objects.values('form_id','approved_status','project_id','created_on','approved_on','project__name').filter(form_id=9, approved_status=2, project_id=item['project_id']).exclude(status='3')

            (bs).append({
                "project_id":result[0]['project_id'],
                "project":result[0]['project__name'],
                'start': result[0]['created_on'],
                'start_year':int(result[0]['created_on'].year),
                'start_month':int(result[0]['created_on'].month),
                'start_day':int(result[0]['created_on'].day),
                'end':result[0]['approved_on'],
                'end_year':int(result[0]['approved_on'].year),
                'end_month':int(result[0]['approved_on'].month),
                'end_day':int(result[0]['approved_on'].day),
                })

            response['bs'] = bs

        return Response({"status":error.context['success_code'], "data": response}, status=status.HTTP_200_OK)


class getDashBoardModulePSR(APIView):

    def get(self,request, pk = None):

        response = {}
        initiation = []
        formulation = []
        presentation = []
        input_sr = []
        concept = []
        incorporation = []
        rfi = []

        form1 = models.GlobalTransaction.objects.values('form_id','approved_status','project_id').filter(form_id=1, approved_status=2).exclude(status='3').distinct('project_id')

        form2 = models.GlobalTransaction.objects.values('form_id','approved_status','project_id').filter(form_id=2, approved_status=2).exclude(status='3').distinct('project_id')

        form3 = models.GlobalTransaction.objects.values('form_id','approved_status','project_id').filter(form_id=3, approved_status=2).exclude(status='3').distinct('project_id')

        form4 = models.GlobalTransaction.objects.values('form_id','approved_status','project_id').filter(form_id=4, approved_status=2).exclude(status='3').distinct('project_id')

        form5 = models.GlobalTransaction.objects.values('form_id','approved_status','project_id').filter(form_id=5, approved_status=2).exclude(status='3').distinct('project_id')

        form6 = models.GlobalTransaction.objects.values('form_id','approved_status','project_id').filter(form_id=6, approved_status=2).exclude(status='3').distinct('project_id')

        form7 = models.GlobalTransaction.objects.values('form_id','approved_status','project_id').filter(form_id=7, approved_status=2).exclude(status='3').distinct('project_id')

        for item in form1:
            result = models.GlobalTransaction.objects.values('form_id','approved_status','project_id','created_on','approved_on','project__name').filter(form_id=1, approved_status=2, project_id=item['project_id']).exclude(status='3')

            (initiation).append({
                "form_id":result[0]['form_id'],
                "project_id":result[0]['project_id'],
                "project":result[0]['project__name'],
                'start': result[0]['created_on'],
                'start_year':int(result[0]['created_on'].year),
                'start_month':int(result[0]['created_on'].month),
                'start_day':int(result[0]['created_on'].day),
                'end':result[0]['approved_on'],
                'end_year':int(result[0]['approved_on'].year),
                'end_month':int(result[0]['approved_on'].month),
                'end_day':int(result[0]['approved_on'].day),
                })

            response['initiation'] = initiation

        for item in form2:
            result = models.GlobalTransaction.objects.values('form_id','approved_status','project_id','created_on','approved_on','project__name').filter(form_id=2, approved_status=2, project_id=item['project_id']).exclude(status='3')

            (formulation).append({
                "form_id":result[0]['form_id'],
                "project_id":result[0]['project_id'],
                "project":result[0]['project__name'],
                'start': result[0]['created_on'],
                'start_year':int(result[0]['created_on'].year),
                'start_month':int(result[0]['created_on'].month),
                'start_day':int(result[0]['created_on'].day),
                'end':result[0]['approved_on'],
                'end_year':int(result[0]['approved_on'].year),
                'end_month':int(result[0]['approved_on'].month),
                'end_day':int(result[0]['approved_on'].day),
                })

            response['formulation'] = formulation

        for item in form3:
            result = models.GlobalTransaction.objects.values('form_id','approved_status','project_id','created_on','approved_on','project__name').filter(form_id=3, approved_status=2, project_id=item['project_id']).exclude(status='3')

            (presentation).append({
                "form_id":result[0]['form_id'],
                "project_id":result[0]['project_id'],
                "project":result[0]['project__name'],
                'start': result[0]['created_on'],
                'start_year':int(result[0]['created_on'].year),
                'start_month':int(result[0]['created_on'].month),
                'start_day':int(result[0]['created_on'].day),
                'end':result[0]['approved_on'],
                'end_year':int(result[0]['approved_on'].year),
                'end_month':int(result[0]['approved_on'].month),
                'end_day':int(result[0]['approved_on'].day),
                })

            response['presentation'] = presentation

        for item in form4:
            result = models.GlobalTransaction.objects.values('form_id','approved_status','project_id','created_on','approved_on','project__name').filter(form_id=4, approved_status=2, project_id=item['project_id']).exclude(status='3')

            (input_sr).append({
                "form_id":result[0]['form_id'],
                "project_id":result[0]['project_id'],
                "project":result[0]['project__name'],
                'start': result[0]['created_on'],
                'start_year':int(result[0]['created_on'].year),
                'start_month':int(result[0]['created_on'].month),
                'start_day':int(result[0]['created_on'].day),
                'end':result[0]['approved_on'],
                'end_year':int(result[0]['approved_on'].year),
                'end_month':int(result[0]['approved_on'].month),
                'end_day':int(result[0]['approved_on'].day),
                })

            response['input_sr'] = input_sr

        for item in form5:
            result = models.GlobalTransaction.objects.values('form_id','approved_status','project_id','created_on','approved_on','project__name').filter(form_id=5, approved_status=2, project_id=item['project_id']).exclude(status='3')

            (concept).append({
                "form_id":result[0]['form_id'],
                "project_id":result[0]['project_id'],
                "project":result[0]['project__name'],
                'start': result[0]['created_on'],
                'start_year':int(result[0]['created_on'].year),
                'start_month':int(result[0]['created_on'].month),
                'start_day':int(result[0]['created_on'].day),
                'end':result[0]['approved_on'],
                'end_year':int(result[0]['approved_on'].year),
                'end_month':int(result[0]['approved_on'].month),
                'end_day':int(result[0]['approved_on'].day),
                })

            response['concept'] = concept

        for item in form6:
            result = models.GlobalTransaction.objects.values('form_id','approved_status','project_id','created_on','approved_on','project__name').filter(form_id=6, approved_status=2, project_id=item['project_id']).exclude(status='3')

            (incorporation).append({
                "form_id":result[0]['form_id'],
                "project_id":result[0]['project_id'],
                "project":result[0]['project__name'],
                'start': result[0]['created_on'],
                'start_year':int(result[0]['created_on'].year),
                'start_month':int(result[0]['created_on'].month),
                'start_day':int(result[0]['created_on'].day),
                'end':result[0]['approved_on'],
                'end_year':int(result[0]['approved_on'].year),
                'end_month':int(result[0]['approved_on'].month),
                'end_day':int(result[0]['approved_on'].day),
                })

            response['incorporation'] = incorporation

        for item in form7:
            result = models.GlobalTransaction.objects.values('form_id','approved_status','project_id','created_on','approved_on','project__name').filter(form_id=7, approved_status=2, project_id=item['project_id']).exclude(status='3')

            (rfi).append({
                "form_id":result[0]['form_id'],
                "project_id":result[0]['project_id'],
                "project":result[0]['project__name'],
                'start': result[0]['created_on'],
                'start_year':int(result[0]['created_on'].year),
                'start_month':int(result[0]['created_on'].month),
                'start_day':int(result[0]['created_on'].day),
                'end':result[0]['approved_on'],
                'end_year':int(result[0]['approved_on'].year),
                'end_month':int(result[0]['approved_on'].month),
                'end_day':int(result[0]['approved_on'].day),
                })

            response['rfi'] = rfi

        return Response({"status":error.context['success_code'], "data": response}, status=status.HTTP_200_OK)


class GetFormLevelHierarchy(APIView):

    def get(self, request, pk = None):

        if 'form_id' in request.GET:
            response = {}
            form_id = request.GET.get('form_id')

            response['recommender'] = models.FormLevelRecommenderHierarchy.objects.values('id', 'form_id','recommender_id','recommender_level').filter(form_id=form_id).exclude(status='3').order_by('recommender_level')

            response['approver'] = models.FormLevelApproverHierarchy.objects.values('id', 'form_id','approver_id','approver_level').filter(form_id=form_id).exclude(status='3').order_by('approver_level')

            return Response({"status":error.context['success_code'] , "data": response}, status=status.HTTP_200_OK)

class FormLevelHierarchyViews(APIView):

    def post(self,request, pk = None):

        #print(request.data)
        if 'form_hierarchy' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "Form id" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)

        if len(request.data['recommender'])==0 and len(request.data['approver'])==0:
            return Response({"status":error.context['error_code'], "message" : "Both recommender and approver" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)

        else: 
            if 'id' in request.data:
                pk = request.data['id']
                created_ip = Common.get_client_ip(request)
                request.data['created_ip'] = created_ip
                form_id = request.data['form_hierarchy']

                rec_count = models.FormLevelRecommenderHierarchy.objects.filter(form_id = form_id).count()
                app_count = models.FormLevelApproverHierarchy.objects.filter(form_id = form_id).count()
               
                models.FormLevelRecommenderHierarchy.objects.filter(form_id = form_id).delete()
                models.FormLevelApproverHierarchy.objects.filter(form_id = form_id).delete()

                #print(pk,"PK")
                pk = None

                if not pk:

                    for index, access in enumerate((request.data['recommender'])):

                        saveserialize = cSerializer.FormsLevelRecommenderHierarchySerializer(data={
                            'form' : form_id,
                            "recommender": access['recommender'] if 'recommender' in access else None,
                            "recommender_level": index+1,
                            "recommender_level_status": -1 if index==0 else 2,
                            "status" : 1,
                            "created_ip" : created_ip,
                            "created_by" : request.user.id
                            })

                        if saveserialize.is_valid():
                            saveserialize.save()

                    for index, access in enumerate((request.data['approver'])):

                        saveserialize = cSerializer.FormsLevelApproverHierarchySerializer(data={
                            'form' : form_id,
                            "approver": access['approver'] if 'approver' in access else None,
                            "approver_level": index+1,
                            "approver_level_status": -1 if index==0 else 2,
                            "status" : 1,
                            "created_ip" : created_ip,
                            "created_by" : request.user.id
                            })

                        if saveserialize.is_valid():
                            saveserialize.save()
                        

                    if rec_count > 0 or app_count > 0:
                        return Response({"status" :error.context['success_code'], "message":"Form level" +language.context[language.defaultLang]['update'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                    else:
                        return Response({"status" :error.context['success_code'], "message":"Form level" +language.context[language.defaultLang]['insert'], "data":saveserialize.data}, status=status.HTTP_200_OK)

                else:                                       

                    list = self.get_object(pk)

                    saveserialize = cSerializer.TemplateConfigSerializer(list, data = request.data, partial= True)
                    if saveserialize.is_valid():
                        saveserialize.save()
                        return Response({"status" :error.context['success_code'], "message":"Template config" +language.context[language.defaultLang]['update'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                    else:
                        return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)
            else: 
                return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)

class GetProjectLevelHierarchy(APIView):

    def get(self, request, pk = None):

        if 'project_id' in request.GET and 'form_id' in request.GET:

            response = {}
            project_id = request.GET.get('project_id')
            form_id = request.GET.get('form_id')

            response['recommender'] = models.ProjectLevelRecommenderHierarchy.objects.values('id', 'form_id','recommender_id','recommender_level').filter(form_id=form_id,project_id=project_id).exclude(status='3').order_by('recommender_level')

            response['approver'] = models.ProjectLevelApproverHierarchy.objects.values('id', 'form_id','approver_id','approver_level').filter(form_id=form_id,project_id=project_id).exclude(status='3').order_by('approver_level')

            return Response({"status":error.context['success_code'] , "data": response}, status=status.HTTP_200_OK)

class ProjectLevelHierarchyViews(APIView):

    def post(self,request, pk = None):

        #print(request.data)
        if 'form_hierarchy' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "Form id" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        if len(request.data['recommender'])==0 and len(request.data['approver'])==0:
            return Response({"status":error.context['error_code'], "message" : "Both recommender and approver" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        else: 
            if 'id' in request.data:
                pk = request.data['id']
                created_ip = Common.get_client_ip(request)
                request.data['created_ip'] = created_ip
                form_id = request.data['form_hierarchy']
                project_id = request.data['project_id']

                rec_count = models.ProjectLevelRecommenderHierarchy.objects.filter(form_id = form_id, project_id = project_id).count()
                app_count = models.ProjectLevelApproverHierarchy.objects.filter(form_id = form_id, project_id = project_id).count()
               
                models.ProjectLevelRecommenderHierarchy.objects.filter(form_id = form_id, project_id = project_id).delete()
                models.ProjectLevelApproverHierarchy.objects.filter(form_id = form_id, project_id = project_id).delete()

                #print(pk,"PK")
                pk = None

                if not pk:

                    for index, access in enumerate((request.data['recommender'])):

                        saveserialize = cSerializer.ProjectLevelRecommenderHierarchySerializer(data={
                            'form' : form_id,
                            'project' : project_id,
                            "recommender" : access['recommender'] if 'recommender' in access else None,
                            "recommender_level": index+1,
                            "recommender_level_status": -1 if index==0 else 2,
                            "status" : 1,
                            "created_ip" : created_ip,
                            "created_by" : request.user.id
                            })

                        if saveserialize.is_valid():
                            saveserialize.save()

                    for index, access in enumerate((request.data['approver'])):

                        saveserialize = cSerializer.ProjectLevelApproverHierarchySerializer(data={
                            'form' : form_id,
                            'project' : project_id,
                            "approver": access['approver'] if 'approver' in access else None,
                            "approver_level": index+1,
                            "approver_level_status": -1 if index==0 else 2,
                            "status" : 1,
                            "created_ip" : created_ip,
                            "created_by" : request.user.id
                            })

                        if saveserialize.is_valid():
                            saveserialize.save()

                    if rec_count > 0 or app_count > 0:
                        return Response({"status" :error.context['success_code'], "message":"Project level" +language.context[language.defaultLang]['update'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                    else:
                        return Response({"status" :error.context['success_code'], "message":"Project level" +language.context[language.defaultLang]['insert'], "data":saveserialize.data}, status=status.HTTP_200_OK)

                else:                                       

                    list = self.get_object(pk)

                    saveserialize = cSerializer.TemplateConfigSerializer(list, data = request.data, partial= True)
                    if saveserialize.is_valid():
                        saveserialize.save()
                        return Response({"status" :error.context['success_code'], "message":"Template config" +language.context[language.defaultLang]['update'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                    else:
                        return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)
            else: 
                return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)



class SaveHierarchyLevel_OLD(APIView):

    def post(self,request, pk = None):

        #print(request.data)
        if 'form_id' not in request.data:
            return Response({"status":error.context['error_code'], "message" : "Form id" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)
        else:
            if 'project_id' in request.data:
                transaction_id = request.data['transaction_id']
                form_id = request.data['form_id']
                project_id = request.data['project_id']
                role_id = request.headers["Authorized-By"]


                secondLevel = accessModels.ProcessFlow.objects.values('process_id','user_role_id').filter(level=2).first()

                # 1st Project level
                rec_level = models.ProjectLevelRecommenderHierarchy.objects.values('recommender_id').filter(form_id = form_id, project_id = project_id).count()

                if rec_level==0:
                    rec_level = models.FormLevelRecommenderHierarchy.objects.values('recommender_id').filter(form_id = form_id).count()

                    if rec_level==0:
                        return Response({"status" :error.context['success_code'], "message":"Add recommender", "data":""}, status = status.HTTP_200_OK)

                    else: # Form Level

                        models.GlobalTransactionApproval.objects.create(
                            approved_level = 1,
                            approved_ip = Common.get_client_ip(request),
                            approved_by_id = request.user.id,
                            transaction_id = transaction_id,
                            comments = "Initiated",
                            status = 1,
                            approved_role_id = role_id,
                            type=-1,
                        )

                        # Getting Recommender Level 1 at Form Level Hierarchy
                        rec_level_1 = models.FormLevelRecommenderHierarchy.objects.values('recommender_id').filter(form_id = form_id, recommender_level = 1).first()

                        res = models.GlobalTransaction.objects.filter(id = transaction_id).update(recommender_level = rec_level, approver_level=1, approved_level=2)

                        formDet = models.Forms.objects.values('id','name').filter(id=request.data["form_id"]).first()

                        projectDet = models.GlobalTransaction.objects.values('id','project__name','created_by__first_name','created_by__last_name').filter(id=request.data["transaction_id"]).first()

                        notificationMessage = (
                            "Project "+projectDet['project__name']
                            + " is initiated by "
                            + projectDet['created_by__first_name']+" "+projectDet['created_by__last_name']+" in "+formDet['name']
                            + " is waiting for your next level recommendations "
                        )

                        notificationModels.NotificationUser.objects.create(
                            form_id = request.data["form_id"],
                            transaction_id = request.data["transaction_id"],
                            project_id = request.data["project_id"],
                            message = notificationMessage,
                            user_role_id = 2,
                            process_id = 1,
                            hierarchy_type = 2,
                            from_user_id = request.user.id,
                            to_user_id = rec_level_1['recommender_id']
                        )


                        return Response({"status" :error.context['success_code'], "message":"Data exist at form level", "data":""}, status=status.HTTP_200_OK)

                else: # Project level.

                    models.GlobalTransactionApproval.objects.create(
                        approved_level = 1,
                        approved_ip = Common.get_client_ip(request),
                        approved_by_id = request.user.id,
                        transaction_id = transaction_id,
                        comments = "Initiated",
                        status = 1,
                        approved_role_id = role_id,
                        type=-1,
                    )

                    # Getting Recommender Level 1 at Form Level Hierarchy
                    rec_level_1 = models.ProjectLevelRecommenderHierarchy.objects.values('recommender_id').filter(form_id = form_id, project_id = project_id, recommender_level = 1).first()

                    res = models.GlobalTransaction.objects.filter(id = transaction_id).update(recommender_level = rec_level, approver_level=1, approved_level=2)

                    ###
                    formDet = models.Forms.objects.values('id','name').filter(id=request.data["form_id"]).first()

                    projectDet = models.GlobalTransaction.objects.values('id','project__name','created_by__first_name','created_by__last_name').filter(id=request.data["transaction_id"]).first()
                    ###

                    notificationMessage = (
                        "Project "+projectDet['project__name']
                        + " is initiated by "
                        + projectDet['created_by__first_name']+" "+projectDet['created_by__last_name']+" in "+formDet['name']
                        + " is waiting for your next level recommendations "
                    )

                    notificationModels.NotificationUser.objects.create(
                        form_id = request.data["form_id"],
                        transaction_id = request.data["transaction_id"],
                        project_id = request.data["project_id"],
                        message = notificationMessage,
                        user_role_id = 2,
                        process_id = 1,
                        hierarchy_type = 2,
                        from_user_id = request.user.id,
                        to_user_id = rec_level_1['recommender_id']
                    )


                    return Response({"status" :error.context['success_code'], "message":"Data exist at project level", "data":""}, status=status.HTTP_200_OK)

            else: 
                return Response({"status" : {"project_id" : ['project_id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)


class SaveHierarchyLevel(APIView):

    def post(self,request, pk = None):

        #print(request.data)
        if 'form_id' not in request.data:
            return Response({"status":error.context['error_code'], "message" : "Form id" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)

        else:
            if 'project_id' in request.data:
                transaction_id = request.data['transaction_id']
                form_id = request.data['form_id']
                project_id = request.data['project_id']
                role_id = request.headers["Authorized-By"]
                created_ip = Common.get_client_ip(request)


                # 1st Project level
                rec_level = models.ProjectLevelRecommenderHierarchy.objects.values('recommender_id').filter(form_id = form_id, project_id = project_id).count()

                if rec_level==0:
                    rec_level = models.FormLevelRecommenderHierarchy.objects.values('recommender_id').filter(form_id = form_id).count()

                    if rec_level==0:
                        return Response({"status" :error.context['success_code'], "message":"Add recommender", "data":""}, status = status.HTTP_200_OK)

                    else: # Form Level

                        # Gather information from FORM LEVEL.
                        form_recommender = models.FormLevelRecommenderHierarchy.objects.values('form_id','recommender_id','recommender_level','recommender_level_status').filter(form_id = form_id).order_by('recommender_level')


                        form_approver = models.FormLevelApproverHierarchy.objects.values('form_id','approver_id','approver_level','approver_level_status').filter(form_id = form_id).order_by('approver_level')

                        #print(form_recommender)
                        #print(form_approver)

                        # Create the hierarchy to project level.
                        for access in form_recommender:

                            saveserialize = cSerializer.ProjectLevelRecommenderHierarchySerializer(data={
                                'form' : form_id,
                                'project' : project_id,
                                "recommender" : access['recommender_id'],
                                "recommender_level": access['recommender_level'],
                                "recommender_level_status": access['recommender_level_status'],
                                "status" : 1,
                                "created_ip" : created_ip,
                                "created_by" : request.user.id
                                })

                            if saveserialize.is_valid():
                                saveserialize.save()

                        for access in form_approver:

                            saveserialize = cSerializer.ProjectLevelApproverHierarchySerializer(data={
                                'form' : form_id,
                                'project' : project_id,
                                "approver": access['approver_id'],
                                "approver_level": access['approver_level'],
                                "approver_level_status": access['approver_level_status'],
                                "status" : 1,
                                "created_ip" : created_ip,
                                "created_by" : request.user.id
                                })

                            if saveserialize.is_valid():
                                saveserialize.save()



                        models.GlobalTransactionApproval.objects.create(
                            approved_level = 1,
                            approved_ip = Common.get_client_ip(request),
                            approved_by_id = request.user.id,
                            transaction_id = transaction_id,
                            comments = "Initiated",
                            status = 1,
                            approved_role_id = role_id,
                            type=-1,
                        )

                        # Getting Recommender Level 1 at Project Level Hierarchy
                        rec_level_1 = models.ProjectLevelRecommenderHierarchy.objects.values('recommender_id').filter(form_id = form_id, project_id = project_id, recommender_level = 1).first()

                        res = models.GlobalTransaction.objects.filter(id = transaction_id).update(recommender_level = rec_level, approver_level=1, approved_level=2)

                        ###
                        formDet = models.Forms.objects.values('id','name').filter(id=request.data["form_id"]).first()

                        projectDet = models.GlobalTransaction.objects.values('id','project__name','created_by__first_name','created_by__last_name').filter(id=request.data["transaction_id"]).first()
                        ###

                        notificationMessage = (
                            "Project "+projectDet['project__name']
                            + " is initiated by "
                            + projectDet['created_by__first_name']+" "+projectDet['created_by__last_name']+" in "+formDet['name']
                            + " is waiting for your next level recommendations "
                        )

                        notificationModels.NotificationUser.objects.create(
                            form_id = request.data["form_id"],
                            transaction_id = request.data["transaction_id"],
                            project_id = request.data["project_id"],
                            message = notificationMessage,
                            user_role_id = 2,
                            process_id = 1,
                            hierarchy_type = 2,
                            from_user_id = request.user.id,
                            to_user_id = rec_level_1['recommender_id']
                        )


                        #return Response({"status" :error.context['success_code'], "message":"Data exist at project level", "data":""}, status=status.HTTP_200_OK)
                        return Response({"status" :error.context['success_code'], "message":"Transaction initiated successfully", "data":""}, status=status.HTTP_200_OK)

                else: # Project level.

                    models.GlobalTransactionApproval.objects.create(
                        approved_level = 1,
                        approved_ip = Common.get_client_ip(request),
                        approved_by_id = request.user.id,
                        transaction_id = transaction_id,
                        comments = "Initiated",
                        status = 1,
                        approved_role_id = role_id,
                        type=-1,
                    )

                    # Getting Recommender Level 1 at Form Level Hierarchy
                    rec_level_1 = models.ProjectLevelRecommenderHierarchy.objects.values('recommender_id').filter(form_id = form_id, project_id = project_id, recommender_level = 1).first()

                    res = models.GlobalTransaction.objects.filter(id = transaction_id).update(recommender_level = rec_level, approver_level=1, approved_level=2)

                    ###
                    formDet = models.Forms.objects.values('id','name').filter(id=request.data["form_id"]).first()

                    projectDet = models.GlobalTransaction.objects.values('id','project__name','created_by__first_name','created_by__last_name').filter(id=request.data["transaction_id"]).first()
                    ###

                    notificationMessage = (
                        "Project "+projectDet['project__name']
                        + " is initiated by "
                        + projectDet['created_by__first_name']+" "+projectDet['created_by__last_name']+" in "+formDet['name']
                        + " is waiting for your next level recommendations "
                    )

                    notificationModels.NotificationUser.objects.create(
                        form_id = request.data["form_id"],
                        transaction_id = request.data["transaction_id"],
                        project_id = request.data["project_id"],
                        message = notificationMessage,
                        user_role_id = 2,
                        process_id = 1,
                        hierarchy_type = 2,
                        from_user_id = request.user.id,
                        to_user_id = rec_level_1['recommender_id']
                    )


                    # return Response({"status" :error.context['success_code'], "message":"Data exist at project level", "data":""}, status=status.HTTP_200_OK)

                    return Response({"status" :error.context['success_code'], "message":"Transaction initiated successfully", "data":""}, status=status.HTTP_200_OK)

            else: 
                return Response({"status" : {"project_id" : ['project_id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)


class ShipDetailsViews(APIView):

    def post(self,request, pk = None):

        #print(request.data)
        if 'project_id' not in request.data and request.data['status'] != 3:
            return Response({"status":error.context['error_code'], "message" : "Project id" +language.context[language.defaultLang]['missing']},status=status.HTTP_200_OK)

        else: 
            if 'id' in request.data:
                pk = request.data['id']
                #form_id = request.data['form_id']
                #form_id = 8
                project_id = request.data['project_id']
                created_ip = Common.get_client_ip(request)
                created_id = request.user.id 

                #ship_count = models.GlobalTransactionShipDetail.objects.filter(form_id = form_id, project_id = project_id).count()
                #models.GlobalTransactionShipDetail.objects.filter(form_id = form_id, project_id = project_id).delete()

                ship_count = models.GlobalTransactionShipDetail.objects.filter(project_id = project_id).count()
                models.GlobalTransactionShipDetail.objects.filter(project_id = project_id).delete()

                pk = None
                if not pk:

                    for index, access in enumerate((request.data['ship'])):

                        saveserialize = cSerializer.ShipDetailsSerializer(data={
                            #'form' : form_id,
                            'project' : project_id,
                            "name": access['name'] if 'name' in access else None,
                            "code": access['code'] if 'code' in access else None,
                            "ship_id": access['ship_id'] if 'ship_id' in access else None,
                            "command": access['command'] if 'command' in access else None,
                            "status" : 1,
                            "created_ip" : created_ip,
                            "created_by" : created_id
                            })

                        if saveserialize.is_valid():
                            saveserialize.save()                        

                    if ship_count > 0:
                        return Response({"status" :error.context['success_code'], "message":"Ship details" +language.context[language.defaultLang]['update'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                    else:
                        return Response({"status" :error.context['success_code'], "message":"Ship details" +language.context[language.defaultLang]['insert'], "data":saveserialize.data}, status=status.HTTP_200_OK)

                else:                                       

                    list = self.get_object(pk)

                    saveserialize = cSerializer.TemplateConfigSerializer(list, data = request.data, partial= True)
                    if saveserialize.is_valid():
                        saveserialize.save()
                        return Response({"status" :error.context['success_code'], "message":"Template config" +language.context[language.defaultLang]['update'], "data":saveserialize.data}, status=status.HTTP_200_OK)
                    else:
                        return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)
            else: 
                return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)


class GetShipDetails(APIView):

    def get(self, request, pk = None):

        #if 'project_id' in request.GET and 'form_id' in request.GET:
        if 'project_id' in request.GET:
            response = {}
            project_id = request.GET.get('project_id')
            #form_id = request.GET.get('form_id')
            #form_id = 8

            response['ship'] = models.GlobalTransactionShipDetail.objects.values('id','project_id','name','code','ship_id','command_id').filter(project_id=project_id).exclude(status='3')

            return Response({"status":error.context['success_code'] , "data": response}, status=status.HTTP_200_OK)

### Excel ####

class TestExcelExport(APIView):

    def post(self,request, pk = None): 
        from pathlib import Path
        BASE_DIR = Path(__file__).resolve().parent.parent
        request_file = request.FILES['excel_file_upload']

        dir_storage='static/import_excel'
        fs = FileSystemStorage(location=dir_storage)
        filename = fs.save(request_file.name, request_file)
        if os.path.splitext(request_file.name)[1] == ".xls" or  os.path.splitext(request_file.name)[1] == ".xlsx":
            excel_folder = os.path.join(BASE_DIR, 'media/Excel/test/')
            read_file = pd.read_excel(request_file)
            read_file.to_csv(excel_folder +'import_excel_file.csv')
            fhand = open('media/Excel/test/import_excel_file.csv')
        else:
             return Response({"status":error.context['error_code'],"message" : "File format not supported (Xls and Xlsx only allowed)" })    
        reader = csv.reader(fhand)
        next(reader)
        #print(reader,"Reader 12312312")
        created_ip = Common.get_client_ip(request)
        created_by = request.user.id
        
        for row in reader:

            request_data = {
                'code' : row[1],
                'name' : row[2],
                'equipment_ship_id' : row[3],
                'equipment_type_name' : row[1],
                'status':1,
                'created_ip': created_ip,
                'created_by': created_by
            }
            #print(row[1],'--',row[2],'--',row[3])
            #print(request_data,"request_data")
            saveserialize = cSerializer.EquipmentSerializer(data = request_data)
            if saveserialize.is_valid():
                saveserialize.save()
        return Response("ok")


            # saveserialize = cSerializer.EquipmentSerializer(data = request_data)
            # if saveserialize.is_valid():
            #     saveserialize.save()
            #     return Response({"status" :error.context['success_code'], "message":"New Equipment" +language.context[language.defaultLang]['insert'], "data":saveserialize.data}, status=status.HTTP_200_OK)
            # else:
            #     return Response({"status" :error.context['error_code'], "message": error.serializerError(saveserialize)}, status=status.HTTP_200_OK)




### TEST ####

# class TestFormsMappingDetailsViews(APIView):

#     def get_object(self,pk):

#         try:
#             return models.Forms.objects.get(pk = pk)
#         except models.Forms.DoesNotExist:
#             raise Http404


#     def post(self,request, pk = None):


#         ss = cSerializer.AllFormsMappingSerializer(models.FormsMapping.objects.filter(form_id=4), many=True)

#         for index, access in enumerate((ss.data)):
#             saveserialize = cSerializer.FormsMappingSerializer(data={
#                 'form' : 8,
#                 "order": access['order'],
#                 "Class": access['Class'],
#                 "module" : access['module'],
#                 "sub_module" : access['sub_module'],
#                 "section" : access['section'],
#                 "sub_section" : access['sub_section'],
#                 "sub_sub_section" : access['sub_sub_section'],
#                 "status" : access['status'],
#                 'created_on' : access['created_on'],
#                 'created_by' : access['created_by'],
#                 'created_ip' : access['created_ip'],
#                 'modified_on' : access['modified_on'],
#                 'modified_by' : access['modified_by'],
#                 'modified_ip' : access['modified_ip']

#                 })

#             if saveserialize.is_valid():
#                 saveserialize.save()
                        
#         return Response({"status" :error.context['success_code'], "message":"Form mapping" +language.context[language.defaultLang]['insert'], "data":saveserialize.data}, status=status.HTTP_200_OK)









### Dock Yard ###

class ImportExcelDart(APIView):

    def post(self,request, pk = None):

        from pathlib import Path
        BASE_DIR = Path(__file__).resolve().parent.parent
        created_ip = Common.get_client_ip(request)
        request_file = request.FILES['excel_file_upload']
        created_by = request.data['created_by']

        dir_storage='static/import_excel'
        fs = FileSystemStorage(location=dir_storage)
        filename = fs.save(request_file.name, request_file)
        if os.path.splitext(request_file.name)[1] == ".xls" or  os.path.splitext(request_file.name)[1] == ".xlsx":
            excel_folder = os.path.join(BASE_DIR, 'media/Excel/Dart/')
            read_file = pd.read_excel(request_file)
            read_file.to_csv(excel_folder +'import_excel_file.csv')
            fhand = open('media/Excel/Dart/import_excel_file.csv')
        else:
             return Response({"status":error.context['error_code'],"message" : "File format not supported (Xls and Xlsx only allowed)" })    
        reader = csv.reader(fhand)
        next(reader)
        #print(reader)


        for row in reader:
            #print(row[1])
            #print(type(int(row[10])),"TYYGHHHHV", row[10])
            if not request_file:
                return Response({"status":error.context['error_code'],"message" : "Upload file is required" })

            department = masterModels.Department.objects.filter(code=row[4]).first()

            if not department:
                department_id = None
            else:
                department_id = department.id

            request_data = {
                'SrNo' : row[1],
                'ShipSrNo' : row[2],

                'DartDate' : pd.to_datetime(row[3]),

                'DepartmentID' : department_id,
                'ExDept' : row[5],
                'ExDeptID' : row[6],
                'EquipmentShipID' : row[7],
                'EquipmentCode' : row[8],
                'SeverityID' : row[9],
                'SeverityCode' : row[10],
                'DiagnosticID' : row[11],
                'DiagnosticCode': row[12],
                #'RectifiedDate' : row[13],
                #'RectifiedDate' : datetime.strptime(row[13], "%Y-%m-%d"),
                'RectifiedDate' : pd.to_datetime(row[13]),
                'RepairID' : row[14],
                'RepairCode' : row[15],
                'RepairAgencyID' : row[16],
                'AgencyCode': row[17],
                'DelayID' : row[18],
                'DelayCode' : row[19],
                'Remarks' : row[20],
                'OpdefSrNo' : row[21],
                'XdueRefitType': row[22],
                'XdueRefitRemarks' : row[23],
                #'CancelDate' : row[24],
                #'CancelDate' : datetime.strptime(row[24], "%Y-%m-%d"),
                'CancelDate' : pd.to_datetime(row[24]),
                'NILDart': row[25],
                'Active' : row[26],
                'CreatedBy': created_by,
                'Source': row[27],

                'status' : 1,
                'created_ip': created_ip,
            }

            #print(request_data,"GGGGG")

            saveserialize = cSerializer.DartSerializer(data = request_data)
            if saveserialize.is_valid():
                saveserialize.save()
            else:
                return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)

        excel_upload_obj = models.ExcelFileDartUpload.objects.create(
        excel_file_upload = request.data['excel_file_upload'],
        created_ip =  created_ip
        )

        if excel_upload_obj:
            return Response({"status" :error.context['success_code'], "message":"File imported successfully"}, status=status.HTTP_200_OK)
        else:
            return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)

class ImportExcelRA(APIView):

    def post(self,request, pk = None):

        from pathlib import Path
        BASE_DIR = Path(__file__).resolve().parent.parent
        created_ip = Common.get_client_ip(request)
        request_file = request.FILES['excel_file_upload']
        created_by = request.data['created_by']

        dir_storage='static/import_excel'
        fs = FileSystemStorage(location=dir_storage)
        filename = fs.save(request_file.name, request_file)
        if os.path.splitext(request_file.name)[1] == ".xls" or  os.path.splitext(request_file.name)[1] == ".xlsx":
            excel_folder = os.path.join(BASE_DIR, 'media/Excel/RA/')
            read_file = pd.read_excel(request_file)
            read_file.to_csv(excel_folder +'import_excel_file.csv')
            fhand = open('media/Excel/RA/import_excel_file.csv')
        else:
             return Response({"status":error.context['error_code'],"message" : "File format not supported (Xls and Xlsx only allowed)" })    
        reader = csv.reader(fhand)
        next(reader)
        #print(reader)


        for row in reader:
            #print(row[1])
            #print(type(int(row[10])),"TYYGHHHHV", row[10])
            if not request_file:
                return Response({"status":error.context['error_code'],"message" : "Upload file is required" })

            department = masterModels.Department.objects.filter(code=row[4]).first()

            if not department:
                department_id = None
            else:
                department_id = department.id

            request_data = {
                'SrNo' : row[1],
                'ShipSrNo' : row[2],

                'DartDate' : pd.to_datetime(row[3]),

                'DepartmentID' : department_id,
                'ExDept' : row[5],
                'ExDeptID' : row[6],
                'EquipmentShipID' : row[7],
                'EquipmentCode' : row[8],
                'SeverityID' : row[9],
                'SeverityCode' : row[10],
                'DiagnosticID' : row[11],
                'DiagnosticCode': row[12],
                #'RectifiedDate' : row[13],
                #'RectifiedDate' : datetime.strptime(row[13], "%Y-%m-%d"),
                'RectifiedDate' : pd.to_datetime(row[13]),
                'RepairID' : row[14],
                'RepairCode' : row[15],
                'RepairAgencyID' : row[16],
                'AgencyCode': row[17],
                'DelayID' : row[18],
                'DelayCode' : row[19],
                'Remarks' : row[20],
                'OpdefSrNo' : row[21],
                'XdueRefitType': row[22],
                'XdueRefitRemarks' : row[23],
                #'CancelDate' : row[24],
                #'CancelDate' : datetime.strptime(row[24], "%Y-%m-%d"),
                'CancelDate' : pd.to_datetime(row[24]),
                'NILDart': row[25],
                'Active' : row[26],
                'CreatedBy': created_by,
                'Source': row[27],

                'status' : 1,
                'created_ip': created_ip,
            }

            #print(request_data,"GGGGG")

            saveserialize = cSerializer.RASerializer(data = request_data)
            if saveserialize.is_valid():
                saveserialize.save()
            else:
                return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)

        excel_upload_obj = models.ExcelFileRAUpload.objects.create(
        excel_file_upload = request.data['excel_file_upload'],
        created_ip =  created_ip
        )

        if excel_upload_obj:
            return Response({"status" :error.context['success_code'], "message":"File imported successfully"}, status=status.HTTP_200_OK)
        else:
            return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)

class ImportExcelOPDEF(APIView):

    def post(self,request, pk = None):

        from pathlib import Path
        BASE_DIR = Path(__file__).resolve().parent.parent
        created_ip = Common.get_client_ip(request)
        request_file = request.FILES['excel_file_upload']
        created_by = request.data['created_by']

        dir_storage='static/import_excel'
        fs = FileSystemStorage(location=dir_storage)
        filename = fs.save(request_file.name, request_file)
        if os.path.splitext(request_file.name)[1] == ".xls" or  os.path.splitext(request_file.name)[1] == ".xlsx":
            excel_folder = os.path.join(BASE_DIR, 'media/Excel/OPDEF/')
            read_file = pd.read_excel(request_file)
            read_file.to_csv(excel_folder +'import_excel_file.csv')
            fhand = open('media/Excel/OPDEF/import_excel_file.csv')
        else:
             return Response({"status":error.context['error_code'],"message" : "File format not supported (Xls and Xlsx only allowed)" })    
        reader = csv.reader(fhand)
        next(reader)
        #print(reader)


        for row in reader:
            #print(row[1])
            #print(type(int(row[10])),"TYYGHHHHV", row[10])
            if not request_file:
                return Response({"status":error.context['error_code'],"message" : "Upload file is required" })

            department = masterModels.Department.objects.filter(code=row[5]).first()

            if not department:
                department_id = None
            else:
                department_id = department.id

            request_data = {
                'OpdefDate' : pd.to_datetime(row[1]),
                'OpdefNo' : row[2],
                'ShipID' : row[3],
                'Command':row[4],
                'DepartmentID' : department_id,
                'DartNo' : row[6],
                'STA' : row[7],
                'EquipmentID' : row[8],
                'LocationCode' : row[9],
                'Unit' : row[10],
                'Defect' : row[11],
                'Severity' : row[12],
                'AssistanceRequired': row[13],
                'EffectOnOperation': row[14],
                'DowngradeDate' : pd.to_datetime(row[15]),
                'STADowngradeDate' : pd.to_datetime(row[16]),
                'STAUpgradeDate' : pd.to_datetime(row[17]),
                'CancelDtg' : row[18],
                'CancelDate' : pd.to_datetime(row[19]),
                'RepairDate' : pd.to_datetime(row[20]),
                'RepairID' : row[21],
                'RepairParts': row[22],
                'HoldupCode' : row[23],
                'HoldupDetail' : row[24],
                'INSMAEffort' : row[25],
                'Refit' : row[26],
                'Active': row[27],
                'Remarks' : row[28],
                'StoredemSrNo' : row[29],
                'OpDefStatus': row[30],
                'EquipmentSerialNo' : row[31],
                'OfarDate' : pd.to_datetime(row[32]),
                'status' : 1,
                'created_ip': created_ip,
            }

            #print(request_data,"GGGGG")

            saveserialize = cSerializer.OPDEFSerializer(data = request_data)
            if saveserialize.is_valid():
                saveserialize.save()
            else:
                return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)

        excel_upload_obj = models.ExcelFileOPDEFUpload.objects.create(
        excel_file_upload = request.data['excel_file_upload'],
        created_ip =  created_ip
        )

        if excel_upload_obj:
            return Response({"status" :error.context['success_code'], "message":"File imported successfully"}, status=status.HTTP_200_OK)
        else:
            return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)

# Dart
class DartList(APIView):
    
    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.Dart.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.ListDartSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.Dart.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Initiation Notes" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        #lists = models.Dart.objects
        lists = models.Dart.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListDartSerializer(lists, many=True)
        print('serializer serializer',serializer.data)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class DartCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.Dart.objects.get(pk = pk)
            except models.Dart.DoesNotExist:
                raise Http404

    def post(self,request, pk = None):

        #print(request.data)
        #pass
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']
            if not request.data['DepartmentID'] and request.data['status'] != 3:
                return Response({"status":error.context['error_code'], "message": "Department ID missing"}, status=status.HTTP_200_OK)
            if not pk:

                request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.DartSerializer(data = request.data)

                if saveserialize.is_valid():
                    saveserialize.save()

                    return Response({"status" : error.context['success_code'], "message":"Dart" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:

                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id


                list = self.get_object(pk)
                
                saveserialize = cSerializer.DartSerializer(list, data = request.data, partial= True)                
                if saveserialize.is_valid():
                    saveserialize.save()
                    # logData=request.data

                    # if 'approved_status' in request.data:
                    #     models.ProjectModuleStatus.objects.filter(project_id=request.data['project'], project_module_master_id=1).delete()

                    #     models.ProjectModuleStatus.objects.create(
                    #         project_module_master_id = 1,
                    #         project_id = request.data['project'],
                    #         status  = request.data['approved_status']
                    #     )

                    return Response({"status" :error.context['success_code'], "message":"Dart" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)


# RA
class RAList(APIView):
    
    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.RA.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.ListRASerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.RA.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Initiation Notes" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        #lists = models.RA.objects
        lists = models.RA.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListRASerializer(lists, many=True)
        print('serializer serializer',serializer.data)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class RACRUD(APIView):
    def get_object(self, pk):
            try:
                return models.RA.objects.get(pk = pk)
            except models.RA.DoesNotExist:
                raise Http404

    def post(self,request, pk = None):

        #print(request.data)
        #pass
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']
            if not request.data['DepartmentID'] and request.data['status'] != 3:
                return Response({"status":error.context['error_code'], "message": "Department ID missing"}, status=status.HTTP_200_OK)
            if not pk:

                request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.RASerializer(data = request.data)

                if saveserialize.is_valid():
                    saveserialize.save()

                    return Response({"status" : error.context['success_code'], "message":"RA" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:

                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id


                list = self.get_object(pk)
                
                saveserialize = cSerializer.RASerializer(list, data = request.data, partial= True)                
                if saveserialize.is_valid():
                    saveserialize.save()
                    # logData=request.data

                    # if 'approved_status' in request.data:
                    #     models.ProjectModuleStatus.objects.filter(project_id=request.data['project'], project_module_master_id=1).delete()

                    #     models.ProjectModuleStatus.objects.create(
                    #         project_module_master_id = 1,
                    #         project_id = request.data['project'],
                    #         status  = request.data['approved_status']
                    #     )

                    return Response({"status" :error.context['success_code'], "message":"RA" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)


# OPDEF
class OPDEFList(APIView):
    
    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.OPDEF.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.ListOPDEFSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.OPDEF.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"OPDEF" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.OPDEF.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListOPDEFSerializer(lists, many=True)
        print('serializer serializer',serializer.data)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class OPDEFCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.OPDEF.objects.get(pk = pk)
            except models.OPDEF.DoesNotExist:
                raise Http404

    def post(self,request, pk = None):

        #print(request.data)
        #pass
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']
            if not request.data['DepartmentID'] and request.data['status'] != 3:
                return Response({"status":error.context['error_code'], "message": "Department ID missing"}, status=status.HTTP_200_OK)
            if not pk:

                request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.OPDEFSerializer(data = request.data)

                if saveserialize.is_valid():
                    saveserialize.save()

                    return Response({"status" : error.context['success_code'], "message":"OPDEF" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:

                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id


                list = self.get_object(pk)
                
                saveserialize = cSerializer.OPDEFSerializer(list, data = request.data, partial= True)                
                if saveserialize.is_valid():
                    saveserialize.save()
                    # logData=request.data

                    # if 'approved_status' in request.data:
                    #     models.ProjectModuleStatus.objects.filter(project_id=request.data['project'], project_module_master_id=1).delete()

                    #     models.ProjectModuleStatus.objects.create(
                    #         project_module_master_id = 1,
                    #         project_id = request.data['project'],
                    #         status  = request.data['approved_status']
                    #     )

                    return Response({"status" :error.context['success_code'], "message":"OPDEF" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)


# Work Instruction
class WorkInstructionList(APIView):
    
    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.WorkInstruction.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.ListWorkInstructionSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.WorkInstruction.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Initiation Notes" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.WorkInstruction.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListWorkInstructionSerializer(lists, many=True)
        print('serializer serializer',serializer.data)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class WorkInstructionCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.WorkInstruction.objects.get(pk = pk)
            except models.WorkInstruction.DoesNotExist:
                raise Http404

    def post(self,request, pk = None):

        #print(request.data)
        #pass
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']
            if not request.data['dart'] and request.data['status'] != 3:
                return Response({"status":error.context['error_code'], "message": "Dart id missing"}, status=status.HTTP_200_OK)
            if not pk:

                request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.WorkInstructionSerializer(data = request.data)

                if saveserialize.is_valid():
                    saveserialize.save()

                    return Response({"status" : error.context['success_code'], "message":"Work instruction" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:

                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id


                list = self.get_object(pk)
                
                saveserialize = cSerializer.WorkInstructionSerializer(list, data = request.data, partial= True)                
                if saveserialize.is_valid():
                    saveserialize.save()

                    return Response({"status" :error.context['success_code'], "message":"Work instruction" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)

class QCWorkInstructionCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.WorkInstructionQCCheck.objects.get(pk = pk)
            except models.WorkInstructionQCCheck.DoesNotExist:
                raise Http404

    def post(self,request, pk = None):

        #print(request.data)
        #pass
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:

            request.data['work_instruction'] = request.data['id']
            request.data['created_ip'] = Common.get_client_ip(request)
            request.data['created_by'] = request.user.id
            request.data['status'] = 1
            saveserialize = cSerializer.WorkInstructionQCCheckSerializer(data = request.data)

            if saveserialize.is_valid():
                saveserialize.save()

                return Response({"status" : error.context['success_code'], "message":"QC work instruction" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
            else:
                return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)

class QCWorkInstructionHistory(APIView):
    def post(self, request, pk=None):
        if "work_instruction_id" in request.data:
            history = (
                models.WorkInstructionQCCheck.objects.values(
                    "id",
                    "qc_status",
                    "remarks",
                    "created_on",
                    "created_by__first_name",
                    "created_by__last_name"
                )
                .filter(work_instruction_id=request.data["work_instruction_id"])
                .order_by("id")
            )
            return Response(
                {"status": error.context["success_code"], "data": history},
                status=status.HTTP_200_OK,
            )
        else:
            return Response(
                {"status": error.context["success_code"], "data": "No data"},
                status=status.HTTP_200_OK,
            )


# Job Card
class JobCardList(APIView):
    
    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.JobCard.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.ListJobCardSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.JobCard.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Job card" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.JobCard.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListJobCardSerializer(lists, many=True)
        #print('serializer serializer',serializer.data)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class JobCardCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.JobCard.objects.get(pk = pk)
            except models.JobCard.DoesNotExist:
                raise Http404

    def post(self,request, pk = None):

        #print(request.data)
        #pass
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']
            if not request.data['work_instruction'] and request.data['status'] != 3:
                return Response({"status":error.context['error_code'], "message": "Work instruction id missing"}, status=status.HTTP_200_OK)
            if not pk:

                request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.JobCardSerializer(data = request.data)

                if saveserialize.is_valid():
                    saveserialize.save()

                    return Response({"status" : error.context['success_code'], "message":"Job card" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:

                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id


                list = self.get_object(pk)
                
                saveserialize = cSerializer.JobCardSerializer(list, data = request.data, partial= True)                
                if saveserialize.is_valid():
                    saveserialize.save()

                    return Response({"status" :error.context['success_code'], "message":"Job card" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)

class QCJobCardCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.JobCardQCCheck.objects.get(pk = pk)
            except models.JobCardQCCheck.DoesNotExist:
                raise Http404

    def post(self,request, pk = None):

        #print(request.data)
        #pass
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:

            request.data['job_card'] = request.data['id']
            request.data['created_ip'] = Common.get_client_ip(request)
            request.data['created_by'] = request.user.id
            request.data['status'] = 1
            saveserialize = cSerializer.JobCardQCCheckSerializer(data = request.data)

            if saveserialize.is_valid():
                saveserialize.save()

                return Response({"status" : error.context['success_code'], "message":"QC job card" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
            else:
                return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)

class QCJobCardHistory(APIView):
    def post(self, request, pk=None):
        if "job_card_id" in request.data:
            history = (
                models.JobCardQCCheck.objects.values(
                    "id",
                    "qc_status",
                    "remarks",
                    "created_on",
                    "created_by__first_name",
                    "created_by__last_name"
                )
                .filter(job_card_id=request.data["job_card_id"])
                .order_by("id")
            )
            return Response(
                {"status": error.context["success_code"], "data": history},
                status=status.HTTP_200_OK,
            )
        else:
            return Response(
                {"status": error.context["success_code"], "data": "No data"},
                status=status.HTTP_200_OK,
            )



# Attendance
class AttendanceList(APIView):
    
    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.Attendance.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.ListAttendanceSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.Attendance.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Attendance" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.Attendance.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListAttendanceSerializer(lists, many=True)
        print('serializer serializer',serializer.data)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class AttendanceCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.Attendance.objects.get(pk = pk)
            except models.Attendance.DoesNotExist:
                raise Http404

    def post(self,request, pk = None):

        #print(request.data)
        #pass
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']
            if not request.data['user'] and request.data['status'] != 3:
                return Response({"status":error.context['error_code'], "message": "Employee missing"}, status=status.HTTP_200_OK)
            if not pk:

                request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.AttendanceSerializer(data = request.data)

                if saveserialize.is_valid():
                    saveserialize.save()

                    return Response({"status" : error.context['success_code'], "message":"Attendance" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:

                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id


                list = self.get_object(pk)
                
                saveserialize = cSerializer.AttendanceSerializer(list, data = request.data, partial= True)                
                if saveserialize.is_valid():
                    saveserialize.save()

                    return Response({"status" :error.context['success_code'], "message":"Attendance" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)



class CheckInOutOLD(APIView):
    def post(self, request, pk=None):
        if "attendance_id" in request.data:

            check_type = request.data['check_type']

            if check_type == 'check_in':
                models.Attendance.objects.filter(id=request.data["attendance_id"]).update(check_in = datetime.now())
                return Response(
                    {"status": error.context["success_code"], "message": "Check in updated"},
                    status=status.HTTP_200_OK,
                )

            if check_type == 'check_out':
                models.Attendance.objects.filter(id=request.data["attendance_id"]).update(check_out = datetime.now())
                return Response(
                    {"status": error.context["success_code"], "message": "Check out updated"},
                    status=status.HTTP_200_OK,
                )

        else:
            return Response(
                {"status": error.context["error_code"], "message": "Attendance id missing"},
                status=status.HTTP_200_OK,
            )


class CheckInOut(APIView):
    def post(self, request, pk=None):
        if "check_type" in request.data:

            check_type = request.data['check_type']
            ipAddress = Common.get_client_ip(request)

            center = models.User.objects.values('center').filter(id=request.user.id)

            if check_type == 'check_in':

                models.Attendance.objects.create(
                    user_id = request.user.id,
                    center_id = center[0]['center'],
                    attendance_date = datetime.now(),
                    check_in = datetime.now(), 
                    attendance_status = 1,
                    status = 1,
                    created_ip = ipAddress,
                    created_by_id = request.user.id
                )

                return Response(
                    {"status": error.context["success_code"], "message": "Checked in successfully"},
                    status=status.HTTP_200_OK,
                )

            if check_type == 'check_out':
                
                models.Attendance.objects.filter(user_id=request.user.id).update(
                    check_out = datetime.now(), 
                    modified_ip = ipAddress,
                    modified_by = request.user.id,
                    modified_on = datetime.now(),
                    )

                return Response(
                    {"status": error.context["success_code"], "message": "Checked out successfully"},
                    status=status.HTTP_200_OK,
                )

        else:
            return Response(
                {"status": error.context["error_code"], "message": "Check type is missing"},
                status=status.HTTP_200_OK,
            )


class getAttendance(APIView):
    def post(self, request, pk=None):
        if request.user.id:
            attendance =  models.Attendance.objects.values(
                    "id",
                    "user_id",
                    "center_id",
                    "attendance_date",
                    "check_in",
                    "check_out",
                    "attendance_status",
                    "center",
                    "status"
                ).filter(user_id=request.user.id, attendance_date=date.today())
            
            return Response(
                {"status": error.context["success_code"], "data": attendance},
                status=status.HTTP_200_OK,
            )
        else:
            return Response(
                {"status": error.context["success_code"], "data": "Not logged in"},
                status=status.HTTP_200_OK,
            )




# Monthly Credits Debits
class MonthlyCreditsDebitsList(APIView):
    
    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.MonthlyCreditsDebits.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.ListMonthlyCreditsDebitsSerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.MonthlyCreditsDebits.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"Monthly credits debits" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.MonthlyCreditsDebits.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListMonthlyCreditsDebitsSerializer(lists, many=True)
        print('serializer serializer',serializer.data)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class MonthlyCreditsDebitsCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.MonthlyCreditsDebits.objects.get(pk = pk)
            except models.MonthlyCreditsDebits.DoesNotExist:
                raise Http404

    def post(self,request, pk = None):

        #print(request.data)
        #pass
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']
            if not request.data['user'] and request.data['status'] != 3:
                return Response({"status":error.context['error_code'], "message": "Employee missing"}, status=status.HTTP_200_OK)
            if not pk:

                request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.MonthlyCreditsDebitsSerializer(data = request.data)

                if saveserialize.is_valid():
                    saveserialize.save()

                    return Response({"status" : error.context['success_code'], "message":"MonthlyCreditsDebits" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:

                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id


                list = self.get_object(pk)
                
                saveserialize = cSerializer.MonthlyCreditsDebitsSerializer(list, data = request.data, partial= True)                
                if saveserialize.is_valid():
                    saveserialize.save()

                    return Response({"status" :error.context['success_code'], "message":"MonthlyCreditsDebits" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)
# MonthlySalary
class MonthlySalaryList(APIView):
    
    def get(self, request, pk = None):       
        filter_values = dict(request.GET.items())
        search_string=order_type=order_column=limit_start=limit_end=''
        normal_values=dict()
        array_values=dict()
        if filter_values:
            for key,values in filter_values.items():
                if values.find("[") !=-1 and values.find("]") !=-1:
                    res = values.strip('][').split(',')
                    array_values[key]=(res)
                else:
                    normal_values[key]=(values)

            strings = []
            search_string = dict((k, normal_values[k]) for k in strings
                                            if k in normal_values)  
            order_column =  request.GET.get('order_column')
            order_type = request.GET.get('order_type')  
            limit_start = request.GET.get('limit_start')
            limit_end = request.GET.get('limit_end')  


            if order_column is not None:                                      
                normal_values.pop('order_column')
            if order_type is not None: 
                normal_values.pop('order_type')  
            if limit_start is not None: 
                normal_values.pop('limit_start')
            if limit_end is not None: 
                normal_values.pop('limit_end')     

            for key in strings:
                if key in normal_values:
                    normal_values.pop(key)

            if search_string:       
                filter_string = None
                for field in search_string:
                    q = Q(**{"%s__contains" % field: search_string[field] })
                    if filter_string:
                        filter_string = filter_string & q
                    else:
                        filter_string = q
        try:
            if pk:
                list = models.MonthlySalary.objects.filter(pk=pk).exclude(status='3').get()
                serializeobj = serializer.ListMonthlySalarySerializer(list)
                return Response({"status":error.context['success_code'], "data": serializeobj.data}, status=status.HTTP_200_OK)
       
        except models.MonthlySalary.DoesNotExist:
            return Response({"status" :error.context['error_code'], "message":"MonthlySalary" +language.context[language.defaultLang]['dataNotFound']}, status = status.HTTP_200_OK)

        lists = models.MonthlySalary.objects.exclude(status='3')
        if normal_values:
            lists = lists.filter(reduce(operator.and_, 
                               (Q(**d) for d in [dict([i]) for i in normal_values.items()])))
        if array_values:
            for key,values in array_values.items():
                queries= [Q(**{"%s__contains" % key: value }) for value in values]
                query=queries.pop()
                for item in queries:
                    query |= item
                lists = lists.filter(query)

        if search_string:
            lists = lists.filter(filter_string)

        if order_type is None: 
            if order_column:
                lists = lists.order_by(order_column)  

        elif order_type in 'asc':
            if order_column:
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('id')   

        elif order_type in 'desc':
            if order_column:
                order_column = '-' + str(order_column)
                lists = lists.order_by(order_column)
            else: 
                lists = lists.order_by('-id') 

        if limit_start and limit_end:
                lists = lists[int(limit_start):int(limit_end)]

        elif limit_start:
                lists = lists[int(limit_start):]

        elif limit_end:
                lists = lists[0:int(limit_end)]           
        
        serializer = cSerializer.ListMonthlySalarySerializer(lists, many=True)
        print('serializer serializer',serializer.data)
        return Response({"status":error.context['success_code'], "data": serializer.data}, status=status.HTTP_200_OK)

class MonthlySalaryCRUD(APIView):
    def get_object(self, pk):
            try:
                return models.MonthlySalary.objects.get(pk = pk)
            except models.MonthlySalary.DoesNotExist:
                raise Http404

    def post(self,request, pk = None):

        #print(request.data)
        #pass
        if 'id' not in request.data:
            return Response({"status" : {"id" : ['id' +language.context[language.defaultLang]['missing']]}}, status=status.HTTP_200_OK)
        else:
            pk = request.data['id']
            if not request.data['user'] and request.data['status'] != 3:
                return Response({"status":error.context['error_code'], "message": "Employee missing"}, status=status.HTTP_200_OK)
            if not pk:

                request.data['created_ip'] = Common.get_client_ip(request)
                saveserialize = cSerializer.MonthlySalarySerializer(data = request.data)

                if saveserialize.is_valid():
                    saveserialize.save()

                    return Response({"status" : error.context['success_code'], "message":"MonthlySalary" +language.context[language.defaultLang]['insert'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'], "message":error.serializerError(saveserialize)}, status=status.HTTP_200_OK)
            else:

                request.data['modified_ip'] = Common.get_client_ip(request)
                request.data['modified_by'] = request.user.id


                list = self.get_object(pk)
                
                saveserialize = cSerializer.MonthlySalarySerializer(list, data = request.data, partial= True)                
                if saveserialize.is_valid():
                    saveserialize.save()

                    return Response({"status" :error.context['success_code'], "message":"MonthlySalary" +language.context[language.defaultLang]['update'], "data":''}, status=status.HTTP_200_OK)
                else:
                    return Response({"status" :error.context['error_code'],"message":error.serializerError(saveserialize)}, status = status.HTTP_200_OK)



